from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


DOCX = Path(r"C:\CODING\Thesis\Doc\CareerFit-Thesis-Report(5).docx")
BACKUP = Path(
    r"C:\CODING\Thesis\Doc\working\CareerFit-Thesis-Report(5)-backup-before-table-indent-fix-20260812.docx"
)


if not BACKUP.exists():
    shutil.copy2(DOCX, BACKUP)

document = Document(DOCX)
changed = 0

for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                formatting = paragraph.paragraph_format
                formatting.left_indent = Pt(0)
                formatting.right_indent = Pt(0)
                formatting.first_line_indent = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                changed += 1

document.save(DOCX)
print(f"Updated {changed} table paragraphs.")
print(f"Backup: {BACKUP}")
