from __future__ import annotations

import hashlib
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(r"C:\CODING\Thesis")
OLD = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-backup-before-Ch4-Conclusion-compression-20260812.docx"
NEW = ROOT / "Doc" / "CareerFit-Thesis-Report(5).docx"


def clean(text: str) -> str:
    return " ".join(text.replace("\u00a0", " ").split())


def word_count(text: str) -> int:
    return len(re.findall(r"[A-Za-z0-9_@'–/-]+", text))


def bounds(document: Document, start: str, end: str):
    starts = [i for i, p in enumerate(document.paragraphs) if clean(p.text) == start]
    ends = [i for i, p in enumerate(document.paragraphs) if clean(p.text) == end]
    if len(starts) != 1 or len(ends) != 1:
        raise RuntimeError((start, starts, end, ends))
    return starts[0], ends[0]


def narrative_words(document: Document, start: str, end: str) -> int:
    first, last = bounds(document, start, end)
    values = []
    for paragraph in document.paragraphs[first + 1 : last]:
        text = clean(paragraph.text)
        if not text or paragraph.style.name.startswith("Heading"):
            continue
        if paragraph.style.name in {"Table Caption", "Figure Caption"}:
            continue
        values.append(text)
    return word_count(" ".join(values))


def range_text(document: Document, start: int, end: int) -> str:
    return "\n".join(f"{p.style.name}|{clean(p.text)}" for p in document.paragraphs[start:end])


def sha(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def caption_position_errors(document: Document):
    body = list(document.element.body.iterchildren())
    paragraphs = {p._p: p for p in document.paragraphs}
    errors = []
    for index, element in enumerate(body):
        if element.tag != qn("w:p") or element not in paragraphs:
            continue
        paragraph = paragraphs[element]
        text = clean(paragraph.text)
        if paragraph.style.name == "Table Caption" and text.startswith(("Table 4.", "Table Con.")):
            if index + 1 >= len(body) or body[index + 1].tag != qn("w:tbl"):
                errors.append(f"table caption detached: {text}")
        if paragraph.style.name == "Figure Caption" and text.startswith("Figure 4."):
            if index == 0 or not body[index - 1].xpath(".//w:drawing | .//w:pict"):
                errors.append(f"figure caption detached: {text}")
    return errors


def main() -> None:
    old = Document(OLD)
    new = Document(NEW)

    old_ch4, old_conclusion = bounds(old, "CHAPTER 4. TESTING AND EVALUATION", "PART 3. CONCLUSION")[0], bounds(old, "PART 3. CONCLUSION", "REFERENCES")[0]
    new_ch4, new_conclusion = bounds(new, "CHAPTER 4. TESTING AND EVALUATION", "PART 3. CONCLUSION")[0], bounds(new, "PART 3. CONCLUSION", "REFERENCES")[0]
    old_refs = bounds(old, "PART 3. CONCLUSION", "REFERENCES")[1]
    new_refs = bounds(new, "PART 3. CONCLUSION", "REFERENCES")[1]

    print("old_ch4_words", narrative_words(old, "CHAPTER 4. TESTING AND EVALUATION", "PART 3. CONCLUSION"))
    print("new_ch4_words", narrative_words(new, "CHAPTER 4. TESTING AND EVALUATION", "PART 3. CONCLUSION"))
    print("old_conclusion_words", narrative_words(old, "PART 3. CONCLUSION", "REFERENCES"))
    print("new_conclusion_words", narrative_words(new, "PART 3. CONCLUSION", "REFERENCES"))
    print("tables", len(old.tables), "->", len(new.tables))
    print("drawings", len(old.inline_shapes), "->", len(new.inline_shapes))

    print("pre_ch4_text_unchanged", sha(range_text(old, 0, old_ch4)) == sha(range_text(new, 0, new_ch4)))
    old_intro = bounds(old, "PART 1. INTRODUCTION", "CHAPTER 4. TESTING AND EVALUATION")[0]
    new_intro = bounds(new, "PART 1. INTRODUCTION", "CHAPTER 4. TESTING AND EVALUATION")[0]
    print("introduction_chapters1_3_text_unchanged", sha(range_text(old, old_intro, old_ch4)) == sha(range_text(new, new_intro, new_ch4)))
    print("references_appendices_text_unchanged", sha(range_text(old, old_refs, len(old.paragraphs))) == sha(range_text(new, new_refs, len(new.paragraphs))))

    chapter_captions = [clean(p.text) for p in new.paragraphs if p.style.name == "Table Caption" and clean(p.text).startswith("Table 4.")]
    conclusion_captions = [clean(p.text) for p in new.paragraphs if p.style.name == "Table Caption" and clean(p.text).startswith("Table Con.")]
    figure_captions = [clean(p.text) for p in new.paragraphs if p.style.name == "Figure Caption" and clean(p.text).startswith("Figure 4.")]
    print("chapter4_table_captions", chapter_captions)
    print("conclusion_table_captions", conclusion_captions)
    print("chapter4_figure_captions", figure_captions)
    print("caption_position_errors", caption_position_errors(new))

    complete_text = "\n".join(clean(p.text) for p in new.paragraphs)
    for term in [
        "46 tests", "46 Chrome", "46-test", "46/49", "Frontend build artifacts",
        "Consolidated result status", "Principal limitations and impact",
        "Figure 4.2", "Table Con.2", "Table Con.3", "288 modified",
        "temporary proxy", "port 8080",
    ]:
        matches = [clean(p.text) for p in new.paragraphs if term.lower() in clean(p.text).lower()]
        print(f"stale[{term}]", matches)

    errors = []
    if chapter_captions != [
        "Table 4.1. Evaluation environment",
        "Table 4.2. Backend automated test results",
        "Table 4.3. Baseline and Rocchio results on the synthetic holdout scenario",
        "Table 4.4. Integrated Chrome workflow results",
        "Table 4.5. API authorization observations",
        "Table 4.6. Runtime observations",
        "Table 4.7. Answers supported by current evidence",
    ]:
        errors.append("Chapter 4 table captions are not consecutive")
    if conclusion_captions != ["Table Con.1. Objective assessment"]:
        errors.append("Conclusion table caption is not Table Con.1")
    if figure_captions != ["Figure 4.1. Baseline and Rocchio benchmark metrics at K = 5"]:
        errors.append("Chapter 4 figure captions are incorrect")
    errors.extend(caption_position_errors(new))
    print("errors", errors)
    if errors:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
