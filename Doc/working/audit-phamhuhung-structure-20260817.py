from collections import Counter
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE

PATH = r"C:\CODING\Thesis\Doc\PhamHuuHung__B2203557.docx"
doc = Document(PATH)

def cm(x):
    return None if x is None else round(x.cm, 3)

print("SECTIONS")
for i, s in enumerate(doc.sections, 1):
    print(i, {
        "orientation": "landscape" if s.orientation == WD_ORIENT.LANDSCAPE else "portrait",
        "page_w_cm": cm(s.page_width), "page_h_cm": cm(s.page_height),
        "top_cm": cm(s.top_margin), "bottom_cm": cm(s.bottom_margin),
        "left_cm": cm(s.left_margin), "right_cm": cm(s.right_margin),
        "header_cm": cm(s.header_distance), "footer_cm": cm(s.footer_distance),
        "start_type": str(s.start_type),
    })

print("\nSTYLES")
for name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Table Caption", "Figure Caption", "List Paragraph", "toc 1", "toc 2", "toc 3", "toc 4"]:
    st = doc.styles[name]
    pf = st.paragraph_format
    f = st.font
    print(name, {
        "font": f.name, "size_pt": None if f.size is None else f.size.pt,
        "bold": f.bold, "italic": f.italic,
        "alignment": str(pf.alignment),
        "line_spacing": pf.line_spacing,
        "before_pt": None if pf.space_before is None else pf.space_before.pt,
        "after_pt": None if pf.space_after is None else pf.space_after.pt,
        "first_indent_cm": cm(pf.first_line_indent),
        "left_indent_cm": cm(pf.left_indent),
        "keep_with_next": pf.keep_with_next,
        "page_break_before": pf.page_break_before,
    })

empty = []
for i,p in enumerate(doc.paragraphs):
    if p.style.name.startswith("Heading") and not p.text.strip():
        empty.append((i,p.style.name))
print("\nEMPTY_HEADINGS", empty)

table_run_sizes = Counter()
table_run_fonts = Counter()
table_para_align = Counter()
table_para_indent = Counter()
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                table_para_align[str(p.alignment)] += 1
                li = p.paragraph_format.left_indent
                fi = p.paragraph_format.first_line_indent
                table_para_indent[(cm(li),cm(fi))] += 1
                for r in p.runs:
                    if r.text:
                        table_run_sizes[None if r.font.size is None else r.font.size.pt] += 1
                        table_run_fonts[r.font.name] += 1
print("\nTABLE_RUN_SIZES", table_run_sizes)
print("TABLE_RUN_FONTS", table_run_fonts)
print("TABLE_PARA_ALIGN", table_para_align)
print("TABLE_PARA_INDENT", table_para_indent)

caption_counts = Counter()
caption_bad_align = []
for i,p in enumerate(doc.paragraphs):
    if p.style.name in ("Table Caption", "Figure Caption"):
        caption_counts[p.style.name] += 1
        if str(p.alignment) not in ("CENTER (1)", "None"):
            caption_bad_align.append((i,p.style.name,str(p.alignment),p.text[:80]))
print("\nCAPTIONS", caption_counts)
print("CAPTION_BAD_ALIGN", caption_bad_align)

print("\nDIRECT_BODY_RUN_SIZES")
c = Counter()
samples = {}
for i,p in enumerate(doc.paragraphs):
    for r in p.runs:
        if r.text and r.font.size is not None:
            key = r.font.size.pt
            c[key] += 1
            samples.setdefault(key, (i,p.style.name,p.text[:100]))
print(c)
print(samples)

print("\nTABLE_GRID_WIDTHS")
for i,t in enumerate(doc.tables, 1):
    grid = t._tbl.tblGrid
    widths = []
    if grid is not None:
        for col in grid.gridCol_lst:
            w = col.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
            if w:
                widths.append(int(w))
    total_cm = round(sum(widths) / 1440 * 2.54, 3) if widths else None
    first = " | ".join(c.text.replace("\n", " ")[:40] for c in t.rows[0].cells) if t.rows else ""
    print(i, total_cm, first)
