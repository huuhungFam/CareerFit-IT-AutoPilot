from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"C:\CODING\Thesis")
SOURCE = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
DRAFT = ROOT / "Doc" / "working" / "CH4_CONCLUSION_COMPRESSION_REVIEW_DRAFT.md"
OUTPUT = ROOT / "Doc" / "CareerFit-Thesis-Report(5).docx"


TABLES = {
    "Table 4.1": ("Table 4.1. Evaluation environment", "Table 4.1. Evaluation environment"),
    "Table 4.2": ("Table 4.2. Backend automated test results", "Table 4.2. Backend automated test results"),
    "Table 4.3": (
        "Table 4.3. Baseline and Rocchio results on the synthetic holdout scenario",
        "Table 4.3. Baseline and Rocchio results on the synthetic holdout scenario",
    ),
    "Table 4.5": ("Table 4.5. Integrated Chrome workflow results", "Table 4.4. Integrated Chrome workflow results"),
    "Table 4.6": ("Table 4.6. API authorization observations", "Table 4.5. API authorization observations"),
    "Table 4.7": ("Table 4.7. Runtime observations", "Table 4.6. Runtime observations"),
    "Table 4.8": ("Table 4.8. Answers supported by current evidence", "Table 4.7. Answers supported by current evidence"),
    "Table Con.2": ("Table Con.2. Objective assessment", "Table Con.1. Objective assessment"),
}

FIGURE_CAPTION = "Figure 4.1. Baseline and Rocchio benchmark metrics at K = 5"


def clean(text: str) -> str:
    return " ".join(text.replace("\u00a0", " ").split())


def strip_markdown(text: str) -> str:
    return text.replace("`", "").strip()


def body_blocks(document: Document):
    paragraphs = {p._p: p for p in document.paragraphs}
    tables = {t._tbl: t for t in document.tables}
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p") and child in paragraphs:
            yield child, "P", paragraphs[child]
        elif child.tag == qn("w:tbl") and child in tables:
            yield child, "T", tables[child]


def find_paragraph_element(document: Document, text: str):
    matches = [p._p for p in document.paragraphs if clean(p.text) == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph {text!r}, found {len(matches)}")
    return matches[0]


def capture_table(document: Document, caption: str):
    blocks = list(body_blocks(document))
    for index, (element, kind, item) in enumerate(blocks):
        if kind == "P" and clean(item.text) == caption:
            for next_element, next_kind, _ in blocks[index + 1 :]:
                if next_kind == "T":
                    return deepcopy(next_element)
                if next_kind == "P" and clean(document.paragraphs[0].text):
                    # Captions and tables are adjacent in the current thesis; continue
                    # past an empty paragraph only.
                    paragraph_text = "" if next_element.tag != qn("w:p") else clean("".join(next_element.itertext()))
                    if paragraph_text:
                        break
    raise RuntimeError(f"Table not found after caption: {caption}")


def capture_figure_paragraph(document: Document, caption: str):
    blocks = list(body_blocks(document))
    for index, (element, kind, item) in enumerate(blocks):
        if kind == "P" and clean(item.text) == caption:
            for previous_element, previous_kind, _ in reversed(blocks[:index]):
                if previous_kind != "P":
                    continue
                if previous_element.xpath(".//w:drawing | .//w:pict"):
                    return deepcopy(previous_element)
                if clean("".join(previous_element.itertext())):
                    break
    raise RuntimeError(f"Drawing paragraph not found before caption: {caption}")


def new_paragraph(document: Document, text: str, style: str):
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(strip_markdown(text))
    element = paragraph._p
    document.element.body.remove(element)

    if style == "Normal":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif style in {"Table Caption", "Figure Caption"}:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        if style == "Normal":
            run.font.size = Pt(13)
        elif style in {"Table Caption", "Figure Caption"}:
            run.font.size = Pt(12)
            run.font.italic = True
    return element


def create_numbering_instance(document: Document) -> int:
    numbering = document.part.numbering_part.element
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), "0")
    num.append(abstract)
    numbering.append(num)
    return num_id


def new_list_paragraph(document: Document, text: str, num_id: int):
    paragraph = document.add_paragraph(style="List Paragraph")
    paragraph.add_run(strip_markdown(text))
    paragraph.paragraph_format.line_spacing = 1.2
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    paragraph._p.get_or_add_pPr().append(num_pr)
    element = paragraph._p
    document.element.body.remove(element)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
    return element


def remove_between(body, start_element, end_element) -> None:
    children = list(body.iterchildren())
    start = children.index(start_element)
    end = children.index(end_element)
    for child in children[start + 1 : end]:
        body.remove(child)


def insert_before(target, elements) -> None:
    for element in elements:
        target.addprevious(element)


def parse_draft(document: Document, table_clones, figure_clone):
    lines = DRAFT.read_text(encoding="utf-8").splitlines()
    chapter_elements = []
    conclusion_elements = []
    current = None
    list_num_id = create_numbering_instance(document)

    for raw_line in lines:
        line = raw_line.strip()
        if line == "## CHAPTER 4. TESTING AND EVALUATION":
            current = chapter_elements
            continue
        if line == "## PART 3. CONCLUSION":
            current = conclusion_elements
            continue
        if current is None or not line or line.startswith("# Review Draft") or line.startswith("This file"):
            continue

        if line.startswith("#### "):
            current.append(new_paragraph(document, line[5:], "Heading 3"))
            continue
        if line.startswith("### "):
            current.append(new_paragraph(document, line[4:], "Heading 2"))
            continue

        if line.startswith("[Retain Table"):
            key_match = re.match(r"\[Retain (Table (?:4\.\d+|Con\.2))", line)
            if not key_match:
                raise RuntimeError(f"Unrecognized table directive: {line}")
            key = key_match.group(1)
            _, new_caption = TABLES[key]
            current.append(new_paragraph(document, new_caption, "Table Caption"))
            current.append(deepcopy(table_clones[key]))
            continue

        if line.startswith("[Retain Figure 4.1"):
            current.append(deepcopy(figure_clone))
            current.append(new_paragraph(document, FIGURE_CAPTION, "Figure Caption"))
            continue

        if line.startswith("["):
            # Approved removal or editorial instruction; it is not document text.
            continue

        number_match = re.match(r"^\d+\.\s+(.*)$", line)
        if number_match and current is conclusion_elements:
            current.append(new_list_paragraph(document, number_match.group(1), list_num_id))
            continue

        current.append(new_paragraph(document, line, "Normal"))

    return chapter_elements, conclusion_elements


def main() -> None:
    document = Document(SOURCE)
    body = document.element.body

    table_clones = {
        key: capture_table(document, old_caption)
        for key, (old_caption, _) in TABLES.items()
    }
    figure_clone = capture_figure_paragraph(document, FIGURE_CAPTION)

    chapter_heading = find_paragraph_element(document, "CHAPTER 4. TESTING AND EVALUATION")
    conclusion_heading = find_paragraph_element(document, "PART 3. CONCLUSION")
    references_heading = find_paragraph_element(document, "REFERENCES")

    chapter_elements, conclusion_elements = parse_draft(document, table_clones, figure_clone)

    remove_between(body, chapter_heading, conclusion_heading)
    remove_between(body, conclusion_heading, references_heading)
    insert_before(conclusion_heading, chapter_elements)
    insert_before(references_heading, conclusion_elements)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(f"source={SOURCE}")
    print(f"output={OUTPUT}")
    print(f"chapter_elements={len(chapter_elements)}")
    print(f"conclusion_elements={len(conclusion_elements)}")


if __name__ == "__main__":
    main()
