import sys
from docx import Document

PATH = r"C:\CODING\Thesis\Doc\PhamHuuHung__B2203557.docx"
doc = Document(PATH)

print("=== RELEVANT PARAGRAPHS ===")
ranges = [(int(sys.argv[1]), int(sys.argv[2]))] if len(sys.argv) == 3 else [(544, 627), (767, 837), (857, 916)]
for start, end in ranges:
    for i in range(start, min(end + 1, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        text = " ".join(p.text.split())
        if text:
            print(f"P{i} [{p.style.name}] {text}")

if len(sys.argv) != 3:
 print("\n=== RELEVANT TABLES ===")
for table_no in ([] if len(sys.argv) == 3 else [33, 34, 38, 39, 40, 43, 44]):
    t = doc.tables[table_no - 1]
    print(f"TABLE {table_no}")
    for row in t.rows:
        print(" || ".join(" / ".join(c.text.split()) for c in row.cells))
