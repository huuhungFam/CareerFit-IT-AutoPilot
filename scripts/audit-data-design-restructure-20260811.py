from pathlib import Path
import re

from docx import Document
from docx.oxml.ns import qn


REPORT = Path("Doc/CareerFit-Thesis-Report.docx")
document = Document(REPORT)


def clean(text):
    return " ".join(text.split())


def paragraph_has_drawing(paragraph):
    return bool(paragraph._p.xpath(".//w:drawing"))


figure_captions = [
    p for p in document.paragraphs
    if p.style.name == "Figure Caption" and clean(p.text).startswith("Figure ")
]
table_captions = [
    p for p in document.paragraphs
    if p.style.name == "Table Caption" and clean(p.text).startswith("Table ")
]

chapter_3_numbers = []
for paragraph in figure_captions:
    match = re.match(r"Figure 3\.(\d+)\.", clean(paragraph.text))
    if match:
        chapter_3_numbers.append(int(match.group(1)))

assert sorted(chapter_3_numbers) == list(range(1, 23)), chapter_3_numbers

appendix_c_captions = [
    p for p in table_captions if clean(p.text).startswith("Table App.C.")
]
assert len(appendix_c_captions) == 25, len(appendix_c_captions)

body = document.element.body
children = list(body.iterchildren())
paragraph_by_element = {p._p: p for p in document.paragraphs}
table_by_element = {t._tbl: t for t in document.tables}

dictionary_column_count = 0
dictionary_table_count = 0
for index, child in enumerate(children):
    if child.tag != qn("w:p"):
        continue
    paragraph = paragraph_by_element.get(child)
    if paragraph is None:
        continue
    text = clean(paragraph.text)

    if paragraph.style.name == "Figure Caption" and text.startswith("Figure "):
        preceding_drawing = None
        for preceding in reversed(children[:index]):
            if preceding.tag == qn("w:tbl"):
                break
            if preceding.tag == qn("w:p"):
                candidate = paragraph_by_element.get(preceding)
                if candidate is None:
                    continue
                if paragraph_has_drawing(candidate):
                    preceding_drawing = candidate
                    break
                if clean(candidate.text):
                    break
        assert preceding_drawing is not None, text

    match = re.match(r"Table App\.C\.(\d+)\.", text) if paragraph.style.name == "Table Caption" else None
    if match and int(match.group(1)) >= 2:
        following_table = None
        for following in children[index + 1:]:
            if following.tag == qn("w:p"):
                candidate = paragraph_by_element.get(following)
                if candidate is not None and clean(candidate.text):
                    break
            if following.tag == qn("w:tbl"):
                following_table = table_by_element.get(following)
                break
        assert following_table is not None, text
        dictionary_table_count += 1
        dictionary_column_count += len(following_table.rows) - 1

assert dictionary_table_count == 24, dictionary_table_count
assert dictionary_column_count == 293, dictionary_column_count

bad_table_runs = []
for table_index, table in enumerate(document.tables, start=1):
    for row_index, row in enumerate(table.rows, start=1):
        for cell_index, cell in enumerate(row.cells, start=1):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.text.strip() and run.font.size is not None and abs(run.font.size.pt - 13.0) > 0.01:
                        bad_table_runs.append((table_index, row_index, cell_index, run.font.size.pt, run.text[:30]))
assert not bad_table_runs, bad_table_runs[:10]

all_text = "\n".join(p.text for p in document.paragraphs)
for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            all_text += "\n" + cell.text
assert "\x00" not in all_text
assert "\ufffd" not in all_text
assert "Error! Bookmark not defined" not in all_text

for section in document.sections:
    assert abs(section.page_width.cm - 21.0) < 0.02
    assert abs(section.page_height.cm - 29.7) < 0.02
    assert abs(section.top_margin.cm - 3.0) < 0.02
    assert abs(section.bottom_margin.cm - 3.0) < 0.02
    assert abs(section.left_margin.cm - 3.5) < 0.02
    assert abs(section.right_margin.cm - 2.0) < 0.02

print(f"paragraphs={len(document.paragraphs)}")
print(f"tables={len(document.tables)}")
print(f"drawings={sum(1 for p in document.paragraphs if paragraph_has_drawing(p))}")
print(f"figure_captions={len(figure_captions)}")
print(f"chapter_3_figures={len(chapter_3_numbers)}")
print(f"table_captions={len(table_captions)}")
print(f"appendix_c_captions={len(appendix_c_captions)}")
print(f"dictionary_tables={dictionary_table_count}")
print(f"dictionary_columns={dictionary_column_count}")
print("a4_margins=PASS")
print("caption_placement=PASS")
print("table_font_size_13=PASS")
print("text_integrity=PASS")
