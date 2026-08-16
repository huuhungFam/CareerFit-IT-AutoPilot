from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"


def find_one(doc: Document, text: str):
    matches = [p for p in doc.paragraphs if p.text.strip() == text]
    if len(matches) != 1:
        raise ValueError(f"Expected one paragraph {text!r}, found {len(matches)}")
    return matches[0]


def set_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


doc = Document(DOCX)

theory_summary = find_one(
    doc,
    "This chapter presented the main ideas used by CareerFit. TF-IDF and cosine similarity "
    "provide an inspectable lexical baseline, while Rocchio uses explicit feedback to adjust "
    "ranking. Ranking metrics measure ordered results, and Human-in-the-Loop policies keep "
    "similarity evidence separate from recruitment actions. Chapters 1 and 2 turn these ideas "
    "into system requirements and design decisions.",
)
set_text(
    theory_summary,
    "This chapter presented the main ideas used by CareerFit. TF-IDF and cosine similarity "
    "provide an inspectable lexical baseline, while Rocchio uses explicit feedback to adjust "
    "ranking. Ranking metrics measure ordered results, and Human-in-the-Loop policies keep "
    "similarity evidence separate from recruitment actions. Chapter 1 turns these ideas into "
    "system requirements, and Chapter 3 applies them in the system design and implementation.",
)

redundant_heading = find_one(doc, "2.11 Chapter Summary")
chapter_bridge = find_one(
    doc,
    "This chapter presented the theoretical foundations used by CareerFit. It distinguished "
    "matching, recommendation, and recruitment action; described text preprocessing, TF-IDF, "
    "cosine similarity, Rocchio feedback, and ranking metrics; and discussed Human-in-the-Loop "
    "control, explainability, security, auditability, related work, and the research gap. These "
    "foundations guide the system design and implementation presented in Chapter 3.",
)
chapter_three = find_one(doc, "CHAPTER 3. SYSTEM DESIGN AND IMPLEMENTATION")

redundant_heading._p.getparent().remove(redundant_heading._p)
chapter_three._p.addnext(chapter_bridge._p)
set_text(
    chapter_bridge,
    "This chapter applies the foundations from Chapter 2 to the CareerFit solution. It first "
    "presents the architecture, modules, data model, security and consistency design, and "
    "deployment topology. It then explains how these decisions are implemented in the backend, "
    "frontend, database, automation, email-action, reporting, and monitoring workflows.",
)

doc.save(DOCX)
print(f"updated={DOCX}")
print("duplicate_chapter_2_summary_removed=true")
print("chapter_3_bridge_preserved=true")
