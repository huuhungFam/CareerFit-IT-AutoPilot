from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
ASSETS = ROOT / "Doc" / "working" / "functional-design-assets-20260811"

CAPTION_TO_IMAGE = {
    "Figure 3.12.": "sequence-explore-jobs.png",
    "Figure 3.14.": "sequence-autofit.png",
    "Figure 3.16.": "sequence-cv-upload.png",
    "Figure 3.18.": "sequence-recruiter-jobs.png",
    "Figure 3.20.": "sequence-talent-pool.png",
    "Figure 3.22.": "sequence-admin-audit.png",
}


def clear_paragraph_content(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def main():
    document = Document(DOCX)
    replaced = []

    for index, paragraph in enumerate(document.paragraphs):
        caption = paragraph.text.strip()
        matched_prefix = next(
            (prefix for prefix in CAPTION_TO_IMAGE if caption.startswith(prefix)), None
        )
        if matched_prefix is None:
            continue
        # Ignore generated List of Figures entries that repeat the same visible
        # caption text. Only an actual Caption paragraph immediately following
        # a drawing is eligible for replacement.
        if paragraph.style.name != "Figure Caption":
            continue
        if index == 0:
            raise RuntimeError(f"Caption has no preceding image paragraph: {caption}")

        image_paragraph = document.paragraphs[index - 1]
        if not image_paragraph._p.xpath(".//w:drawing | .//w:pict"):
            raise RuntimeError(f"Caption does not follow a drawing: {caption}")
        asset = ASSETS / CAPTION_TO_IMAGE[matched_prefix]
        if not asset.exists():
            raise FileNotFoundError(asset)

        clear_paragraph_content(image_paragraph)
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.paragraph_format.keep_with_next = True
        image_paragraph.add_run().add_picture(str(asset), width=Inches(6.0))

        drawing = image_paragraph._p.xpath(".//wp:docPr")
        if drawing:
            drawing[0].set("title", caption)
            drawing[0].set("descr", caption)
        replaced.append((matched_prefix, asset.name))

    missing = set(CAPTION_TO_IMAGE) - {prefix for prefix, _ in replaced}
    if missing:
        raise RuntimeError(f"Missing expected captions: {sorted(missing)}")

    document.save(DOCX)
    print(f"Replaced {len(replaced)} sequence diagrams in {DOCX}")
    for prefix, image_name in replaced:
        print(f"  {prefix} <- {image_name}")


if __name__ == "__main__":
    main()
