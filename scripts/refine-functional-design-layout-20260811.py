from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"

CAPTIONS = [
    "Table 3.3. Interface components for Explore Jobs",
    "Table 3.4. Interface components for Manage AutoFit",
    "Table 3.5. Interface components for CV upload",
    "Table 3.6. Interface components for Job and applicant management",
    "Table 3.7. Interface components for Talent Pool and invitations",
    "Table 3.8. Interface components for administrative audit activity",
]


def find_exact(document: Document, text: str):
    matches = [p for p in document.paragraphs if p.text.strip() == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph for {text!r}; found {len(matches)}")
    return matches[0]


def next_table(document: Document, caption_text: str) -> Table:
    caption = find_exact(document, caption_text)
    table_map = {table._tbl: table for table in document.tables}
    node = caption._p.getnext()
    while node is not None:
        if node.tag == qn("w:tbl"):
            return table_map[node]
        if node.tag == qn("w:p") and "Heading" in node.xml:
            break
        node = node.getnext()
    raise RuntimeError(f"No table found after {caption_text!r}")


def set_cell_margins(cell, top=55, start=80, bottom=55, end=80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table: Table, widths_cm: list[float]) -> None:
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    total_twips = round(sum(widths_cm) / 2.54 * 1440)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_twips))

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_cm in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(round(width_cm / 2.54 * 1440)))
        grid.append(grid_col)

    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for cell, width_cm in zip(row.cells, widths_cm):
            twips = round(width_cm / 2.54 * 1440)
            cell.width = Cm(width_cm)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(twips))
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(13)


def remove_note_column(table: Table) -> None:
    if len(table.columns) != 4:
        raise RuntimeError(f"Expected a 4-column component table, found {len(table.columns)} columns")
    for row in table.rows:
        last_cell = row._tr.tc_lst[-1]
        row._tr.remove(last_cell)
    grid = table._tbl.tblGrid
    if len(grid.gridCol_lst) >= 4:
        grid.remove(grid.gridCol_lst[-1])
    set_table_widths(table, [1.2, 3.4, 10.9])


def suppress_heading4_auto_numbering(document: Document) -> None:
    for paragraph in document.paragraphs:
        if paragraph.style.name != "Heading 4" or not paragraph.text.strip().startswith("3.4."):
            continue
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            p_pr.append(num_pr)
        for child in list(num_pr):
            num_pr.remove(child)
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), "0")
        num_pr.append(num_id)


def include_heading4_in_toc(document: Document) -> None:
    changed = 0
    for instruction in document.element.xpath(".//w:instrText"):
        text = instruction.text or ""
        if 'TOC \\o "1-3"' in text:
            instruction.text = text.replace('TOC \\o "1-3"', 'TOC \\o "1-4"')
            changed += 1
    if changed != 1:
        raise RuntimeError(f"Expected one main TOC field update, changed {changed}")


def request_field_refresh(document: Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main() -> None:
    document = Document(DOCX)
    for caption in CAPTIONS:
        table = next_table(document, caption)
        remove_note_column(table)
    suppress_heading4_auto_numbering(document)
    include_heading4_in_toc(document)
    request_field_refresh(document)
    document.save(DOCX)
    print(f"updated={DOCX}")
    print(f"component_tables={len(CAPTIONS)}")
    print("component_table_columns=3")
    print("heading4_auto_numbering=suppressed")
    print("toc_levels=1-4")


if __name__ == "__main__":
    main()
