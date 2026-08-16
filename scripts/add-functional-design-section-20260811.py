from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-before-20260811-functional-design.docx"
ASSETS = ROOT / "Doc" / "working" / "functional-design-assets-20260811"


FUNCTIONS = [
    {
        "role": "candidate",
        "heading": "3.4.1.1 Explore Jobs",
        "purpose": (
            "This function allows a Candidate or Guest to search the active Job catalogue, apply visible filters, "
            "open a selected Job, and review the information needed before deciding whether to apply. Candidate-only "
            "guidance may also be shown when relevant Matching context is available."
        ),
        "image": "interface-explore-jobs-annotated.png",
        "interface_caption": "Figure 3.11. Annotated interface for Explore Jobs",
        "table_caption": "Table 3.3. Interface components for Explore Jobs",
        "components": [
            ("1", "Navigation", "Open Candidate pages such as Jobs, Recommendations, Applications, and AutoFit.", "Candidate session"),
            ("2", "Select", "Choose the Job location scope.", "Optional filter"),
            ("3", "Input", "Enter a Job title, skill, or company keyword.", "Optional search"),
            ("4", "Button", "Submit the current search criteria.", ""),
            ("5", "Filter / Sort", "Refine results by urgent state, level, work mode, salary, domain, and ordering.", "Combinable"),
            ("6", "Information panel", "Show match-quality guidance or a limited-data state for the current result set.", "Candidate context"),
        ],
        "sequence": "sequence-explore-jobs.png",
        "sequence_caption": "Figure 3.12. Sequence diagram for Explore Jobs",
    },
    {
        "role": "candidate",
        "heading": "3.4.1.2 Manage AutoFit",
        "purpose": (
            "This function allows a Candidate to review and save AutoFit configuration and, when requested, start a "
            "manual AutoFit run. The saved configuration is also available to the system-configured automatic execution."
        ),
        "image": "interface-autofit-annotated.png",
        "interface_caption": "Figure 3.13. Annotated interface for Manage AutoFit",
        "table_caption": "Table 3.4. Interface components for Manage AutoFit",
        "components": [
            ("1", "Navigation", "Open the Candidate workspace and AutoFit page.", "Candidate only"),
            ("2", "Summary section", "Present the most important current AutoFit states.", "Read-only overview"),
            ("3", "Status card", "Show whether automatic application is enabled and its threshold.", ""),
            ("4", "Status card", "Show the next system-configured AutoFit scan.", ""),
            ("5", "Status card", "Show the high-match email state and threshold.", "Notification setting"),
            ("6", "Status card", "Show the configured quiet-hours window.", "Notification setting"),
            ("7", "Configuration section", "Change automatic-application consent and the minimum accepted Matching score.", "Saved to policy"),
        ],
        "sequence": "sequence-autofit.png",
        "sequence_caption": "Figure 3.14. Sequence diagram for Manage AutoFit",
    },
    {
        "role": "candidate",
        "heading": "3.4.1.3 Upload and Confirm a CV",
        "purpose": (
            "This function allows a Candidate to upload a supported CV document or use manual entry, review the "
            "extracted content, correct it when needed, and confirm the CV before scoring and Matching continue."
        ),
        "image": "interface-cv-upload-annotated.png",
        "interface_caption": "Figure 3.15. Annotated interface for CV upload",
        "table_caption": "Table 3.5. Interface components for CV upload",
        "components": [
            ("1", "Navigation", "Open Candidate pages and the CV upload entry point.", "Candidate only"),
            ("2", "Tab", "Choose document analysis or manual CV entry.", "Two input paths"),
            ("3", "Upload area", "Select or drop a supported CV file for processing.", "File rules are validated"),
            ("4", "Status panel", "Show file acceptance, processing, review, or failure information.", "Observable processing state"),
        ],
        "sequence": "sequence-cv-upload.png",
        "sequence_caption": "Figure 3.16. Sequence diagram for CV upload, review, and confirmation",
    },
    {
        "role": "recruiter",
        "heading": "3.4.2.1 Manage Job Postings and Applicants",
        "purpose": (
            "This function supports the Recruiter's main Job workspace. It provides Job search and status filtering, "
            "draft or publication actions, selected-Job details, and access to the Applications submitted for an owned Job."
        ),
        "image": "interface-recruiter-jobs-annotated.png",
        "interface_caption": "Figure 3.17. Annotated interface for Job and applicant management",
        "table_caption": "Table 3.6. Interface components for Job and applicant management",
        "components": [
            ("1", "Navigation", "Open Recruiter pages such as Jobs, Talent Pool, and analytics.", "Recruiter only"),
            ("2", "Button", "Export the Recruiter's Job list as CSV.", "Supporting action"),
            ("3", "Input", "Search within the Recruiter's Jobs.", ""),
            ("4", "Select", "Filter Jobs by lifecycle status.", ""),
            ("5", "Button", "Open the form for a new Job posting.", "Draft or publication flow"),
            ("6", "Job list", "Display owned Jobs and allow one Job to be selected.", ""),
            ("7", "Action buttons", "Edit or delete the selected Job when the current state permits it.", "Ownership required"),
            ("8", "Detail panel", "Show the selected Job and provide access to its applicant workflow.", ""),
        ],
        "sequence": "sequence-recruiter-jobs.png",
        "sequence_caption": "Figure 3.18. Sequence diagram for Job publication and applicant review",
    },
    {
        "role": "recruiter",
        "heading": "3.4.2.2 Manage Talent Pool and Invitations",
        "purpose": (
            "This function allows a Recruiter to inspect Matching and Potential CV groups for an owned Job, bookmark "
            "Candidates for later review, inspect visible CV details, and create or track recruitment invitations."
        ),
        "image": "interface-talent-pool-annotated.png",
        "interface_caption": "Figure 3.19. Annotated interface for Talent Pool and invitations",
        "table_caption": "Table 3.7. Interface components for Talent Pool and invitations",
        "components": [
            ("1", "Navigation", "Open the Recruiter Talent Pool.", "Recruiter only"),
            ("2", "Job list", "Select an owned Job whose Candidate results will be reviewed.", "Ownership required"),
            ("3", "Tabs", "Switch among all visible CVs, bookmarked CVs, and invited CVs.", ""),
            ("4", "Result groups", "Separate high Matching and high-Potential CV results.", "Decision support"),
            ("5", "Candidate summary", "Show the Candidate identity summary and relevant score context.", "Visibility rules apply"),
            ("6", "Button", "Save or remove the Candidate bookmark for the selected Job.", "No Candidate notification"),
            ("7", "State badge", "Show whether an Application or invitation already exists.", "Prevents duplicate action"),
            ("8", "Button", "Open the visible CV detail.", "Visibility rules apply"),
        ],
        "sequence": "sequence-talent-pool.png",
        "sequence_caption": "Figure 3.20. Sequence diagram for Talent Pool bookmarking and invitation",
    },
    {
        "role": "administrator",
        "heading": "3.4.3.1 Review Administrative Audit Activity",
        "purpose": (
            "This supporting function allows an Administrator to review recorded administrative and business events. "
            "It helps verify who performed an action, which target was affected, and whether the recorded result succeeded."
        ),
        "image": "interface-admin-audit-annotated.png",
        "interface_caption": "Figure 3.21. Annotated interface for administrative audit activity",
        "table_caption": "Table 3.8. Interface components for administrative audit activity",
        "components": [
            ("1", "Navigation", "Open the administrative overview.", "Administrator only"),
            ("2", "Navigation", "Open user access administration.", "Supporting admin function"),
            ("3", "Navigation", "Open Job visibility administration.", "Supporting admin function"),
            ("4", "Navigation", "Open the audit-log view.", "Selected view"),
            ("5", "Navigation", "Open email-action monitoring.", "Supporting admin function"),
            ("6", "Audit table", "Review event time, actor, action, target, and result.", "Paginated records"),
        ],
        "sequence": "sequence-admin-audit.png",
        "sequence_caption": "Figure 3.22. Sequence diagram for reviewing administrative audit activity",
    },
]


def find_exact(document: Document, text: str) -> Paragraph:
    matches = [p for p in document.paragraphs if p.text.strip() == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph for {text!r}; found {len(matches)}")
    return matches[0]


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def configure_font(run, *, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_body(paragraph: Paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.2
    for run in paragraph.runs:
        configure_font(run)


def configure_heading(paragraph: Paragraph, *, page_break: bool = False) -> None:
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.page_break_before = page_break


def configure_caption(paragraph: Paragraph, style: str) -> None:
    paragraph.style = style
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = style == "Table Caption"
    for run in paragraph.runs:
        configure_font(run, italic=True)


def insert_label(anchor: Paragraph, label: str, *, page_break: bool = False) -> Paragraph:
    paragraph = anchor.insert_paragraph_before()
    paragraph.style = "List Paragraph"
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = page_break
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(label)
    configure_font(run, bold=True)
    return paragraph


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str = "E7E7E7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_widths(table, widths_cm: list[float]) -> None:
    table.autofit = False
    table.allow_autofit = False
    total_twips = round(sum(widths_cm) / 2.54 * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "0")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_cm in widths_cm:
        twips = round(width_cm / 2.54 * 1440)
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(twips))
        grid.append(col)

    for row in table.rows:
        for cell, width_cm in zip(row.cells, widths_cm):
            twips = round(width_cm / 2.54 * 1440)
            cell.width = Cm(width_cm)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(twips))
            set_cell_margins(cell)


def insert_component_table(anchor: Paragraph, caption_text: str, rows: list[tuple[str, str, str, str]]) -> None:
    caption = anchor.insert_paragraph_before(caption_text)
    configure_caption(caption, "Table Caption")

    table = anchor._parent.add_table(rows=len(rows) + 1, cols=4, width=Cm(15.5))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    anchor._p.addprevious(table._tbl)
    set_table_widths(table, [1.2, 3.0, 8.2, 3.1])
    header = ("No.", "Control Type", "Description", "Note")
    values = [header, *rows]
    for row_index, row_values in enumerate(values):
        for cell, value in zip(table.rows[row_index].cells, row_values):
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(value)
            configure_font(run, bold=row_index == 0)
            if row_index == 0:
                shade_cell(cell)
    repeat_header(table.rows[0])


def insert_image(anchor: Paragraph, image_path: Path, caption_text: str, *, page_break: bool = False) -> None:
    paragraph = anchor.insert_paragraph_before()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.page_break_before = page_break
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    shape = paragraph.add_run().add_picture(str(image_path), width=Inches(6.0))
    shape._inline.docPr.set("title", caption_text)
    shape._inline.docPr.set("descr", caption_text)
    caption = anchor.insert_paragraph_before(caption_text)
    configure_caption(caption, "Figure Caption")


def remove_old_screenshot(document: Document, caption_text: str) -> None:
    caption = find_exact(document, caption_text)
    image_node = caption._p.getprevious()
    if image_node is None or not image_node.xpath(".//w:drawing | .//w:pict"):
        raise RuntimeError(f"Expected an image before {caption_text!r}")
    image_node.getparent().remove(image_node)
    caption._p.getparent().remove(caption._p)


def replace_caption_numbers_after_anchor(document: Document, anchor: Paragraph) -> None:
    node = anchor._p
    chapter4 = find_exact(document, "CHAPTER 4. TESTING AND EVALUATION")._p
    while node is not None and node is not chapter4:
        if node.tag == qn("w:p"):
            paragraph = Paragraph(node, anchor._parent)
            text = paragraph.text.strip()
            if paragraph.style.name in {"Figure Caption", "Table Caption"}:
                figure_match = re.match(r"Figure 3\.(\d+)(\..*)", text)
                table_match = re.match(r"Table 3\.(\d+)(\..*)", text)
                if figure_match and int(figure_match.group(1)) >= 11:
                    set_paragraph_text(paragraph, f"Figure 3.{int(figure_match.group(1)) + 12}{figure_match.group(2)}")
                elif table_match and int(table_match.group(1)) >= 3:
                    set_paragraph_text(paragraph, f"Table 3.{int(table_match.group(1)) + 6}{table_match.group(2)}")
        node = node.getnext()


def renumber_headings_after_anchor(document: Document, anchor: Paragraph) -> None:
    node = anchor._p
    chapter4 = find_exact(document, "CHAPTER 4. TESTING AND EVALUATION")._p
    while node is not None and node is not chapter4:
        if node.tag == qn("w:p"):
            paragraph = Paragraph(node, anchor._parent)
            if paragraph.style.name in {"Heading 2", "Heading 3", "Heading 4"}:
                match = re.match(r"3\.(\d+)(.*)", paragraph.text.strip())
                if match and int(match.group(1)) >= 4:
                    set_paragraph_text(paragraph, f"3.{int(match.group(1)) + 1}{match.group(2)}")
        node = node.getnext()


def request_field_refresh(document: Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def update_chapter_prose(document: Document) -> None:
    replacements = {
        (
            "This chapter applies the foundations from Chapter 2 to the CareerFit solution. It first presents the architecture, modules, data model, security and consistency design, and deployment topology. It then explains how these decisions are implemented in the backend, frontend, database, automation, email-action, reporting, and monitoring workflows."
        ): (
            "This chapter applies the foundations from Chapter 2 to the CareerFit solution. It first presents the architecture, modules, data model, functional design, security and consistency design, and deployment topology. It then explains how these decisions are implemented in the backend, frontend, database, automation, email-action, reporting, and monitoring workflows."
        ),
        (
            "App.tsx defines public, Candidate, Recruiter, and Administrator routes. Protected routes check the loaded account role, while session restoration calls /api/auth/me and clears invalid state. Candidate pages include Job reporting. Recruiter applicant and Talent Pool views can report a visible CV with the owned Job ID. The Administrator workspace includes a Report moderation tab for the pending queue, case detail, ban, and dismiss actions."
        ): (
            "Section 3.4 presents selected role-based interfaces and their actor-visible sequences. At the integration level, App.tsx defines public, Candidate, Recruiter, and Administrator routes. Protected routes check the loaded account role, while session restoration calls /api/auth/me and clears invalid state. Candidate Job reporting, Recruiter CV reporting, and Administrator Report moderation remain protected by role, ownership, and target-visibility checks."
        ),
        (
            "This chapter presented the CareerFit architecture, modules, data model, security and consistency design, deployment topology, and implementation in Spring Boot, React, PostgreSQL, and Flyway. It covered JWT security, CV review, direct matching, the Potential assessment, feedback learning, applications, Talent Pool, AutoFit, email actions, Job/CV reporting, administrator moderation, frontend integration, and monitoring. Chapter 4 evaluates the refreshed system."
        ): (
            "This chapter presented the CareerFit architecture, modules, data model, functional design, security and consistency design, deployment topology, and implementation in Spring Boot, React, PostgreSQL, and Flyway. It covered JWT security, CV review, direct matching, the Potential assessment, feedback learning, applications, Talent Pool, AutoFit, email actions, Job/CV reporting, administrator moderation, frontend integration, and monitoring. Chapter 4 evaluates the refreshed system."
        ),
    }
    by_text = {p.text.strip(): p for p in document.paragraphs}
    missing = [text for text in replacements if text not in by_text]
    if missing:
        raise RuntimeError(f"Missing synchronized Chapter 3 prose: {missing}")
    for old, new in replacements.items():
        set_paragraph_text(by_text[old], new)


def add_function_section(anchor: Paragraph, item: dict, *, page_break: bool) -> None:
    heading = anchor.insert_paragraph_before(item["heading"])
    heading.style = "Heading 4"
    configure_heading(heading, page_break=page_break)

    purpose = anchor.insert_paragraph_before()
    purpose.style = "List Paragraph"
    label = purpose.add_run("Purpose: ")
    configure_font(label, bold=True)
    content = purpose.add_run(item["purpose"])
    configure_font(content)
    purpose.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    purpose.paragraph_format.line_spacing = 1.2

    insert_label(anchor, "Interface:")
    insert_image(anchor, ASSETS / item["image"], item["interface_caption"])
    insert_label(anchor, "Interface Components:")
    insert_component_table(anchor, item["table_caption"], item["components"])
    insert_label(anchor, "Processing Logic:", page_break=True)
    insert_image(anchor, ASSETS / item["sequence"], item["sequence_caption"])


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)
    missing_assets = [item[name] for item in FUNCTIONS for name in ("image", "sequence") if not (ASSETS / item[name]).exists()]
    if missing_assets:
        raise FileNotFoundError(f"Missing functional-design assets: {missing_assets}")

    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)

    document = Document(DOCX)
    if any(p.text.strip() == "3.4 Functional Design" for p in document.paragraphs):
        raise RuntimeError("Functional Design section already exists; refusing to duplicate it")

    for caption in (
        "Screen 3.1. Candidate urgent-job catalogue",
        "Screen 3.2. Candidate AutoFit policy settings",
        "Screen 3.3. Candidate CV upload entry point",
        "Screen 3.4. Recruiter Jobs and applicant workspace",
        "Screen 3.5. Recruiter Talent Pool and Potential CV",
        "Screen 3.6. Administrator audit logs",
    ):
        remove_old_screenshot(document, caption)

    old_anchor = find_exact(document, "3.4 Security, Failure, and Consistency Design")
    replace_caption_numbers_after_anchor(document, old_anchor)
    renumber_headings_after_anchor(document, old_anchor)
    anchor = find_exact(document, "3.5 Security, Failure, and Consistency Design")

    heading = anchor.insert_paragraph_before("3.4 Functional Design")
    heading.style = "Heading 2"
    configure_heading(heading, page_break=True)
    intro = anchor.insert_paragraph_before(
        "This section follows the role-oriented Functional Design structure used by the thesis template. Each selected function is presented through its purpose, an annotated implementation screenshot, a compact interface-component table, and a sequence diagram. Data tables are not repeated here because Section 3.3 and Appendix C already provide the complete data models and data dictionary."
    )
    configure_body(intro)

    candidate = anchor.insert_paragraph_before("3.4.1 Functional Design for Candidate")
    candidate.style = "Heading 3"
    configure_heading(candidate)
    first_for_role = {"candidate": True, "recruiter": True, "administrator": True}
    current_role = "candidate"
    for item in FUNCTIONS:
        if item["role"] != current_role:
            current_role = item["role"]
            role_number = {"recruiter": "3.4.2", "administrator": "3.4.3"}[current_role]
            role_name = {"recruiter": "Recruiter", "administrator": "Administrator"}[current_role]
            role_heading = anchor.insert_paragraph_before(f"{role_number} Functional Design for {role_name}")
            role_heading.style = "Heading 3"
            configure_heading(role_heading, page_break=True)
        add_function_section(anchor, item, page_break=not first_for_role[item["role"]])
        first_for_role[item["role"]] = False

    update_chapter_prose(document)
    request_field_refresh(document)
    document.save(DOCX)

    print(f"updated={DOCX}")
    print(f"backup={BACKUP}")
    print(f"functions={len(FUNCTIONS)}")
    print(f"new_figures={len(FUNCTIONS) * 2}")
    print(f"new_tables={len(FUNCTIONS)}")


if __name__ == "__main__":
    main()
