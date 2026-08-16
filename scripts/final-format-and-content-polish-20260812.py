from __future__ import annotations

import copy
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-before-final-format-20260812.docx"
TEMP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-final-format-20260812.docx"


def set_rfonts(element, name: str = "Times New Roman") -> None:
    rpr = element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def format_run(run, *, size: float, bold=None, italic=None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    set_rfonts(run._element)


def set_style_font(style, *, size: float, bold=None, italic=None) -> None:
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic
    set_rfonts(style.element)


def replace_in_runs(paragraph, old: str, new: str) -> bool:
    changed = False
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            changed = True
    return changed


def set_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        flag = OxmlElement("w:tblHeader")
        flag.set(qn("w:val"), "true")
        tr_pr.append(flag)


def insert_table_after(doc, paragraph, rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    paragraph._p.addnext(table._tbl)
    for r_idx, values in enumerate(rows):
        row = table.rows[r_idx]
        set_cant_split(row)
        if r_idx == 0:
            set_repeat_header(row)
        for c_idx, value in enumerate(values):
            cell = row.cells[c_idx]
            cell.text = value
            if widths:
                cell.width = Inches(widths[c_idx])
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 1.2
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    format_run(run, size=13, bold=(r_idx == 0), italic=False)
    return table


def insert_paragraph_after_table(table, text: str):
    p = OxmlElement("w:p")
    table._tbl.addnext(p)
    paragraph = table._parent.add_paragraph()
    paragraph._p.getparent().remove(paragraph._p)
    p.addnext(paragraph._p)
    paragraph.add_run(text)
    return paragraph


def find_paragraph(doc, exact: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == exact:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact}")


def clear_and_write(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def main() -> None:
    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)

    doc = Document(DOCX)

    # Normalize the hierarchy to the latest faculty template.
    normal = doc.styles["Normal"]
    set_style_font(normal, size=13)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.2

    heading_rules = {
        "Heading 1": (14, True, False, WD_ALIGN_PARAGRAPH.CENTER),
        "Heading 2": (14, True, False, WD_ALIGN_PARAGRAPH.LEFT),
        "Heading 3": (13, True, False, WD_ALIGN_PARAGRAPH.LEFT),
        "Heading 4": (13, True, True, WD_ALIGN_PARAGRAPH.LEFT),
    }
    for name, (size, bold, italic, alignment) in heading_rules.items():
        style = doc.styles[name]
        set_style_font(style, size=size, bold=bold, italic=italic)
        style.paragraph_format.alignment = alignment

    for name in ("toc 1", "toc 2", "toc 3"):
        style = doc.styles[name]
        set_style_font(style, size=13, bold=False, italic=False)
        style.paragraph_format.line_spacing = 1.0
    set_style_font(doc.styles["toc 4"], size=11, bold=False, italic=False)
    doc.styles["toc 4"].paragraph_format.line_spacing = 1.0

    for name in ("Caption", "Table Caption", "Figure Caption"):
        style = doc.styles[name]
        set_style_font(style, size=12, bold=False, italic=True)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.line_spacing = 1.0

    # Remove direct formatting that contradicted the normalized styles.
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name
        if style_name in heading_rules:
            size, bold, italic, alignment = heading_rules[style_name]
            paragraph.alignment = alignment
            for run in paragraph.runs:
                format_run(run, size=size, bold=bold, italic=italic)
        elif style_name.lower() in {"toc 1", "toc 2", "toc 3"}:
            for run in paragraph.runs:
                format_run(run, size=13, bold=False, italic=False)
        elif style_name.lower() == "toc 4":
            for run in paragraph.runs:
                format_run(run, size=11, bold=False, italic=False)
        elif style_name in {"Caption", "Table Caption", "Figure Caption"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                format_run(run, size=12, bold=False, italic=True)

    # Use existing core references more explicitly instead of padding the bibliography.
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Prior work can be grouped into lexical retrieval"):
            clear_and_write(
                paragraph,
                "Prior work can be grouped into lexical retrieval, learned resume–job representation, "
                "explainable recommendation, and human-centered governance. The classical vector-space "
                "and relevance-feedback literature provides transparent mathematical foundations [1], [4]. "
                "Contemporary resume–job matching methods such as ConFit v2 use trained dense representations "
                "and hard-negative strategies to improve semantic matching [7]. Real-world observational "
                "research comparing human and LLM resume ratings found only minor correlation between the two, "
                "indicating that automated and human judgments should not be treated as interchangeable [8]. "
                "Explainable recommendation research seeks to accompany scores with understandable, preferably "
                "grounded reasons [14]. Risk-management and algorithmic-hiring studies further support explicit "
                "human oversight, documented controls, and careful treatment of bias claims [9], [10].",
            )
            for run in paragraph.runs:
                format_run(run, size=13)
            break

    # Disambiguate Conclusion table numbering from Appendix C.
    for paragraph in doc.paragraphs:
        for n in (1, 2, 3):
            replace_in_runs(paragraph, f"Table C.{n}.", f"Table Con.{n}.")

    # Normalize quotation marks in the five web references.
    for paragraph in doc.paragraphs:
        if re.match(r"^\[(?:1[5-9])\]", paragraph.text.strip()):
            text = paragraph.text
            first = text.find('"')
            last = text.rfind('"')
            if first >= 0 and last > first:
                text = text[:first] + "“" + text[first + 1:last] + "”" + text[last + 1:]
                clear_and_write(paragraph, text)

    # References: left aligned with an IEEE-style hanging indent.
    refs = find_paragraph(doc, "REFERENCES")
    appendices = find_paragraph(doc, "APPENDICES")
    in_references = False
    for paragraph in doc.paragraphs:
        if paragraph._p is refs._p:
            in_references = True
            continue
        if paragraph._p is appendices._p:
            in_references = False
        if in_references and paragraph.text.strip():
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.5)
            paragraph.paragraph_format.line_spacing = 1.2
            for run in paragraph.runs:
                format_run(run, size=13)

    # Replace the dense Appendix A paragraph with a readable traceability table.
    appendix_a_rows = [
        ["Use Case", "Business goal", "Verified workflow or test scope"],
        ["UC-01", "Manage Career Profile", "Candidate Profile, Portfolio, CV review and confirmation, default CV, and deletion"],
        ["UC-02", "Explore Jobs", "Public Job catalogue, search, detail, employer information, urgent Jobs, and similar Jobs"],
        ["UC-03", "Manage Job Applications", "Application submission, history, withdrawal, and invitation response"],
        ["UC-04", "Provide Matching Feedback", "Feedback submission and learned-ranking behavior"],
        ["UC-05", "Review Personalized Career Insights", "CV–Job Matching, profile-based recommendation, and Candidate analytics"],
        ["UC-06", "Manage AutoFit", "AutomationPolicy configuration, manual execution, and AutoApply eligibility"],
        ["UC-07", "Respond Through Actionable Email", "Confirmation, feedback or invitation redemption, replay rejection, and expiry handling"],
        ["UC-08", "Manage Employer Profile and Job Postings", "Employer Profile and Job lifecycle operations"],
        ["UC-09", "Review and Process Applicants", "Applicant review and Application decision changes"],
        ["UC-10", "Manage Talent Pool and Invitations", "Bookmark, Talent Pool, invitation, and invitation withdrawal"],
        ["UC-11", "Review Recruitment Analytics", "Recruiter analytics and supported recruitment summaries"],
        ["UC-12", "Report Suspicious Recruitment Content", "Candidate Job-reporting and Recruiter visible-CV-reporting paths"],
        ["UC-13", "Administer Platform Access and Job Visibility", "Administrative User access and Job visibility controls"],
        ["UC-14", "Review and Resolve Content Reports", "Report review, dismissal, and supported moderation actions"],
    ]
    old_a = next(p for p in doc.paragraphs if p.text.startswith("The core traceability chain is:"))
    clear_and_write(old_a, "Table App.A.1. Requirement–test traceability for the approved Use Cases")
    old_a.style = doc.styles["Table Caption"]
    table_a = insert_table_after(doc, old_a, appendix_a_rows, [0.75, 2.35, 3.65])

    # Replace the dense Appendix B endpoint paragraph with a compact contract table.
    appendix_b_rows = [
        ["Method", "Endpoint", "Business purpose"],
        ["POST", "/api/auth/login", "Authenticate a User and issue the supported signed JWT response"],
        ["GET", "/api/jobs/search", "Explore the public Job catalogue using supported search criteria"],
        ["POST", "/api/cv/upload", "Upload a Candidate CV for validation and review"],
        ["GET", "/api/cv/{cvId}/review", "View the extracted CV review draft"],
        ["PATCH", "/api/cv/{cvId}/review", "Confirm or correct the supported CV review data"],
        ["POST", "/api/matches/{matchingId}/feedback", "Submit Candidate feedback for a Matching result"],
        ["POST", "/api/applications", "Submit a Job Application"],
        ["POST", "/api/reports/jobs/{jobId}", "Report suspicious Job content"],
        ["POST", "/api/reports/cvs/{cvId}", "Report a visible Candidate CV"],
        ["GET", "/api/admin/reports", "Review submitted content reports"],
        ["POST", "/api/admin/reports/{type}/{targetId}/ban", "Apply the supported moderation action to reported content"],
        ["POST", "/api/admin/reports/{type}/{targetId}/dismiss", "Dismiss a reviewed content report"],
        ["PATCH", "/api/automation/policy", "Save Candidate AutoFit configuration"],
        ["GET", "/actuator/health", "Read the configured aggregate runtime health status"],
    ]
    old_b = next(p for p in doc.paragraphs if p.text.startswith("Representative endpoints are POST"))
    clear_and_write(old_b, "Table App.B.1. Selected CareerFit API contracts")
    old_b.style = doc.styles["Table Caption"]
    table_b = insert_table_after(doc, old_b, appendix_b_rows, [0.75, 2.65, 3.35])
    note_b = insert_paragraph_after_table(
        table_b,
        "Except for intentionally public contracts, protected routes require a signed JWT together with "
        "backend role and ownership checks. The table describes actor-visible contracts rather than controller, "
        "repository, transaction, or persistence details.",
    )
    note_b.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note_b.paragraph_format.line_spacing = 1.2
    for run in note_b.runs:
        format_run(run, size=13)

    # All substantive table text uses 13 pt and avoids full justification.
    header_terms = {
        "no.", "use case", "business goal", "method", "endpoint", "module", "field",
        "metric", "risk", "objective", "test group", "layer", "component", "column",
        "postgresql type", "null", "key", "default / constraint", "description",
    }
    for table_index, table in enumerate(doc.tables, start=1):
        if table_index <= 2:  # cover-page layout tables
            continue
        for row_index, row in enumerate(table.rows):
            set_cant_split(row)
            first_row_text = {cell.text.strip().lower() for cell in row.cells}
            header = row_index == 0 and bool(first_row_text & header_terms)
            if header:
                set_repeat_header(row)
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.line_spacing = 1.2
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        format_run(run, size=13, bold=(True if header else None))

    # Long machine-readable tokens should never be stretched by full justification.
    technical_pattern = re.compile(r"(?:https?://|/api/|/actuator/|[A-Fa-f0-9]{32,}|`[^`]+`)")
    for paragraph in doc.paragraphs:
        if technical_pattern.search(paragraph.text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    TEMP.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TEMP)
    shutil.copy2(TEMP, DOCX)
    print(f"Updated: {DOCX}")
    print(f"Backup:  {BACKUP}")
    print(f"Tables:  {len(doc.tables)}")


if __name__ == "__main__":
    main()
