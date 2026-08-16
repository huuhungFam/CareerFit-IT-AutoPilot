from __future__ import annotations

import hashlib
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-before-20260811-standard-flowcharts.docx"
FIGURES = ROOT / "Doc" / "figures"


def table_digest(document: Document) -> str:
    payload = []
    for table in document.tables:
        payload.append([[cell.text for cell in row.cells] for row in table.rows])
    return hashlib.sha256(repr(payload).encode("utf-8")).hexdigest()


def find_exact(document: Document, text: str):
    matches = [p for p in document.paragraphs if p.text.strip() == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph for {text!r}; found {len(matches)}")
    return matches[0]


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def previous_paragraph(paragraph):
    previous = paragraph._p.getprevious()
    while previous is not None and not previous.tag.endswith("}p"):
        previous = previous.getprevious()
    if previous is None:
        raise RuntimeError(f"No previous paragraph before {paragraph.text!r}")
    from docx.text.paragraph import Paragraph

    return Paragraph(previous, paragraph._parent)


def configure_image_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.page_break_before = True
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0


def configure_caption(paragraph) -> None:
    paragraph.style = "Figure Caption"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = False


def add_picture(paragraph, image_path: Path, alt_text: str) -> None:
    run = paragraph.add_run()
    inline = run.add_picture(str(image_path), width=Inches(5.2))
    inline._inline.docPr.set("title", alt_text)
    inline._inline.docPr.set("descr", alt_text)


def replace_figure(document: Document, old_caption: str, new_caption: str, image_name: str) -> None:
    caption = find_exact(document, old_caption)
    old_image = previous_paragraph(caption)
    if not old_image._p.xpath(".//w:drawing"):
        raise RuntimeError(f"Expected an image immediately before {old_caption!r}")
    remove_paragraph(old_image)

    image_paragraph = caption.insert_paragraph_before()
    configure_image_paragraph(image_paragraph)
    add_picture(image_paragraph, FIGURES / image_name, new_caption)
    caption.text = new_caption
    configure_caption(caption)


def remove_chapter_one_figure(document: Document, caption_text: str) -> None:
    caption = find_exact(document, caption_text)
    image = previous_paragraph(caption)
    if not image._p.xpath(".//w:drawing"):
        raise RuntimeError(f"Expected an image immediately before {caption_text!r}")
    remove_paragraph(image)
    remove_paragraph(caption)


def insert_autofit_figure(document: Document) -> None:
    heading = find_exact(document, "3.13.3 Notification Guard and Delivery")
    image_paragraph = heading.insert_paragraph_before()
    configure_image_paragraph(image_paragraph)
    caption_text = "Figure 3.11. AutoFit eligibility and application decision flowchart"
    add_picture(image_paragraph, FIGURES / "flowchart-autofit-20260811.png", caption_text)

    caption = heading.insert_paragraph_before(caption_text, style="Figure Caption")
    configure_caption(caption)


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)
    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(DOCX, BACKUP)

    document = Document(DOCX)
    before_tables = table_digest(document)

    for caption in (
        "Figure 1.4. CV review, confirmation, and matching sequence",
        "Figure 1.5. Feedback learning and recomputation sequence",
        "Figure 1.6. AutoFit decision flow",
    ):
        remove_chapter_one_figure(document, caption)

    chapter_one_renames = {
        "Figure 1.7. Recruiter use-case diagram": "Figure 1.4. Recruiter use-case diagram",
        "Figure 1.8. Shared Candidate\u2013Recruiter reporting use-case diagram": "Figure 1.5. Shared Candidate-Recruiter reporting use-case diagram",
        "Figure 1.9. Administrator use-case diagram": "Figure 1.6. Administrator use-case diagram",
    }
    for old, new in chapter_one_renames.items():
        paragraph = find_exact(document, old)
        paragraph.text = new
        configure_caption(paragraph)

    replace_figure(
        document,
        "Figure 3.6. CV ingestion, review, and confirmation pipeline",
        "Figure 3.6. CV ingestion, review, confirmation, and matching flowchart",
        "flowchart-cv-processing-20260811.png",
    )
    replace_figure(
        document,
        "Figure 3.8. Direct score and Potential assessment flow",
        "Figure 3.8. CV-Job matching and Potential assessment flowchart",
        "flowchart-matching-potential-20260811.png",
    )
    replace_figure(
        document,
        "Figure 3.9. Feedback processing and post-commit learning",
        "Figure 3.9. Rocchio feedback-learning and recomputation flowchart",
        "flowchart-rocchio-feedback-20260811.png",
    )

    insert_autofit_figure(document)

    chapter_three_renames = {
        "Figure 3.11. Per-account notification policy guard": "Figure 3.12. Per-account notification policy guard",
        "Figure 3.12. Hashed email-action token and confirm-then-POST flow": "Figure 3.13. Hashed email-action token and confirm-then-POST flow",
        "Figure 3.13. Frontend routes, server-side catalogue, and API data flow": "Figure 3.14. Frontend routes, server-side catalogue, and API data flow",
    }
    for old, new in chapter_three_renames.items():
        paragraph = find_exact(document, old)
        paragraph.text = new
        configure_caption(paragraph)

    after_tables = table_digest(document)
    if before_tables != after_tables:
        raise RuntimeError("Table contents changed during the figure-only edit")

    document.save(DOCX)
    print(f"Updated: {DOCX}")
    print(f"Backup:  {BACKUP}")
    print(f"Tables unchanged: {before_tables}")


if __name__ == "__main__":
    main()
