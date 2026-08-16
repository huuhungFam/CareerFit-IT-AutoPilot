from collections import Counter
from pathlib import Path
import re

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"


def words(text):
    return re.findall(r"[A-Za-z][A-Za-z'–-]*", text.lower())


doc = Document(DOCX)
body = []
chapter = "Front matter"
for index, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()
    if text.startswith("CHAPTER ") and paragraph.style.name.startswith("Heading"):
        chapter = text
    if paragraph.style.name == "Normal" and len(words(text)) >= 25:
        body.append((index, chapter, text))

ngrams = Counter()
for _, _, text in body:
    tokens = words(text)
    for n in (4, 5, 6):
        ngrams.update((n, " ".join(tokens[i:i+n])) for i in range(len(tokens)-n+1))

print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} images={len(doc.inline_shapes)}")
print("\nLONG PARAGRAPHS")
for index, chapter, text in sorted(body, key=lambda x: len(words(x[2])), reverse=True)[:35]:
    print(f"{index}\t{len(words(text))}\t{chapter}\t{text[:180]}")

print("\nREPEATED PHRASES")
blocked = {"the system is designed", "this chapter presents", "in this thesis the", "as a result the"}
shown = 0
for (n, phrase), count in ngrams.most_common():
    if count < 3 or phrase in blocked:
        continue
    if any(phrase in longer for longer in blocked):
        continue
    print(f"{count}x\t{n}-gram\t{phrase}")
    shown += 1
    if shown >= 50:
        break

print("\nNEAR-DUPLICATE PARAGRAPHS")
sets = [(index, chapter, set(words(text)), text) for index, chapter, text in body]
pairs = []
for a in range(len(sets)):
    ia, ca, sa, ta = sets[a]
    if len(sa) < 20:
        continue
    for b in range(a + 1, len(sets)):
        ib, cb, sb, tb = sets[b]
        score = len(sa & sb) / max(1, len(sa | sb))
        if score >= 0.38:
            pairs.append((score, ia, ib, ca, cb, ta, tb))
for score, ia, ib, ca, cb, ta, tb in sorted(pairs, reverse=True)[:35]:
    print(f"{score:.2f}\t{ia}/{ib}\t{ca} <> {cb}\n  A {ta[:160]}\n  B {tb[:160]}")
