from __future__ import annotations

from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-before-20260803-project-sync.docx"


def set_run_font(run, size=13):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Times New Roman")
    rpr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rpr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def set_paragraph_text(paragraph, text):
    paragraph.text = text
    for run in paragraph.runs:
        set_run_font(run)


def replace_start(document, prefix, text):
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting {prefix!r}, found {len(matches)}")
    set_paragraph_text(matches[0], text)


def normalize_cell(cell, text, bold=False, size=11):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.bold = bold
    set_run_font(r, size)


def set_table(table, rows):
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    if any(len(row) != len(table.columns) for row in rows):
        raise RuntimeError("Table shape mismatch")
    for ri, row in enumerate(rows):
        for ci, value in enumerate(row):
            normalize_cell(table.cell(ri, ci), value, bold=(ri == 0))


def image_paragraph_before(document, caption_prefix):
    captions = [i for i, p in enumerate(document.paragraphs) if p.text.strip().startswith(caption_prefix)]
    if not captions:
        raise RuntimeError(f"Expected caption {caption_prefix!r}, found none")
    # The first match can be the generated List of Figures entry; the last is the body caption.
    caption_index = captions[-1]
    for index in range(caption_index - 1, max(-1, caption_index - 5), -1):
        p = document.paragraphs[index]
        if p._p.xpath(".//a:blip"):
            return p
    raise RuntimeError(f"No image found before {caption_prefix}")


def replace_picture(document, caption_prefix, image_path, alt_text, max_width_cm=15.0, max_height_cm=9.4):
    image_path = Path(image_path)
    p = image_paragraph_before(document, caption_prefix)
    blips = p._p.xpath(".//a:blip")
    if len(blips) != 1:
        raise RuntimeError(f"Expected one picture before {caption_prefix}, found {len(blips)}")
    rid = blips[0].get(qn("r:embed"))
    document.part.related_parts[rid]._blob = image_path.read_bytes()

    with Image.open(image_path) as im:
        ratio = im.width / im.height
    width_cm = min(max_width_cm, max_height_cm * ratio)
    height_cm = width_cm / ratio
    cx, cy = int(Cm(width_cm)), int(Cm(height_cm))
    for extent in p._p.xpath(".//wp:extent"):
        extent.set("cx", str(cx)); extent.set("cy", str(cy))
    for extent in p._p.xpath(".//a:xfrm/a:ext"):
        extent.set("cx", str(cx)); extent.set("cy", str(cy))
    for prop in p._p.xpath(".//wp:docPr"):
        prop.set("descr", alt_text)
        prop.set("title", alt_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(REPORT, BACKUP)

    document = Document(REPORT)

    replacements = {
        "Methods: Recruitment text": "Methods: Recruitment text is normalized and represented with Term Frequency-Inverse Document Frequency vectors. Cosine similarity ranks CV-job pairs and candidate-job recommendations, while Rocchio relevance feedback updates learned representations from explicit feedback. A separate, versioned skill-transfer model checks aliases, skill families, shared foundations, and seniority to explain Potential candidates without changing the direct cosine score. AutoFit checks consent, thresholds, interaction history, quotas, cooldowns, time zones, and quiet hours before choosing an action. CareerFit uses a Spring Boot backend, a React and TypeScript frontend, and PostgreSQL persistence managed through Flyway migrations.",
        "Results: The refreshed backend suite": "Results: On August 3, 2026, the refreshed backend suite passed 131 of 131 tests across 33 suites. TypeScript checking, ESLint, the production frontend build, and the bundle check also passed. All 41 Chrome workflow, contract, and resilience tests passed against the integrated application. In the controlled synthetic benchmark, nDCG@5 increased from 0.037737 to 0.837737 after Rocchio feedback. These results support a well-tested academic prototype, but they do not prove production readiness or real-world hiring effectiveness.",
        "This thesis develops CareerFit IT AutoPilot": "This thesis develops CareerFit IT AutoPilot, an IT-focused web platform that combines job discovery, reviewed CV and JD processing, matching, recommendation, Rocchio feedback, AutoFit policies, secure email actions, recruiter talent discovery, and audit records. The system helps users make decisions but does not replace recruitment judgment.",
        "• Provide candidate profile and CV management": "• Provide candidate profile and CV management, including upload, manual drafts, OCR-assisted extraction, section review, user correction, confirmation, validation, and quality feedback before matching.",
        "• Present normalized scores": "• Present normalized scores, relevance labels, matching reasons, validation signals, and a separate Potential explanation based on transferable skills, skill families, shared foundations, and seniority compatibility.",
        "This thesis focuses on an academic prototype": "This thesis focuses on an academic prototype for IT recruitment. Its functional scope includes a public job portal; authenticated Candidate, Recruiter, and Administrator workspaces; reviewed CV and profile management; recruiter company profiles; Job drafts, quality preview, publishing, urgent and deadline fields; CV–JD matching; personalized recommendation; applications; Talent Pool bookmarks and invitations; explicit feedback; AutoFit policies; notifications; analytics; and audit records.",
        "The data-processing scope includes extraction": "The data-processing scope includes extraction from supported CV documents, image preprocessing and OCR fallback, OCR cleanup, section review, bilingual-oriented text normalization, TF-IDF vectorization, cosine-similarity scoring, a versioned skill-transfer Potential assessment, relevance labeling, and Rocchio feedback updates. PostgreSQL is the main data store, and Flyway manages database changes. The React and TypeScript frontend calls REST APIs provided by the Spring Boot backend.",
        "• An integrated recruitment platform": "• An integrated recruitment platform that combines public job discovery, reviewed CV processing, CV–JD matching, personalized recommendation, application management, Recruiter Jobs, and a Talent Pool workspace.",
        "• A transparent text-scoring pipeline": "• A transparent scoring design that keeps the TF-IDF/cosine direct-match score separate from the versioned Potential assessment and explains transferable skills, shared foundations, and seniority checks.",
        "The implementation exposes these groups": "The implementation exposes these groups through REST-oriented endpoints. Examples include /api/jobs/search, /api/cv/upload, /api/cv/{cvId}/review, /api/matches/me/cards, /api/recommendations/jobs, /api/applications, /api/recruiter/jobs/{jobId}/applicants, /api/recruiter/talent/jobs/{jobId}/bookmarks, /api/automation/policy, and /api/admin/audit-logs. Chapter 3 explains the main behavior.",
        "user_account stores identity": "user_account stores identity, role, activation, verification, language, and password hash. A Candidate owns a profile, multiple CVs, reviewable CV sections, extracted skills, portfolio links, and portfolio projects. A Recruiter owns an employer profile and Jobs. A Job stores its JD, structured recruitment fields, draft/publish state, urgent flag, application deadline, application count, source details, and vector data.",
        "matching is the unique association": "matching is the unique association between one CV and one Job. It stores the direct score, label, separate Potential flag and explanation, reasons, and recomputation state. application links a Candidate and Job and may reference the CV and Matching used at that time. feedback stores a judgment about a Matching. recruiter_cv_bookmark keeps a Recruiter's saved Candidate/CV for an owned Job, while recommendation_interaction records events such as viewed, skipped, applied, saved, not interested, or show similar.",
        "automation_policy stores one policy": "automation_policy stores one policy per account, including enable/pause state, thresholds, quotas, cooldowns, quiet hours, category guards, and notification preferences. email_action stores expiring one-time actions using a token hash. Notification delivery records support deduplication and policy decisions. audit_log records the actor, action, target, result, source channel, request details, and structured metadata. Analytics tables support aggregate dashboards without changing Matching records.",
        "The backend package root is": "The backend package root is com.careerfit.backend. Domain packages include auth, candidate, cv, skill, job, employer, matching, recommendation, application, feedback, automation, notification, analytics, admin, audit, and settings. Shared responses, exceptions, text utilities, validation, security filters, and configuration are placed in common or configuration packages.",
        "AuthController exposes registration": "AuthController exposes registration, email/password login, and current-account operations. AuthService checks account state and credentials and delegates JWT creation to JwtService. Passwords are stored with BCryptPasswordEncoder. A successful login returns an access token and a user DTO with the ID, email, full name, role, and verification state. Passwordless login was removed; signed one-time tokens remain only for explicit email actions.",
        "CvController provides multipart upload": "CvController provides upload, manual creation, manual drafts, CV listing, details, processing status, review read/update, review confirmation, default selection, and deletion. CvIngestionService coordinates the workflow. Uploaded content is extracted and converted into reviewable sections. Manual entry can be saved as a DRAFT. The Candidate can correct sections and confirm them before vectorization and Job scoring start.",
        "PdfExtractionService accepts PDF": "PdfExtractionService accepts PDF, PNG, JPG/JPEG, and DOCX. It checks extension, MIME type, and leading magic bytes. PDFBox extracts embedded PDF text, Apache POI extracts DOCX text, and images use OCR. Scanned pages are preprocessed before Tesseract runs, and OCR cleanup removes common noise before the text is split into review sections. The default OCR languages are vie+eng, with configured page and time limits.",
        "Text shorter than 50 characters": "Text shorter than 50 characters after extraction and cleanup is a hard failure. Text from 50 to 199 characters is accepted with a sparse-content warning. Encrypted PDFs are rejected, image dimensions are limited, and temporary OCR files are deleted. The review step also lets the Candidate fix extraction mistakes before the CV is used for scoring. Production use would still need malware scanning and stronger isolation.",
        "QualityValidationService produces": "QualityValidationService produces structured signals for questionable CV or JD content, such as missing sections, seniority/experience conflicts, salary/content conflicts, and invalid deadlines. Blocking errors are separated from warnings so users can fix the data instead of losing their work.",
        "Although configuration contains low": "The direct-match label uses the configured medium boundary (70) and high boundary (90), producing LOW, MEDIUM, or HIGH. Potential is a separate boolean and does not replace that label. The current Potential model is loaded from a versioned JSON knowledge base and is evaluated separately from cosine similarity.",
        "Match reasons are the five": "Direct-match reasons use the highest-weight Job terms that also appear in the CV, with the Job domain added when available. Potential uses normalized aliases, direct skill-transfer links, role-family transfer, shared foundations, and seniority compatibility. Its explanation can list transferable skills and shared foundations. This remains a rule-based explanation, not a learned claim that the Candidate will succeed.",
        "After CV vectorization commits": "After the Candidate confirms the CV review, vectorization is committed and AfterCommitExecutor submits MatchingService.scoreAllJobsForCv. The service reloads the CV, scores eligible active Jobs, and creates or updates the unique Matching row. Reasons and Potential metadata are stored as JSON. The same post-commit boundary is used when a published Job is scored against existing confirmed CVs.",
        "When a Recruiter creates or updates": "A Recruiter first completes a company profile, then can save a Job as DRAFT, check JD quality, and publish it with the required deadline. Published Jobs can be marked urgent and are scored after commit. MatchingQueryService supports Candidate and Recruiter views with server-side filters, pagination, sorting, label/Potential filters, and clear empty-state metadata. Portfolio details remain protected by the Candidate's visibility setting and application state.",
        "ApplicationService.submit resolves": "ApplicationService.submit resolves the authenticated Candidate, requires an active Job before its application deadline, rejects an existing Candidate–Job application, and uses the requested or default confirmed CV. If a Matching exists, it is attached to keep the score context. The service saves an audit event, updates the Job application count through the database rule, and requests Candidate and Recruiter notifications.",
        "A Candidate lists owned applications": "A Candidate views owned applications through a paginated API and can filter them by status in the interface. A Recruiter can list applicants only for an owned Job, update valid states, or invite a Candidate. The Talent Pool adds Job-based Candidate discovery, a Potential-only view, and CV bookmarks so a Recruiter can save a profile before inviting the Candidate.",
        "The response DTOs combine": "Response DTOs combine application, Candidate, CV, Matching, Job, and bookmark details without exposing persistence entities. Portfolio fields are returned only when the Candidate has enabled visibility and the application state permits it. Server-side pagination, sorting, and filters keep larger Job and Candidate catalogues manageable.",
        "AsyncConfig creates a bounded": "AsyncConfig creates a bounded ThreadPoolTaskExecutor. AfterCommitExecutor submits CV and Job matching only after the related transaction commits. Mail sending and Rocchio learning can also leave the request thread, while persisted CV, Matching, email-action, and delivery states make completion and failures observable.",
        "AutomationScheduler coordinates five": "AutomationScheduler coordinates matching recomputation, digests, high-match checks, auto-apply, token cleanup, and deadline-related reminders using configured schedules. Per-item error handling prevents one failed record from stopping a full scan. Account policies add pause/resume, thresholds, quotas, cooldowns, quiet hours, and category guards before eligible actions or notifications are executed.",
        "AutoApplyService.runForPolicy resolves": "AutoApplyService.runForPolicy resolves the policy owner, Candidate, and confirmed default CV. It checks policy enablement and pause state, eligible categories, active Job state, deadline, score threshold, previous interactions, quota, and cooldown before creating an automatic application. Database uniqueness remains the final concurrency guard. Every created or skipped outcome can be audited and followed by a notification decision.",
        "The current auto-apply implementation": "Auto-apply and notification delivery use related but separate checks. Application creation requires consent and application-policy conditions, while email delivery also checks timing, quota, cooldown, category, and deduplication. Keeping the two decisions separate prevents a notification setting from silently becoming permission to submit an application.",
        "NotificationPolicyGuard evaluates": "NotificationPolicyGuard evaluates whether a message is allowed and stores sent, skipped, or failed results in separate transactions. It supports deduplication, quota, cooldown, quiet hours, category guards, and deadline reminders. SMTP is used when enabled, while NoOpMailService keeps local development usable without claiming that real email delivery was tested.",
        "JPA entities map the domain tables": "JPA entities map the domain tables, and Spring Data repositories provide pagination, sorting, locking, and explicit update operations. Flyway migrations V1–V24 create and update the current schema, including removal of passwordless login, CV review/draft states, the skill catalogue, Potential support, recruiter bookmarks, Job popularity, urgent hiring, account-level automation fields, and application deadlines. Hibernate uses ddl-auto=validate.",
        "App.tsx defines public": "App.tsx defines public, Candidate, Recruiter, and Administrator routes. Protected routes check the loaded account role, while session restoration calls /api/auth/me and clears invalid state. Public pages include Jobs and employer details. Candidate pages cover CV drafts and review, matching, urgent Jobs, recommendations, applications, analytics, AutoFit, and settings. Recruiter pages cover company onboarding, Jobs, applicants, Talent Pool, bookmarks, invitations, and analytics.",
        "API-driven pages no longer substitute": "API-driven pages do not replace failed responses with mock Job data. The current frontend uses server-side filters, pagination, and sorting for larger catalogues; skill/title/location/domain autocomplete; Candidate status tabs; CV review and confirmation; recruiter company onboarding; JD quality preview; Job draft/publish; urgent and deadline controls; Talent Pool bookmarks and invitations; per-account AutoFit settings; and analytics refetch for current-day data.",
        "This chapter explained how the design": "This chapter explained the current implementation in Spring Boot, React, PostgreSQL, and Flyway. It covered password-based JWT security, CV extraction and human review, direct matching, skill-transfer Potential assessment, feedback learning, recruiter company and Talent Pool workflows, applications, account-level AutoFit controls, email actions, frontend integration, and monitoring. Chapter 4 evaluates the refreshed system.",
        "Backend, algorithm, frontend build": "Backend, algorithm, frontend build, and integrated Chrome results were refreshed on August 3, 2026 (ICT). The observed Git HEAD was 242e13a8f7d16fc9ebcab9780264c2c2b2b4ef06, but the worktree contained 197 modified or untracked entries. The results therefore describe this working tree and cannot be reproduced from the commit alone unless the full source state is preserved.",
        "Docker Desktop was initially unavailable": "Docker Desktop and Compose provided PostgreSQL and the backend for integrated browser verification. Testcontainers created isolated PostgreSQL 16 databases for backend integration tests. The browser suite used the local CareerFit database, which contained imported, seeded, and E2E records. The public interface showed 992 open Jobs during the screenshot session. This runtime data is separate from the controlled algorithm dataset.",
        "The original command catalog": "Earlier commands are recorded in evidence/CHAPTER5_EVIDENCE_20260703.md. For the August 3 refresh, Surefire XML, Maven output, TypeScript/ESLint/Vite results, Playwright output, the current screenshots, and evaluation/result.json are the main evidence. The report states the working-tree limitation instead of treating the commit hash as a complete release identifier.",
        "The complete backend suite was executed": "The complete backend suite was executed with .\\mvnw.cmd clean verify. It compiled 142 application source files and 35 test source files, started PostgreSQL 16 Testcontainers, applied Flyway migrations V1–V24, executed 33 Surefire suites, and built the backend JAR.",
        "The suites covered application context": "The suites covered application context, API contracts, security, account login, CV drafts and review, OCR/extraction, skill suggestions and transfer rules, Job quality and lifecycle, server-side catalogues, matching, feedback, automation, deadlines, bookmarks, invitations, analytics, settings, and the controlled AlgorithmEvaluatorTest.",
        "All 72 registered JUnit tests": "All 131 registered JUnit tests passed on August 3, with no failures, errors, or skips. The aggregated Surefire suite time was 102.194 seconds. Negative tests still produced expected handled errors in their own cases, but the Maven build completed successfully and no final test report contained a failure.",
        "The July 18 full-suite run": "The August 3 full-suite run reproduced dataset hash 6e935639ba6d3290dca8ad91a35d714c5e30c7e69a59af23ddbcf89fcc5cc2f2. It produced baseline nDCG@5 of 0.037737056145, Rocchio nDCG@5 of 0.837737056145, and a delta of 0.80. These values describe the controlled dataset and should not be read as expected improvement on real recruitment data.",
        "npm run build performs": "The frontend verification ran TypeScript checking with no output, ESLint, Vite 6.4.3 production build, and the bundle check. The August 3 build completed successfully and transformed 2,361 modules.",
        "Manual Rollup chunking": "Manual Rollup chunking separated React, query, chart, and icon dependencies. The largest generated JavaScript chunk was charts at 387.39 kB (112.17 kB gzip), below Vite's 500 kB warning threshold. This is build evidence only; no browser performance profile was collected.",
        "The Playwright Chromium project ran": "The Playwright suite ran with the Vite frontend, the current backend, PostgreSQL, and the installed desktop Chrome channel. All 41 tests across workflow, backend-contract, account-state, catalogue, and resilience coverage passed in 42.1 seconds.",
        "The result covers the four original": "The result covers Guest search and details; email/password login; Candidate CV draft/review, matching, urgent filters, applications, settings, and AutoFit; Recruiter company onboarding, Job draft/publish, applicants, Talent Pool, bookmarks, and invitations; Administrator operations; session restoration; server-side catalogues; retryable errors; dark mode; and role navigation. It is not an independent UAT study, and Firefox/WebKit were not run in this refresh.",
        "The July 18 evaluation passed": "The August 3 evaluation passed all 131 backend tests, completed TypeScript, ESLint, production-build, and bundle checks, and passed all 41 integrated Chrome tests. The controlled Rocchio benchmark kept its expected synthetic improvement. These results support a functioning academic prototype, but not production readiness or proven hiring effectiveness.",
        "CareerFit IT AutoPilot was implemented": "CareerFit IT AutoPilot was implemented as an IT recruitment prototype with public Job discovery, password-based role workspaces, reviewed CV processing, direct matching, a separate Potential assessment, recommendations, Rocchio feedback, applications, Recruiter Talent Pool features, account-level AutoFit policies, email actions, analytics, and administrative monitoring.",
        "The refreshed evaluation provides evidence": "The refreshed evaluation provides evidence for several technical outcomes. The backend passed 131 tests across 33 suites. TypeScript checking, ESLint, Vite production compilation, and the bundle check passed. All 41 integrated Chrome tests passed across role workflows, backend contracts, account state, catalogues, and resilience, including the new CV review, Job draft/publish, urgent, Talent Pool, bookmark, invitation, and policy-setting flows.",
        "The July 18 refresh retained": "The August 3 refresh verified the larger project state, kept the largest JavaScript chunk below 500 kB, removed stale passwordless claims, updated the CV review and skill-transfer design, expanded browser coverage to 41 tests, and replaced six report interface images with current screens. The result is suitable for a local thesis demonstration, not a production release.",
        "The controlled dataset is synthetic": "The controlled dataset is synthetic and designed to show how feedback changes ranking. It does not include normal language variation, recruiter disagreement, demographic analysis, misleading CVs, or changing market terms. TF-IDF still uses a static IT corpus and simple tokenization. The Potential assessment now uses a versioned skill-transfer knowledge base, but its aliases, transfer weights, family rules, and guard thresholds are still manually defined.",
        "The 72 backend tests": "The 131 backend tests do not prove complete path coverage. Browser evaluation covers 41 automated cases in desktop Chrome only, and no independent users participated. Security checks were targeted rather than a penetration test, while the latency sample used 30 sequential local requests. Concurrency, capacity, cross-browser behavior, accessibility with users, and real hiring outcomes remain unevaluated.",
        "The final remediation pass resolved": "The current implementation improves earlier gaps through after-commit matching, CV review before scoring, versioned Potential rules, server-side catalogues, recruiter company onboarding, Job drafts, bookmarks, policy guards, and broader E2E coverage. Audit records are still written directly by several services, and the evaluated worktree is not a fixed release commit. These limits make exact reproduction and production operation harder.",
        "Local evaluation passed the backend tests": "Local evaluation passed 131 backend tests, the frontend type/lint/build/bundle checks, 41 integrated Chrome tests, health checks, and the controlled Rocchio benchmark. This supports a tested Human-in-the-Loop prototype, not a system ready to make real hiring decisions.",
        "Representative endpoints are POST": "Representative endpoints are POST /api/auth/login; GET /api/jobs/search; POST /api/cv/upload; GET/PATCH /api/cv/{cvId}/review; POST /api/cv/{cvId}/review/confirm; GET /api/matches/me/cards; POST /api/matches/{matchingId}/feedback; POST /api/applications; GET /api/recruiter/talent/jobs/{jobId}/bookmarks; PATCH /api/automation/policy; GET then POST /api/email-action/redeem; and GET /actuator/health. Protected routes require a signed JWT and backend role/ownership checks.",
        "Identity data is centered": "Identity data is centered on UserAccount and role profiles. Recruitment data includes Job, CV with review fields, Matching, Application, RecruiterCvBookmark, Invitation state, Feedback, and the Skill catalogue. Automation and operations use AutomationPolicy, EmailAction, Notification, DeliveryLog, AuditLog, and scheduler state. Flyway V1–V24 records schema changes, and one-time email action tokens are stored as SHA-256 hashes.",
        "The final evidence package contains": "The final evidence package contains the 131-test backend reports, controlled algorithm benchmark, frontend type/lint/build results, 41-test Chrome output, runtime health evidence, updated screenshots, and evaluation/result.json. Chapter 4 reports what this evidence supports and what it does not support.",
        "2. As a Candidate, sign in": "2. As a Candidate, sign in, upload a CV or save a manual draft, review and correct extracted sections, confirm the CV, wait for scoring, inspect match reasons and Potential explanations, filter urgent Jobs, apply or withdraw where permitted, and submit feedback.",
        "3. As a Recruiter, create": "3. As a Recruiter, complete the company profile, save a test Job draft, run the JD quality preview, publish it with an application deadline, inspect applicants and the Talent Pool, bookmark and invite a Candidate, update application status, and remove the test Job after verification.",
        "Sign in, open My CVs": "Sign in, open My CVs, upload a PDF/DOCX/image file or save a manual draft. Review the extracted sections and quality warnings, correct them if needed, and confirm the CV before scoring. Then choose the default CV, review direct-match and Potential explanations, manage applications, use status and urgent filters, provide feedback, inspect recommendations, and configure AutoFit settings.",
        "Sign in to the Recruiter workspace": "Sign in to the Recruiter workspace and complete the company profile before creating a Job. Save the Job as a draft, check JD quality, set the urgent flag and application deadline when needed, then publish it. Use Jobs and applicants for active recruitment, and use Talent Pool to filter Potential CVs, bookmark a Candidate, send an invitation, and update valid application states.",
    }

    for prefix, text in replacements.items():
        replace_start(document, prefix, text)

    # Update screenshot captions before replacing the associated image parts.
    captions = {
        "Screen 3.1.": "Screen 3.1. Candidate urgent-job catalogue",
        "Screen 3.2.": "Screen 3.2. Candidate AutoFit policy settings",
        "Screen 3.3.": "Screen 3.3. Candidate CV upload entry point",
        "Screen 3.4.": "Screen 3.4. Recruiter Jobs and applicant workspace",
        "Screen 3.5.": "Screen 3.5. Recruiter Talent Pool and Potential CV",
        "Screen 3.6.": "Screen 3.6. Administrator audit logs",
    }
    for prefix, text in captions.items():
        matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
        if not matches:
            raise RuntimeError(f"Missing body caption {prefix}")
        set_paragraph_text(matches[-1], text)

    # Functional and implementation tables.
    set_table(document.tables[4], [
        ["Group", "Required capabilities", "Main actor"],
        ["Authentication and account", "Registration, email/password login, current-account lookup, role enforcement, and settings", "Candidate, Recruiter, Administrator"],
        ["Public Job portal", "Server-side search, suggestions, filters, urgent flag, sorting, pagination, Job details, employers, and public analytics", "Guest and authenticated users"],
        ["Candidate profile and CV", "Profile/portfolio, multiple CVs, manual drafts, upload, OCR cleanup, section review/edit/confirm, default CV, and quality signals", "Candidate"],
        ["Job and employer management", "Required company profile; JD quality preview; Job draft, publish, update, urgent/deadline fields, status, export, and counts", "Recruiter"],
        ["Matching and recommendation", "Direct CV–JD score, labels/reasons, separate Potential assessment, Candidate cards, recruiter ranking, recommendations, and similar Jobs", "Candidate and Recruiter"],
        ["Application and Talent Pool", "Apply, history, withdrawal, applicant review, Potential discovery, bookmark, invitation, and status update", "Candidate and Recruiter"],
        ["Feedback learning", "Record explicit feedback, update learned Job vectors with Rocchio, and mark Matchings for recomputation", "Candidate"],
        ["Automation and notification", "Per-account enable/pause, thresholds, categories, quota, cooldown, quiet hours, run-now, auto-apply, reminders, and delivery logs", "Authenticated user and Scheduler"],
        ["Analytics and administration", "Live summaries, event tracking, user/Job moderation, audit logs, notifications, and email-action monitoring", "Candidate, Recruiter, Administrator"],
    ])

    set_table(document.tables[6], [
        ["Field", "Description"],
        ["Use-case ID", "UC-01"],
        ["Primary actor", "Candidate"],
        ["Preconditions", "The Candidate account is active, authenticated, and has permission to manage its own CVs."],
        ["Trigger", "The Candidate uploads a supported document or saves manual CV content."],
        ["Main flow", "Validate and extract the source; preprocess and OCR when needed; create review sections; show warnings; let the Candidate edit and confirm; normalize and vectorize; score active Jobs; persist unique Matchings; return ordered cards."],
        ["Alternative/exception flows", "Save manual work as DRAFT; reject unsupported or unsafe content; record extraction failure; keep REVIEW_REQUIRED until the Candidate confirms; show an empty state when no Job is eligible."],
        ["Postconditions", "The CV remains DRAFT/REVIEW_REQUIRED, reaches SCORING_DONE after confirmation, or ends in FAILED with a visible reason."],
    ])

    set_table(document.tables[8], [
        ["Field", "Description"],
        ["Use-case ID", "UC-03"],
        ["Primary actor", "Recruiter"],
        ["Preconditions", "The Recruiter account is active, authenticated, and has completed its company profile."],
        ["Trigger", "The Recruiter saves a Job draft or opens an owned Job/Talent Pool workspace."],
        ["Main flow", "Validate structured and free-text data; preview JD quality; save DRAFT; set deadline/urgent state; publish; vectorize after commit; review applicants and Potential CVs; bookmark, invite, or update application state."],
        ["Alternative/exception flows", "Reject missing company data, invalid deadline/content, another Recruiter's Job, invalid invitation, or unsupported status transition; keep incomplete work as DRAFT."],
        ["Postconditions", "The Job draft or published Job and permitted recruitment actions are persisted with ownership and audit evidence."],
    ])

    set_table(document.tables[13], [
        ["Module", "Main responsibility", "Representative services"],
        ["Auth and Security", "Registration, password login, JWT validation, account resolution, and role rules", "AuthService, JwtService, security filters"],
        ["Candidate, CV, and Skill", "Profile/portfolio, CV draft/review, storage, OCR cleanup, validation, and skill suggestions", "CandidateProfileService, CvIngestionService, PdfExtractionService, SkillService"],
        ["Job and Employer", "Public catalogue, company onboarding, JD quality, draft/publish, urgent/deadline, and recruiter-owned lifecycle", "JobService, EmployerService, QualityValidationService"],
        ["Matching", "TF-IDF/cosine score, labels/reasons, skill-transfer Potential, batch matching, filters, and ranking", "ScoringService, SkillTransferService, MatchingService, MatchingQueryService"],
        ["Recommendation", "Candidate-profile Job ordering and similar-Job retrieval", "RecommendationService"],
        ["Application and Talent", "Apply/withdraw, applicants, invitations, status changes, Talent Pool, and CV bookmarks", "ApplicationService, RecruiterTalentService"],
        ["Feedback", "Feedback validation, Rocchio update, and recomputation signaling", "FeedbackService, RocchioService"],
        ["Automation", "Per-account policy, run-now, pause/resume, auto-apply, reminders, and scheduler", "AutomationPolicyService, AutoApplyService, AutomationScheduler"],
        ["Notification", "Policy guard, delivery log, SMTP/no-op mail, and signed email actions", "NotificationPolicyGuard, NotificationEmailService, EmailActionService"],
        ["Analytics, Admin, and Audit", "Live analytics, moderation, monitoring, audit persistence, and common API behavior", "Analytics services, administrative services, audit/common components"],
    ])

    set_table(document.tables[14], [
        ["Constraint", "Purpose"],
        ["Unique normalized user email", "Prevent duplicate identities that differ only by case"],
        ["One Candidate/Recruiter profile and one policy per account", "Preserve role ownership and account-level automation settings"],
        ["Partial unique default CV per Candidate", "Ensure at most one active default CV"],
        ["Unique Matching per CV and Job", "Prevent duplicate direct scores for the same pair"],
        ["Unique Application per Candidate and Job", "Prevent duplicate applications or invitations"],
        ["Unique Recruiter bookmark per Job and Candidate", "Prevent duplicate Talent Pool bookmarks"],
        ["Unique feedback per Matching and actor", "Keep one current judgment per actor"],
        ["Check constraints on roles, states, labels, salary, and source channels", "Reject invalid domain values at the database boundary"],
        ["Indexes for active/urgent/popular Jobs, deadlines, ownership, and ranking", "Support server-side catalogue, reminder, and matching queries"],
    ])

    set_table(document.tables[17], [
        ["State", "Meaning", "Typical transition"],
        ["DRAFT", "Manual content is saved but not submitted for scoring", "Draft edit → review/confirm"],
        ["UPLOADED", "Uploaded metadata and source are accepted", "Upload → validation/extraction"],
        ["VALIDATING", "File, content, and extraction checks are running", "Valid source → review; hard error → failed"],
        ["REVIEW_REQUIRED", "Extracted/manual sections are ready for Candidate review", "Edit sections → confirm"],
        ["PROCESSING", "Confirmed content is normalized and vectorized", "Vector saved → asynchronous scoring"],
        ["SCORING_DONE", "The confirmed CV vector is available and Matchings can be queried", "Use cards, feedback, applications, and AutoFit"],
        ["FAILED", "Processing cannot continue", "Failure reason is stored and audited"],
    ])

    set_table(document.tables[18], [
        ["Condition", "Stored/displayed result"],
        ["Direct score below 70", "LOW label"],
        ["Direct score from 70 to below 90", "MEDIUM label"],
        ["Direct score at least 90", "HIGH label; Potential is not added"],
        ["Potential score at least 62 with skill compatibility ≥0.50 and family compatibility ≥0.55", "Potential guard continues"],
        ["Core target and transferable career evidence exist, with no severe seniority gap", "isPotential=true and an explanation is stored"],
        ["Direct score below 20 and skill compatibility below 0.70", "Potential is rejected as weak evidence"],
    ])

    set_table(document.tables[20], [
        ["API operation", "Controller/service responsibility", "Persistence effect"],
        ["POST /api/cv/upload; GET/PATCH /api/cv/{id}/review", "Extract content, expose/edit review sections", "CV, review data, quality signals, audit"],
        ["POST /api/cv/{id}/review/confirm", "Confirm owned CV and start vectorization/matching", "CV vector, Matchings, audit"],
        ["GET /api/matches/me/cards", "Filter, sort, paginate, and map Candidate match cards", "Read-only"],
        ["POST /api/matches/{id}/feedback", "Upsert feedback and trigger Rocchio after commit", "Feedback, audit, learned Job vector, stale flags"],
        ["POST /api/jobs/drafts; POST /api/jobs/{id}/publish", "Validate company/JD/deadline and manage owned Job lifecycle", "Job, vectors, Matchings, audit"],
        ["PUT /api/recruiter/talent/jobs/{jobId}/candidates/{candidateId}/bookmark", "Validate Job ownership and save Talent Pool choice", "Recruiter CV bookmark"],
        ["PATCH /api/automation/policy; POST /api/automation/pause", "Update or pause the account's automation policy", "Automation policy and version"],
    ])

    set_table(document.tables[21], [
        ["Area", "Current limitation", "Consequence or required improvement"],
        ["Email actions", "Hashed token and confirm-then-POST flow", "Add rate limiting, delivery monitoring, and deployment-specific origin controls"],
        ["CV/OCR", "Preprocessing and cleanup are heuristic", "Scanned layouts and low-quality images still need user correction"],
        ["Text processing", "Simple tokenization remains the direct-score baseline", "Technology spelling and Vietnamese phrases may still lose information"],
        ["Potential model", "Versioned aliases and transfer weights are manually defined", "Validate rules with recruiter-labeled data and review model drift"],
        ["Labels", "Direct labels use only the medium/high boundaries", "Align configuration names and UI wording with tested boundaries"],
        ["Automation", "Many policy guards and schedules interact", "Add more time-zone, quota, cooldown, deadline, and concurrency tests"],
        ["Frontend session", "JWT and account summary use sessionStorage", "Evaluate stronger refresh/cookie architecture and XSS controls"],
        ["File storage", "Local path or Docker volume", "Add malware scanning, encryption, backup, retention, and protected object storage"],
        ["Audit", "Several services write audit rows directly", "Centralize redaction, request context, schema rules, and integrity policy"],
    ])

    set_table(document.tables[22], [
        ["Component", "Observed version or configuration"],
        ["Host operating environment", "Windows, PowerShell"],
        ["Java", "21.0.1"],
        ["Backend framework", "Spring Boot 3.2.5"],
        ["Maven", "Repository Maven Wrapper; clean verify"],
        ["Node.js / npm", "Bundled Node runtime with repository npm scripts"],
        ["Frontend build", "Vite 6.4.3, React 18.3, TypeScript 5.9, ESLint"],
        ["Docker / Compose", "Docker Desktop with Compose"],
        ["Test database", "PostgreSQL 16 through Testcontainers; local Compose PostgreSQL for E2E"],
        ["Database migrations", "Flyway V1–V24"],
        ["Browser E2E", "Playwright 1.61 using installed desktop Chrome"],
        ["Integrated backend", "Compose backend on port 8080"],
        ["Integrated frontend", "Vite development server on 127.0.0.1:5173"],
    ])

    set_table(document.tables[23], [
        ["Measure", "Result"],
        ["Test suites", "33"],
        ["Tests", "131"],
        ["Failures", "0"],
        ["Errors", "0"],
        ["Skipped", "0"],
        ["Aggregated Surefire suite time", "102.194 s"],
        ["Maven lifecycle", "clean verify and JAR build passed"],
    ])

    set_table(document.tables[25], [
        ["Artifact", "Uncompressed size", "Gzip size"],
        ["index.html", "1.37 kB", "0.67 kB"],
        ["Main CSS", "132.25 kB", "24.36 kB"],
        ["Icons chunk", "22.10 kB", "5.03 kB"],
        ["Query chunk", "39.78 kB", "12.20 kB"],
        ["React chunk", "181.35 kB", "59.62 kB"],
        ["Application chunk", "310.48 kB", "83.45 kB"],
        ["Charts chunk", "387.39 kB", "112.17 kB"],
    ])

    set_table(document.tables[26], [
        ["Scenario group", "Main verification", "Result"],
        ["Guest and authentication", "Public catalogue/detail, email/password login, session restoration, and role routes", "Passed"],
        ["Candidate CV and matching", "Manual draft, upload/review/confirm, polling, match cards, filters, and feedback", "Passed"],
        ["Candidate recruitment", "Urgent catalogue, apply/withdraw/history, settings, and AutoFit controls", "Passed"],
        ["Recruiter Jobs", "Company onboarding, quality preview, draft/publish, deadline, applicants, and catalogue", "Passed"],
        ["Recruiter Talent Pool", "Potential view, Candidate details, bookmark, invitation, and state changes", "Passed"],
        ["Administration and resilience", "Suspend/reactivate, audit, API errors, navigation, account state, and dark mode", "Passed"],
    ])

    set_table(document.tables[30], [
        ["Evaluation area", "Supported result", "Qualification"],
        ["Backend automated tests", "131/131 tests across 33 suites passed", "clean verify and JAR build completed"],
        ["Controlled feedback benchmark", "Large, deterministic Rocchio improvement", "Synthetic causal design; not production effectiveness"],
        ["Frontend checks", "TypeScript, ESLint, Vite build, and bundle check passed", "Largest JS chunk 387.39 kB; no Vite size warning"],
        ["Browser workflows", "41/41 integrated Chrome tests passed", "No Firefox/WebKit or independent participant UAT"],
        ["Authorization spot checks", "Observed 200/401/403 outcomes matched expectations", "Not a complete security assessment"],
        ["Runtime APIs", "Core Job API and health endpoints returned HTTP 200/UP", "Local profile only; no production monitoring"],
        ["Local latency sample", "Mean 61.79 ms and p95 85.11 ms for 30 sequential Job queries", "No concurrency, controlled load, or production network"],
    ])

    set_table(document.tables[31], [
        ["Objective", "Assessment", "Primary evidence"],
        ["Role-based web platform", "Achieved for prototype scope", "API contracts and 41 integrated Chrome tests"],
        ["Reviewed CV/JD ingestion", "Achieved for supported formats and configured OCR", "Draft/review/confirm, extraction, quality, Job, and integration tests"],
        ["Explainable matching and Potential", "Achieved as rule-based baseline", "TF-IDF/scoring tests and versioned skill-transfer reasons"],
        ["Rocchio feedback adaptation", "Achieved in controlled test; production concurrency unverified", "Synthetic benchmark, after-commit execution, and clean reports"],
        ["AutoFit and communication", "Achieved as configurable prototype", "Account policies, scheduler, reminders, notifications, and signed email actions"],
        ["Recruiter operational workflow", "Achieved for prototype scope", "Company, Job draft/publish, applicants, Talent Pool, bookmark, and invitation tests"],
        ["Production readiness", "Not demonstrated", "Local health passed; scale, real users, and production security remain missing"],
    ])

    set_table(document.tables[32], [
        ["Limitation", "Impact on conclusions"],
        ["Synthetic benchmark", "Supports causal behavior only, not real hiring effectiveness"],
        ["Lexical direct score", "Limits semantic and contextual understanding"],
        ["Manual skill-transfer knowledge base", "Potential rules require labeled validation and ongoing maintenance"],
        ["Production concurrency not evaluated", "Conflict, retry, and recovery behavior remain unproven at scale"],
        ["Local-only health verification", "Does not establish deployed availability"],
        ["Chrome-only 41-test automation suite", "Limits cross-browser and real-user usability conclusions"],
        ["Email delivery and abuse controls not production-tested", "Signed action flow exists, but real delivery and rate limits remain unverified"],
        ["Dirty worktree and mutable E2E database", "Weakens exact experiment reproduction"],
    ])

    # Replace implementation diagrams generated from the current architecture.
    diagram_map = {
        "Figure 1.4.": ROOT / "Doc/figures/fig-3-2.png",
        "Figure 1.5.": ROOT / "Doc/figures/fig-3-3.png",
        "Figure 2.6.": ROOT / "Doc/figures/fig-3-8.png",
        "Figure 3.3.": ROOT / "Doc/figures/fig-4-3.png",
        "Figure 3.5.": ROOT / "Doc/figures/fig-4-5.png",
        "Figure 3.8.": ROOT / "Doc/figures/fig-4-8.png",
        "Figure 3.10.": ROOT / "Doc/figures/fig-4-10.png",
    }
    for caption_prefix, path in diagram_map.items():
        replace_picture(document, caption_prefix, path, f"Current CareerFit diagram for {caption_prefix.rstrip('.')}")

    screenshot_map = {
        "Screen 3.1.": ROOT / "Doc/screenshots/screen-3-3-candidate-urgent-jobs-20260803.png",
        "Screen 3.2.": ROOT / "Doc/screenshots/screen-3-2-candidate-autofit-20260803.png",
        "Screen 3.3.": ROOT / "Doc/screenshots/screen-4-3-cv-upload.png",
        "Screen 3.4.": ROOT / "Doc/screenshots/screen-3-4-recruiter-workspace-20260803.png",
        "Screen 3.5.": ROOT / "Doc/screenshots/screen-3-5-recruiter-talent-pool-20260803.png",
        "Screen 3.6.": ROOT / "Doc/screenshots/screen-3-6-admin-audit-20260803.png",
    }
    for caption_prefix, path in screenshot_map.items():
        replace_picture(document, caption_prefix, path, next(v for k, v in captions.items() if k == caption_prefix))

    # Enforce the faculty margin requirement on every section.
    for section in document.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(3.0)
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2.0)

    # Mark all fields dirty so Word refreshes TOC and lists on open.
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        from docx.oxml import OxmlElement
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    document.core_properties.comments = "Updated against CareerFit working tree and verified on 2026-08-03."
    document.save(REPORT)
    print(f"updated {REPORT}")
    print(f"backup  {BACKUP}")


if __name__ == "__main__":
    main()
