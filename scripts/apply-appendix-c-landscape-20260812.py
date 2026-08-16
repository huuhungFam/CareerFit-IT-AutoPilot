from __future__ import annotations

import copy
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
TEMP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-landscape-20260812.docx"


def paragraph_text(element, parent) -> str:
    return Paragraph(element, parent).text.strip()


def set_section_geometry(sect_pr, *, landscape: bool) -> None:
    pg_sz = sect_pr.find(qn("w:pgSz"))
    if pg_sz is None:
        pg_sz = OxmlElement("w:pgSz")
        sect_pr.append(pg_sz)
    if landscape:
        pg_sz.set(qn("w:w"), "16838")
        pg_sz.set(qn("w:h"), "11906")
        pg_sz.set(qn("w:orient"), "landscape")
    else:
        pg_sz.set(qn("w:w"), "11906")
        pg_sz.set(qn("w:h"), "16838")
        pg_sz.attrib.pop(qn("w:orient"), None)

    pg_mar = sect_pr.find(qn("w:pgMar"))
    if pg_mar is None:
        pg_mar = OxmlElement("w:pgMar")
        sect_pr.append(pg_mar)
    pg_mar.set(qn("w:top"), "1701")
    pg_mar.set(qn("w:bottom"), "1701")
    pg_mar.set(qn("w:left"), "1984")
    pg_mar.set(qn("w:right"), "1134")

    section_type = sect_pr.find(qn("w:type"))
    if section_type is None:
        section_type = OxmlElement("w:type")
    else:
        sect_pr.remove(section_type)
    # CT_SectPr has a strict child order: header/footer references precede w:type,
    # while page size and margins follow it. Keeping this order avoids Word's
    # unreadable-content repair prompt.
    pg_sz_index = list(sect_pr).index(pg_sz)
    sect_pr.insert(pg_sz_index, section_type)
    section_type.set(qn("w:val"), "nextPage")


def insert_section_break_before(paragraph_element, source_sect_pr, *, landscape: bool) -> None:
    section_p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    section_p.append(p_pr)
    sect_pr = copy.deepcopy(source_sect_pr)
    set_section_geometry(sect_pr, landscape=landscape)
    p_pr.append(sect_pr)
    paragraph_element.addprevious(section_p)


def set_table_widths(table: Table, widths: list[float]) -> None:
    table.autofit = False
    tbl_grid = table._tbl.tblGrid
    for grid_col, width in zip(tbl_grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(int(Inches(width).twips)))
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)


def main() -> None:
    doc = Document(DOCX)
    if len(doc.sections) not in {3, 5}:
        raise RuntimeError(f"Expected 3 or 5 sections, found {len(doc.sections)}")

    body = doc._element.body
    source_sect_pr = body.sectPr
    c_heading = None
    d_heading = None
    in_appendix_c = False
    appendix_c_tables: list[Table] = []

    for child in list(body.iterchildren()):
        if child.tag == qn("w:p"):
            text = paragraph_text(child, doc._body)
            if text == "Appendix C. Full Data Dictionary":
                c_heading = child
                in_appendix_c = True
            elif text == "Appendix D. Evaluation Summary":
                d_heading = child
                in_appendix_c = False
        elif child.tag == qn("w:tbl") and in_appendix_c:
            appendix_c_tables.append(Table(child, doc._body))

    if c_heading is None or d_heading is None:
        raise RuntimeError("Appendix C or Appendix D heading was not found")

    # The section property stored before a heading describes the section that ends there.
    # Portrait before Appendix C closes the preceding content; landscape before Appendix D
    # closes Appendix C; the body-level final section remains portrait after Appendix D.
    if len(doc.sections) == 3:
        insert_section_break_before(c_heading, source_sect_pr, landscape=False)
        insert_section_break_before(d_heading, source_sect_pr, landscape=True)

    # Reorder and normalize all section properties, including files produced by
    # an earlier run before the strict OOXML ordering fix.
    section_properties = body.xpath(".//w:sectPr")
    for index, sect_pr in enumerate(section_properties, start=1):
        set_section_geometry(sect_pr, landscape=(index == 4))

    for table in appendix_c_tables:
        if len(table.columns) == 7:
            set_table_widths(table, [0.45, 1.40, 1.40, 0.65, 0.65, 1.80, 3.10])
        else:
            table.autofit = True
        if len(table.rows) >= 3:
            for cell in table.rows[-2].cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    )

    doc.save(TEMP)
    shutil.copy2(TEMP, DOCX)
    check = Document(DOCX)
    print(f"Updated: {DOCX}")
    print(f"Sections: {len(check.sections)}")
    print(f"Appendix C tables: {len(appendix_c_tables)}")
    for i, section in enumerate(check.sections, start=1):
        print(
            f"Section {i}: {section.orientation}, "
            f"{section.page_width.inches:.2f} x {section.page_height.inches:.2f} in"
        )


if __name__ == "__main__":
    main()
