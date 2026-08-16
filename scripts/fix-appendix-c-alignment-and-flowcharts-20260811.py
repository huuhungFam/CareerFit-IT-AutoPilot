from __future__ import annotations

import hashlib
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = (
    ROOT
    / "Doc"
    / "working"
    / "CareerFit-Thesis-Report-before-20260811-flowchart-and-dictionary-fix.docx"
)
FIGURES = ROOT / "Doc" / "figures"
FLOWCHART_WIDTH = 5.30

FLOWCHARTS = {
    "Figure 3.14. CV ingestion, review, confirmation, and matching flowchart": "flowchart-cv-processing-20260811.png",
    "Figure 3.15. Seed-corpus initialization and TF-IDF construction flowchart": "flowchart-tfidf-construction-20260811.png",
    "Figure 3.16. CV-Job matching and Potential assessment flowchart": "flowchart-matching-potential-20260811.png",
    "Figure 3.17. Rocchio feedback-learning and recomputation flowchart": "flowchart-rocchio-feedback-20260811.png",
    "Figure 3.19. AutoFit eligibility and application decision flowchart": "flowchart-autofit-20260811.png",
    "Figure 3.20. Notification policy evaluation and delivery outcome flowchart": "flowchart-notification-policy-20260811.png",
    "Figure 3.21. Actionable-email confirmation and redemption flowchart": "flowchart-email-action-redemption-20260811.png",
}

CATALOGUE_WIDTHS_CM = [0.9, 3.6, 3.4, 7.6]
DICTIONARY_WIDTHS_CM = [1.0, 2.8, 2.8, 1.2, 1.3, 3.0, 3.4]


def digest_tables(document: Document) -> str:
    payload = [
        [
            [cell.text for cell in row.cells]
            for row in table.rows
            if any(cell.text.strip() for cell in row.cells)
        ]
        for table in document.tables
    ]
    return hashlib.sha256(repr(payload).encode("utf-8")).hexdigest()


def digest_body_text(document: Document) -> str:
    return hashlib.sha256(
        "\n".join(paragraph.text for paragraph in document.paragraphs).encode("utf-8")
    ).hexdigest()


def find_caption(document: Document, caption_text: str) -> Paragraph:
    matches = [p for p in document.paragraphs if p.text.strip() == caption_text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one caption {caption_text!r}; found {len(matches)}")
    return matches[0]


def previous_paragraph(paragraph: Paragraph) -> Paragraph:
    element = paragraph._p.getprevious()
    while element is not None and not element.tag.endswith("}p"):
        element = element.getprevious()
    if element is None:
        raise RuntimeError(f"No image paragraph before {paragraph.text!r}")
    return Paragraph(element, paragraph._parent)


def replace_flowchart(document: Document, caption_text: str, image_name: str) -> None:
    caption = find_caption(document, caption_text)
    image_paragraph = previous_paragraph(caption)
    if not image_paragraph._p.xpath(".//w:drawing"):
        raise RuntimeError(f"Expected drawing before {caption_text!r}")

    page_break_before = image_paragraph.paragraph_format.page_break_before
    keep_with_next = image_paragraph.paragraph_format.keep_with_next
    image_paragraph._element.getparent().remove(image_paragraph._element)

    replacement = caption.insert_paragraph_before()
    replacement.alignment = WD_ALIGN_PARAGRAPH.CENTER
    replacement.paragraph_format.page_break_before = page_break_before
    replacement.paragraph_format.keep_with_next = keep_with_next
    replacement.paragraph_format.space_before = Pt(0)
    replacement.paragraph_format.space_after = Pt(0)
    inline = replacement.add_run().add_picture(
        str(FIGURES / image_name), width=Inches(FLOWCHART_WIDTH)
    )
    inline._inline.docPr.set("title", caption_text)
    inline._inline.docPr.set("descr", caption_text)


def next_table(paragraph: Paragraph) -> Table:
    element = paragraph._p.getnext()
    while element is not None:
        if element.tag.endswith("}tbl"):
            return Table(element, paragraph._parent)
        if element.tag.endswith("}p") and Paragraph(element, paragraph._parent).text.strip():
            raise RuntimeError(f"Unexpected paragraph between caption and table: {paragraph.text!r}")
        element = element.getnext()
    raise RuntimeError(f"No table after caption {paragraph.text!r}")


def set_cell_margins(cell, top=80, start=60, bottom=80, end=60) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table: Table, widths_cm: list[float]) -> None:
    if len(table.columns) != len(widths_cm):
        raise RuntimeError(
            f"Expected {len(widths_cm)} columns but found {len(table.columns)}"
        )
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(sum(int(Cm(width).twips) for width in widths_cm)))
    tbl_w.set(qn("w:type"), "dxa")

    grid_columns = table._tbl.tblGrid.gridCol_lst
    for index, width_cm in enumerate(widths_cm):
        width_twips = int(Cm(width_cm).twips)
        table.columns[index].width = Cm(width_cm)
        if index < len(grid_columns):
            grid_columns[index].set(qn("w:w"), str(width_twips))
        for cell in table.columns[index].cells:
            cell.width = Cm(width_cm)
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(width_twips))
            tc_w.set(qn("w:type"), "dxa")


def format_appendix_c_table(table: Table, widths_cm: list[float]) -> None:
    for row in list(table.rows):
        if not any(cell.text.strip() for cell in row.cells):
            table._tbl.remove(row._tr)
    set_table_geometry(table, widths_cm)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.left_indent = Cm(0)
                paragraph.paragraph_format.right_indent = Cm(0)
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(13)


def appendix_c_tables(document: Document) -> list[tuple[int, Table]]:
    result: list[tuple[int, Table]] = []
    pattern = re.compile(r"^Table App\.C\.(\d+)\.")
    for paragraph in document.paragraphs:
        if paragraph.style.name != "Table Caption":
            continue
        match = pattern.match(paragraph.text.strip())
        if match:
            result.append((int(match.group(1)), next_table(paragraph)))
    result.sort(key=lambda item: item[0])
    if [number for number, _ in result] != list(range(1, 26)):
        raise RuntimeError(f"Expected Appendix C tables 1-25; found {[n for n, _ in result]}")
    return result


def main() -> None:
    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)

    document = Document(DOCX)
    before_tables = digest_tables(document)
    before_text = digest_body_text(document)

    for caption_text, image_name in FLOWCHARTS.items():
        replace_flowchart(document, caption_text, image_name)

    targets = appendix_c_tables(document)
    for number, table in targets:
        format_appendix_c_table(
            table,
            CATALOGUE_WIDTHS_CM if number == 1 else DICTIONARY_WIDTHS_CM,
        )

    if before_tables != digest_tables(document):
        raise RuntimeError("Appendix C table text changed during formatting")
    if before_text != digest_body_text(document):
        raise RuntimeError("Body paragraph text changed during formatting")

    document.save(DOCX)
    print(f"Updated flowcharts: {len(FLOWCHARTS)}")
    print(f"Left-aligned Appendix C tables: {len(targets)}")
    print(f"Saved: {DOCX}")
    print(f"Backup: {BACKUP}")


if __name__ == "__main__":
    main()
