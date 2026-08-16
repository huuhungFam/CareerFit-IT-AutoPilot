from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


path = Path(__file__).resolve().parents[1] / "Doc" / "CareerFit-Thesis-Report.docx"
doc = Document(path)

old = (
    "CareerFit is implemented as a role-based web platform for guests, candidates, recruiters, and administrators using a Spring Boot backend, a React and TypeScript frontend, and PostgreSQL persistence managed through Flyway migrations. The final backend suite passed 63 of 63 tests, the production frontend build completed with a largest JavaScript chunk of 375.64 kB, and four Chromium P0 workflows passed with cleanup of generated Job data. On a synthetic causal benchmark containing 50 Jobs and 100 CVs, Rocchio increased nDCG@5 from 0.037737 to 0.817737. The final benchmark log contained no optimistic-lock exception, and aggregate runtime health returned HTTP 200 with status UP. These results demonstrate controlled feedback behavior and integrated prototype workflows, not production hiring effectiveness."
)
new = (
    "CareerFit is implemented as a role-based web platform for guests, candidates, recruiters, and administrators using a Spring Boot backend, a React and TypeScript frontend, and PostgreSQL persistence managed through Flyway migrations. The refreshed backend suite passed 72 of 72 tests, the production frontend build completed with a largest JavaScript chunk of 378.44 kB, and all 20 Chromium workflow, contract, and resilience tests passed. On a synthetic causal benchmark containing 50 Jobs and 100 CVs, Rocchio increased nDCG@5 from 0.037737 to 0.837737. The final benchmark log contained no optimistic-lock exception, and aggregate runtime health returned HTTP 200 with status UP. These results demonstrate controlled feedback behavior and integrated prototype workflows, not production hiring effectiveness."
)
matches = [p for p in doc.paragraphs if p.text == old]
if len(matches) != 1:
    raise RuntimeError(f"Expected one abstract paragraph, found {len(matches)}")
matches[0].text = new


def set_cell(table, row, col, text):
    cell = table.rows[row].cells[col]
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
            run.font.size = Pt(11)


set_cell(doc.tables[24], 3, 2, "Largest JavaScript chunk was 378.44 kB; no Vite size warning")
set_cell(doc.tables[25], 1, 2, "API contracts and 20 Chromium workflow, contract, and resilience tests")
doc.save(path)
print(path)
