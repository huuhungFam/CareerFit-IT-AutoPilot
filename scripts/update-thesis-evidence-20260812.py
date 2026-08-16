from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
OUTPUT = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-evidence-20260812.docx"


def set_font(run) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), "Times New Roman")


def replace_paragraph_text(paragraph, replacements: list[tuple[str, str]]) -> bool:
    original = paragraph.text
    updated = original
    for old, new in replacements:
        updated = updated.replace(old, new)
    if updated == original:
        return False
    paragraph.clear()
    run = paragraph.add_run(updated)
    set_font(run)
    return True


def main() -> None:
    status = subprocess.run(
        ["git", "status", "--short"], cwd=ROOT, check=True, capture_output=True, text=True
    ).stdout.splitlines()
    dirty_count = len(status)

    doc = Document(DOCX)
    replacements = [
        ("August 7, 2026", "August 12, 2026"),
        ("August 7", "August 12"),
        ("209 modified or untracked entries", f"{dirty_count} modified or untracked entries"),
        ("102.453 seconds", "164.516 seconds"),
        ("102.453 s", "164.516 s"),
        ("All 46 Chrome", "All 49 Chrome"),
        ("46 integrated Chrome", "49 integrated Chrome"),
        ("46-test Chrome", "49-test Chrome"),
        ("46/46 integrated Chrome", "49/49 integrated Chrome"),
        ("The 46 Chrome tests", "The 49 Chrome tests"),
    ]

    changed_paragraphs = 0
    for paragraph in doc.paragraphs:
        if replace_paragraph_text(paragraph, replacements):
            changed_paragraphs += 1

    changed_cells = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if replace_paragraph_text(paragraph, replacements):
                        changed_cells += 1

    # Record the only extra harness condition needed by the refreshed E2E run.
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Earlier commands are recorded in evidence/CHAPTER5_EVIDENCE_20260703.md"):
            paragraph.add_run(
                " The isolated stack exposes the application at localhost:18080, while one P0 request helper "
                "still targets localhost:8080. The final 49-test run therefore used a temporary loopback proxy "
                "from port 8080 to 18080; no frontend or backend implementation code was changed for this run."
            )
            for run in paragraph.runs:
                set_font(run)
            break

    doc.save(OUTPUT)
    DOCX.write_bytes(OUTPUT.read_bytes())
    print(f"dirty_count={dirty_count}")
    print(f"changed_paragraphs={changed_paragraphs}")
    print(f"changed_table_paragraphs={changed_cells}")
    print(DOCX)


if __name__ == "__main__":
    main()
