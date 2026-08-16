from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"C:\CODING\Thesis\demo\CV_Candidate_CF_Demo_Matching.docx")

NAVY = RGBColor(15, 42, 68)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(85, 96, 110)
LIGHT_BLUE = "DCE6F1"


def set_run_font(run, *, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_bottom_border(paragraph, color="2E74B5", size="16", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(25, 31, 40)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_1 = doc.styles["Heading 1"]
    heading_1.font.name = "Calibri"
    heading_1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading_1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading_1.font.size = Pt(16)
    heading_1.font.bold = True
    heading_1.font.color.rgb = BLUE
    heading_1.paragraph_format.space_before = Pt(16)
    heading_1.paragraph_format.space_after = Pt(8)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = doc.styles["Heading 2"]
    heading_2.font.name = "Calibri"
    heading_2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading_2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading_2.font.size = Pt(13)
    heading_2.font.bold = True
    heading_2.font.color.rgb = BLUE
    heading_2.paragraph_format.space_before = Pt(12)
    heading_2.paragraph_format.space_after = Pt(6)
    heading_2.paragraph_format.keep_with_next = True

    heading_3 = doc.styles["Heading 3"]
    heading_3.font.name = "Calibri"
    heading_3._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading_3._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading_3.font.size = Pt(12)
    heading_3.font.bold = True
    heading_3.font.color.rgb = RGBColor(31, 77, 120)
    heading_3.paragraph_format.space_before = Pt(8)
    heading_3.paragraph_format.space_after = Pt(4)
    heading_3.paragraph_format.keep_with_next = True


def add_body_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.10
    paragraph.paragraph_format.widow_control = True
    run = paragraph.add_run(text)
    set_run_font(run, size=10.8, color=RGBColor(25, 31, 40))
    return paragraph


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.right_margin = Inches(0.82)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    # Named page-geometry override for a one-page CV: 0.72/0.82-inch margins.
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(2)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("PHẠM HỮU HƯNG")
    set_run_font(run, size=26, color=NAVY, bold=True)

    role = doc.add_paragraph()
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(5)
    role.paragraph_format.keep_with_next = True
    run = role.add_run("FRONTEND ENGINEER")
    set_run_font(run, size=13.5, color=BLUE, bold=True)

    metadata = doc.add_paragraph()
    metadata.paragraph_format.space_before = Pt(0)
    metadata.paragraph_format.space_after = Pt(12)
    metadata.paragraph_format.keep_with_next = True
    run = metadata.add_run("CAREERFIT DEMO LAB   |   CẦN THƠ, VIỆT NAM")
    set_run_font(run, size=10.5, color=MUTED, bold=True)
    add_bottom_border(metadata, color="2E74B5", size="12", space="7")

    profile = doc.add_paragraph()
    profile.paragraph_format.space_before = Pt(0)
    profile.paragraph_format.space_after = Pt(7)
    profile.paragraph_format.keep_with_next = True
    shade_paragraph(profile, LIGHT_BLUE)
    run = profile.add_run("HỒ SƠ MỤC TIÊU")
    set_run_font(run, size=11.2, color=NAVY, bold=True)

    paragraphs = [
        "Kỹ sư Frontend có 4 năm kinh nghiệm phát triển ứng dụng web bằng React và TypeScript. Thành thạo JavaScript, Redux, Next.js, Vite, HTML, CSS, thiết kế responsive, accessibility, Jest, REST API, Git, Agile và Scrum.",
        "Xây dựng giao diện tìm kiếm việc làm bằng React và TypeScript, phát triển thư viện UI component tái sử dụng, quản lý trạng thái Redux, tích hợp REST API và kiểm thử tự động bằng Jest. Tối ưu hiệu năng tải trang 35 phần trăm bằng Vite, code splitting và caching.",
        "Phát triển CareerFit IT AutoPilot, dashboard phân tích dữ liệu, biểu đồ tương tác, form xác thực, thông báo thời gian thực và trải nghiệm người dùng trên desktop lẫn mobile. Phối hợp cùng backend, product và designer trong quy trình Agile Scrum.",
        "Tốt nghiệp Kỹ sư Công nghệ Thông tin. Có khả năng đọc hiểu tài liệu tiếng Anh, giao tiếp tốt, giải quyết vấn đề và làm việc độc lập.",
    ]
    for text in paragraphs:
        add_body_paragraph(doc, text)

    experience = doc.add_paragraph()
    experience.paragraph_format.space_before = Pt(5)
    experience.paragraph_format.space_after = Pt(7)
    experience.paragraph_format.keep_with_next = True
    shade_paragraph(experience, LIGHT_BLUE)
    run = experience.add_run("KINH NGHIỆM")
    set_run_font(run, size=11.2, color=NAVY, bold=True)

    exp = doc.add_paragraph()
    exp.paragraph_format.space_before = Pt(0)
    exp.paragraph_format.space_after = Pt(0)
    exp.paragraph_format.line_spacing = 1.10
    run = exp.add_run("Frontend Engineer tại CareerFit Demo Lab từ 2022 đến 2026.")
    set_run_font(run, size=10.8, color=RGBColor(25, 31, 40), bold=True)

    doc.core_properties.title = "CV Candidate CareerFit Demo Matching"
    doc.core_properties.subject = "CV dùng cho kịch bản demo đối sánh 100-96-94"
    doc.core_properties.author = "CareerFit Demo"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
