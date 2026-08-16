from __future__ import annotations

from collections import Counter
from pathlib import Path
import json
import re
import unicodedata
import zipfile

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
OUT = ROOT / "Doc" / "working" / "review-20260718.json"


def iter_items(doc: Document):
    heading = "Front matter"
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        if paragraph.style.name.startswith("Heading") and text.strip():
            heading = text.strip()
        yield {
            "kind": "paragraph",
            "location": f"P{index}",
            "heading": heading,
            "style": paragraph.style.name,
            "text": text,
        }
    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            text = " | ".join(cell.text.replace("\n", " / ") for cell in row.cells)
            yield {
                "kind": "table",
                "location": f"T{table_index}R{row_index}",
                "heading": "Table content",
                "style": "TABLE",
                "text": text,
            }


def is_suspicious_character(ch: str) -> bool:
    code = ord(ch)
    category = unicodedata.category(ch)
    if ch in "\n\r\t":
        return False
    if category in {"Cc", "Cf", "Cs", "Co", "Cn"}:
        return True
    if ch == "\ufffd":
        return True
    if code in {0x00AD, 0x200B, 0x200C, 0x200D, 0x2060, 0xFEFF}:
        return True
    return False


def main() -> None:
    doc = Document(DOCX)
    items = list(iter_items(doc))
    suspicious = []
    repeated_words = []
    spacing = []
    non_ascii = Counter()
    terminology = []

    term_patterns = {
        "raw email-action token claim": re.compile(r"random token strings|raw token.*persist", re.I),
        "mock/fallback claim": re.compile(r"mock data remains|mapper fallback|mock login|fallback development", re.I),
        "recruiter feedback claim": re.compile(r"recruiter.{0,50}feedback|feedback.{0,50}recruiter", re.I),
        "localStorage claim": re.compile(r"localStorage", re.I),
        "portfolio claim": re.compile(r"portfolio", re.I),
        "async matching claim": re.compile(r"async|asynchronous|after commit", re.I),
    }

    for item in items:
        text = item["text"]
        for ch in text:
            if ord(ch) > 127:
                non_ascii[f"U+{ord(ch):04X} {unicodedata.name(ch, 'UNKNOWN')} {ch!r}"] += 1
            if is_suspicious_character(ch):
                suspicious.append({**item, "character": f"U+{ord(ch):04X}"})
        for match in re.finditer(r"\b([A-Za-z][A-Za-z'-]*)\s+\1\b", text, re.I):
            repeated_words.append({**item, "match": match.group(0)})
        if re.search(r" {2,}|\s+[,.!?;:]|[,.!?;:][A-Za-z]", text):
            spacing.append(item)
        for label, pattern in term_patterns.items():
            if pattern.search(text):
                terminology.append({"label": label, **item})

    table_headers = []
    for index, table in enumerate(doc.tables):
        first = table.rows[0]._tr if table.rows else None
        has_header = first is not None and first.find(qn("w:trPr")) is not None and first.find(qn("w:trPr")).find(qn("w:tblHeader")) is not None
        table_headers.append({"table": index, "rows": len(table.rows), "columns": len(table.columns), "has_header": has_header, "first_row": " | ".join(c.text for c in table.rows[0].cells) if table.rows else ""})

    drawings = []
    with zipfile.ZipFile(DOCX) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        for match in re.finditer(r"<wp:docPr\b([^>]*)/>", document_xml):
            attrs = match.group(1)
            def attr(name: str) -> str:
                found = re.search(fr'{name}="([^"]*)"', attrs)
                return found.group(1) if found else ""
            drawings.append({"id": attr("id"), "name": attr("name"), "title": attr("title"), "descr": attr("descr")})

    result = {
        "document": str(DOCX),
        "counts": {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "inline_images": len(doc.inline_shapes),
            "suspicious_characters": len(suspicious),
            "repeated_words": len(repeated_words),
            "spacing_flags": len(spacing),
            "drawings_missing_alt": sum(not (d["title"] or d["descr"]) for d in drawings),
            "tables_missing_header_flag": sum(not t["has_header"] for t in table_headers),
        },
        "suspicious_characters": suspicious,
        "repeated_words": repeated_words,
        "spacing_flags": spacing,
        "non_ascii": non_ascii.most_common(),
        "terminology": terminology,
        "table_headers": table_headers,
        "drawings": drawings,
    }
    OUT.parent.mkdir(parents=True, exist_ok=True)
    OUT.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(result["counts"], ensure_ascii=False, indent=2))
    print("\nSUSPICIOUS")
    for row in suspicious[:50]:
        print(row)
    print("\nREPEATED WORDS")
    for row in repeated_words:
        print(row)
    print("\nTABLES WITHOUT HEADER FLAG")
    for row in table_headers:
        if not row["has_header"]:
            print(row)


if __name__ == "__main__":
    main()
