from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
LOGO = ROOT / "Doc" / "working" / "ctu-template-logo.png"
_fallback_list_number = 0


def element_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t"))).strip()


def is_paragraph(element) -> bool:
    return element.tag == qn("w:p")


def is_table(element) -> bool:
    return element.tag == qn("w:tbl")


def paragraph(element, document: Document) -> Paragraph:
    return Paragraph(element, document._body)


def has_drawing(element) -> bool:
    return bool(element.findall(".//" + qn("w:drawing")))


def has_section_break(element) -> bool:
    return (
        is_paragraph(element)
        and element.find("./" + qn("w:pPr") + "/" + qn("w:sectPr")) is not None
    )


def find_index(elements, exact_text: str) -> int:
    for index, element in enumerate(elements):
        if element_text(element) == exact_text:
            return index
    raise RuntimeError(f"Could not find block marker: {exact_text}")


def detach(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    return element


def make_paragraph(
    document: Document,
    text: str = "",
    style: str | None = None,
    *,
    alignment=None,
    bold: bool | None = None,
    italic: bool | None = None,
    size: float | None = None,
    space_before: float | None = None,
    space_after: float | None = None,
    line_spacing: float | None = None,
    first_indent_cm: float | None = None,
    page_break_before: bool | None = None,
):
    global _fallback_list_number
    if style == "Heading 3":
        _fallback_list_number = 0
    available_styles = {item.name for item in document.styles}
    resolved_style = style
    if style in {"List Number", "List Bullet"} and style not in available_styles:
        resolved_style = "List Paragraph"
        if style == "List Number":
            _fallback_list_number += 1
            text = f"{_fallback_list_number}. {text}"
        else:
            text = f"• {text}"
    p = document.add_paragraph(style=resolved_style)
    run = p.add_run(text)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if alignment is not None:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing
    if first_indent_cm is not None:
        p.paragraph_format.first_line_indent = Cm(first_indent_cm)
    if page_break_before is not None:
        p.paragraph_format.page_break_before = page_break_before
    return detach(p._p)


def make_page_break(document: Document):
    p = document.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    p.paragraph_format.space_after = Pt(0)
    return detach(p._p)


def set_paragraph_text(element, document: Document, text: str):
    p = paragraph(element, document)
    p.text = text
    for run in p.runs:
        run.font.name = "Times New Roman"
        run._element.get_or_add_rPr().rFonts.set(
            qn("w:eastAsia"), "Times New Roman"
        )


def make_cover_line(
    document: Document,
    text: str = "",
    *,
    size: float = 13,
    bold: bool = False,
    after: float = 0,
    before: float = 0,
):
    return make_paragraph(
        document,
        text,
        "Normal",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=bold,
        size=size,
        space_before=before,
        space_after=after,
        line_spacing=1.0,
        first_indent_cm=0,
    )


def make_logo_paragraph(document: Document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(12)
    p.add_run().add_picture(str(LOGO), width=Inches(1.05))
    return detach(p._p)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 11):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def make_inner_cover_table(document: Document):
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(7.2)
    table.columns[1].width = Cm(7.2)
    remove_table_borders(table)
    set_cell_text(
        table.cell(0, 0),
        "Supervisor\nPh.D. Nguyen Thanh Khoa",
        bold=False,
        size=13,
    )
    set_cell_text(
        table.cell(0, 1),
        "Student\nPham Huu Hung - B2203557",
        bold=False,
        size=13,
    )
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                if run.text in {"Supervisor", "Student"}:
                    run.bold = True
    return detach(table._tbl)


def build_cover(document: Document, include_supervisor: bool) -> list:
    result = [
        make_cover_line(
            document, "MINISTRY OF EDUCATION AND TRAINING", size=13, after=0
        ),
        make_cover_line(document, "CAN THO UNIVERSITY", size=14, bold=True),
        make_cover_line(
            document,
            "COLLEGE OF INFORMATION AND COMMUNICATION TECHNOLOGY",
            size=14,
            bold=True,
        ),
        make_cover_line(
            document, "FACULTY OF SOFTWARE ENGINEERING", size=14, bold=True
        ),
        make_logo_paragraph(document),
        make_cover_line(
            document,
            "SOFTWARE DEVELOPMENT THESIS PROJECT",
            size=15,
            bold=True,
            before=8,
        ),
        make_cover_line(document, "COURSE: CT250H", size=14, bold=True, after=18),
        make_cover_line(document, "Thesis title", size=13, bold=True, after=6),
        make_cover_line(
            document,
            "DESIGN AND IMPLEMENTATION OF A HUMAN-IN-THE-LOOP",
            size=15,
            bold=True,
        ),
        make_cover_line(
            document,
            "AI-ASSISTED RECRUITMENT AUTOMATION PLATFORM FOR",
            size=15,
            bold=True,
        ),
        make_cover_line(
            document,
            "CV-JD EVALUATION AND RECOMMENDATION IN IT",
            size=15,
            bold=True,
            after=24,
        ),
    ]
    if include_supervisor:
        result.append(make_inner_cover_table(document))
    else:
        result.extend(
            [
                make_cover_line(
                    document, "Student", size=13, bold=True, before=4
                ),
                make_cover_line(
                    document, "Pham Huu Hung - B2203557", size=13, bold=False
                ),
            ]
        )
    result.append(
        make_cover_line(
            document, "Can Tho, August 2026", size=13, bold=True, before=28
        )
    )
    return result


def make_signature_block(document: Document, label: str, name: str):
    return [
        make_paragraph(
            document,
            "Can Tho, August 2026",
            "Normal",
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            space_before=18,
            space_after=0,
            line_spacing=1.0,
            first_indent_cm=0,
        ),
        make_paragraph(
            document,
            label,
            "Normal",
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            bold=True,
            space_after=28,
            line_spacing=1.0,
            first_indent_cm=0,
        ),
        make_paragraph(
            document,
            name,
            "Normal",
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            space_after=0,
            line_spacing=1.0,
            first_indent_cm=0,
        ),
    ]


def build_declaration(document: Document) -> list:
    result = [
        make_paragraph(
            document,
            "DECLARATION OF ORIGINALITY",
            "Heading 1",
            page_break_before=True,
        ),
        make_paragraph(
            document,
            (
                "I hereby declare that this thesis, titled “Design and "
                "Implementation of a Human-in-the-Loop AI-Assisted Recruitment "
                "Automation Platform for CV-JD Evaluation and Recommendation in "
                "IT,” is my original work completed under the supervision of "
                "Ph.D. Nguyen Thanh Khoa. The implementation, experiments, data "
                "analysis, and conclusions presented in this report are reported "
                "truthfully within the stated scope."
            ),
            "Normal",
        ),
        make_paragraph(
            document,
            (
                "All ideas, publications, standards, software documentation, "
                "figures, and other materials obtained from external sources are "
                "acknowledged through citations or source notes. I accept full "
                "responsibility for the accuracy, academic integrity, and "
                "originality of the submitted work."
            ),
            "Normal",
        ),
    ]
    result.extend(make_signature_block(document, "Student", "Pham Huu Hung"))
    return result


def build_supervisor_comments(document: Document) -> list:
    result = [
        make_paragraph(
            document,
            "SUPERVISOR'S COMMENTS",
            "Heading 1",
            page_break_before=True,
        )
    ]
    for _ in range(17):
        result.append(
            make_paragraph(
                document,
                "....................................................................................................................",
                "Normal",
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_after=3,
                line_spacing=1.0,
                first_indent_cm=0,
            )
        )
    result.extend(
        make_signature_block(
            document, "Supervisor", "Ph.D. Nguyen Thanh Khoa"
        )
    )
    return result


def build_abstract(document: Document) -> list:
    return [
        make_paragraph(
            document, "ABSTRACT", "Heading 1", page_break_before=True
        ),
        make_paragraph(
            document,
            (
                "Background: Recruitment in the information technology sector "
                "requires candidates and recruiters to compare many types of "
                "information from curricula vitae and job descriptions. Most job "
                "portals support posting, searching, and manual applications, "
                "while automated screening tools may not clearly explain how a "
                "score is calculated or how an automated action is controlled."
            ),
            "Normal",
        ),
        make_paragraph(
            document,
            (
                "Objectives: This thesis designs and implements CareerFit IT "
                "AutoPilot, a web-based recruitment platform that combines a job "
                "portal, CV-job description matching, personalized job "
                "recommendations, policy-based automation, and Human-in-the-Loop "
                "interaction in one auditable workflow."
            ),
            "Normal",
        ),
        make_paragraph(
            document,
            (
                "Methods: Recruitment text is normalized and represented with "
                "Term Frequency-Inverse Document Frequency vectors. Cosine "
                "similarity ranks CV-job pairs and candidate-job recommendations, "
                "while Rocchio relevance feedback updates learned "
                "representations from explicit feedback. AutoFit checks scores, "
                "consent, thresholds, interaction history, quotas, cooldowns, "
                "time zones, and quiet hours before choosing an action. CareerFit "
                "uses a Spring Boot backend, a React and TypeScript frontend, and "
                "PostgreSQL persistence managed through Flyway migrations."
            ),
            "Normal",
        ),
        make_paragraph(
            document,
            (
                "Results: The refreshed backend suite passed 72 of 72 tests, the "
                "production frontend build completed successfully, and all 20 "
                "Chromium workflow, contract, and resilience tests passed. In the "
                "controlled synthetic benchmark, nDCG@5 increased from 0.037737 "
                "to 0.837737 after Rocchio feedback. These results support a "
                "well-tested academic prototype, but they do not establish "
                "production readiness or real-world hiring effectiveness."
            ),
            "Normal",
        ),
        make_paragraph(
            document,
            (
                "Keywords: recruitment automation; CV-job matching; "
                "recommendation system; Human-in-the-Loop; Rocchio feedback; "
                "explainable automation."
            ),
            "Normal",
            italic=True,
            first_indent_cm=0,
        ),
    ]


def make_table_caption(document: Document, text: str):
    return make_paragraph(
        document,
        text,
        "Table Caption",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
        size=11,
        first_indent_cm=0,
    )


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_cell_width(cell, width_cm: float):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).emu / 635)))
    tc_w.set(qn("w:type"), "dxa")


def make_use_case_table(document: Document, rows: list[tuple[str, str]]):
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(3.7)
    table.columns[1].width = Cm(11.3)
    set_cell_text(table.cell(0, 0), "Field", bold=True, size=10.5)
    set_cell_text(table.cell(0, 1), "Description", bold=True, size=10.5)
    set_repeat_table_header(table.rows[0])
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, bold=True, size=10.5)
        set_cell_text(cells[1], value, bold=False, size=10.5)
        set_table_cell_width(cells[0], 3.7)
        set_table_cell_width(cells[1], 11.3)
    return detach(table._tbl)


def replace_text_nodes(elements: list, mapping: dict[str, str]):
    tokens = {key: f"[[[MAP_{index:04d}]]]" for index, key in enumerate(mapping)}
    for element in elements:
        for text_node in element.iter(qn("w:t")):
            text = text_node.text or ""
            for key, token in tokens.items():
                if key in text:
                    text = text.replace(key, token)
            for key, token in tokens.items():
                if token in text:
                    text = text.replace(token, mapping[key])
            text_node.text = text


def style_table(table, *, cover: bool = False):
    if cover:
        return
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run._element.get_or_add_rPr().rFonts.set(
                        qn("w:eastAsia"), "Times New Roman"
                    )
                    run.font.size = Pt(10.5)
                    if row_index == 0:
                        run.bold = True


def add_page_border(section):
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:pgBorders"))
    if existing is not None:
        sect_pr.remove(existing)
    borders = OxmlElement("w:pgBorders")
    borders.set(qn("w:offsetFrom"), "page")
    borders.set(qn("w:display"), "allPages")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "12")
        border.set(qn("w:space"), "24")
        border.set(qn("w:color"), "000000")
        borders.append(border)
    sect_pr.append(borders)


def style_document(document: Document):
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.first_line_indent = Cm(0.75)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    heading_1 = document.styles["Heading 1"]
    heading_1.font.name = "Times New Roman"
    heading_1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading_1.font.size = Pt(16)
    heading_1.font.bold = True
    heading_1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_1.paragraph_format.page_break_before = True
    heading_1.paragraph_format.space_before = Pt(4)
    heading_1.paragraph_format.space_after = Pt(6)
    heading_1.paragraph_format.first_line_indent = Cm(0)

    heading_2 = document.styles["Heading 2"]
    heading_2.font.name = "Times New Roman"
    heading_2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading_2.font.size = Pt(13)
    heading_2.font.bold = True
    heading_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_2.paragraph_format.line_spacing = 1.2
    heading_2.paragraph_format.space_before = Pt(6)
    heading_2.paragraph_format.space_after = Pt(0)
    heading_2.paragraph_format.first_line_indent = Cm(0)

    heading_3 = document.styles["Heading 3"]
    heading_3.font.name = "Times New Roman"
    heading_3._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading_3.font.size = Pt(13)
    heading_3.font.bold = True
    heading_3.font.italic = True
    heading_3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_3.paragraph_format.line_spacing = 1.2
    heading_3.paragraph_format.space_before = Pt(2)
    heading_3.paragraph_format.space_after = Pt(4)
    heading_3.paragraph_format.first_line_indent = Cm(0)

    if "Heading 4" in [style.name for style in document.styles]:
        heading_4 = document.styles["Heading 4"]
        heading_4.font.name = "Times New Roman"
        heading_4._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        heading_4.font.size = Pt(13)
        heading_4.font.bold = True
        heading_4.paragraph_format.line_spacing = 1.2
        heading_4.paragraph_format.first_line_indent = Cm(0)

    for caption_name in ("Figure Caption", "Table Caption"):
        style = document.styles[caption_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(11)
        style.font.italic = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.first_line_indent = Cm(0)

    for index, section in enumerate(document.sections):
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.orientation = WD_ORIENT.PORTRAIT
        if index == 0:
            section.top_margin = Cm(3)
            section.bottom_margin = Cm(3)
            section.left_margin = Cm(3.5)
            section.right_margin = Cm(2)
            add_page_border(section)
        else:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2)
        section.header_distance = Cm(1.27)
        section.footer_distance = Cm(1.27)


def rebuild_appendices(document: Document) -> list:
    elements = [
        make_paragraph(document, "APPENDICES", "Heading 1"),
        make_paragraph(
            document,
            "Appendix A. Requirement-Test Traceability Matrix",
            "Heading 2",
        ),
        make_paragraph(
            document,
            (
                "The core traceability chain is: authentication and role controls "
                "to SecurityHardeningTest and P0 login flows; CV ingestion to "
                "service and integration tests plus the Candidate upload flow; "
                "matching and feedback to AlgorithmEvaluatorTest; Job lifecycle "
                "to the Recruiter create/verify/delete flow; administration to "
                "suspend/activate workflows; and operations to Actuator health, "
                "build, and runtime evidence."
            ),
            "Normal",
        ),
        make_paragraph(
            document, "Appendix B. Selected API Contracts", "Heading 2"
        ),
        make_paragraph(
            document,
            (
                "Representative endpoints are POST /api/auth/login; GET and POST "
                "/api/jobs; POST /api/cvs/upload; GET /api/matchings; POST "
                "/api/feedback; GET and PATCH /api/automation/policies/me; GET "
                "then POST /api/email-action/redeem; and GET /actuator/health. "
                "Protected endpoints require a signed JWT and enforce role or "
                "ownership rules in the backend."
            ),
            "Normal",
        ),
        make_paragraph(document, "Appendix C. Data Model Summary", "Heading 2"),
        make_paragraph(
            document,
            (
                "Identity data is centered on Account and role-specific profiles. "
                "Recruitment data includes Job, CV, Matching, Application, "
                "Invitation, and Feedback. Automation and operations use "
                "AutomationPolicy, EmailAction, EmailToken, Notification, "
                "AuditLog, and scheduler state. Flyway migrations provide the "
                "reproducible schema history, and action tokens are persisted as "
                "SHA-256 hashes."
            ),
            "Normal",
        ),
        make_paragraph(document, "Appendix D. Evaluation Summary", "Heading 2"),
        make_paragraph(
            document,
            (
                "The final evidence package contains the backend test log, "
                "controlled algorithm benchmark, frontend production build, "
                "Chromium workflow results, runtime health response, screenshots, "
                "and evaluation/result.json. Chapter 4 reports the observed "
                "results and limitations; the artifacts support demonstration "
                "and thesis-defense readiness rather than a production deployment "
                "claim."
            ),
            "Normal",
        ),
        make_paragraph(
            document,
            "Appendix E. UAT and Demonstration Script",
            "Heading 2",
            page_break_before=True,
        ),
        make_paragraph(document, "Preparation", "Heading 3"),
        make_paragraph(
            document,
            "Start PostgreSQL, the backend, and the frontend; verify aggregate health and the public Job API; and prepare separate Candidate, Recruiter, and Administrator accounts.",
            "List Number",
            first_indent_cm=0,
        ),
        make_paragraph(
            document,
            "Record the current Git commit, working-tree status, configuration profile, database source, and evaluation artifact locations.",
            "List Number",
            first_indent_cm=0,
        ),
        make_paragraph(
            document,
            "Use clearly marked test records and note the cleanup action for every Job, CV, application, invitation, or policy created during the demonstration.",
            "List Number",
            first_indent_cm=0,
        ),
        make_paragraph(document, "Execution", "Heading 3"),
        make_paragraph(
            document,
            "As a Guest, search for an IT Job, apply filters, and open Job and employer details.",
            "List Number",
            first_indent_cm=0,
        ),
        make_paragraph(
            document,
            "As a Candidate, sign in, upload or select a CV, wait for processing, inspect ranked matches and reasons, apply to a Job, withdraw where permitted, and submit feedback.",
            "List Number",
            first_indent_cm=0,
        ),
        make_paragraph(
            document,
            "As a Recruiter, create a test Job, inspect applicants and discovered candidates, invite a Candidate, update application status, and remove the test Job after verification.",
            "List Number",
            first_indent_cm=0,
        ),
        make_paragraph(
            document,
            "Configure Candidate AutoFit thresholds, cooldown, quota, notification options, quiet hours, and human-control settings; then use run-now only for a controlled verification.",
            "List Number",
            first_indent_cm=0,
        ),
        make_paragraph(
            document,
            "As an Administrator, inspect user, Job, audit, notification, and email-action records; suspend and reactivate a designated test account.",
            "List Number",
            first_indent_cm=0,
        ),
        make_paragraph(
            document,
            "Show Actuator health, the archived backend and Chromium test results, the benchmark artifact, and the final cleanup state.",
            "List Number",
            first_indent_cm=0,
        ),
        make_paragraph(document, "Acceptance and cleanup", "Heading 3"),
        make_paragraph(
            document,
            "Record each expected and observed result, preserve any failure evidence, and do not convert an unverified step into a passing claim.",
            "List Number",
            first_indent_cm=0,
        ),
        make_paragraph(
            document,
            "Delete or deactivate demonstration data, restore account state, and confirm that background logs contain no unhandled exception.",
            "List Number",
            first_indent_cm=0,
        ),
        make_paragraph(
            document,
            "Archive the final evidence with a clean commit or a complete source-state package.",
            "List Number",
            first_indent_cm=0,
        ),
        make_paragraph(
            document,
            "Appendix F. Local Deployment Instructions",
            "Heading 2",
            page_break_before=True,
        ),
        make_paragraph(document, "Prerequisites", "Heading 3"),
        make_paragraph(
            document,
            "Install Java 21, Node.js with npm, Docker Desktop with Docker Compose, and Git.",
            "List Bullet",
            first_indent_cm=0,
        ),
        make_paragraph(
            document,
            "Provide environment-specific database, JWT, CORS, application URL, mail, OCR, and storage configuration. Do not commit production secrets.",
            "List Bullet",
            first_indent_cm=0,
        ),
        make_paragraph(document, "Startup procedure", "Heading 3"),
        make_paragraph(
            document,
            "From the repository root, start PostgreSQL with `docker compose up -d postgres` and wait for the container health check to pass.",
            "List Number",
            first_indent_cm=0,
        ),
        make_paragraph(
            document,
            "From `Backend/careerfit-backend`, run `mvnw.cmd spring-boot:run`. The host backend connects to Compose PostgreSQL at `localhost:5433`.",
            "List Number",
            first_indent_cm=0,
        ),
        make_paragraph(
            document,
            "From `Frontend`, run `npm install` when dependencies are not present, followed by `npm run dev`. The local Vite application is available at `http://localhost:5173`.",
            "List Number",
            first_indent_cm=0,
        ),
        make_paragraph(document, "Verification", "Heading 3"),
        make_paragraph(
            document,
            "Open `http://localhost:8080/actuator/health` and verify the expected aggregate, liveness, and readiness responses for the selected profile.",
            "List Number",
            first_indent_cm=0,
        ),
        make_paragraph(
            document,
            "Open the frontend, verify public Job search, sign in with a designated test account, and confirm authenticated API access.",
            "List Number",
            first_indent_cm=0,
        ),
        make_paragraph(
            document,
            "If OCR, mail, or another optional dependency is disabled, record the limitation instead of treating the missing external service as a verified production integration.",
            "List Number",
            first_indent_cm=0,
        ),
        make_paragraph(document, "Shutdown and cleanup", "Heading 3"),
        make_paragraph(
            document,
            "Stop the frontend and backend processes, then use `docker compose down` to stop the local containers. Remove volumes only when test data loss is intended.",
            "List Number",
            first_indent_cm=0,
        ),
        make_paragraph(
            document,
            "Appendix G. Role-Based User Guide",
            "Heading 2",
            page_break_before=True,
        ),
        make_paragraph(document, "Guest", "Heading 3"),
        make_paragraph(
            document,
            "Open the Jobs page, search by title, skill, or company, select filters, and open a Job card to view its public details. Authentication is required before application or personalized scoring.",
            "Normal",
        ),
        make_paragraph(document, "Candidate", "Heading 3"),
        make_paragraph(
            document,
            "Sign in, open My CVs, upload a PDF/DOCX/image file or create a CV manually, and wait until processing reaches a terminal status. Select the default CV, review matching cards and reasons, manage applications, provide feedback, inspect recommendations, and configure AutoFit settings.",
            "Normal",
        ),
        make_paragraph(document, "Recruiter", "Heading 3"),
        make_paragraph(
            document,
            "Sign in to the Recruiter workspace, create a Job with the required structured fields and description, inspect applicants and discovered candidates, send invitations, and update application status only for owned Jobs. Remove demonstration Jobs after testing.",
            "Normal",
        ),
        make_paragraph(document, "Administrator", "Heading 3"),
        make_paragraph(
            document,
            "Use the administrative workspace to inspect users, Jobs, audit records, notification and email-action state, and supported maintenance views. Suspend or reactivate only a designated test account and verify the resulting audit record.",
            "Normal",
        ),
    ]
    return elements


document = Document(REPORT)
if any(
    p.text.strip() == "CHAPTER 1. PROBLEM DESCRIPTION AND REQUIREMENTS"
    for p in document.paragraphs
):
    raise RuntimeError("The report already uses the July 26 four-chapter structure.")

body = document._body._element
children = list(body)
body_final_sectpr = children[-1]
if body_final_sectpr.tag != qn("w:sectPr"):
    raise RuntimeError("Expected final body section properties.")

markers = {
    name: find_index(children, name)
    for name in (
        "ACKNOWLEDGEMENTS",
        "ABSTRACT",
        "TABLE OF CONTENTS",
        "LIST OF FIGURES",
        "LIST OF TABLES",
        "LIST OF ABBREVIATIONS",
        "CHAPTER 1. INTRODUCTION",
        "CHAPTER 2. THEORETICAL BACKGROUND AND RELATED WORK",
        "CHAPTER 3. SYSTEM ANALYSIS AND DESIGN",
        "CHAPTER 4. SYSTEM IMPLEMENTATION",
        "CHAPTER 5. EXPERIMENTAL EVALUATION",
        "CHAPTER 6. RESULTS, DISCUSSION AND CONCLUSION",
        "REFERENCES",
        "APPENDICES",
    )
}

cover_old = children[: markers["ACKNOWLEDGEMENTS"]]
ack_block = children[markers["ACKNOWLEDGEMENTS"] : markers["ABSTRACT"]]
toc_block = children[
    markers["TABLE OF CONTENTS"] : markers["LIST OF FIGURES"]
]
figure_list_block = children[
    markers["LIST OF FIGURES"] : markers["LIST OF TABLES"]
]
table_list_block = children[
    markers["LIST OF TABLES"] : markers["LIST OF ABBREVIATIONS"]
]
abbreviation_block = children[
    markers["LIST OF ABBREVIATIONS"] : markers["CHAPTER 1. INTRODUCTION"]
]
intro_block = children[
    markers["CHAPTER 1. INTRODUCTION"] :
    markers["CHAPTER 2. THEORETICAL BACKGROUND AND RELATED WORK"]
]
theory_block = children[
    markers["CHAPTER 2. THEORETICAL BACKGROUND AND RELATED WORK"] :
    markers["CHAPTER 3. SYSTEM ANALYSIS AND DESIGN"]
]
analysis_block = children[
    markers["CHAPTER 3. SYSTEM ANALYSIS AND DESIGN"] :
    markers["CHAPTER 4. SYSTEM IMPLEMENTATION"]
]
implementation_block = children[
    markers["CHAPTER 4. SYSTEM IMPLEMENTATION"] :
    markers["CHAPTER 5. EXPERIMENTAL EVALUATION"]
]
evaluation_block = children[
    markers["CHAPTER 5. EXPERIMENTAL EVALUATION"] :
    markers["CHAPTER 6. RESULTS, DISCUSSION AND CONCLUSION"]
]
conclusion_block = children[
    markers["CHAPTER 6. RESULTS, DISCUSSION AND CONCLUSION"] :
    markers["REFERENCES"]
]
references_block = children[markers["REFERENCES"] : markers["APPENDICES"]]

cover_section_break = next(
    element for element in cover_old if has_section_break(element)
)
front_section_break = next(
    element for element in abbreviation_block if has_section_break(element)
)
abbreviation_block = [
    element for element in abbreviation_block if element is not front_section_break
]

design_start = find_index(analysis_block, "3.6 System Architecture")
requirements_source = analysis_block[:design_start]
design_block = analysis_block[design_start:]

# Build covers and new front matter.
cover_elements = build_cover(document, include_supervisor=False)
cover_elements.append(make_page_break(document))
cover_elements.extend(build_cover(document, include_supervisor=True))
cover_elements.append(cover_section_break)

declaration_block = build_declaration(document)
supervisor_comments_block = build_supervisor_comments(document)
abstract_block = build_abstract(document)

# Restructure Introduction.
intro_renames = {
    "CHAPTER 1. INTRODUCTION": "INTRODUCTION",
    "1.1 Problem Statement": "Problem Statement",
    "1.2 Motivation": "Background and Related Work",
    "1.3 Objectives": "Research Objectives",
    "1.3.1 Overall Objective": "Overall Objective",
    "1.3.2 Specific Objectives": "Specific Objectives",
    "1.4 Scope of the Thesis": "Research Subjects and Scope",
    "1.5 Contributions of the Thesis": "Key Contributions",
    "1.6 Thesis Organization": "Thesis Structure",
}
for element in intro_block:
    text = element_text(element)
    if text in intro_renames:
        set_paragraph_text(element, document, intro_renames[text])

motivation_end = find_index(intro_block, "Research Objectives")
intro_block.insert(
    motivation_end,
    make_paragraph(
        document,
        (
            "Related work includes classical lexical retrieval, learned "
            "resume-job representations, explainable recommendation, and "
            "human-centered governance. Lexical methods provide inspectable "
            "weights and matching reasons, while recent dense approaches improve "
            "semantic representation but require stronger data and evaluation "
            "assumptions [6]-[8]. Research on algorithmic hiring and explainable "
            "recommendation also shows that model performance, fairness, user "
            "control, and explanation quality must be evaluated separately "
            "[9], [10], [14]. CareerFit addresses a narrower integration gap by "
            "combining portal workflows, explainable scoring, feedback learning, "
            "policy-controlled actions, and audit records in one academic "
            "prototype."
        ),
        "Normal",
    ),
)

scope_heading = find_index(intro_block, "Research Subjects and Scope")
intro_block.insert(
    scope_heading + 1,
    make_paragraph(document, "Research Subject", "Heading 3"),
)
intro_block.insert(
    scope_heading + 2,
    make_paragraph(
        document,
        (
            "The research subject is the IT recruitment workflow supported by a "
            "web platform, including CV and Job-description text, lexical ranking "
            "and recommendation, explicit feedback signals, policy-controlled "
            "automation, and auditable user and system actions."
        ),
        "Normal",
    ),
)
intro_block.insert(
    scope_heading + 3,
    make_paragraph(document, "Research Scope", "Heading 3"),
)

contribution_index = find_index(intro_block, "Key Contributions")
method_content = [
    make_paragraph(document, "Research Methods and Content", "Heading 2"),
    make_paragraph(document, "Research Methods", "Heading 3"),
    make_paragraph(
        document,
        (
            "This thesis follows an applied software-engineering research "
            "approach. The work combines requirement analysis, architecture and "
            "data design, incremental implementation, a controlled algorithm "
            "experiment, automated backend and browser testing, runtime "
            "observation, and threats-to-validity analysis. Implementation claims "
            "are checked against source code, migrations, API contracts, logs, "
            "test reports, and generated evaluation artifacts."
        ),
        "Normal",
    ),
    make_paragraph(document, "Research Content", "Heading 3"),
    make_paragraph(
        document,
        (
            "The research models CareerFit as a Perception-Decision-Action-"
            "Learning-Audit workflow. It first identifies recruitment roles, "
            "requirements, and control boundaries; then designs and implements "
            "the data, matching, recommendation, feedback, automation, email "
            "action, and audit modules. Finally, the study evaluates algorithm "
            "behavior, software correctness, browser workflows, authorization, "
            "runtime health, and local latency within the stated experimental "
            "scope."
        ),
        "Normal",
    ),
]
intro_block[contribution_index:contribution_index] = method_content

structure_heading_index = find_index(intro_block, "Thesis Structure")
intro_block = intro_block[: structure_heading_index + 1] + [
    make_paragraph(
        document,
        (
            "The report contains an Introduction, four main chapters, and a "
            "separate Conclusion. Chapter 1 defines the problem, stakeholders, "
            "requirements, and use cases. Chapter 2 presents the theoretical "
            "background and the solution design. Chapter 3 explains the "
            "implementation in Spring Boot, React, PostgreSQL, and Flyway. "
            "Chapter 4 presents the evaluation method, verified evidence, and "
            "threats to validity. The Conclusion discusses achievements, "
            "limitations, future work, and the final contribution."
        ),
        "Normal",
    )
]

# Build Chapter 1 from the former requirements/use-case part of Chapter 3.
requirements_block = requirements_source[1:]
chapter_1_heading = make_paragraph(
    document,
    "CHAPTER 1. PROBLEM DESCRIPTION AND REQUIREMENTS",
    "Heading 1",
    page_break_before=False,
)
requirements_heading_map = {
    "3.1 System Analysis Context": "1.1 System Analysis Context",
    "3.2 Stakeholders and Actors": "1.2 Stakeholders and Actors",
    "3.3 Functional Requirements": "1.3 Functional Requirements",
    "3.4 Non-Functional Requirements": "1.4 Non-Functional Requirements",
    "3.5 Use-Case Analysis": "1.5 Use-Case Analysis",
    "3.5.1 Candidate Uploads a CV and Receives Matches":
        "1.5.1 Candidate Uploads a CV and Receives Matches",
    "3.5.2 Candidate Searches and Applies for a Job":
        "1.5.2 Candidate Searches, Applies, and Withdraws",
    "3.5.3 Recruiter Creates a Job and Reviews Candidates":
        "1.5.3 Recruiter Creates a Job and Reviews Candidates",
    "3.5.4 User Feedback Changes Later Ranking":
        "1.5.4 Candidate Feedback Changes Later Ranking",
    "3.5.5 AutoFit Executes a Policy-Governed Action":
        "1.5.5 AutoFit Executes a Policy-Governed Action",
    "3.5.6 Email Feedback Action":
        "1.5.6 User Completes an Email Feedback Action",
}
for element in requirements_block:
    text = element_text(element)
    if text in requirements_heading_map:
        set_paragraph_text(element, document, requirements_heading_map[text])

use_cases = [
    (
        "1.5.1 Candidate Uploads a CV and Receives Matches",
        "Table 1.5. Use case - Candidate uploads a CV and receives matches",
        [
            ("Use-case ID", "UC-01"),
            ("Primary actor", "Candidate"),
            ("Preconditions", "The Candidate account is active and authenticated."),
            ("Trigger", "The Candidate uploads a supported document or submits manual CV data."),
            ("Main flow", "Validate the source; store the CV; extract and normalize text; derive terms and skills; create or update the vector; score eligible active Jobs; persist unique CV-Job Matching records; return ordered match cards."),
            ("Alternative/exception flows", "Reject unsupported or oversized files and invalid fields; record OCR or extraction failure; allow soft quality warnings; return an empty result when no active Job is eligible."),
            ("Postconditions", "The CV ends in SCORING_DONE or FAILED, with a visible status and failure reason where applicable."),
        ],
    ),
    (
        "1.5.2 Candidate Searches, Applies, and Withdraws",
        "Table 1.6. Use case - Candidate searches, applies, and withdraws",
        [
            ("Use-case ID", "UC-02"),
            ("Primary actor", "Guest or Candidate"),
            ("Preconditions", "Public search needs no account; application and personalized scoring require an authenticated Candidate."),
            ("Trigger", "The user searches Jobs, opens details, or selects Apply/Withdraw."),
            ("Main flow", "Search and filter public Jobs; open Job details; resolve the Candidate and selected/default CV; verify the Job; prevent duplicate Candidate-Job applications; create an application; list owned applications; withdraw when the current state permits."),
            ("Alternative/exception flows", "Reject missing authentication, missing/default CV, invalid Job, duplicate application, unrelated ownership, or invalid withdrawal state."),
            ("Postconditions", "A valid application is created or withdrawn, and the resulting state is returned to the owning Candidate."),
        ],
    ),
    (
        "1.5.3 Recruiter Creates a Job and Reviews Candidates",
        "Table 1.7. Use case - Recruiter creates a Job and reviews candidates",
        [
            ("Use-case ID", "UC-03"),
            ("Primary actor", "Recruiter"),
            ("Preconditions", "The Recruiter account is active and authenticated."),
            ("Trigger", "The Recruiter creates a Job or opens an owned Job workspace."),
            ("Main flow", "Validate structured and free-text Job data; normalize and vectorize the Job; persist ownership; expose eligible Jobs to matching; rank applicants and discovered candidates; invite a Candidate or update application state."),
            ("Alternative/exception flows", "Reject invalid salary/content rules, missing ownership, another Recruiter's Job identifier, invalid invitation, or unsupported application transition."),
            ("Postconditions", "The owned Job and permitted recruitment actions are persisted with audit evidence."),
        ],
    ),
    (
        "1.5.4 Candidate Feedback Changes Later Ranking",
        "Table 1.8. Use case - Candidate feedback changes later ranking",
        [
            ("Use-case ID", "UC-04"),
            ("Primary actor", "Candidate"),
            ("Preconditions", "The Candidate owns the CV and Matching associated with the feedback."),
            ("Trigger", "The Candidate submits GOOD_MATCH, POTENTIAL, BAD_MATCH, or NOT_INTERESTED feedback."),
            ("Main flow", "Resolve the Matching and actor; verify ownership and role; upsert one current judgment; commit the transaction; start Rocchio learning after commit; mark affected Matchings for recomputation; rescore and clear the marker."),
            ("Alternative/exception flows", "Reject another role, another Candidate's Matching, missing evidence, or invalid feedback; retain retry visibility when asynchronous recomputation fails."),
            ("Postconditions", "The feedback and audit state are stored, and later rankings use the rebuilt learned vector."),
        ],
    ),
    (
        "1.5.5 AutoFit Executes a Policy-Governed Action",
        "Table 1.9. Use case - AutoFit executes a policy-governed action",
        [
            ("Use-case ID", "UC-05"),
            ("Primary actor", "Candidate and background scheduler"),
            ("Preconditions", "An enabled owned policy, an eligible default CV, completed scoring, and an allowed action state exist."),
            ("Trigger", "A scheduled scan or controlled run-now request evaluates the policy."),
            ("Main flow", "Resolve the user, Candidate, CV, matches, thresholds, consent, quota, cooldown, quiet hours, time zone, and prior interactions; select eligible actions; create notifications or applications; write audit evidence."),
            ("Alternative/exception flows", "Skip disabled, paused, ineligible, duplicate, quota-exceeded, cooldown, quiet-hour, stale, or invalid-state items; isolate per-item failures."),
            ("Postconditions", "Only policy-allowed actions are recorded, and the remaining items stay unchanged for later evaluation."),
        ],
    ),
    (
        "1.5.6 User Completes an Email Feedback Action",
        "Table 1.10. Use case - User completes an email feedback action",
        [
            ("Use-case ID", "UC-06"),
            ("Primary actor", "Email recipient"),
            ("Preconditions", "A pending, unexpired EmailAction exists and only its SHA-256 token hash is stored."),
            ("Trigger", "The recipient opens the link and deliberately confirms the action."),
            ("Main flow", "GET hashes and validates the token and displays a confirmation page without changing recruitment data; POST repeats validation, performs the supported action, marks the token redeemed, and records the outcome."),
            ("Alternative/exception flows", "Reject missing, expired, redeemed, malformed, or unsupported actions; mark expired pending records; do not execute state changes from GET."),
            ("Postconditions", "The requested action is executed at most once under normal validation and remains auditable."),
        ],
    ),
]

for heading, caption_text, rows in use_cases:
    heading_index = find_index(requirements_block, heading)
    cursor = heading_index + 1
    while cursor < len(requirements_block):
        element = requirements_block[cursor]
        if not is_paragraph(element):
            break
        p = paragraph(element, document)
        if p.style.name != "Normal" or has_drawing(element):
            break
        requirements_block.pop(cursor)
    requirements_block[heading_index + 1 : heading_index + 1] = [
        make_table_caption(document, caption_text),
        make_use_case_table(document, rows),
    ]

requirements_block.extend(
    [
        make_paragraph(document, "1.6 Chapter Summary", "Heading 2"),
        make_paragraph(
            document,
            (
                "This chapter defined CareerFit's actors, functional and "
                "non-functional requirements, and six representative use cases. "
                "The requirements separate matching evidence, application state, "
                "user policy, automated actions, and audit records so that later "
                "design and implementation decisions can be traced to expected "
                "behavior."
            ),
            "Normal",
        ),
    ]
)

# Build Chapter 2 from the former theory chapter and design portion.
set_paragraph_text(
    theory_block[0],
    document,
    "CHAPTER 2. THEORETICAL BACKGROUND AND SOLUTION DESIGN",
)
for element in theory_block:
    if element_text(element) == "2.10 Chapter Summary":
        set_paragraph_text(
            element, document, "2.10 Theoretical Background Summary"
        )

design_heading_map = {
    "3.6 System Architecture": "2.11 System Architecture",
    "3.7 Module Design": "2.12 Module Design",
    "3.8 Data Design": "2.13 Data Design",
    "3.8.1 Core Identity and Recruitment Data":
        "2.13.1 Core Identity and Recruitment Data",
    "3.8.2 Automation, Communication, and Operational Data":
        "2.13.2 Automation, Communication, and Operational Data",
    "3.9 Security, Failure, and Consistency Design":
        "2.14 Security, Failure, and Consistency Design",
    "3.9.1 Authentication and Authorization":
        "2.14.1 Authentication and Authorization",
    "3.9.2 Failure Handling": "2.14.2 Failure Handling",
    "3.9.3 Identified Design Risks": "2.14.3 Identified Design Risks",
    "3.10 Deployment Architecture": "2.15 Deployment Architecture",
    "3.11 Chapter Summary": "2.16 Chapter Summary",
}
for element in design_block:
    text = element_text(element)
    if text in design_heading_map:
        set_paragraph_text(element, document, design_heading_map[text])

summary_index = find_index(design_block, "2.16 Chapter Summary")
design_block = design_block[: summary_index + 1] + [
    make_paragraph(
        document,
        (
            "This chapter connected the theoretical foundations to the CareerFit "
            "solution design. TF-IDF, cosine similarity, Rocchio feedback, "
            "ranking metrics, and Human-in-the-Loop principles explain the core "
            "decision-support behavior. The architecture, modules, data model, "
            "security boundaries, consistency rules, and deployment topology "
            "define how those ideas are separated into implementable components."
        ),
        "Normal",
    )
]
chapter_2_block = theory_block + design_block

# Renumber implementation, evaluation, and conclusion headings.
set_paragraph_text(
    implementation_block[0], document, "CHAPTER 3. SYSTEM IMPLEMENTATION"
)
for element in implementation_block[1:]:
    if not is_paragraph(element):
        continue
    p = paragraph(element, document)
    if p.style.name in {"Heading 2", "Heading 3", "Heading 4"}:
        text = element_text(element)
        if re.match(r"^4\.", text):
            set_paragraph_text(element, document, "3." + text[2:])

set_paragraph_text(
    evaluation_block[0], document, "CHAPTER 4. TESTING AND EVALUATION"
)
for element in evaluation_block[1:]:
    if not is_paragraph(element):
        continue
    p = paragraph(element, document)
    if p.style.name in {"Heading 2", "Heading 3", "Heading 4"}:
        text = element_text(element)
        if re.match(r"^5\.", text):
            set_paragraph_text(element, document, "4." + text[2:])

set_paragraph_text(conclusion_block[0], document, "CONCLUSION")
conclusion_heading_map = {
    "6.1 Summary of Results": "1. Summary of Results",
    "6.2 Achievement of Thesis Objectives": "2. Achievement of Thesis Objectives",
    "6.2.1 Role-Based Recruitment Platform": "2.1 Role-Based Recruitment Platform",
    "6.2.2 CV and Job-Description Processing": "2.2 CV and Job-Description Processing",
    "6.2.3 Matching and Recommendation": "2.3 Matching and Recommendation",
    "6.2.4 Feedback Learning": "2.4 Feedback Learning",
    "6.2.5 Policy-Driven Automation and Email Action": "2.5 Policy-Driven Automation and Email Action",
    "6.2.6 Auditability and Evaluation": "2.6 Auditability and Evaluation",
    "6.3 Discussion": "3. Discussion",
    "6.3.1 Value of an Interpretable Baseline": "3.1 Value of an Interpretable Baseline",
    "6.3.2 Human-in-the-Loop as System Structure": "3.2 Human-in-the-Loop as System Structure",
    "6.3.3 Meaning of the Rocchio Improvement": "3.3 Meaning of the Rocchio Improvement",
    "6.3.4 Test Results versus Background Errors": "3.4 Test Results versus Background Errors",
    "6.3.5 Product Positioning": "3.5 Product Positioning",
    "6.4 Limitations": "4. Limitations",
    "6.4.1 Data and Algorithm Limitations": "4.1 Data and Algorithm Limitations",
    "6.4.2 Evaluation Limitations": "4.2 Evaluation Limitations",
    "6.4.3 Security and Privacy Limitations": "4.3 Security and Privacy Limitations",
    "6.4.4 Reliability and Maintainability Limitations": "4.4 Reliability and Maintainability Limitations",
    "6.5 Future Work": "5. Future Work",
    "6.6 Conclusion": "6. Closing Remarks",
}
for element in conclusion_block:
    text = element_text(element)
    if text in conclusion_heading_map:
        set_paragraph_text(element, document, conclusion_heading_map[text])

# References and rebuilt appendices.
appendix_block = rebuild_appendices(document)

main_content_heading = make_paragraph(
    document, "MAIN CONTENT", "Heading 1", page_break_before=True
)

new_elements = (
    cover_elements
    + ack_block
    + declaration_block
    + supervisor_comments_block
    + toc_block
    + table_list_block
    + figure_list_block
    + abbreviation_block
    + abstract_block
    + [front_section_break]
    + intro_block
    + [main_content_heading, chapter_1_heading]
    + requirements_block
    + chapter_2_block
    + implementation_block
    + evaluation_block
    + conclusion_block
    + references_block
    + appendix_block
)

# Renumber figures, screens, tables, and in-text references without cascading.
caption_mapping = {}
for old, new in {
    1: 3, 2: 4, 3: 5, 4: 6, 5: 7, 6: 8
}.items():
    caption_mapping[f"Figure 3.{old}"] = f"Figure 1.{new}"
for old, new in {7: 5, 8: 6, 9: 7}.items():
    caption_mapping[f"Figure 3.{old}"] = f"Figure 2.{new}"
for number in range(1, 11):
    caption_mapping[f"Figure 4.{number}"] = f"Figure 3.{number}"
for number in range(1, 7):
    caption_mapping[f"Screen 4.{number}"] = f"Screen 3.{number}"
for number in range(1, 5):
    caption_mapping[f"Figure 5.{number}"] = f"Figure 4.{number}"

table_number_mapping = {
    "Table 3.1": "Table 1.2",
    "Table 3.2": "Table 1.3",
    "Table 3.3": "Table 1.4",
    "Table 3.4": "Table 2.2",
    "Table 3.5": "Table 2.3",
    "Table 3.6": "Table 2.4",
}
for number in range(1, 7):
    table_number_mapping[f"Table 4.{number}"] = f"Table 3.{number}"
for number in range(1, 9):
    table_number_mapping[f"Table 5.{number}"] = f"Table 4.{number}"
for number in range(1, 4):
    table_number_mapping[f"Table 6.{number}"] = f"Table C.{number}"

replace_text_nodes(new_elements, caption_mapping)
replace_text_nodes(new_elements, table_number_mapping)

chapter_reference_mapping = {
    "Chapter 6": "the Conclusion",
    "Chapter 5": "Chapter 4",
    "Chapter 4": "Chapter 3",
    "Chapter 3": "Chapter 2",
}
replace_text_nodes(new_elements, chapter_reference_mapping)

# Targeted corrections after broad chapter-reference renumbering.
targeted_text = {
    "Chapter 2 turns these ideas into system requirements and design decisions.":
        "Chapters 1 and 2 turn these ideas into system requirements and design decisions.",
    "Chapter 2 defines these responsibilities at the design level, and Chapter 3 describes":
        "Chapter 2 defines these responsibilities at the design level, and Chapter 3 describes",
    "This chapter uses the endpoint groups only to establish component boundaries and traceability.":
        "This chapter uses endpoint groups only to establish functional boundaries and traceability.",
}
replace_text_nodes(new_elements, targeted_text)

# Reassemble the document body while preserving the final section properties.
for child in list(body):
    body.remove(child)
for element in new_elements:
    detach(element)
    body.append(element)
body.append(body_final_sectpr)

style_document(document)

# Avoid an extra page break immediately after section breaks or MAIN CONTENT.
for p in document.paragraphs:
    if p.text.strip() in {
        "ACKNOWLEDGEMENTS",
        "INTRODUCTION",
        "CHAPTER 1. PROBLEM DESCRIPTION AND REQUIREMENTS",
    }:
        p.paragraph_format.page_break_before = False

# Format all tables consistently, excluding the inner-cover layout table.
for table in document.tables:
    text = " ".join(cell.text for row in table.rows for cell in row.cells)
    style_table(
        table,
        cover=("Supervisor" in text and "Pham Huu Hung - B2203557" in text),
    )

document.save(REPORT)
print(f"Restructured report saved to {REPORT}")
print(
    f"paragraphs={len(document.paragraphs)} tables={len(document.tables)} "
    f"inline_images={len(document.inline_shapes)} sections={len(document.sections)}"
)
