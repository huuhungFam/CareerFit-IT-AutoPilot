from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"


def set_text(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)


def find_by_start(doc, start):
    matches = [p for p in doc.paragraphs if p.text.startswith(start)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting with {start!r}, found {len(matches)}")
    return matches[0]


def remove(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


doc = Document(DOCX)

replacements = {
    "Recruitment in the information technology sector involves matching candidates":
        "IT recruitment requires candidates and recruiters to compare many details, including technical skills, experience, language, location, and working conditions. Candidates must identify suitable vacancies, while recruiters must review CVs and decide which applicants deserve closer attention. As the number of jobs and profiles grows, doing this manually becomes slow and inconsistent [10].",
    "Most job portals effectively support vacancy publication":
        "Most job portals support posting, searching, and applying, but they do not fully explain why a job fits a candidate or why one applicant ranks above another. CareerFit therefore treats CV–JD matching and personalized recommendation as related but separate tasks. The first compares a specific CV and job; the second prioritizes jobs from a candidate's broader profile and preferences.",
    "Automation can reduce repetitive work, but fully automatic decisions":
        "Automation can reduce repetitive work, but a similarity score should not directly trigger a consequential action. Consent, current application state, previous interactions, thresholds, quotas, and exceptional conditions also matter. CareerFit uses a Human-in-the-Loop design so that users can review important actions and administrators can trace what happened.",
    "To address these problems, this thesis develops CareerFit IT AutoPilot":
        "This thesis develops CareerFit IT AutoPilot, an IT-focused web platform that combines job discovery, CV and JD processing, matching, recommendation, Rocchio feedback, AutoFit policies, email actions, and audit records. The aim is not to replace recruitment judgment, but to provide understandable evidence and controlled assistance throughout the workflow.",
    "The first motivation for this work is to reduce fragmentation":
        "The first motivation is to reduce fragmentation. Candidates often move between job search, CV management, application history, and email, while recruiters switch between vacancy management, applicant review, candidate discovery, and reporting. CareerFit brings these activities into one role-based system and uses email as an additional action channel.",
    "The second motivation is explainability":
        "The second motivation is explainability. TF-IDF and cosine similarity do not capture meaning as deeply as modern embedding models, but their terms and weights can be inspected. This makes them suitable for an academic baseline in which the scoring process, limitations, and effect of feedback must be demonstrated clearly.",
    "The third motivation is controlled automation":
        "The third motivation is controlled adaptation and automation. AutoFit places policy checks between a score and an action, considering consent, thresholds, previous interactions, cooldown, quota, time zone, and quiet hours. Rocchio feedback then provides a transparent way to adjust later rankings from explicit positive and negative judgments without overwriting the original representation.",
    "• Support expiring email-action links, action-state tracking":
        "• Support expiring email-action links with hashed token storage, a non-mutating confirmation page, POST execution, action-state tracking, and audit logging.",
    "This thesis is organized into six chapters":
        "The thesis contains six chapters. Chapter 1 introduces the problem, objectives, scope, and contributions. Chapter 2 presents the theoretical background and related work, including TF-IDF, cosine similarity, Rocchio feedback, ranking metrics, Human-in-the-Loop control, and relevant security foundations.",
    "Chapter 3 analyzes the system requirements":
        "Chapter 3 covers requirements, use cases, architecture, data design, security, and deployment. Chapter 4 explains the implementation. Chapter 5 presents the evaluation method and verified results, and Chapter 6 discusses achievements, limitations, future work, and the final conclusion.",
    "This chapter established the foundations for CareerFit":
        "This chapter presented the main ideas used by CareerFit. TF-IDF and cosine similarity provide an inspectable lexical baseline, while Rocchio uses explicit feedback to adjust ranking. Ranking metrics measure ordered results, and Human-in-the-Loop policies keep similarity evidence separate from recruitment actions. Chapter 3 turns these ideas into system requirements and design decisions.",
    "CareerFit IT AutoPilot is analyzed as a role-based recruitment information system":
        "CareerFit is designed as a role-based system for IT recruitment. Its boundary includes the React frontend, Spring Boot backend, PostgreSQL database, CV storage, background scheduler, and optional email provider. External job boards, interview scheduling, offers, payroll, and unrestricted automated hiring remain outside the implemented scope.",
    "This chapter translated the CareerFit objectives into actors":
        "This chapter defined the actors, requirements, use cases, architecture, data model, security boundaries, failure handling, and deployment design. The main design rule is to keep matching evidence, application state, user policy, and automated actions separate. Chapter 4 explains how these decisions are implemented.",
    "This chapter described how CareerFit implements the architecture":
        "This chapter explained how the design is implemented in Spring Boot, React, PostgreSQL, and Flyway. It covered security, CV processing, TF-IDF scoring, Rocchio feedback, application workflows, AutoFit, email actions, audit records, frontend integration, and runtime monitoring. Chapter 5 evaluates these parts using fresh test and runtime evidence.",
    "`AutomationScheduler` coordinates five tasks documented in Chapter 3":
        "AutomationScheduler coordinates five tasks described in Chapter 3. Their delays, cron expressions, and time zone are read from app.scheduler properties. Per-item exception handling prevents one failed candidate or matching record from stopping an entire scan, while logs record the processing outcome.",
    "The zero-failure result means the registered JUnit assertions passed":
        "All 63 registered JUnit tests passed in the final run. Flyway 9.22.3 still warned that PostgreSQL 16.14 is newer than the highest version it reports as tested, and some negative tests intentionally logged handled exceptions. The final benchmark log contained no StaleObjectStateException or build failure.",
    "This mixed state must not be summarized as monitoring health passing":
        "The final runtime check returned HTTP 200 and status UP for aggregate health, liveness, and readiness. Mail health is disabled when application mail is disabled, so the absence of a local mail provider no longer makes aggregate health misleading. Production monitoring would still require protected component details, alerting, and an external mail check when email is enabled.",
    "The benchmark test runs within a Spring context where schedulers":
        "The benchmark runs inside a Spring context, so asynchronous work and persisted matching records can affect repeatability if they are not controlled. The final test registers learning after commit, clears previous matching state before recomputation, and checks the logs for background failures. The browser tests also avoid relying only on the first seeded job and delete the job created by the recruiter flow.",
    "Fresh evaluation showed that all 63 registered backend tests passed":
        "The final evaluation passed all 63 backend tests, completed the frontend production build, and passed four Chromium P0 workflows. The controlled Rocchio benchmark improved the holdout ranking deterministically on the synthetic dataset, and its final log contained no optimistic-lock exception. Aggregate health returned HTTP 200 with status UP. These results support a functioning academic prototype, but not production readiness or proven effectiveness on real recruitment data.",
    "This thesis designed, implemented, and evaluated CareerFit IT AutoPilot":
        "CareerFit IT AutoPilot was implemented as an IT-focused recruitment prototype with public job discovery, role-based workspaces, CV processing, matching, recommendation, Rocchio feedback, applications, AutoFit policies, email actions, analytics, and administrative monitoring. The backend follows a modular-monolith structure and the user interface is implemented in React.",
    "The controlled benchmark strongly supports mathematical behavior":
        "The controlled benchmark confirms the intended Rocchio behavior in the planted latent-skill scenario. The final run completed without the earlier optimistic-lock exception after learning was moved to an after-commit boundary and benchmark state was cleared before recomputation. This supports deterministic controlled behavior, although it does not prove reliability under production concurrency.",
    "Audit coverage is not yet formally proven for every state-changing endpoint":
        "Audit coverage has not been formally proven for every state-changing endpoint, and direct repository calls can still produce inconsistent metadata conventions. The evaluation nevertheless checked logs and side effects in addition to test exit codes, which helped detect and correct the earlier concurrency and health problems.",
    "The backend suite contains 63 tests but does not prove exhaustive path coverage":
        "The 63 backend tests do not prove exhaustive path coverage. Browser evaluation covers four P0 cases in Chromium only, and no independent users participated. Security checks were targeted API observations rather than penetration testing, while the latency sample used 30 sequential warm requests on one workstation. Aggregate health passed in the final run, but concurrency, capacity, cross-browser behavior, and real-user usability remain unevaluated.",
    "The first priority is correctness and security hardening":
        "Future security work should add rate limiting, stronger session handling, protected management endpoints, malware scanning, encrypted CV storage, retention rules, and tested backup recovery. Email delivery should be tested with a real provider, including expiry, replay, scanner behavior, and failure recovery. CI should continue checking background logs and asynchronous completion rather than relying only on foreground test status.",
    "CareerFit IT AutoPilot demonstrates an end-to-end approach":
        "CareerFit demonstrates an end-to-end approach to controlled recruitment automation in the IT domain. It connects job discovery, interpretable matching, profile-based recommendation, feedback learning, policy-driven actions, and audit records while keeping the relationship between score, business state, and user decision visible.",
    "Fresh evaluation supports the operation of the academic prototype":
        "The final evidence shows that the prototype works in the evaluated local environment: backend tests passed, the frontend built successfully, selected browser workflows completed, aggregate health was UP, and Rocchio produced deterministic improvement in the synthetic scenario. The main contribution is therefore a working and critically evaluated Human-in-the-Loop workflow, not a claim that the model is universally superior or ready to make production hiring decisions.",
}

for start, new_text in replacements.items():
    set_text(find_by_start(doc, start), new_text)

for start in [
    "Feedback creates another practical challenge",
    "The final motivation is continuous but bounded adaptation",
    "Chapter 2 presents the theoretical background and related work",
    "Chapter 4 explains the system implementation",
    "Chapter 5 defines the experimental and system evaluation methodology",
    "Chapter 6 summarizes the verified results",
    "The most important contribution is therefore not a claim",
]:
    remove(find_by_start(doc, start))

# Remove Markdown notation that leaked from the source drafts into Word.
for paragraph in doc.paragraphs:
    if "`" in paragraph.text:
        set_text(paragraph, paragraph.text.replace("`", ""))
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if "`" in paragraph.text:
                    set_text(paragraph, paragraph.text.replace("`", ""))

settings = doc.settings._element
update = settings.find(qn("w:updateFields"))
if update is None:
    update = OxmlElement("w:updateFields")
    settings.append(update)
update.set(qn("w:val"), "true")

doc.save(DOCX)
print(DOCX)
