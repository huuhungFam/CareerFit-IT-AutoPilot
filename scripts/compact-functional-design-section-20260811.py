from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"


TABLE_ROWS = {
    "Table 3.3. Interface components for Explore Jobs": [
        ("1", "Navigation", "Open Candidate pages."),
        ("2", "Select", "Choose the Job location scope."),
        ("3", "Input", "Enter a Job title, skill, or company keyword."),
        ("4", "Button", "Submit the current search criteria."),
        ("5", "Filter / Sort", "Refine the results and sort order."),
        ("6", "Information panel", "Show result-quality guidance or a limited-data state."),
    ],
    "Table 3.6. Interface components for Job and applicant management": [
        ("1", "Navigation", "Open Recruiter pages."),
        ("2", "Button", "Export the Recruiter's Job list as CSV."),
        ("3", "Input", "Search within the Recruiter's Jobs."),
        ("4", "Select", "Filter Jobs by lifecycle status."),
        ("5", "Button", "Open the form for a new Job posting."),
        ("6", "Job list", "Display owned Jobs and select one Job."),
        ("7", "Action buttons", "Edit or delete the selected Job when permitted."),
        ("8", "Detail panel", "Show the selected Job and open applicant review."),
    ],
    "Table 3.7. Interface components for Talent Pool and invitations": [
        ("1", "Navigation", "Open the Recruiter Talent Pool."),
        ("2", "Job list", "Select an owned Job for Candidate results."),
        ("3", "Tabs", "Switch between all, bookmarked, and invited CVs."),
        ("4", "Result groups", "Separate Matching and Potential CV groups."),
        ("5", "Candidate summary", "Show the Candidate summary and score context."),
        ("6", "Button", "Save or remove the Candidate bookmark."),
        ("7", "State badge", "Show an existing Application or invitation."),
        ("8", "Button", "Open the visible CV detail."),
    ],
}


def find_exact(document: Document, text: str):
    matches = [p for p in document.paragraphs if p.text.strip() == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph for {text!r}; found {len(matches)}")
    return matches[0]


def next_table(document: Document, caption_text: str) -> Table:
    caption = find_exact(document, caption_text)
    table_map = {table._tbl: table for table in document.tables}
    node = caption._p.getnext()
    while node is not None:
        if node.tag == qn("w:tbl"):
            return table_map[node]
        node = node.getnext()
    raise RuntimeError(f"No table after {caption_text}")


def set_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)


def main() -> None:
    document = Document(DOCX)
    functional_heading = find_exact(document, "3.4 Functional Design")
    functional_heading.paragraph_format.page_break_before = False

    for caption, rows in TABLE_ROWS.items():
        table = next_table(document, caption)
        if len(table.columns) != 3 or len(table.rows) != len(rows) + 1:
            raise RuntimeError(f"Unexpected table structure for {caption}")
        for col, value in enumerate(("No.", "Control Type", "Description")):
            set_cell(table.cell(0, col), value, bold=True)
        for row_index, values in enumerate(rows, start=1):
            for col, value in enumerate(values):
                set_cell(table.cell(row_index, col), value)

    document.save(DOCX)
    print(f"updated={DOCX}")
    print("functional_design_page_break=removed")
    print(f"compacted_tables={len(TABLE_ROWS)}")


if __name__ == "__main__":
    main()
