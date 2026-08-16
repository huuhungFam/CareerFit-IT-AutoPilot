from __future__ import annotations

from pathlib import Path
import hashlib
import re
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-before-20260811-shortened-usecases.docx"


def join(*lines: str) -> str:
    return "\n".join(lines)


def sections(*blocks: str) -> str:
    return "\n\n".join(blocks)


SHORT = {
    "UC-01": {
        "Description": "This use case allows a Candidate to maintain their Profile, Portfolio, and CV collection, including CV review and confirmation for later matching.",
        "Preconditions": join(
            "P1. The Candidate account is active and authenticated.",
            "P2. The Candidate may manage only their own Profile, Portfolio, and CVs.",
        ),
        "Trigger": "The Candidate opens Career Profile and selects a Profile, Portfolio, or CV action.",
        "Main Flow": join(
            "1. The Candidate opens Career Profile.",
            "2. The System displays the Profile, Portfolio, and owned CVs with their statuses.",
            "3. The Candidate selects Upload CV and submits a supported file.",
            "4. The System validates the file and extracts reviewable CV sections.",
            "5. The System marks the CV REVIEW_REQUIRED and displays extracted content and warnings.",
            "6. The Candidate reviews and corrects the content.",
            "7. The Candidate confirms the CV.",
            "8. The System validates and accepts it for matching processing.",
            "9. After successful processing, the System marks the CV SCORING_DONE and makes it default if no default CV exists.",
            "10. The System displays the updated CV collection; Matching results are viewed through UC-05.",
        ),
        "Alternative Flows": sections(
            join(
                "A1 – Create or save a manual CV (from Step 3)",
                "1. The Candidate enters CV information manually.",
                "2. The System saves unfinished content as DRAFT or completed content as REVIEW_REQUIRED.",
                "3. A REVIEW_REQUIRED CV continues at Step 5.",
            ),
            join(
                "A2 – Update Profile or Portfolio (from Step 3)",
                "1. The Candidate changes editable Profile information or adds, updates, or removes a Portfolio item.",
                "2. The System validates, saves, and displays the change.",
            ),
            join(
                "A3 – Resume an unconfirmed CV (from Step 3)",
                "1. The Candidate opens an owned DRAFT or REVIEW_REQUIRED CV and updates it.",
                "2. Confirmation continues at Step 7; otherwise, the CV remains unconfirmed.",
            ),
            join(
                "A4 – Select a default CV (from Step 3)",
                "1. The Candidate selects an owned SCORING_DONE CV.",
                "2. The System makes it the only default CV and updates the collection.",
            ),
            join(
                "A5 – Delete an eligible CV (from Step 3)",
                "1. The Candidate selects an owned non-default CV and confirms deletion.",
                "2. The System verifies eligibility and removes the CV.",
            ),
        ),
        "Exception Flows": sections(
            join(
                "E1 – File or extraction is invalid (from Step 4)",
                "1. The System rejects an unsupported file or marks the CV FAILED when extraction cannot complete.",
                "2. No invalid Matching result is presented.",
            ),
            join(
                "E2 – Confirmation or processing fails (from Step 8)",
                "1. The System displays validation issues or marks processing as FAILED.",
                "2. The CV remains unconfirmed or failed and may be corrected when allowed.",
            ),
            join(
                "E3 – CV is unavailable or unauthorized",
                "1. The System denies the operation and changes no CV information.",
            ),
            join(
                "E4 – Default or deletion request is ineligible (from A4 or A5)",
                "1. The System rejects a non-SCORING_DONE default request or deletion of a default or BANNED CV, or a CV with pending reports.",
                "2. Existing CV states remain unchanged.",
            ),
        ),
        "Postconditions": join(
            "Success Postconditions",
            "S1. Valid Profile or Portfolio changes are stored.",
            "S2. A CV remains DRAFT/REVIEW_REQUIRED or reaches SCORING_DONE according to its progress.",
            "S3. The Candidate has at most one default CV.",
            "S4. A CV may reach SCORING_DONE even when no eligible Job produces a Matching result.",
            "",
            "Minimal Guarantees",
            "F1. Invalid operations do not replace valid career information.",
            "F2. The Candidate cannot change another Candidate's data or view invalid Matching results.",
        ),
        "Related Use Cases": join("UC-03 – Manage Job Applications", "UC-05 – Review Personalized Career Insights", "UC-06 – Manage AutoFit"),
    },
    "UC-02": {
        "Description": "This use case allows a Candidate or Guest to search and review publicly available Jobs; protected Candidate actions still require authentication.",
        "Preconditions": join(
            "P1. Public Job exploration requires no authentication.",
            "P2. Protected actions require an active, authenticated Candidate account.",
        ),
        "Trigger": "The Candidate or Guest opens Job exploration or submits supported criteria.",
        "Main Flow": join(
            "1. The Candidate opens Job exploration.",
            "2. The System displays public ACTIVE Jobs and supported search, filter, sorting, and pagination controls.",
            "3. The Candidate enters criteria.",
            "4. The System displays the matching Job summaries.",
            "5. The Candidate selects a Job.",
            "6. The System verifies public availability and displays Job and employer details.",
            "7. The System displays similar Jobs when available.",
        ),
        "Alternative Flows": sections(
            join(
                "A1 – Guest explores Jobs (from Step 1)",
                "1. A Guest searches, filters, and views public Job details.",
                "2. The System requests authentication before a protected action.",
            ),
            join(
                "A2 – Change the result view (from Step 3 or Step 4)",
                "1. The Candidate selects urgent Jobs, clears criteria, changes sorting, or changes page.",
                "2. The System displays the corresponding ACTIVE Job set.",
            ),
            join(
                "A3 – No Jobs match (from Step 4)",
                "1. The System displays an empty state and allows the criteria to be changed.",
            ),
            join(
                "A4 – Open a similar Job (from Step 7)",
                "1. The Candidate selects a similar Job, and the flow returns to Step 6.",
            ),
        ),
        "Exception Flows": sections(
            join(
                "E1 – Selected Job is unavailable (from Step 6)",
                "1. The System does not expose a missing or non-public Job and returns the user to available results.",
            ),
            join(
                "E2 – Protected action lacks authentication",
                "1. The System does not perform the action and requests Candidate authentication.",
            ),
        ),
        "Postconditions": join(
            "Success Postconditions",
            "S1. Current public Job information is displayed according to the selected criteria.",
            "S2. Exploration alone creates no Application.",
            "",
            "Minimal Guarantees",
            "F1. Non-public Job details and protected actions are not exposed without authorization.",
            "F2. Exploration changes no Job or Candidate data.",
        ),
        "Related Use Cases": join("UC-03 – Manage Job Applications", "UC-05 – Review Personalized Career Insights", "UC-12 – Report Suspicious Recruitment Content"),
    },
    "UC-03": {
        "Description": "This use case allows a Candidate to apply for Jobs, review Application history, withdraw eligible Applications, and respond to Recruiter invitations.",
        "Preconditions": join(
            "P1. The Candidate account is active and authenticated.",
            "P2. The Candidate may manage only their own Applications.",
            "P3. A new Application requires an ACTIVE Job and an owned selected or default CV.",
        ),
        "Trigger": "The Candidate applies for a Job or opens Application management.",
        "Main Flow": join(
            "1. The Candidate opens an ACTIVE Job and selects Apply.",
            "2. The System verifies the Candidate and Job.",
            "3. The System checks for an existing Candidate–Job Application.",
            "4. The System identifies the selected owned CV or owned default CV.",
            "5. The System creates a PENDING Application.",
            "6. The System confirms submission.",
            "7. The Application appears in Candidate history and the Recruiter's applicant list.",
        ),
        "Alternative Flows": sections(
            join(
                "A1 – Select the CV (from Step 4)",
                "1. The System uses an explicitly selected owned CV or, when none is supplied, the owned default CV.",
                "2. The flow continues at Step 5.",
            ),
            join(
                "A2 – Review Application history (alternative from the Trigger)",
                "1. The System displays owned Applications and statuses for filtering and review.",
            ),
            join(
                "A3 – Withdraw an eligible Application (from A2)",
                "1. The Candidate requests withdrawal of an Application not APPROVED or REJECTED.",
                "2. The System verifies eligibility and changes it to NOT_INTERESTED.",
            ),
            join(
                "A4 – Respond to an invitation (alternative from the Trigger)",
                "1. The Candidate selects an INVITED Application.",
                "2. The System changes it to PENDING when accepted or NOT_INTERESTED when declined.",
            ),
        ),
        "Exception Flows": sections(
            join(
                "E1 – Job is unavailable (from Step 2)",
                "1. The System rejects a missing or non-ACTIVE Job; no Application is created.",
            ),
            join(
                "E2 – Application already exists (from Step 3)",
                "1. The System rejects the duplicate request and keeps the existing Application.",
            ),
            join(
                "E3 – No eligible owned CV is available (from Step 4)",
                "1. The System rejects a missing, unauthorized, or unavailable CV; no Application is created.",
            ),
            join(
                "E4 – Withdrawal or invitation response is invalid (from A3 or A4)",
                "1. The Application is unavailable, unauthorized, or in an ineligible status.",
                "2. The System rejects the request and preserves its status.",
            ),
        ),
        "Postconditions": join(
            "Success Postconditions",
            "S1. A submitted or accepted Application is PENDING and visible to both sides.",
            "S2. A withdrawn or declined Application is NOT_INTERESTED.",
            "",
            "Minimal Guarantees",
            "F1. Failed requests create no duplicate or unauthorized Application.",
            "F2. A Candidate cannot change another Candidate's Application.",
        ),
        "Related Use Cases": join("UC-01 – Manage Career Profile", "UC-02 – Explore Jobs", "UC-06 – Manage AutoFit", "UC-07 – Respond Through Actionable Email", "UC-09 – Review and Process Applicants"),
    },
    "UC-04": {
        "Description": "This use case allows a Candidate to evaluate an owned CV–Job Matching result so supported Feedback can influence later personalized rankings.",
        "Preconditions": join(
            "P1. The Candidate account is active and authenticated.",
            "P2. The Matching belongs to a CV owned by the Candidate.",
        ),
        "Trigger": "The Candidate selects a Feedback option for a Matching result.",
        "Main Flow": join(
            "1. The Candidate opens a Matching result.",
            "2. The System displays the Job, Matching information, and supported Feedback options.",
            "3. The Candidate selects a Feedback type.",
            "4. The System verifies ownership and validates the type.",
            "5. The System records the Feedback and confirms receipt.",
            "6. When applicable, the Feedback is used for later personalized rankings.",
        ),
        "Alternative Flows": sections(
            join(
                "A1 – Change previous Feedback (from Step 3)",
                "1. The Candidate selects another supported type, and the System replaces the previous Feedback.",
            ),
            join(
                "A2 – Mark the Job NOT_INTERESTED (from Step 3)",
                "1. The System records the choice without using it as a preference-learning signal.",
            ),
        ),
        "Exception Flows": sections(
            join(
                "E1 – Matching is unavailable or unauthorized (from Step 4)",
                "1. The System rejects the request and changes no Feedback.",
            ),
            join(
                "E2 – Feedback type is unsupported (from Step 4)",
                "1. The System rejects the value and preserves existing Feedback.",
            ),
        ),
        "Postconditions": join(
            "Success Postconditions",
            "S1. The Matching stores the Candidate's latest supported Feedback.",
            "S2. Eligible Feedback may influence later rankings; NOT_INTERESTED is stored without preference learning.",
            "",
            "Minimal Guarantees",
            "F1. Invalid or unauthorized requests do not change existing Feedback.",
        ),
        "Related Use Cases": join("UC-05 – Review Personalized Career Insights", "UC-07 – Respond Through Actionable Email"),
    },
    "UC-05": {
        "Description": "This use case allows a Candidate to review CV–Job Matching results, independent profile/preference-based Job recommendations, and available career analytics.",
        "Preconditions": "P1. The Candidate account is active and authenticated.",
        "Trigger": "The Candidate opens personalized career insights and selects a view.",
        "Main Flow": join(
            "1. The Candidate opens personalized career insights.",
            "2. The System displays the available Matching, recommendation, and analytics views.",
            "3. The Candidate opens CV–Job Matching results.",
            "4. The System displays available results for the Candidate's eligible CVs.",
            "5. The Candidate opens Job recommendations.",
            "6. The System independently displays recommendations based on supported Candidate profile or preference information.",
            "7. The Candidate reviews the available insights.",
        ),
        "Alternative Flows": sections(
            join(
                "A1 – Review Candidate analytics (from Step 2)",
                "1. The Candidate selects analytics, and the System displays available summaries and trends.",
            ),
            join(
                "A2 – Limited data in a selected view (from Step 4, Step 6, or A1)",
                "1. The System displays available information and an empty or limited-data state for the missing section.",
                "2. Other insight views remain available.",
            ),
        ),
        "Exception Flows": join(
            "E1 – Matching information is unavailable or unauthorized (from Step 4)",
            "1. The System withholds the unavailable result and changes no Candidate data.",
        ),
        "Postconditions": join(
            "Success Postconditions",
            "S1. Available Matching, recommendation, or analytics information is displayed in its corresponding view.",
            "S2. Reviewing insights changes no Application or Job state.",
            "",
            "Minimal Guarantees",
            "F1. Missing data in one view does not incorrectly hide other available insights.",
            "F2. Unauthorized Matching information is not exposed.",
        ),
        "Related Use Cases": join("UC-01 – Manage Career Profile", "UC-02 – Explore Jobs", "UC-04 – Provide Matching Feedback", "UC-06 – Manage AutoFit"),
    },
    "UC-06": {
        "Description": "This use case allows a Candidate to configure AutoFit and request a controlled run; the System may also execute it automatically on its configured schedule.",
        "Preconditions": "P1. The Candidate account is active and authenticated.",
        "Trigger": "The Candidate changes AutoFit configuration or requests a run; alternatively, the system-configured AutoFit schedule initiates execution.",
        "Main Flow": join(
            "1. The Candidate opens AutoFit settings.",
            "2. The System displays enablement, pause state, and threshold.",
            "3. The Candidate changes the configuration or requests a manual run.",
            "4. The System validates and saves a threshold from 50 to 100.",
            "5. For execution, the System verifies an owned default CV in SCORING_DONE status.",
            "6. The System selects Matchings for ACTIVE Jobs at or above the threshold.",
            "7. The System skips existing Candidate–Job Applications.",
            "8. The System creates at most three AUTO_APPLIED Applications and displays the run summary.",
        ),
        "Alternative Flows": sections(
            join(
                "A1 – Automatic execution (alternative from the Trigger)",
                "1. The configured schedule initiates execution for an enabled, unpaused policy.",
                "2. The flow continues at Step 5 without treating the internal scheduler as an actor.",
            ),
            join(
                "A2 – Pause or resume AutoFit (from Step 3)",
                "1. The Candidate changes the pause state, and the System saves it.",
            ),
            join(
                "A3 – No eligible Matching (from Step 6)",
                "1. The System creates no Application and displays a zero-result summary.",
            ),
            join(
                "A4 – Existing Application is skipped (from Step 7)",
                "1. The System skips the duplicate and continues with other eligible Matchings.",
            ),
            join(
                "A5 – No eligible default CV (from Step 5)",
                "1. The System creates no Application and informs the Candidate.",
            ),
            join(
                "A6 – Save configuration without running AutoFit (from Step 4)",
                "1. The Candidate does not request a manual run.",
                "2. The System retains the configuration for future manual or automatic execution, and the Use Case ends.",
            ),
        ),
        "Exception Flows": sections(
            join(
                "E1 – Threshold is outside 50–100 (from Step 4)",
                "1. The System rejects the value and preserves the previous valid configuration.",
            ),
            join(
                "E2 – Policy is unavailable or unauthorized",
                "1. The System denies the operation and creates no Application.",
            ),
        ),
        "Postconditions": join(
            "Success Postconditions",
            "S1. Valid AutoFit configuration is stored.",
            "S2. An execution creates zero to three non-duplicate AUTO_APPLIED Applications for eligible ACTIVE Jobs.",
            "S3. Notification quota, cooldown, quiet hours, and preferences are not AutoApply eligibility rules.",
            "",
            "Minimal Guarantees",
            "F1. Invalid configuration does not replace the last valid configuration.",
            "F2. Missing eligibility creates no Application.",
        ),
        "Related Use Cases": join("UC-01 – Manage Career Profile", "UC-03 – Manage Job Applications", "UC-05 – Review Personalized Career Insights"),
    },
    "UC-07": {
        "Description": "This use case allows a Candidate to confirm and complete supported Matching Feedback or Recruiter-invitation actions from an actionable email.",
        "Preconditions": join(
            "P1. The Candidate has received a CareerFit actionable email.",
            "P2. The email contains a link associated with a CareerFit email action.",
        ),
        "Trigger": "The Candidate opens the actionable link.",
        "Main Flow": join(
            "1. The Candidate opens the link.",
            "2. The System identifies the email action and verifies its current state and expiry.",
            "3. The System displays a non-mutating confirmation page.",
            "4. The Candidate reviews and confirms the action.",
            "5. The System validates the action and referenced record again.",
            "6. The System performs the supported action.",
            "7. The System marks the email action REDEEMED.",
            "8. The System displays the result.",
        ),
        "Alternative Flows": sections(
            join(
                "A1 – Submit Matching Feedback (from Step 6)",
                "1. The System submits the corresponding GOOD_MATCH, POTENTIAL, BAD_MATCH, or NOT_INTERESTED Feedback.",
            ),
            join(
                "A2 – Respond to a Recruiter invitation (from Step 6)",
                "1. For an INVITED Application, acceptance changes it to PENDING and decline changes it to NOT_INTERESTED.",
            ),
            join(
                "A3 – Leave without confirming (from Step 4)",
                "1. No business action occurs; the email action remains pending until redemption or expiry.",
            ),
        ),
        "Exception Flows": sections(
            join(
                "E1 – Action link is invalid (from Step 2)",
                "1. The System displays an invalid-link message and changes no business data.",
            ),
            join(
                "E2 – Action was already redeemed (from Step 2 or Step 5)",
                "1. The System informs the Candidate that it was processed previously and does not execute it again.",
            ),
            join(
                "E3 – Action has expired (from Step 2 or Step 5)",
                "1. The System rejects the action without changing the referenced state.",
                "2. When expiry is detected during confirmation, the System records EXPIRED.",
            ),
            join(
                "E4 – Referenced Matching or invitation is unavailable (from Step 5)",
                "1. The System rejects the action and preserves existing Feedback or Application state.",
            ),
        ),
        "Postconditions": join(
            "Success Postconditions",
            "S1. The supported action is applied to the referenced Matching or Application.",
            "S2. The email action is REDEEMED and cannot be executed again.",
            "",
            "Minimal Guarantees",
            "F1. Opening confirmation changes no business data.",
            "F2. Invalid, expired, redeemed, or failed actions do not change the referenced state.",
        ),
        "Related Use Cases": join("UC-03 – Manage Job Applications", "UC-04 – Provide Matching Feedback"),
    },
    "UC-08": {
        "Description": "This use case allows a Recruiter to maintain Employer Profile information and manage the lifecycle of owned Job postings.",
        "Preconditions": join(
            "P1. The Recruiter account is active and authenticated.",
            "P2. Existing-Job actions apply only to Jobs owned by the Recruiter.",
        ),
        "Trigger": "The Recruiter opens Employer Profile or Job management.",
        "Main Flow": join(
            "1. The Recruiter opens Job management.",
            "2. The System displays owned Jobs and available actions.",
            "3. The Recruiter selects Create Job.",
            "4. The System displays the Job form.",
            "5. The Recruiter enters Job information and submits it for publication.",
            "6. The System validates the information and Recruiter ownership.",
            "7. The System creates the Job as ACTIVE.",
            "8. The System displays the published Job in Job management.",
        ),
        "Alternative Flows": sections(
            join(
                "A1 – Update Employer Profile (alternative from the Trigger)",
                "1. The Recruiter changes editable employer information, and the System validates and saves it.",
            ),
            join(
                "A2 – Save as draft (from Step 5)",
                "1. The Recruiter saves the Job without publication, and the System stores it as DRAFT.",
            ),
            join(
                "A3 – Publish an existing draft (from Step 2)",
                "1. The Recruiter selects an owned DRAFT Job, completes valid information, and publishes it as ACTIVE.",
            ),
            join(
                "A4 – Update or close an existing Job (from Step 2)",
                "1. The Recruiter edits an owned Job or closes it through an allowed transition.",
                "2. The System validates, saves, and displays the result.",
            ),
            join(
                "A5 – Delete an eligible Job (from Step 2)",
                "1. The Recruiter confirms deletion of an owned Job with no blocking Applications or pending reports.",
                "2. The System deletes it and updates the list.",
            ),
        ),
        "Exception Flows": sections(
            join(
                "E1 – Job is unavailable or unauthorized",
                "1. The System denies access and changes no Job.",
            ),
            join(
                "E2 – Job information is invalid (from Step 6 or A3)",
                "1. The System displays validation issues and does not publish the Job.",
            ),
            join(
                "E3 – Job action is not allowed (from A4 or A5)",
                "1. The System rejects an invalid transition or deletion blocked by Applications or pending reports.",
                "2. The Job remains unchanged.",
            ),
        ),
        "Postconditions": join(
            "Success Postconditions",
            "S1. Valid Employer Profile changes are stored.",
            "S2. A published Job is ACTIVE; a draft or closed Job retains its selected valid state.",
            "S3. An eligible deleted Job is unavailable.",
            "",
            "Minimal Guarantees",
            "F1. A Recruiter cannot modify another Recruiter's Job.",
            "F2. Invalid or blocked actions do not publish, transition, or delete the Job.",
        ),
        "Related Use Cases": join("UC-02 – Explore Jobs", "UC-09 – Review and Process Applicants", "UC-10 – Manage Talent Pool and Invitations"),
    },
    "UC-09": {
        "Description": "This use case allows a Recruiter to review Applications for an owned Job and record a supported recruitment decision.",
        "Preconditions": join(
            "P1. The Recruiter account is active and authenticated.",
            "P2. The selected Job belongs to the Recruiter.",
        ),
        "Trigger": "The Recruiter opens the applicant workspace for an owned Job.",
        "Main Flow": join(
            "1. The Recruiter selects an owned Job.",
            "2. The System displays its Applications and supported filters.",
            "3. The Recruiter selects an applicant.",
            "4. The System displays available Candidate, CV, Portfolio, Application, Matching, and Potential information.",
            "5. The Recruiter reviews the information and selects Approve or Reject.",
            "6. The System verifies ownership, CV status, and the requested decision.",
            "7. The System updates the Application and displays the result.",
        ),
        "Alternative Flows": sections(
            join(
                "A1 – Review without a decision (from Step 5)",
                "1. The Recruiter closes the detail, and the Application remains unchanged.",
            ),
            join(
                "A2 – No applicants match (from Step 2)",
                "1. The System displays an empty state and allows filters or the selected Job to change.",
            ),
        ),
        "Exception Flows": sections(
            join(
                "E1 – Job or Application is unavailable or unauthorized",
                "1. The System denies access or rejects the decision without changing status.",
            ),
            join(
                "E2 – Associated CV is BANNED (from Step 6)",
                "1. The System rejects the decision and preserves the Application status.",
            ),
            join(
                "E3 – Requested status is unsupported (from Step 6)",
                "1. The System rejects the update and preserves the existing status.",
            ),
        ),
        "Postconditions": join(
            "Success Postconditions",
            "S1. The Application reflects the supported Recruiter decision and is visible to the Candidate.",
            "",
            "Minimal Guarantees",
            "F1. Review alone changes nothing.",
            "F2. Unauthorized, BANNED-CV, or unsupported decisions do not change status.",
        ),
        "Related Use Cases": join("UC-03 – Manage Job Applications", "UC-08 – Manage Employer Profile and Job Postings", "UC-10 – Manage Talent Pool and Invitations", "UC-12 – Report Suspicious Recruitment Content"),
    },
    "UC-10": {
        "Description": "This use case allows a Recruiter to discover suitable Candidates for an owned Job, manage Job-specific CV bookmarks, and manage invitations.",
        "Preconditions": join(
            "P1. The Recruiter account is active and authenticated.",
            "P2. The selected Job belongs to the Recruiter.",
        ),
        "Trigger": "The Recruiter opens the Talent Pool for an owned Job.",
        "Main Flow": join(
            "1. The Recruiter selects an owned Job.",
            "2. The System displays Candidates with available Matchings, separated into High and Potential groups.",
            "3. The Recruiter filters and selects a Candidate.",
            "4. The System displays the available CV, score, label, and explanation.",
            "5. The Recruiter selects Invite.",
            "6. The System verifies the ACTIVE Job, Candidate default CV, and existing Application state.",
            "7. The System creates an INVITED Application.",
            "8. The System displays the Candidate in the invited group.",
        ),
        "Alternative Flows": sections(
            join(
                "A1 – Add or remove a bookmark (from Step 4)",
                "1. The System updates the Job-specific bookmark and bookmark view.",
            ),
            join(
                "A2 – Withdraw a pending invitation (from Step 8)",
                "1. The Recruiter confirms withdrawal of an INVITED Application.",
                "2. The System removes it and updates the invited group.",
            ),
            join(
                "A3 – An Application already exists (from Step 6)",
                "1. The System displays the existing state instead of creating a duplicate.",
            ),
            join(
                "A4 – Talent Pool is empty (from Step 2)",
                "1. The System displays an empty state; no bookmark or invitation is created.",
            ),
        ),
        "Exception Flows": sections(
            join(
                "E1 – Job is unavailable, unauthorized, or not ACTIVE (from Step 6)",
                "1. The System rejects the invitation and creates nothing.",
            ),
            join(
                "E2 – Candidate has no default CV (from Step 6)",
                "1. The System rejects the invitation and creates nothing.",
            ),
            join(
                "E3 – Invitation cannot be withdrawn (from A2)",
                "1. The Application is unavailable, unauthorized, or no longer INVITED.",
                "2. The System rejects withdrawal and preserves it.",
            ),
        ),
        "Postconditions": join(
            "Success Postconditions",
            "S1. Bookmarks reflect the Recruiter's Job-specific selections.",
            "S2. A new invitation is INVITED; a withdrawn pending invitation is removed.",
            "S3. Existing Applications are not duplicated.",
            "",
            "Minimal Guarantees",
            "F1. Invalid Job, CV, ownership, or withdrawal state changes nothing.",
        ),
        "Related Use Cases": join("UC-03 – Manage Job Applications", "UC-08 – Manage Employer Profile and Job Postings", "UC-09 – Review and Process Applicants", "UC-12 – Report Suspicious Recruitment Content"),
    },
    "UC-11": {
        "Description": "This use case allows a Recruiter to review available analytics for owned Jobs, Applications, recruitment trends, and skill gaps.",
        "Preconditions": join(
            "P1. The Recruiter account is active and authenticated.",
            "P2. Job-specific analytics apply only to an owned Job.",
        ),
        "Trigger": "The Recruiter opens recruitment analytics.",
        "Main Flow": join(
            "1. The Recruiter opens recruitment analytics and selects a period.",
            "2. The System displays Recruiter-level Job and Application statistics.",
            "3. The Recruiter selects an owned Job.",
            "4. The System verifies ownership.",
            "5. The System displays the Job funnel, skill gaps, and available trends.",
            "6. The Recruiter reviews the results.",
        ),
        "Alternative Flows": sections(
            join(
                "A1 – Review aggregate analytics (from Step 3)",
                "1. No Job is selected, so the System continues to display Recruiter-level analytics.",
            ),
            join(
                "A2 – Limited analytics data (from Step 5)",
                "1. The System displays available values and a limited-data state for missing sections.",
            ),
        ),
        "Exception Flows": join(
            "E1 – Job is unavailable or unauthorized (from Step 4)",
            "1. The System denies Job-specific analytics while keeping the Recruiter overview available.",
        ),
        "Postconditions": join(
            "Success Postconditions",
            "S1. Available aggregate or Job-specific analytics are displayed.",
            "",
            "Minimal Guarantees",
            "F1. Another Recruiter's analytics are not exposed, and review changes no operational data.",
        ),
        "Related Use Cases": join("UC-08 – Manage Employer Profile and Job Postings", "UC-09 – Review and Process Applicants", "UC-10 – Manage Talent Pool and Invitations"),
    },
    "UC-12": {
        "Description": "This use case allows a Candidate to report a suspicious Job and a Recruiter to report a suspicious visible CV for Administrator review.",
        "Preconditions": join(
            "P1. The reporting account is active and authenticated.",
            "P2. A Candidate reports an ACTIVE Job.",
            "P3. A Recruiter reports a CV visible through an Application or Matching for an owned Job.",
        ),
        "Trigger": "A Candidate or Recruiter selects Report for eligible recruitment content.",
        "Main Flow": join(
            "1. The Candidate opens an ACTIVE Job and selects Report.",
            "2. The System displays supported reasons and a comment field.",
            "3. The Candidate selects a reason and submits.",
            "4. The System verifies actor, target, reason, and duplicate state.",
            "5. The System creates a PENDING report.",
            "6. The System confirms submission for Administrator review.",
        ),
        "Alternative Flows": join(
            "A1 – Recruiter reports a CV (alternative from the Trigger)",
            "1. The Recruiter opens a CV visible through an owned Job and submits a reason, comment, and Job.",
            "2. The System verifies ownership and visibility, creates a PENDING report, and confirms submission.",
        ),
        "Exception Flows": sections(
            join(
                "E1 – Target is unavailable or ineligible (from Step 4 or A1)",
                "1. The System rejects a missing/non-ACTIVE Job or a CV outside the Recruiter's visible owned-Job context.",
                "2. No report is created.",
            ),
            join(
                "E2 – A pending report already exists (from Step 4 or A1)",
                "1. The System rejects the duplicate and preserves the existing pending report.",
            ),
            join(
                "E3 – Reason is unsupported",
                "1. The System rejects the request and creates no report.",
            ),
        ),
        "Postconditions": join(
            "Success Postconditions",
            "S1. A valid report is PENDING and available to Administrator moderation.",
            "S2. The target's pending-report count reflects the report.",
            "",
            "Minimal Guarantees",
            "F1. Invalid, unauthorized, or duplicate requests create no additional report.",
        ),
        "Related Use Cases": join("UC-02 – Explore Jobs", "UC-09 – Review and Process Applicants", "UC-10 – Manage Talent Pool and Invitations", "UC-14 – Review and Resolve Content Reports"),
    },
    "UC-13": {
        "Description": "This use case allows an Administrator to control User-account access and public Job visibility.",
        "Preconditions": "P1. The Administrator account is active and authenticated with the Administrator role.",
        "Trigger": "The Administrator opens User or Job administration and selects an action.",
        "Main Flow": join(
            "1. The Administrator opens User administration and selects a User.",
            "2. The System displays account information and active state.",
            "3. The Administrator selects Suspend.",
            "4. The System verifies that the target is not the current Administrator.",
            "5. The System sets the account inactive.",
            "6. The System confirms suspension.",
        ),
        "Alternative Flows": sections(
            join(
                "A1 – Activate a User (from Step 3)",
                "1. The Administrator selects an inactive User, and the System activates the account.",
            ),
            join(
                "A2 – Hide an ACTIVE Job (alternative from the Trigger)",
                "1. The Administrator selects an ACTIVE Job, and the System changes it to HIDDEN_BY_ADMIN.",
            ),
            join(
                "A3 – Restore a hidden Job (from A2)",
                "1. The Administrator selects a HIDDEN_BY_ADMIN Job, and the System restores it to ACTIVE.",
            ),
            join(
                "A4 – No records match",
                "1. The System displays an empty state and allows filters to change.",
            ),
        ),
        "Exception Flows": sections(
            join(
                "E1 – Administrator attempts self-suspension (from Step 4)",
                "1. The System rejects the request; the Administrator remains active.",
            ),
            join(
                "E2 – Target is unavailable",
                "1. The System rejects the action and changes no state.",
            ),
            join(
                "E3 – Job transition is not allowed (from A2 or A3)",
                "1. The System rejects Hide for a non-ACTIVE Job or Restore for a Job not HIDDEN_BY_ADMIN.",
                "2. The Job remains unchanged.",
            ),
        ),
        "Postconditions": join(
            "Success Postconditions",
            "S1. A suspended User is inactive; an activated User is active.",
            "S2. A hidden Job is HIDDEN_BY_ADMIN; a restored Job is ACTIVE.",
            "",
            "Minimal Guarantees",
            "F1. The Administrator cannot suspend their own account.",
            "F2. Invalid requests do not change target state.",
        ),
        "Related Use Cases": join("UC-08 – Manage Employer Profile and Job Postings", "UC-14 – Review and Resolve Content Reports"),
    },
    "UC-14": {
        "Description": "This use case allows an Administrator to review reports grouped by target and resolve them by dismissal or by banning the reported Job or CV.",
        "Preconditions": "P1. The Administrator account is active and authenticated with the Administrator role.",
        "Trigger": "The Administrator opens the content-report queue.",
        "Main Flow": join(
            "1. The System displays pending reports grouped by Job or CV target.",
            "2. The Administrator selects a case.",
            "3. The System displays the target, reasons, comments, count, and status.",
            "4. The Administrator reviews the evidence and selects Dismiss.",
            "5. The System verifies that pending reports remain.",
            "6. The System marks them DISMISSED and clears the target count.",
            "7. The System displays the resolved case.",
        ),
        "Alternative Flows": sections(
            join(
                "A1 – Ban the reported target (from Step 4)",
                "1. The Administrator selects Ban and may add a note.",
                "2. The System sets the Job or CV to BANNED; for a CV, it removes the default designation.",
                "3. The System marks pending reports ACTIONED, clears the count, and displays the case.",
            ),
            join(
                "A2 – Report queue is empty (from Step 1)",
                "1. The System displays an empty queue; no moderation action occurs.",
            ),
        ),
        "Exception Flows": sections(
            join(
                "E1 – Report target is unavailable (from Step 3)",
                "1. The System cannot resolve the case and records no decision.",
            ),
            join(
                "E2 – No pending reports remain (from Step 5)",
                "1. The System rejects repeated resolution and preserves existing states.",
            ),
        ),
        "Postconditions": join(
            "Success Postconditions",
            "S1. Dismissed reports are DISMISSED; reports for a banned target are ACTIONED.",
            "S2. A banned target is BANNED and has no pending-report count.",
            "",
            "Minimal Guarantees",
            "F1. Review alone changes no target state.",
            "F2. Failed or repeated resolution does not partially change the case.",
        ),
        "Related Use Cases": join("UC-12 – Report Suspicious Recruitment Content", "UC-13 – Administer Platform Access and Job Visibility"),
    },
}


EDITABLE_FIELDS = tuple(next(iter(SHORT.values())).keys())


def text_hash(items):
    return hashlib.sha256("\n\u241e\n".join(items).encode("utf-8")).hexdigest()


def table_values(table):
    return {row.cells[0].text.strip(): row.cells[1].text.strip() for row in table.rows if len(row.cells) >= 2}


def use_case_tables(document):
    found = {}
    for index, table in enumerate(document.tables):
        values = table_values(table)
        use_case_id = values.get("Use Case ID")
        if use_case_id in SHORT:
            found[use_case_id] = (index, table)
    if set(found) != set(SHORT):
        raise RuntimeError(f"Use Case table mismatch: {sorted(found)}")
    return found


def set_cell_text(cell, text):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), "Times New Roman")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def main():
    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(REPORT, BACKUP)

    before = Document(REPORT)
    before_paragraph_hash = text_hash([paragraph.text for paragraph in before.paragraphs])
    before_tables = use_case_tables(before)
    use_case_indexes = {index for index, _ in before_tables.values()}
    before_other_table_hash = text_hash([
        cell.text
        for index, table in enumerate(before.tables) if index not in use_case_indexes
        for row in table.rows for cell in row.cells
    ])
    before_fixed = {
        uc_id: {key: values[key] for key in ("Use Case ID", "Use Case Name", "Primary Actor(s)", "Secondary Actor(s)", "Priority")}
        for uc_id, (_, table) in before_tables.items()
        for values in [table_values(table)]
    }
    before_words = sum(len(table_values(table).get(field, "").split()) for _, table in before_tables.values() for field in EDITABLE_FIELDS)

    for uc_id, (_, table) in before_tables.items():
        row_by_field = {row.cells[0].text.strip(): row for row in table.rows if len(row.cells) >= 2}
        for field, text in SHORT[uc_id].items():
            if field not in row_by_field:
                raise RuntimeError(f"Missing {field} in {uc_id}")
            set_cell_text(row_by_field[field].cells[1], text)

    before.save(REPORT)

    after = Document(REPORT)
    after_tables = use_case_tables(after)
    after_paragraph_hash = text_hash([paragraph.text for paragraph in after.paragraphs])
    after_other_table_hash = text_hash([
        cell.text
        for index, table in enumerate(after.tables) if index not in use_case_indexes
        for row in table.rows for cell in row.cells
    ])
    after_fixed = {
        uc_id: {key: values[key] for key in ("Use Case ID", "Use Case Name", "Primary Actor(s)", "Secondary Actor(s)", "Priority")}
        for uc_id, (_, table) in after_tables.items()
        for values in [table_values(table)]
    }
    after_words = sum(len(table_values(table).get(field, "").split()) for _, table in after_tables.values() for field in EDITABLE_FIELDS)

    if before_paragraph_hash != after_paragraph_hash:
        raise RuntimeError("A paragraph outside the tables changed")
    if before_other_table_hash != after_other_table_hash:
        raise RuntimeError("A non-Use-Case table changed")
    if before_fixed != after_fixed:
        raise RuntimeError("A fixed Use Case field changed")
    for uc_id, (_, table) in after_tables.items():
        if len(table.rows) != 14:
            raise RuntimeError(f"{uc_id} no longer has the 13-field structure")

    reduction = (before_words - after_words) / before_words * 100
    print(f"before_words={before_words}")
    print(f"after_words={after_words}")
    print(f"reduction_percent={reduction:.1f}")
    print("outside_paragraphs_unchanged=true")
    print("other_tables_unchanged=true")
    print("fixed_use_case_fields_unchanged=true")
    print(REPORT)
    print(BACKUP)


if __name__ == "__main__":
    main()
