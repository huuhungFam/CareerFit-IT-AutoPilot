from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def iter_block_items(document: Document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield "P", child
        elif child.tag == qn("w:tbl"):
            yield "T", child


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx")
    parser.add_argument("--out", required=True)
    args = parser.parse_args()

    doc = Document(args.docx)
    paragraph_by_element = {p._p: p for p in doc.paragraphs}
    table_by_element = {t._tbl: t for t in doc.tables}

    lines: list[str] = []
    p_index = 0
    t_index = 0
    for kind, element in iter_block_items(doc):
        if kind == "P":
            paragraph = paragraph_by_element.get(element)
            if paragraph is None:
                continue
            p_index += 1
            text = " ".join(paragraph.text.split())
            drawings = len(element.xpath(".//w:drawing | .//w:pict"))
            image_targets: list[str] = []
            for blip in element.xpath(".//a:blip"):
                rel_id = blip.get(qn("r:embed"))
                if rel_id and rel_id in doc.part.rels:
                    image_targets.append(doc.part.rels[rel_id].target_ref)
            if text or drawings:
                lines.append(
                    f"P{p_index:04d}\tSTYLE={paragraph.style.name!r}\tDRAWINGS={drawings}"
                    f"\tIMAGES={','.join(image_targets)}\t{text}"
                )
        else:
            table = table_by_element.get(element)
            if table is None:
                continue
            t_index += 1
            lines.append(f"TABLE {t_index}: rows={len(table.rows)} cols={len(table.columns)}")
            for r_idx, row in enumerate(table.rows, 1):
                cells = [" ".join(cell.text.split()) for cell in row.cells]
                lines.append(f"  R{r_idx}: " + " || ".join(cells))

    lines.append(f"SUMMARY paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} inline_shapes={len(doc.inline_shapes)}")
    Path(args.out).write_text("\n".join(lines), encoding="utf-8")


if __name__ == "__main__":
    main()
