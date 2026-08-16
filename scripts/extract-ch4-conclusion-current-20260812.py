from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(r"C:\CODING\Thesis")
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"


def clean(text: str) -> str:
    return " ".join(text.replace("\u00a0", " ").split())


def word_count(text: str) -> int:
    return len(re.findall(r"[A-Za-z0-9_@'–/-]+", text))


def iter_blocks(document: Document):
    paragraphs = {p._p: p for p in document.paragraphs}
    tables = {t._tbl: t for t in document.tables}
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p") and child in paragraphs:
            yield "P", paragraphs[child]
        elif child.tag == qn("w:tbl") and child in tables:
            yield "T", tables[child]


def main() -> None:
    document = Document(DOCX)
    active = False
    current_part = ""
    section_text: dict[str, list[str]] = {}

    for kind, item in iter_blocks(document):
        if kind == "P":
            text = clean(item.text)
            style = item.style.name if item.style else ""
            if text == "CHAPTER 4. TESTING AND EVALUATION":
                active = True
                current_part = "CHAPTER 4"
            elif text == "PART 3. CONCLUSION":
                current_part = "CONCLUSION"
            elif text == "REFERENCES":
                break

            if not active:
                continue

            print(f"P|{style}|{text}")
            if style.startswith("Heading") and text:
                section_text.setdefault(text, [])
            elif text and style not in {"Table Caption", "Figure Caption"}:
                headings = [key for key in section_text]
                if headings:
                    section_text[headings[-1]].append(text)
        elif active:
            rows = []
            for row in item.rows:
                rows.append(" || ".join(clean(cell.text) for cell in row.cells))
            print(f"T|rows={len(item.rows)}|cols={len(item.columns)}")
            for row in rows:
                print(f"R|{row}")

    print("=== DIRECT NARRATIVE COUNTS ===")
    for heading, paragraphs in section_text.items():
        print(f"COUNT|{heading}|{word_count(' '.join(paragraphs))}")


if __name__ == "__main__":
    main()
