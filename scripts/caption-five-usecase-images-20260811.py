from __future__ import annotations

from pathlib import Path
import hashlib
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-before-20260811-five-usecase-captions.docx"
DOWNLOADS = Path(r"C:\Users\PhamHuuHwng\Downloads")

SOURCE_FILES = {
    "overall": DOWNLOADS / "UC-overall.png",
    "candidate": DOWNLOADS / "candidate.png",
    "recruiter": DOWNLOADS / "recruiter.png",
    "shared": DOWNLOADS / "uc12.png",
    "administrator": DOWNLOADS / "admin.png",
}

CAPTIONS = {
    "overall": "Figure 1.2. Overall CareerFit use-case diagram",
    "candidate": "Figure 1.3. Candidate use-case diagram",
    "recruiter": "Figure 1.7. Recruiter use-case diagram",
    "shared": "Figure 1.8. Shared Candidate–Recruiter reporting use-case diagram",
    "administrator": "Figure 1.9. Administrator use-case diagram",
}

ALT_TEXT = {
    "overall": "Overall CareerFit use-case diagram showing Guest, Candidate, Recruiter, and Administrator associations with UC-01 through UC-14.",
    "candidate": "Candidate use-case diagram for UC-01 through UC-07, with Guest access to UC-02.",
    "recruiter": "Recruiter use-case diagram for UC-08 through UC-11.",
    "shared": "Temporary shared reporting use-case diagram for UC-12; the left actor must be corrected from Guest to Recruiter.",
    "administrator": "Administrator use-case diagram for UC-13 and UC-14.",
}

SHARED_NOTE = (
    "NOTE: In Figure 1.8, the left actor must be changed from Guest to Recruiter before final "
    "submission. UC-12 is initiated by Candidate and Recruiter, as defined in Table 1.16."
)


def source_hashes():
    result = {}
    for key, path in SOURCE_FILES.items():
        if not path.exists():
            raise RuntimeError(f"Missing source image: {path}")
        result[hashlib.sha256(path.read_bytes()).hexdigest()] = key
    return result


def image_key(document, paragraph, hashes):
    blips = paragraph._p.xpath(".//a:blip")
    if not blips:
        return None
    rid = blips[0].get(qn("r:embed"))
    blob = document.part.related_parts[rid].blob
    return hashes.get(hashlib.sha256(blob).hexdigest())


def set_text(paragraph, text, size=11, italic=True):
    paragraph.clear()
    paragraph.paragraph_format.space_before = Pt(3 if italic and size == 11 else 0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0 if size == 11 else 1.2
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(key), "Times New Roman")


def new_paragraph_after(reference, document, style):
    element = OxmlElement("w:p")
    reference._p.addnext(element)
    paragraph = Paragraph(element, reference._parent)
    paragraph.style = document.styles[style]
    return paragraph


def remove_paragraph(paragraph):
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)


def main():
    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(REPORT, BACKUP)

    document = Document(REPORT)
    hashes = source_hashes()
    found = {}
    for paragraph in document.paragraphs:
        key = image_key(document, paragraph, hashes)
        if key:
            if key in found:
                raise RuntimeError(f"Duplicate inserted image: {key}")
            found[key] = paragraph
    if set(found) != set(CAPTIONS):
        raise RuntimeError(f"Expected five inserted images; found {sorted(found)}")

    # Remove redundant typed titles above the diagrams.
    redundant_titles = {
        "Candidate Use Case Diagram",
        "Recruiter Use Case Diagram",
        "Administrator Use Case Diagram",
    }
    for paragraph in list(document.paragraphs):
        if paragraph.text.strip() in redundant_titles:
            remove_paragraph(paragraph)

    # Remove the obsolete redraw note under the newly replaced overall diagram.
    for paragraph in list(document.paragraphs):
        if paragraph.text.strip().startswith("NOTE: Figure 1.2 is temporarily retained without redrawing"):
            remove_paragraph(paragraph)

    # Clean only empty paragraphs immediately surrounding the newly added role diagrams.
    for key in ("candidate", "recruiter", "shared", "administrator"):
        image_paragraph = found[key]
        previous = image_paragraph._p.getprevious()
        while previous is not None:
            wrapped = next((p for p in document.paragraphs if p._p is previous), None)
            if wrapped is None or wrapped.text.strip() or wrapped._p.xpath(".//a:blip"):
                break
            candidate = previous.getprevious()
            remove_paragraph(wrapped)
            previous = candidate
        following = image_paragraph._p.getnext()
        while following is not None:
            wrapped = next((p for p in document.paragraphs if p._p is following), None)
            if wrapped is None or wrapped.text.strip() or wrapped._p.xpath(".//a:blip"):
                break
            candidate = following.getnext()
            remove_paragraph(wrapped)
            following = candidate

    # Normalize image placement and accessible descriptions.
    for key, paragraph in found.items():
        paragraph.style = document.styles["Normal"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.keep_together = True
        for prop in paragraph._p.xpath(".//wp:docPr"):
            prop.set("title", CAPTIONS[key].split(". ", 1)[1])
            prop.set("descr", ALT_TEXT[key])

    # Update the caption already attached to the overall diagram.
    overall = found["overall"]
    overall_next = overall._p.getnext()
    overall_caption = next((p for p in document.paragraphs if p._p is overall_next), None)
    if overall_caption is None or overall_caption.style.name != "Figure Caption":
        raise RuntimeError("Overall diagram is not followed by its existing caption")
    set_text(overall_caption, CAPTIONS["overall"])
    overall_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    overall_caption.paragraph_format.keep_together = True

    # Insert captions below the four newly added role/shared diagrams.
    for key in ("candidate", "recruiter", "shared", "administrator"):
        image_paragraph = found[key]
        caption = new_paragraph_after(image_paragraph, document, "Figure Caption")
        set_text(caption, CAPTIONS[key])
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.keep_together = True
        if key == "shared":
            note = new_paragraph_after(caption, document, "Normal")
            set_text(note, SHARED_NOTE, size=13, italic=True)
            note.alignment = WD_ALIGN_PARAGRAPH.LEFT
            note.paragraph_format.keep_together = True

    # Renumber the pre-existing Chapter 1 workflow captions around the new diagrams.
    renames = {
        "Figure 1.3. CV review, confirmation, and matching sequence": "Figure 1.4. CV review, confirmation, and matching sequence",
        "Figure 1.4. Feedback learning and recomputation sequence": "Figure 1.5. Feedback learning and recomputation sequence",
        "Figure 1.5. AutoFit decision flow": "Figure 1.6. AutoFit decision flow",
    }
    for old, new in renames.items():
        matches = [p for p in document.paragraphs if p.text.strip() == old]
        if len(matches) != 1:
            raise RuntimeError(f"Expected one caption {old}; found {len(matches)}")
        set_text(matches[0], new)
        matches[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        matches[0].paragraph_format.keep_together = True

    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    document.save(REPORT)

    checked = Document(REPORT)
    captions = [p.text for p in checked.paragraphs if p.style.name == "Figure Caption" and p.text.startswith("Figure 1.")]
    expected = [
        "Figure 1.1. CareerFit system context",
        CAPTIONS["overall"],
        CAPTIONS["candidate"],
        "Figure 1.4. CV review, confirmation, and matching sequence",
        "Figure 1.5. Feedback learning and recomputation sequence",
        "Figure 1.6. AutoFit decision flow",
        CAPTIONS["recruiter"],
        CAPTIONS["shared"],
        CAPTIONS["administrator"],
    ]
    if captions != expected:
        raise RuntimeError(f"Unexpected Chapter 1 captions: {captions}")
    if any(p.style.name.startswith("Heading") and not p.text.strip() for p in checked.paragraphs):
        raise RuntimeError("An empty heading remains after cleanup")

    print("chapter_1_figure_captions=9")
    print("five_inserted_images_centered=true")
    print("empty_headings_removed=true")
    print("uc12_actor_note_added=true")
    print(REPORT)
    print(BACKUP)


if __name__ == "__main__":
    main()
