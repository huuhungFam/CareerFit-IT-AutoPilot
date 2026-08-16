from pathlib import Path
import shutil

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-before-20260809-caption-placement.docx"


def body_items(document):
    paragraphs = {id(paragraph._p): paragraph for paragraph in document.paragraphs}
    tables = {id(table._tbl): table for table in document.tables}
    result = []
    for element in document._element.body.iterchildren():
        if id(element) in paragraphs:
            result.append(("paragraph", paragraphs[id(element)]))
        elif id(element) in tables:
            result.append(("table", tables[id(element)]))
    return result


def has_picture(paragraph):
    return bool(paragraph._p.xpath(".//a:blip"))


def main():
    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(REPORT, BACKUP)

    document = Document(REPORT)

    captions = [paragraph for paragraph in document.paragraphs
                if paragraph.style.name == "Table Caption"
                and paragraph.text.strip().startswith("Table 1.1.")]
    if len(captions) != 1:
        raise RuntimeError(f"Expected exactly one Table 1.1 caption, found {len(captions)}")
    caption = captions[0]

    items = body_items(document)
    caption_index = next(index for index, item in enumerate(items)
                         if item[0] == "paragraph" and item[1]._p is caption._p)
    if caption_index == 0 or items[caption_index - 1][0] != "table":
        raise RuntimeError("Table 1.1 caption is not immediately below a table")
    table = items[caption_index - 1][1]
    table._tbl.addprevious(caption._p)

    # Captions above tables should remain with the first table row.
    for paragraph in document.paragraphs:
        if paragraph.style.name == "Table Caption":
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True

    # Images/screens should remain with their captions below them.
    items = body_items(document)
    for index, (kind, item) in enumerate(items):
        if kind != "paragraph" or item.style.name != "Figure Caption":
            continue
        if index == 0 or items[index - 1][0] != "paragraph" or not has_picture(items[index - 1][1]):
            raise RuntimeError(f"Figure caption is not immediately below an image: {item.text}")
        items[index - 1][1].paragraph_format.keep_with_next = True
        item.paragraph_format.keep_together = True

    document.save(REPORT)

    # Structural verification after saving.
    checked = Document(REPORT)
    items = body_items(checked)
    table_caption_count = 0
    figure_caption_count = 0
    for index, (kind, item) in enumerate(items):
        if kind != "paragraph":
            continue
        if item.style.name == "Table Caption":
            table_caption_count += 1
            if index + 1 >= len(items) or items[index + 1][0] != "table":
                raise RuntimeError(f"Table caption is not immediately above a table: {item.text}")
        elif item.style.name == "Figure Caption":
            figure_caption_count += 1
            if index == 0 or items[index - 1][0] != "paragraph" or not has_picture(items[index - 1][1]):
                raise RuntimeError(f"Figure caption is not immediately below an image: {item.text}")

    print(f"table_captions_above={table_caption_count}")
    print(f"figure_captions_below={figure_caption_count}")
    print(REPORT)
    print(BACKUP)


if __name__ == "__main__":
    main()
