from __future__ import annotations

from collections import Counter
from pathlib import Path
import re

from docx import Document

import hashlib


ROOT = Path(__file__).resolve().parents[1]
CURRENT = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-before-20260812-content-reduction.docx"
PROTECTED = (
    "2.2.2 Bag-of-Words Representation",
    "2.3.1 Term Frequency and Inverse Document Frequency",
    "2.3.2 Cosine Similarity",
)


def section_payload(document: Document, heading_text: str) -> tuple[tuple[str, str], ...]:
    matches = [p for p in document.paragraphs if " ".join(p.text.split()) == heading_text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one protected heading {heading_text!r}; found {len(matches)}")
    heading = matches[0]
    level = int(heading.style.name.split()[-1])
    paragraphs = document.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p._p is heading._p)
    result = []
    for paragraph in paragraphs[start + 1 :]:
        if paragraph.style.name.startswith("Heading "):
            next_level = int(paragraph.style.name.split()[-1])
            if next_level <= level:
                break
        result.append((paragraph.style.name, paragraph.text))
    return tuple(result)


def payload_hash(payload: tuple[tuple[str, str], ...]) -> str:
    return hashlib.sha256(repr(payload).encode("utf-8")).hexdigest()


def main() -> None:
    current = Document(CURRENT)
    backup = Document(BACKUP)

    for title in PROTECTED:
        before = payload_hash(section_payload(backup, title))
        after = payload_hash(section_payload(current, title))
        if before != after:
            raise RuntimeError(f"Protected section changed: {title}")

    body_text = "\n".join(p.text for p in current.paragraphs)
    for marker in ("\ufffd", "\x00", "Error! Reference source not found", "Error! Bookmark not defined"):
        if marker in body_text:
            raise RuntimeError(f"Invalid marker found: {marker!r}")

    drawing_count = len(current.element.body.xpath(".//w:drawing"))
    caption_counts = Counter(
        p.style.name for p in current.paragraphs if p.style.name in {"Figure Caption", "Table Caption", "Caption"}
    )
    use_case_headings = [
        p.text
        for p in current.paragraphs
        if p.style.name.startswith("Heading") and re.match(r"^1\.5\.(?:[1-9]|1[0-4])\s", p.text.strip())
    ]

    if len(current.tables) != 72:
        raise RuntimeError(f"Expected 72 tables, found {len(current.tables)}")
    if drawing_count != 44:
        raise RuntimeError(f"Expected 44 drawings, found {drawing_count}")
    if len(use_case_headings) != 14:
        raise RuntimeError(f"Expected 14 Use Case headings, found {len(use_case_headings)}")

    print(f"paragraphs={len(current.paragraphs)}")
    print(f"tables={len(current.tables)}")
    print(f"drawings={drawing_count}")
    print(f"use_case_headings={len(use_case_headings)}")
    print(f"caption_styles={dict(caption_counts)}")
    print("protected_sections=unchanged")
    print("field_errors=none")


if __name__ == "__main__":
    main()
