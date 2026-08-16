from __future__ import annotations

import re
from collections import Counter
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"


def clean(text: str) -> str:
    return " ".join(text.split())


def words(text: str) -> list[str]:
    return re.findall(r"[A-Za-z0-9_@'/-]+", text)


document = Document(REPORT)
headings: list[tuple[int, int, str]] = []
for index, paragraph in enumerate(document.paragraphs):
    match = re.fullmatch(r"Heading (\d+)", paragraph.style.name)
    if match and clean(paragraph.text):
        headings.append((index, int(match.group(1)), clean(paragraph.text)))

print("CURRENT STRUCTURE WITH NARRATIVE WORD COUNTS")
for position, (index, level, title) in enumerate(headings):
    if index < 379 or index >= 869:
        continue
    next_heading_index = headings[position + 1][0] if position + 1 < len(headings) else len(document.paragraphs)
    own_text = " ".join(
        clean(paragraph.text)
        for paragraph in document.paragraphs[index + 1:next_heading_index]
        if clean(paragraph.text)
    )
    subtree_end = len(document.paragraphs)
    for following_index, following_level, _ in headings[position + 1:]:
        if following_level <= level:
            subtree_end = following_index
            break
    subtree_text = " ".join(
        clean(paragraph.text)
        for paragraph in document.paragraphs[index + 1:subtree_end]
        if clean(paragraph.text)
    )
    print(
        f"{level}\t{title}\town={len(words(own_text))}\t"
        f"subtree={len(words(subtree_text))}\tpreview={own_text[:220]}"
    )

sentences: list[tuple[int, str]] = []
for index, paragraph in enumerate(document.paragraphs[379:869], start=379):
    if paragraph.style.name.startswith("Heading"):
        continue
    for sentence in re.split(r"(?<=[.!?])\s+", clean(paragraph.text)):
        sentence = clean(sentence)
        if len(words(sentence)) >= 12:
            sentences.append((index, sentence))

normalized = Counter(
    re.sub(r"[^a-z0-9]+", " ", sentence.lower()).strip()
    for _, sentence in sentences
)
print("\nEXACT REPEATED SENTENCES")
for key, count in normalized.most_common():
    if count < 2:
        break
    examples = [(index, sentence) for index, sentence in sentences if re.sub(r"[^a-z0-9]+", " ", sentence.lower()).strip() == key]
    print(count, examples)

heading_for_paragraph: dict[int, str] = {}
current_heading = ""
for index, paragraph in enumerate(document.paragraphs):
    if paragraph.style.name.startswith("Heading") and clean(paragraph.text):
        current_heading = clean(paragraph.text)
    heading_for_paragraph[index] = current_heading

stop = {
    "a", "an", "and", "are", "as", "at", "be", "because", "by", "can", "for",
    "from", "has", "have", "in", "is", "it", "of", "on", "or", "that", "the",
    "their", "this", "to", "was", "were", "while", "with", "without",
}

def token_set(sentence: str) -> set[str]:
    return {token.lower() for token in words(sentence) if token.lower() not in stop and len(token) > 2}


near_duplicates: list[tuple[float, int, str, int, str]] = []
for left in range(len(sentences)):
    left_index, left_sentence = sentences[left]
    left_tokens = token_set(left_sentence)
    if len(left_tokens) < 8:
        continue
    for right in range(left + 1, len(sentences)):
        right_index, right_sentence = sentences[right]
        if heading_for_paragraph[left_index] == heading_for_paragraph[right_index]:
            continue
        right_tokens = token_set(right_sentence)
        union = left_tokens | right_tokens
        score = len(left_tokens & right_tokens) / len(union) if union else 0
        if score >= 0.50:
            near_duplicates.append((score, left_index, left_sentence, right_index, right_sentence))

print("\nNEAR-DUPLICATE SENTENCES ACROSS SECTIONS")
for item in sorted(near_duplicates, reverse=True)[:30]:
    score, left_index, left_sentence, right_index, right_sentence = item
    print(
        f"score={score:.2f}\t{heading_for_paragraph[left_index]} [{left_index}] {left_sentence}\t"
        f"<->\t{heading_for_paragraph[right_index]} [{right_index}] {right_sentence}"
    )

print("\nDOCUMENT COUNTS")
print(f"paragraphs={len(document.paragraphs)}")
print(f"tables={len(document.tables)}")
print(f"body_narrative_words={len(words(' '.join(clean(p.text) for p in document.paragraphs[379:869])))}")
print(f"chapter_2_words={len(words(' '.join(clean(p.text) for p in document.paragraphs[495:569])))}")
print(f"chapter_3_words={len(words(' '.join(clean(p.text) for p in document.paragraphs[569:745])))}")
print(f"chapter_4_words={len(words(' '.join(clean(p.text) for p in document.paragraphs[745:809])))}")
