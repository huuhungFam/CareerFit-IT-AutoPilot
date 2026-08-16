from __future__ import annotations

import hashlib
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-before-20260811-additional-flowcharts.docx"
FIGURES = ROOT / "Doc" / "figures"


def table_digest(document: Document) -> str:
    payload = []
    for table in document.tables:
        payload.append([[cell.text for cell in row.cells] for row in table.rows])
    return hashlib.sha256(repr(payload).encode("utf-8")).hexdigest()


def find_exact(document: Document, text: str):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph for {text!r}; found {len(matches)}")
    return matches[0]


def previous_paragraph(paragraph):
    previous = paragraph._p.getprevious()
    while previous is not None and not previous.tag.endswith("}p"):
        previous = previous.getprevious()
    if previous is None:
        raise RuntimeError(f"No previous paragraph before {paragraph.text!r}")
    return Paragraph(previous, paragraph._parent)


def remove_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def configure_image_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.page_break_before = True
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0


def configure_caption(paragraph) -> None:
    paragraph.style = "Figure Caption"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = False


def replace_figure(document: Document, old_caption: str, new_caption: str, image_name: str) -> None:
    caption = find_exact(document, old_caption)
    image = previous_paragraph(caption)
    if not image._p.xpath(".//w:drawing"):
        raise RuntimeError(f"Expected a drawing immediately before {old_caption!r}")
    remove_paragraph(image)

    image_paragraph = caption.insert_paragraph_before()
    configure_image_paragraph(image_paragraph)
    inline = image_paragraph.add_run().add_picture(str(FIGURES / image_name), width=Inches(5.2))
    inline._inline.docPr.set("title", new_caption)
    inline._inline.docPr.set("descr", new_caption)

    caption.text = new_caption
    configure_caption(caption)


def main() -> None:
    shutil.copy2(DOCX, BACKUP)
    document = Document(DOCX)
    before_tables = table_digest(document)

    replacements = (
        (
            "Figure 3.7. Seed-corpus initialization and TF-IDF construction",
            "Figure 3.7. Seed-corpus initialization and TF-IDF construction flowchart",
            "flowchart-tfidf-construction-20260811.png",
        ),
        (
            "Figure 3.12. Per-account notification policy guard",
            "Figure 3.12. Notification policy evaluation and delivery outcome flowchart",
            "flowchart-notification-policy-20260811.png",
        ),
        (
            "Figure 3.13. Hashed email-action token and confirm-then-POST flow",
            "Figure 3.13. Actionable-email confirmation and redemption flowchart",
            "flowchart-email-action-redemption-20260811.png",
        ),
    )
    for old_caption, new_caption, image_name in replacements:
        replace_figure(document, old_caption, new_caption, image_name)

    if before_tables != table_digest(document):
        raise RuntimeError("Table contents changed during the figure-only update")

    document.save(DOCX)
    print(f"Updated: {DOCX}")
    print(f"Backup:  {BACKUP}")
    print(f"Tables unchanged: {before_tables}")


if __name__ == "__main__":
    main()
