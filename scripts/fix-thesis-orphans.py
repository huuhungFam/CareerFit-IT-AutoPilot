from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
doc = Document(DOCX)

# Section breaks inserted at the start of a heading can leave a form-feed-only
# paragraph that inherits Heading 1. Keep the section break but remove heading formatting.
for paragraph in doc.paragraphs:
    if paragraph.text.strip() in {"", "\x0c"} and paragraph.style.name == "Heading 1":
        paragraph.style = doc.styles["Normal"]
        paragraph.paragraph_format.page_break_before = False
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0

# Keep each heading with the complete first explanatory paragraph. This avoids
# a heading plus one or two body lines being stranded at the bottom of a page.
paragraphs = doc.paragraphs
for index, paragraph in enumerate(paragraphs[:-1]):
    if not paragraph.style.name.startswith("Heading"):
        continue
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    cursor = index + 1
    while cursor < len(paragraphs) and paragraphs[cursor].style.name.startswith("Heading"):
        paragraphs[cursor].paragraph_format.keep_with_next = True
        paragraphs[cursor].paragraph_format.keep_together = True
        cursor += 1
    if cursor < len(paragraphs):
        first_body = paragraphs[cursor]
        if first_body.style.name in {"Normal", "List Paragraph"} and len(first_body.text.split()) <= 150:
            first_body.paragraph_format.keep_together = True

settings = doc.settings._element
update = settings.find(qn("w:updateFields"))
if update is None:
    update = OxmlElement("w:updateFields")
    settings.append(update)
update.set(qn("w:val"), "true")

doc.save(DOCX)
print(DOCX)
