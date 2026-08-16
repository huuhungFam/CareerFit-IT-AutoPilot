from pathlib import Path

from docx import Document


REPORT = Path("Doc/CareerFit-Thesis-Report.docx")
TARGET_CAPTIONS = {
    "Figure 3.2. CareerFit conceptual data model",
    "Figure 3.3. CareerFit logical data model - identity and Career Profile",
    "Figure 3.7. CareerFit physical data model - identity and Career Profile tables",
}

APPENDIX_HEADINGS_WITHOUT_EXTRA_BREAK = {
    "C.2.1 Identity, Candidate, CV, and Portfolio",
    "user_account",
    "employer_profile",
    "automation_policy",
    "content_report",
}


def paragraph_text(paragraph):
    return " ".join(paragraph.text.split())


document = Document(REPORT)
changed = []

for index, paragraph in enumerate(document.paragraphs):
    text = paragraph_text(paragraph)
    if text not in TARGET_CAPTIONS:
        continue

    image_index = index - 1
    while image_index >= 0 and not document.paragraphs[image_index].text.strip() and not document.paragraphs[image_index]._p.xpath(".//w:drawing"):
        image_index -= 1

    if image_index < 0 or not document.paragraphs[image_index]._p.xpath(".//w:drawing"):
        raise RuntimeError(f"Could not find the image paragraph preceding {text}")

    image_paragraph = document.paragraphs[image_index]
    image_paragraph.paragraph_format.page_break_before = False
    changed.append(text)

missing = TARGET_CAPTIONS.difference(changed)
if missing:
    raise RuntimeError(f"Missing captions: {sorted(missing)}")

appendix_changed = []
for paragraph in document.paragraphs:
    text = paragraph_text(paragraph)
    if text in APPENDIX_HEADINGS_WITHOUT_EXTRA_BREAK:
        paragraph.paragraph_format.page_break_before = False
        appendix_changed.append(text)

appendix_missing = APPENDIX_HEADINGS_WITHOUT_EXTRA_BREAK.difference(appendix_changed)
if appendix_missing:
    raise RuntimeError(f"Missing Appendix C headings: {sorted(appendix_missing)}")

document.save(REPORT)
print(f"Updated {len(changed)} figure paragraphs in {REPORT}")
for caption in changed:
    print(caption)
print(f"Removed redundant page breaks from {len(appendix_changed)} Appendix C headings")
