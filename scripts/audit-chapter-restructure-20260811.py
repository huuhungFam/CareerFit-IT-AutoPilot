from __future__ import annotations

import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = (
    ROOT
    / "Doc"
    / "working"
    / "CareerFit-Thesis-Report-before-20260811-chapter-restructure.docx"
)

doc = Document(DOCX)
backup = Document(BACKUP)
errors: list[str] = []


def heading_texts(document: Document, start: str, end: str) -> list[str]:
    active = False
    result = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == start:
            active = True
        if active and paragraph.style.name.startswith("Heading"):
            result.append(text)
        if text == end:
            break
    return result


expected_chapter_2 = [
    "CHAPTER 2. THEORETICAL BACKGROUND",
    "2.1 Recruitment Information and CV–Job Description Matching",
    "2.2 Text Representation and Preprocessing",
    "2.2.1 Document Normalization",
    "2.2.2 Bag-of-Words Representation",
    "2.3 TF-IDF and Cosine Similarity",
    "2.3.1 Term Frequency and Inverse Document Frequency",
    "2.3.2 Cosine Similarity",
    "2.4 Matching, Recommendation, and Ranking",
    "2.5 Relevance Feedback and the Rocchio Method",
    "2.5.1 Relevance Feedback",
    "2.5.2 Feedback Semantics in CareerFit",
    "2.6 Evaluation Metrics for Ranked Results",
    "2.6.1 Precision@K, Recall@K, and HitRate@K",
    "2.6.2 Mean Reciprocal Rank",
    "2.6.3 Discounted Cumulative Gain",
    "2.7 Human-in-the-Loop and Explainable Automation",
    "2.7.1 Levels of Human Involvement",
    "2.7.2 Explainability and Auditability",
    "2.8 Web Architecture, Security, and Audit Foundations",
    "2.8.1 Client–Server and REST",
    "2.8.2 Authentication, Authorization, and Action Tokens",
    "2.8.3 Audit Logging",
    "2.9 Related Work and Research Gap",
    "2.10 Theoretical Background Summary",
    "CHAPTER 3. SYSTEM DESIGN AND IMPLEMENTATION",
]
expected_chapter_3 = [
    "CHAPTER 3. SYSTEM DESIGN AND IMPLEMENTATION",
    "3.1 System Architecture",
    "3.2 Module Design",
    "3.3 Data Design",
    "3.3.1 Core Identity and Recruitment Data",
    "3.3.2 Automation, Communication, and Operational Data",
    "3.4 Security, Failure, and Consistency Design",
    "3.4.1 Authentication and Authorization",
    "3.4.2 Failure Handling",
    "3.4.3 Identified Design Risks",
    "3.5 Deployment Architecture",
    "3.6 Implementation Overview",
    "3.7 Authentication and Authorization Implementation",
    "3.7.1 Login and JWT Processing",
    "3.7.2 URL Rules and Ownership Checks",
    "3.8 CV Ingestion and Validation",
    "3.8.1 Upload and Manual Creation",
    "3.8.2 File Validation and Extraction",
    "3.9 Text Normalization and TF-IDF Implementation",
    "3.9.1 Text Normalization",
    "3.9.2 Static Corpus and Vector Construction",
    "3.10 Matching and Recommendation Implementation",
    "3.10.1 Scoring Service",
    "3.10.2 Matching Orchestration and Persistence",
    "3.10.3 Recommendation Service",
    "3.11 Feedback Learning with Rocchio",
    "3.12 Application and Recruiter Workflow",
    "3.13 AutoFit and Background Processing",
    "3.13.1 Async Executor and Scheduler",
    "3.13.2 Auto-Apply",
    "3.13.3 Notification Guard and Delivery",
    "3.14 Email Action and Audit Implementation",
    "3.15 Persistence and API Implementation",
    "3.16 Frontend Integration",
    "3.17 Deployment and Observability Implementation",
    "3.18 Implementation Limitations",
    "3.19 Chapter Summary",
    "CHAPTER 4. TESTING AND EVALUATION",
]
actual_chapter_2 = heading_texts(
    doc, "CHAPTER 2. THEORETICAL BACKGROUND", "CHAPTER 3. SYSTEM DESIGN AND IMPLEMENTATION"
)
actual_chapter_3 = heading_texts(
    doc, "CHAPTER 3. SYSTEM DESIGN AND IMPLEMENTATION", "CHAPTER 4. TESTING AND EVALUATION"
)
if actual_chapter_2 != expected_chapter_2:
    errors.append(f"chapter_2_heading_sequence={actual_chapter_2!r}")
if actual_chapter_3 != expected_chapter_3:
    errors.append(f"chapter_3_heading_sequence={actual_chapter_3!r}")

expected_figure_2 = [
    "Figure 2.1. Distinction between matching, recommendation, and recruitment action",
    "Figure 2.2. TF-IDF vectorization and cosine-similarity pipeline",
    "Figure 2.3. Rocchio relevance-feedback vector update",
    "Figure 2.4. Human-in-the-Loop control cycle in CareerFit",
]
expected_figure_3 = [
    "Figure 3.1. CareerFit container and component architecture",
    "Figure 3.2. CareerFit logical entity relationships",
    "Figure 3.3. Local and containerized deployment topology",
    "Figure 3.4. Backend module structure and request flow",
    "Figure 3.5. JWT authentication and authorization boundaries",
    "Figure 3.6. CV ingestion, review, and confirmation pipeline",
    "Figure 3.7. Seed-corpus initialization and TF-IDF construction",
    "Figure 3.8. Direct score and Potential assessment flow",
    "Figure 3.9. Feedback processing and post-commit learning",
    "Figure 3.10. Application and invitation state transitions",
    "Figure 3.11. Per-account notification policy guard",
    "Figure 3.12. Hashed email-action token and confirm-then-POST flow",
    "Figure 3.13. Frontend routes, server-side catalogue, and API data flow",
]
expected_table_2 = ["Table 2.1. Positioning of CareerFit against major approach categories"]
expected_table_3 = [
    "Table 3.1. Backend modules and responsibilities",
    "Table 3.2. Key integrity constraints",
    "Table 3.3. Design risks and required treatment",
    "Table 3.4. Main implementation technologies",
    "Table 3.5. CV processing states",
    "Table 3.6. Implemented score interpretation",
    "Table 3.7. Feedback handling",
    "Table 3.8. Representative API-to-service mapping",
    "Table 3.9. Verified implementation limitations",
]

for prefix, expected, style in [
    ("Figure 2.", expected_figure_2, "Figure Caption"),
    ("Figure 3.", expected_figure_3, "Figure Caption"),
    ("Table 2.", expected_table_2, "Table Caption"),
    ("Table 3.", expected_table_3, "Table Caption"),
]:
    actual = [p.text.strip() for p in doc.paragraphs if p.style.name == style and p.text.startswith(prefix)]
    if actual != expected:
        errors.append(f"caption_sequence_{prefix.strip()}={actual!r}")

for caption in expected_figure_2 + expected_figure_3 + expected_table_2 + expected_table_3:
    if not any(p.style.name == "toc 1" and p.text.strip().startswith(caption) for p in doc.paragraphs):
        errors.append(f"missing_updated_list_entry={caption}")

toc_text = "\n".join(p.text for p in doc.paragraphs if p.style.name.lower().startswith("toc"))
for stale in (
    "CHAPTER 2. THEORETICAL BACKGROUND AND SOLUTION DESIGN",
    "CHAPTER 3. SYSTEM IMPLEMENTATION",
    "PART 4. REFERENCES",
    "PART 5. APPENDICES",
    "2.11 System Architecture",
    "2.16 Chapter Summary",
):
    if stale in toc_text:
        errors.append(f"stale_toc_entry={stale}")

def use_case_tables(document: Document) -> list[list[list[str]]]:
    result = []
    for table in document.tables:
        cells = [[cell.text for cell in row.cells] for row in table.rows]
        if any(cell.strip() == "Use Case ID" for row in cells for cell in row):
            result.append(cells)
    return result


current_use_cases = use_case_tables(doc)
backup_use_cases = use_case_tables(backup)
if len(current_use_cases) != 14:
    errors.append(f"current_use_case_table_count={len(current_use_cases)}")
if current_use_cases != backup_use_cases:
    errors.append("use_case_table_content_changed")

paragraphs = doc.paragraphs
for index, paragraph in enumerate(paragraphs):
    if paragraph.style.name != "Figure Caption":
        continue
    if index == 0 or "w:drawing" not in paragraphs[index - 1]._p.xml:
        errors.append(f"figure_without_image={paragraph.text.strip()}")

all_text = "\n".join(p.text for p in paragraphs)
for marker in ("Error! Bookmark not defined.", "\ufffd", "\x00"):
    if marker in all_text:
        errors.append(f"invalid_marker={marker!r}")

if not any(p.text.strip() == "REFERENCES" and p.style.name == "Heading 1" for p in paragraphs):
    errors.append("references_heading_missing")
if not any(p.text.strip() == "APPENDICES" and p.style.name == "Heading 1" for p in paragraphs):
    errors.append("appendices_heading_missing")

reference_entries = [p.text.strip() for p in paragraphs if re.match(r"^\[\d+\]", p.text.strip())]
if len(reference_entries) != 14:
    errors.append(f"reference_count={len(reference_entries)}")
reference_heading_index = next(
    i for i, paragraph in enumerate(paragraphs)
    if paragraph.text.strip() == "REFERENCES" and paragraph.style.name == "Heading 1"
)
body_text = "\n".join(paragraph.text for paragraph in paragraphs[:reference_heading_index])
cited = {int(number) for number in re.findall(r"\[(\d+)\]", body_text)}
if cited != set(range(1, 15)):
    errors.append(f"cited_reference_numbers={sorted(cited)}")

for section_index, section in enumerate(doc.sections):
    geometry = (
        round(section.page_width.cm, 1),
        round(section.page_height.cm, 1),
        round(section.top_margin.cm, 1),
        round(section.bottom_margin.cm, 1),
        round(section.left_margin.cm, 1),
        round(section.right_margin.cm, 1),
    )
    if geometry != (21.0, 29.7, 3.0, 3.0, 3.5, 2.0):
        errors.append(f"section_{section_index}_geometry={geometry}")

normal = doc.styles["Normal"]
if normal.font.name != "Times New Roman" or round(normal.font.size.pt, 1) != 13.0:
    errors.append(f"normal_style={normal.font.name}/{normal.font.size.pt}")
if normal.paragraph_format.line_spacing != 1.2:
    errors.append(f"normal_line_spacing={normal.paragraph_format.line_spacing}")

print(f"docx={DOCX}")
print(f"paragraphs={len(paragraphs)} tables={len(doc.tables)} images={len(doc.inline_shapes)}")
print(f"use_case_tables={len(current_use_cases)} unchanged={current_use_cases == backup_use_cases}")
print(f"references={len(reference_entries)} cited_numbers={len(cited)}")
print(f"errors={len(errors)}")
for error in errors:
    print(error)
