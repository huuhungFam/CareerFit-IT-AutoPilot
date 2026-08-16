from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"C:\CODING\Thesis")
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
SOURCE = ROOT / "Doc" / "working" / "CHAPTER_6_SOURCE.md"


def font(run, size=13, bold=None, italic=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def find(document, text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(text)


def remove(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def insert_p(document, anchor, text, level=None):
    paragraph = document.add_paragraph(style=f"Heading {level}" if level else "Normal")
    font(paragraph.add_run(text), 14 if level == 2 else 13, bool(level) if level else None)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    if level:
        paragraph.paragraph_format.first_line_indent = Inches(0)
    else:
        paragraph.paragraph_format.first_line_indent = Inches(0.5)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    anchor._element.addprevious(paragraph._element)


def insert_table(document, anchor, lines):
    rows = [[cell.strip().replace("`", "") for cell in line.strip().strip("|").split("|")] for line in lines]
    rows = [rows[0], *rows[2:]]
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, values in enumerate(rows):
        cells = table.rows[0].cells if row_index == 0 else table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Inches(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    font(run, 9.5, row_index == 0)
    anchor._element.addprevious(table._element)


def populate(document, anchor):
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    pending = []

    def flush():
        nonlocal pending
        if pending:
            insert_p(document, anchor, " ".join(line.strip() for line in pending))
            pending = []

    while index < len(lines):
        line = lines[index]
        if not line.strip():
            flush(); index += 1; continue
        if line.startswith("## "):
            flush(); insert_p(document, anchor, line[3:], level=2); index += 1; continue
        if line.startswith("### "):
            flush(); insert_p(document, anchor, line[4:], level=3); index += 1; continue
        if line.startswith("Table "):
            flush(); insert_p(document, anchor, line); index += 1
            table_lines = []
            while index < len(lines) and (not lines[index].strip() or lines[index].startswith("|")):
                if lines[index].startswith("|"):
                    table_lines.append(lines[index])
                index += 1
            insert_table(document, anchor, table_lines); continue
        pending.append(line); index += 1
    flush()


def update_abstract(document):
    old = "The evaluation is designed to separate controlled algorithm experiments from functional, integration, acceptance, performance, and security testing. At the current writing stage, quantitative results are intentionally omitted and will be inserted only after the final version is evaluated using reproducible runtime evidence."
    new = "The evaluation separates controlled algorithm experiments from functional, integration, browser, security, and runtime evidence. The backend executed 63 registered tests without JUnit failures, the frontend production build succeeded, and four Chromium P0 workflows passed. On a synthetic causal benchmark containing 50 jobs and 100 CVs, Rocchio increased nDCG@5 from 0.037737 to 0.817737; three repeated runs produced the same metric values. These results demonstrate controlled feedback behavior and integrated prototype workflows, not production hiring effectiveness. Background optimistic-lock exceptions and an aggregate health status of DOWN remain unresolved limitations."
    for paragraph in document.paragraphs:
        if old in paragraph.text:
            updated = paragraph.text.replace(old, new)
            paragraph.clear()
            font(paragraph.add_run(updated))
            return
    raise ValueError("Abstract evaluation paragraph not found")


def main():
    document = Document(DOCX)
    if any(p.text.strip() == "6.1 Summary of Results" for p in document.paragraphs):
        raise SystemExit("Chapter 6 is already populated; no changes made.")
    placeholder = "NOTE: [Chapter 6 will be completed after Chapter 5 results are finalized.]"
    for paragraph in list(document.paragraphs):
        if paragraph.text.strip() == placeholder:
            remove(paragraph)
    anchor = find(document, "REFERENCES")
    populate(document, anchor)
    update_abstract(document)
    document.save(DOCX)
    print(f"Updated {DOCX} ({DOCX.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
