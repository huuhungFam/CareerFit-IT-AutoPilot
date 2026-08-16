from __future__ import annotations

import hashlib
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-before-20260811-figure-refinement.docx"
FIGURES = ROOT / "Doc" / "figures"

FLOWCHART_WIDTH = 5.30
FULL_WIDTH = 6.10

FLOWCHARTS = {
    "Figure 3.14. CV ingestion, review, confirmation, and matching flowchart": "flowchart-cv-processing-20260811.png",
    "Figure 3.15. Seed-corpus initialization and TF-IDF construction flowchart": "flowchart-tfidf-construction-20260811.png",
    "Figure 3.16. CV-Job matching and Potential assessment flowchart": "flowchart-matching-potential-20260811.png",
    "Figure 3.17. Rocchio feedback-learning and recomputation flowchart": "flowchart-rocchio-feedback-20260811.png",
    "Figure 3.19. AutoFit eligibility and application decision flowchart": "flowchart-autofit-20260811.png",
    "Figure 3.20. Notification policy evaluation and delivery outcome flowchart": "flowchart-notification-policy-20260811.png",
    "Figure 3.21. Actionable-email confirmation and redemption flowchart": "flowchart-email-action-redemption-20260811.png",
}


def digest_tables(document: Document) -> str:
    payload = [
        [[cell.text for cell in row.cells] for row in table.rows]
        for table in document.tables
    ]
    return hashlib.sha256(repr(payload).encode("utf-8")).hexdigest()


def digest_non_drawing_text(document: Document) -> str:
    return hashlib.sha256(
        "\n".join(paragraph.text for paragraph in document.paragraphs).encode("utf-8")
    ).hexdigest()


def find_caption(document: Document, caption_text: str) -> Paragraph:
    matches = [p for p in document.paragraphs if p.text.strip() == caption_text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one caption {caption_text!r}; found {len(matches)}")
    return matches[0]


def previous_paragraph(paragraph: Paragraph) -> Paragraph:
    element = paragraph._p.getprevious()
    while element is not None and not element.tag.endswith("}p"):
        element = element.getprevious()
    if element is None:
        raise RuntimeError(f"No image paragraph before {paragraph.text!r}")
    return Paragraph(element, paragraph._parent)


def set_inline_width(inline, width_inches: float) -> None:
    extent = inline._inline.extent
    old_cx = int(extent.cx)
    old_cy = int(extent.cy)
    new_cx = int(Inches(width_inches))
    new_cy = int(round(old_cy * new_cx / old_cx))
    extent.cx = new_cx
    extent.cy = new_cy
    graphic_extent = inline._inline.graphic.graphicData.pic.spPr.xfrm.ext
    graphic_extent.cx = new_cx
    graphic_extent.cy = new_cy


def replace_flowchart(document: Document, caption_text: str, image_name: str) -> None:
    caption = find_caption(document, caption_text)
    image_paragraph = previous_paragraph(caption)
    if not image_paragraph._p.xpath(".//w:drawing"):
        raise RuntimeError(f"Expected drawing before {caption_text!r}")

    page_break_before = image_paragraph.paragraph_format.page_break_before
    keep_with_next = image_paragraph.paragraph_format.keep_with_next
    image_paragraph._element.getparent().remove(image_paragraph._element)

    replacement = caption.insert_paragraph_before()
    replacement.alignment = WD_ALIGN_PARAGRAPH.CENTER
    replacement.paragraph_format.page_break_before = page_break_before
    replacement.paragraph_format.keep_with_next = keep_with_next
    replacement.paragraph_format.space_before = 0
    replacement.paragraph_format.space_after = 0
    inline = replacement.add_run().add_picture(
        str(FIGURES / image_name), width=Inches(FLOWCHART_WIDTH)
    )
    inline._inline.docPr.set("title", caption_text)
    inline._inline.docPr.set("descr", caption_text)


def enlarge_remaining_figures(document: Document) -> int:
    changed = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if paragraph.style.name != "Figure Caption":
            continue
        if not (
            text.startswith("Figure 3.")
            or text.startswith("Figure 4.")
            or text.startswith("Screen 3.")
        ):
            continue
        if text in FLOWCHARTS:
            continue
        image_paragraph = previous_paragraph(paragraph)
        drawings = image_paragraph._p.xpath(".//w:drawing")
        if not drawings:
            raise RuntimeError(f"Expected drawing before {text!r}")
        for inline_shape in image_paragraph._p.xpath(".//wp:inline"):
            extent = inline_shape.xpath("./wp:extent")[0]
            old_cx = int(extent.get("cx"))
            old_cy = int(extent.get("cy"))
            new_cx = int(Inches(FULL_WIDTH))
            new_cy = int(round(old_cy * new_cx / old_cx))
            extent.set("cx", str(new_cx))
            extent.set("cy", str(new_cy))
            graphic_extent = inline_shape.xpath(".//a:xfrm/a:ext")[0]
            graphic_extent.set("cx", str(new_cx))
            graphic_extent.set("cy", str(new_cy))
            changed += 1
    return changed


def main() -> None:
    shutil.copy2(DOCX, BACKUP)
    document = Document(DOCX)
    before_tables = digest_tables(document)
    before_text = digest_non_drawing_text(document)

    for caption_text, image_name in FLOWCHARTS.items():
        replace_flowchart(document, caption_text, image_name)
    enlarged = enlarge_remaining_figures(document)

    if before_tables != digest_tables(document):
        raise RuntimeError("Table contents changed during the figure-only refinement")
    if before_text != digest_non_drawing_text(document):
        raise RuntimeError("Paragraph text changed during the figure-only refinement")

    document.save(DOCX)
    print(f"Updated flowcharts: {len(FLOWCHARTS)}")
    print(f"Enlarged remaining Chapter 3/4 figures and screens: {enlarged}")
    print(f"Saved: {DOCX}")
    print(f"Backup: {BACKUP}")


if __name__ == "__main__":
    main()
