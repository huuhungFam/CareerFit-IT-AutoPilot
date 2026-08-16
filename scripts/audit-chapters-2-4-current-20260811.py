from __future__ import annotations

import re
from collections import Counter
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
CURRENT = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BEFORE = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-before-20260811-chapter-restructure.docx"


def clean(text: str) -> str:
    return " ".join(text.replace("\u00a0", " ").split())


def words(text: str) -> list[str]:
    return re.findall(r"[A-Za-zÀ-ỹ0-9_@'–/-]+", text)


def sentence_list(text: str) -> list[str]:
    return [clean(value) for value in re.split(r"(?<=[.!?])\s+", clean(text)) if clean(value)]


def chapter_bounds(document: Document, chapter: int) -> tuple[int, int]:
    start_pattern = re.compile(rf"CHAPTER {chapter}\.\s")
    starts = [
        i
        for i, p in enumerate(document.paragraphs)
        if p.style.name == "Heading 1" and start_pattern.match(clean(p.text))
    ]
    if len(starts) != 1:
        raise RuntimeError(f"Chapter {chapter}: expected one start, found {starts}")
    start = starts[0]
    end = len(document.paragraphs)
    for i in range(start + 1, len(document.paragraphs)):
        text = clean(document.paragraphs[i].text)
        if document.paragraphs[i].style.name == "Heading 1" and (
            re.match(r"CHAPTER \d+\.\s", text) or text == "PART 3. CONCLUSION"
        ):
            end = i
            break
    return start, end


def section_rows(document: Document, chapter: int):
    start, end = chapter_bounds(document, chapter)
    headings = []
    for i in range(start, end):
        paragraph = document.paragraphs[i]
        match = re.fullmatch(r"Heading (\d+)", paragraph.style.name)
        if match and clean(paragraph.text):
            headings.append((i, int(match.group(1)), clean(paragraph.text)))

    rows = []
    for pos, (index, level, title) in enumerate(headings):
        next_index = headings[pos + 1][0] if pos + 1 < len(headings) else end
        direct_paragraphs = []
        for paragraph in document.paragraphs[index + 1 : next_index]:
            text = clean(paragraph.text)
            if not text:
                continue
            if paragraph.style.name in {"Figure Caption", "Table Caption"}:
                continue
            direct_paragraphs.append(text)
        text = " ".join(direct_paragraphs)
        citations = re.findall(r"\[(\d+(?:\s*[-,]\s*\d+)*)\]", text)
        rows.append(
            {
                "index": index,
                "level": level,
                "title": title,
                "paragraphs": direct_paragraphs,
                "text": text,
                "words": len(words(text)),
                "sentences": sentence_list(text),
                "citations": citations,
            }
        )
    return rows


def print_document(label: str, path: Path) -> None:
    document = Document(path)
    print(f"\n===== {label}: {path.name} =====")
    all_sentences = []
    for chapter in (2, 3, 4):
        rows = section_rows(document, chapter)
        chapter_text = " ".join(row["text"] for row in rows)
        print(
            f"\nCHAPTER {chapter}: narrative_words={len(words(chapter_text))} "
            f"headings={len(rows)} citations={sum(len(row['citations']) for row in rows)}"
        )
        for row in rows:
            if row["level"] == 1:
                continue
            marker = ""
            if row["words"] < 50:
                marker = " [VERY SHORT]"
            elif row["words"] < 80:
                marker = " [SHORT]"
            print(
                f"L{row['level']} {row['title']} | words={row['words']} "
                f"paras={len(row['paragraphs'])} citations={len(row['citations'])}{marker}"
            )
            for sentence in row["sentences"]:
                all_sentences.append((chapter, row["title"], sentence))

    normalized = Counter(
        re.sub(r"[^a-z0-9]+", " ", sentence.lower()).strip()
        for _, _, sentence in all_sentences
        if len(words(sentence)) >= 10
    )
    repeated = {text: count for text, count in normalized.items() if count > 1}
    print(f"\nExact repeated substantive sentences: {len(repeated)}")
    for text, count in sorted(repeated.items(), key=lambda item: (-item[1], item[0]))[:20]:
        print(f"  {count}x {text[:180]}")

    long_sentences = [
        (len(words(sentence)), chapter, title, sentence)
        for chapter, title, sentence in all_sentences
        if len(words(sentence)) > 45
    ]
    print(f"Long sentences (>45 words): {len(long_sentences)}")
    for count, chapter, title, sentence in sorted(long_sentences, reverse=True)[:25]:
        print(f"  {count} words | Ch{chapter} | {title} | {sentence[:240]}")


if __name__ == "__main__":
    print_document("CURRENT", CURRENT)
    print_document("BEFORE CHAPTER RESTRUCTURE", BEFORE)
