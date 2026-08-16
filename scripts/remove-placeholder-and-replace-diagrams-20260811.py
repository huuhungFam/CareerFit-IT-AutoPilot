from __future__ import annotations

import hashlib
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-before-20260811-remove-placeholder-diagrams.docx"
FIGURES = ROOT / "Doc" / "figures"


def table_digest(document: Document) -> str:
    payload = []
    for table in document.tables:
        payload.append([[cell.text for cell in row.cells] for row in table.rows])
    return hashlib.sha256(repr(payload).encode("utf-8")).hexdigest()


def find_exact(document: Document, text: str) -> Paragraph:
    matches = [p for p in document.paragraphs if p.text.strip() == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph for {text!r}; found {len(matches)}")
    return matches[0]


def previous_paragraph(paragraph: Paragraph) -> Paragraph:
    previous = paragraph._p.getprevious()
    while previous is not None and not previous.tag.endswith("}p"):
        previous = previous.getprevious()
    if previous is None:
        raise RuntimeError(f"No previous paragraph before {paragraph.text!r}")
    return Paragraph(previous, paragraph._parent)


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def configure_image_paragraph(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.page_break_before = True
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0


def configure_caption(paragraph: Paragraph) -> None:
    paragraph.style = "Figure Caption"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = False


def add_picture(paragraph: Paragraph, image_name: str, alt_text: str) -> None:
    image_path = FIGURES / image_name
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    inline = paragraph.add_run().add_picture(str(image_path), width=Inches(5.45))
    inline._inline.docPr.set("title", alt_text)
    inline._inline.docPr.set("descr", alt_text)


def remove_figure(document: Document, caption_text: str) -> None:
    caption = find_exact(document, caption_text)
    image = previous_paragraph(caption)
    if not image._p.xpath(".//w:drawing"):
        raise RuntimeError(f"Expected a drawing immediately before {caption_text!r}")
    remove_paragraph(image)
    remove_paragraph(caption)


def replace_figure(document: Document, old_caption: str, new_caption: str, image_name: str) -> None:
    caption = find_exact(document, old_caption)
    image = previous_paragraph(caption)
    if not image._p.xpath(".//w:drawing"):
        raise RuntimeError(f"Expected a drawing immediately before {old_caption!r}")
    remove_paragraph(image)

    image_paragraph = caption.insert_paragraph_before()
    configure_image_paragraph(image_paragraph)
    add_picture(image_paragraph, image_name, new_caption)
    caption.text = new_caption
    configure_caption(caption)


def request_field_refresh(document: Document) -> None:
    settings = document.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def figure_captions(document: Document):
    return [p.text.strip() for p in document.paragraphs if p.style.name == "Figure Caption"]


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)
    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(DOCX, BACKUP)

    document = Document(DOCX)
    before_tables = table_digest(document)
    before_figure_count = len(figure_captions(document))

    # Remove diagrams that only repeat the surrounding explanation or tables.
    for caption in (
        "Figure 2.1. Distinction between matching, recommendation, and recruitment action",
        "Figure 2.2. TF-IDF vectorization and cosine-similarity pipeline",
        "Figure 2.3. Rocchio relevance-feedback vector update",
        "Figure 2.4. Human-in-the-Loop control cycle in CareerFit",
        "Figure 4.1. Evaluation environments and evidence sources",
        "Figure 4.3. P0 end-to-end workflow coverage",
    ):
        remove_figure(document, caption)

    replacements = (
        (
            "Figure 3.1. CareerFit container and component architecture",
            "Figure 3.1. CareerFit container and component architecture",
            "uml-careerfit-architecture-20260811.png",
        ),
        (
            "Figure 3.2. CareerFit logical entity relationships",
            "Figure 3.2. CareerFit core logical data model",
            "erd-careerfit-core-20260811.png",
        ),
        (
            "Figure 3.3. Local and containerized deployment topology",
            "Figure 3.3. Local and containerized deployment topology",
            "uml-deployment-20260811.png",
        ),
        (
            "Figure 3.4. Backend module structure and request flow",
            "Figure 3.4. Backend module structure and request path",
            "uml-backend-modules-20260811.png",
        ),
        (
            "Figure 3.5. JWT authentication and authorization boundaries",
            "Figure 3.5. JWT authentication, role authorization, and ownership-check sequence",
            "uml-jwt-sequence-20260811.png",
        ),
        (
            "Figure 3.10. Application and invitation state transitions",
            "Figure 3.10. Application and invitation UML state machine",
            "uml-application-state-20260811.png",
        ),
        (
            "Figure 3.14. Frontend routes, server-side catalogue, and API data flow",
            "Figure 3.14. Frontend request and API-response sequence",
            "uml-frontend-sequence-20260811.png",
        ),
        (
            "Figure 4.4. Local Job-search latency distribution",
            "Figure 4.2. Observed local Job-search latency statistics",
            "chart-job-search-latency-20260811.png",
        ),
    )
    for old_caption, new_caption, image_name in replacements:
        replace_figure(document, old_caption, new_caption, image_name)

    benchmark = find_exact(document, "Figure 4.2. Baseline and Rocchio benchmark metrics at K = 5")
    benchmark.text = "Figure 4.1. Baseline and Rocchio benchmark metrics at K = 5"
    configure_caption(benchmark)

    request_field_refresh(document)

    after_tables = table_digest(document)
    if before_tables != after_tables:
        raise RuntimeError("Table contents changed during the figure-only edit")

    after_captions = figure_captions(document)
    if len(after_captions) != before_figure_count - 6:
        raise RuntimeError(
            f"Unexpected figure-count change: before={before_figure_count}, after={len(after_captions)}"
        )
    if any(c.startswith("Figure 2.") for c in after_captions):
        raise RuntimeError("Chapter 2 still contains figure captions")

    document.save(DOCX)
    print(f"Updated: {DOCX}")
    print(f"Backup:  {BACKUP}")
    print(f"Tables unchanged: {before_tables}")
    print(f"Figure captions: {before_figure_count} -> {len(after_captions)}")


if __name__ == "__main__":
    main()
