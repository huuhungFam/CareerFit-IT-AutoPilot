from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PATH = Path(r"C:\CODING\Thesis\Doc\CareerFit-Thesis-Report.docx")


def set_font(run, size=13, bold=None, italic=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def find_paragraph(document, text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def insert_paragraph(document, anchor, text, level=None, italic=False, indent=True):
    paragraph = document.add_paragraph(style=f"Heading {level}" if level else "Normal")
    run = paragraph.add_run(text)
    set_font(run, 14 if level == 2 else 13, bool(level) if level else None, italic)
    if level:
        paragraph.paragraph_format.first_line_indent = Inches(0)
    else:
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Inches(0.5) if indent else Inches(0)
        paragraph.paragraph_format.space_after = Pt(0)
    anchor._element.addprevious(paragraph._element)
    return paragraph


def insert_note(document, anchor, text):
    paragraph = insert_paragraph(document, anchor, "NOTE: " + text, italic=True, indent=False)
    paragraph.paragraph_format.left_indent = Inches(0.25)


def insert_table(document, anchor, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Inches(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_font(run, 10, row_index == 0)
    anchor._element.addprevious(table._element)


CONTENT = [
    ("h2", "2.1 Recruitment Information and CV–Job Description Matching"),
    ("p", "Recruitment matching is an information-retrieval and decision-support problem in which heterogeneous evidence about candidates and vacancies must be compared. A curriculum vitae commonly describes education, employment history, projects, technical skills, certifications, languages, and contact information. A job description describes responsibilities, required and preferred skills, seniority, location, employment conditions, and organizational context. Some attributes are structured, such as years of experience or location, while others appear in free text and may use inconsistent terminology. Consequently, an exact database filter is useful for hard constraints but insufficient for ranking the overall relevance of a candidate and a position."),
    ("p", "The same information supports different user tasks. From the candidate perspective, the system retrieves and ranks jobs that fit a CV or desired profile. From the recruiter perspective, it ranks applicants or discovers potential candidates for a specific vacancy. These directions share document representations and similarity functions, but they are not identical business decisions. A relevant job recommendation does not imply that an application has been submitted, and a high candidate–job similarity does not imply that a recruiter should accept the candidate. The score is therefore evidence for prioritization rather than a final employment decision."),
    ("p", "Recruitment data also has domain-specific ambiguity. A technology may be expressed by an abbreviation, a product name, or a broader skill family; job titles vary across companies; and the same term may have different importance at junior and senior levels. Open recruitment corpora are relatively uncommon because CVs contain personal information. The Djinni Recruitment Dataset illustrates both the scale of the matching problem and the value of anonymized, documented corpora, providing more than 150,000 jobs and 230,000 candidate profiles in English and Ukrainian [6]. CareerFit is narrower: it focuses on IT recruitment and uses explicit fields and normalized text so that matching behavior can be inspected within the project scope."),
    ("note", "[Figure 2.1 – Conceptual distinction among CV–JD matching, profile-based job recommendation, and recruitment action – to be created later.]"),
    ("h2", "2.2 Text Representation and Preprocessing"),
    ("h3", "2.2.1 Document Normalization"),
    ("p", "Before vectorization, text must be converted into a consistent representation. A typical pipeline extracts machine-readable text, normalizes character encoding and case, separates or standardizes tokens, removes formatting noise, and optionally filters stop words or applies stemming or lemmatization. In recruitment documents, preprocessing must preserve discriminative technical tokens such as framework names, programming languages, version-like expressions, and abbreviations. Aggressive removal can make a CV and a JD appear less related even when they share an important technology."),
    ("p", "Bilingual data adds further complexity. Vietnamese contains diacritics and multi-syllable expressions separated by spaces, while English technical terminology frequently appears unchanged inside Vietnamese sentences. A practical pipeline must apply deterministic normalization to equivalent inputs and maintain a controlled vocabulary. It should also keep structured attributes—such as location, language, seniority, and salary mode—available for validation or filtering instead of forcing every business constraint into a single text vector."),
    ("p", "Preprocessing quality directly affects every later score. Missing OCR text, malformed documents, extremely short descriptions, duplicated boilerplate, and contradictory structured fields can distort term frequencies. For this reason, validation belongs before scoring. Hard validation rejects inputs that cannot support meaningful processing, while soft validation allows processing but records warnings or suggestions. This distinction prevents the ranking engine from silently assigning apparently precise scores to unusable data."),
    ("h3", "2.2.2 Bag-of-Words Representation"),
    ("p", "In a bag-of-words representation, a document is represented by the terms it contains and their weights; word order is not modeled directly. If the vocabulary contains |V| terms, each document is mapped to a vector in R^|V|. This representation is sparse because a particular CV or JD contains only a small portion of the complete vocabulary. Its main advantages are computational simplicity, reproducibility, and the ability to inspect which terms contribute to similarity. Its main limitation is lexical dependence: synonyms and related skills do not match unless normalization or an explicit mapping connects them."),
    ("h2", "2.3 TF-IDF and Cosine Similarity"),
    ("h3", "2.3.1 Term Frequency and Inverse Document Frequency"),
    ("p", "The vector-space model represents documents and queries as weighted term vectors and ranks documents according to their geometric relationship [1]. TF-IDF is a standard weighting family for this model [2]. Term frequency (TF) reflects how strongly a term occurs in a particular document. Inverse document frequency (IDF) reduces the influence of terms that occur in many documents and increases the relative influence of terms that distinguish a smaller subset of the corpus."),
    ("p", "Let tf(t,d) denote the frequency of term t in document d, let N be the number of documents in the reference corpus, and let df(t) be the number of documents containing t. A basic formulation is:"),
    ("eq", "tfidf(t,d) = tf(t,d) × log(N / df(t))."),
    ("p", "Implementations often use logarithmic TF scaling, smoothed IDF, or vector normalization to handle repeated terms and unseen values. The exact formulation changes numeric scores, so a thesis must document the implemented formula rather than referring to TF-IDF as if it were a single universal function. A controlled or static reference corpus can improve score stability because the IDF of a term does not change whenever a user adds one document; however, it may become less representative as technologies and hiring terminology evolve."),
    ("h3", "2.3.2 Cosine Similarity"),
    ("p", "Cosine similarity compares the direction of two vectors rather than their unnormalized magnitude. For vectors x and y, it is defined as:"),
    ("eq", "cos(x,y) = (x · y) / (||x||₂ ||y||₂)."),
    ("p", "For non-negative TF-IDF vectors, cosine similarity usually lies between 0 and 1. A larger value indicates a greater proportion of shared weighted terms. Length normalization is valuable for CV–JD comparison because a long CV should not receive a higher score merely because it contains more words. TF-IDF-weighted cosine distance has also been used as an efficient and reproducible document-alignment method [3]."),
    ("p", "Cosine similarity remains a lexical measure. If a JD uses “container orchestration” while a CV only uses “Kubernetes,” the vectors may not overlap unless the vocabulary or normalization layer establishes that relationship. Conversely, shared generic terms can increase similarity without proving that mandatory requirements are satisfied. CareerFit therefore treats similarity as one component of a broader workflow that also includes structured fields, validation, labels, reasons, user feedback, and policy conditions."),
    ("note", "[Figure 2.2 – TF-IDF vectorization and cosine-similarity pipeline with a small CV/JD example – to be created later.]"),
    ("h2", "2.4 Matching, Recommendation, and Ranking"),
    ("p", "Matching and recommendation can be formalized as ranking tasks. Given a query representation q and a candidate set D, the system computes a relevance score s(q,d) for each d ∈ D and orders the results by decreasing score. For candidate-to-job matching, q may be a CV vector and D the active job vectors. For recruiter-side discovery, q may be the selected JD and D the available candidate or CV vectors. For profile-based recommendation, q represents the candidate’s desired role, skills, location, seniority, and other preferences rather than one uploaded CV."),
    ("p", "This separation matters because a CV describes demonstrated history, whereas a preference profile describes intent. A candidate may have experience in Java but seek a position involving Go; a CV-only ranking and a preference-based recommendation can therefore produce different, legitimate lists. Storing the workflows separately also makes evaluation more defensible: the relevance labels, candidate sets, and user objectives for one task should not be assumed to be valid for the other."),
    ("p", "A raw similarity value is not automatically a calibrated probability of hiring success. Mapping it to a percentage or label is a presentation and business-rule decision, not a statistical guarantee. Thresholds such as Low, Medium, High, or Potential should be documented, tested at boundary values, and kept distinct from application status. Stable secondary sorting is also necessary when multiple items have equal scores so that repeated requests do not appear inconsistent."),
    ("p", "Recent resume–job matching research increasingly uses dense encoders and contrastive learning. ConFit v2, for example, improves resume–job representations through hypothetical resumes and hard-negative mining [7]. Such approaches can capture semantic relations beyond exact terms, but they require training data, model governance, and an evaluation design appropriate to the domain. CareerFit intentionally begins with a transparent lexical baseline whose calculations can be reproduced and explained, while treating semantic or hybrid retrieval as future work rather than an implemented result."),
    ("h2", "2.5 Relevance Feedback and the Rocchio Method"),
    ("h3", "2.5.1 Relevance Feedback"),
    ("p", "Relevance feedback uses judgments about retrieved items to modify a query or profile representation. Instead of assuming that the original text completely expresses the user’s information need, the system observes which results are marked relevant or non-relevant. Explicit feedback is especially useful in recruitment because the importance of related skills may be known to the candidate or recruiter but omitted from the original CV, profile, or JD."),
    ("p", "The classic Rocchio method updates a query vector by combining the original query with the centroids of relevant and non-relevant document vectors [4]. Let q₀ be the original vector, Dᵣ the set of relevant feedback documents, and Dₙᵣ the set of non-relevant documents. A common form is:"),
    ("eq", "qₘ = αq₀ + (β / |Dᵣ|) Σ(d ∈ Dᵣ) d − (γ / |Dₙᵣ|) Σ(d ∈ Dₙᵣ) d."),
    ("p", "The parameters α, β, and γ control retention of the original representation, attraction toward relevant examples, and movement away from non-relevant examples. They are hyperparameters rather than universal constants. Positive and negative feedback may also be weighted differently because an explicit rejection can be ambiguous: a user may reject a job because of location, timing, or salary even when its technical content is relevant."),
    ("h3", "2.5.2 Feedback Semantics in CareerFit"),
    ("p", "CareerFit distinguishes feedback such as GOOD_MATCH, POTENTIAL, BAD_MATCH, and NOT_INTERESTED. These labels should not be collapsed without a documented mapping. GOOD_MATCH provides a strong positive relevance signal. POTENTIAL may represent partial or transferable fit rather than direct relevance. BAD_MATCH is a negative judgment about the match. NOT_INTERESTED can reflect personal intent and may be weaker evidence about technical similarity. The implementation chapter must state exactly how each event affects the learned vector and whether role or channel changes its weight."),
    ("p", "A robust update process retains the immutable base vector and recomputes the learned vector from the relevant feedback history. Repeatedly applying an update to the previously learned vector can create cumulative drift: running the same recomputation more than once may change the output even though no new feedback exists. Idempotent recomputation means that the same base vector, feedback set, and parameters produce the same learned vector. This property is important for retries, scheduled processing, debugging, and auditability."),
    ("p", "Feedback learning must be evaluated causally rather than by reporting only a final score. A controlled experiment records the baseline ranking, introduces a defined feedback event, recomputes the representation, and measures the effect on separate holdout items. If the positive item used to provide feedback is also the sole item used to claim improvement, the result is circular. Chapter 5 therefore separates feedback examples from holdout evaluation and explicitly labels synthetic scenarios."),
    ("note", "[Figure 2.3 – Rocchio update showing the base vector, positive and negative centroids, and the resulting learned vector – to be created later.]"),
    ("h2", "2.6 Evaluation Metrics for Ranked Results"),
    ("p", "Ranking quality cannot be adequately described by ordinary classification accuracy because users observe an ordered list and usually inspect only the first few items. Evaluation therefore requires relevance judgments and metrics that account for the cutoff position K, the location of relevant results, and, where available, graded relevance."),
    ("h3", "2.6.1 Precision@K, Recall@K, and HitRate@K"),
    ("p", "Precision@K is the proportion of the top K results that are relevant. Recall@K is the proportion of all known relevant items that appear in the top K. HitRate@K records whether at least one relevant item appears in the first K positions. Precision emphasizes the usefulness of the visible list, recall emphasizes coverage, and HitRate provides a simple task-success indicator. The metrics answer different questions and should not be substituted for one another."),
    ("h3", "2.6.2 Mean Reciprocal Rank"),
    ("p", "For a query, reciprocal rank is 1/r, where r is the rank of the first relevant result; it is zero if no relevant result is retrieved. Mean Reciprocal Rank (MRR) averages this value across queries. MRR strongly rewards placing the first relevant result near the top but ignores the ordering of additional relevant items. It is therefore appropriate when finding one useful job or candidate is the primary objective, but it should be accompanied by a broader top-K metric when multiple relevant results matter."),
    ("h3", "2.6.3 Discounted Cumulative Gain"),
    ("p", "Discounted Cumulative Gain (DCG) supports graded relevance and discounts useful items that occur at lower ranks. Normalized DCG divides the observed DCG by the ideal ordering for the same relevance judgments, producing nDCG values that are comparable across queries. Järvelin and Kekäläinen introduced cumulated-gain measures to credit systems for ranking highly relevant documents before less relevant ones [5]. One common formulation is:"),
    ("eq", "DCG@K = Σ(i=1..K) (2^relᵢ − 1) / log₂(i + 1),     nDCG@K = DCG@K / IDCG@K."),
    ("p", "Metric values are meaningful only when the relevance labels, candidate set, cutoff K, and aggregation procedure are documented. A large improvement on a synthetic scenario demonstrates behavior under that scenario; it does not by itself demonstrate production hiring quality. Chapter 5 therefore reports dataset provenance, baseline construction, holdout logic, repeated-run behavior, and threats to validity together with the metric values."),
    ("h2", "2.7 Human-in-the-Loop and Explainable Automation"),
    ("h3", "2.7.1 Levels of Human Involvement"),
    ("p", "Human-in-the-Loop is not a single interface feature. Human involvement can occur during data validation, model configuration, review of recommendations, approval of actions, exception handling, and post-action feedback. NIST describes human–AI configurations as ranging from manual decision making through systems that provide an additional opinion or defer to an expert, to more autonomous operation [9]. The appropriate configuration depends on the consequences and failure modes of the task."),
    ("p", "Recruitment decisions can materially affect people, so a similarity engine should assist prioritization rather than act as an unquestioned authority. CareerFit separates perception, decision support, action, learning, and audit. The matching engine produces evidence; the AutoFit policy evaluates whether an action is permitted; a user may review or confirm important actions; feedback changes later ranking; and audit records support investigation. This design does not eliminate bias, but it makes the control points and system state more explicit."),
    ("h3", "2.7.2 Explainability and Auditability"),
    ("p", "Explainability answers why a particular result or action was produced, whereas auditability concerns whether the inputs, rules, actor, channel, time, and outcome can be reconstructed. A list of overlapping skills can provide a local explanation of lexical similarity, but it is not proof that the entire model or business decision is fair. Likewise, an audit log records what occurred but does not guarantee that the recorded policy was appropriate."),
    ("p", "Research on algorithmic hiring warns that claims about bias mitigation and performance must be evaluated against actual practices, data, and institutional context [10]. Explanations can support understanding, but they should be grounded in the features used by the system. Work on explainable recommendation similarly distinguishes recommendation scoring from the generation of user-facing reasons and emphasizes fact-grounded explanations [14]. CareerFit therefore favors inspectable match reasons and policy outcomes over unsupported natural-language claims about candidate suitability."),
    ("p", "The NIST AI Risk Management Framework organizes AI risk work around governance, mapping, measurement, and management [9]. CareerFit is not claimed to be a complete implementation of that framework; however, its design aligns with several practical principles: document scope and limitations, preserve human authority for consequential actions, measure behavior using defined evidence, and maintain records that support review."),
    ("note", "[Figure 2.4 – Human-in-the-Loop control points in the CareerFit Perception–Decision–Action–Learning–Audit cycle – to be created later.]"),
    ("h2", "2.8 Web Architecture, Security, and Audit Foundations"),
    ("h3", "2.8.1 Client–Server and REST"),
    ("p", "CareerFit uses a web client and a server-side application connected through HTTP APIs. The client is responsible for presentation, navigation, local interaction state, and invoking authorized operations. The backend enforces authentication, authorization, validation, business rules, transactions, persistence, and background processing. This separation prevents browser-side presentation logic from becoming the authority for protected actions."),
    ("p", "REST is an architectural style characterized by constraints including client–server separation, stateless interaction, cacheability where appropriate, a uniform interface, layered components, and optional code-on-demand [11]. A REST API models resources and transfers representations of their state. In practice, merely using JSON over HTTP does not guarantee that every REST constraint is satisfied; Chapter 4 therefore describes CareerFit as a REST-oriented API and documents its actual endpoint and state-transition behavior."),
    ("h3", "2.8.2 Authentication, Authorization, and Action Tokens"),
    ("p", "Authentication establishes the identity associated with a request, while authorization determines whether that identity may perform an operation. Role-Based Access Control assigns permissions according to roles such as Candidate, Recruiter, or Administrator, but ownership checks are still required; for example, one candidate must not access another candidate’s private CV merely because both share the same role."),
    ("p", "JSON Web Token is a compact, URL-safe format for transferring claims that can be signed or message-authentication-code protected and may include registered claims such as subject, audience, expiration time, issued-at time, and token identifier [12]. A valid signature alone is not sufficient: the application must validate the expected algorithm and relevant claims and must apply its own account and authorization rules. Short-lived access tokens and purpose-specific action tokens address different workflows and should not be treated as interchangeable credentials."),
    ("p", "An actionable email link can be opened by the intended user, forwarded accidentally, or visited automatically by a mail-security scanner. A safer pattern lets a GET request display a confirmation page without changing business state and requires a deliberate POST request to execute the action. The action token should be random or cryptographically protected, purpose-bound, expiring, and single-use. Server-side consumption state or an equivalent replay-prevention mechanism is needed because expiration alone does not prevent the same valid link from being used repeatedly."),
    ("h3", "2.8.3 Audit Logging"),
    ("p", "Audit logging records security-relevant and business-relevant events so that operations can be monitored and investigated. NIST guidance emphasizes a managed logging process that includes generation, transmission, storage, access, analysis, and disposal rather than treating logs as incidental text files [13]. For recruitment automation, useful fields include the actor, action, target, source channel, relevant policy or score snapshot, result, and timestamp. Sensitive CV content and credentials should not be copied indiscriminately into logs."),
    ("p", "An append-oriented audit history supports traceability, but database permissions, retention, integrity, and access control determine whether that history is trustworthy. Operational logs used for debugging and domain audit records used to explain a business action may have different schemas and retention requirements. Chapter 3 defines these responsibilities at the design level, and Chapter 4 describes the mechanisms actually implemented in CareerFit."),
    ("h2", "2.9 Related Work and Research Gap"),
    ("p", "Prior work can be grouped into lexical retrieval, learned resume–job representation, explainable recommendation, and human-centered governance. The classical vector-space and relevance-feedback literature provides transparent mathematical foundations [1], [4]. Contemporary resume–job matching methods such as ConFit v2 use trained dense representations and hard-negative strategies to improve semantic matching [7]. Real-world observational research comparing human and LLM resume ratings found only minor correlation between the two, indicating that automated and human judgments should not be treated as interchangeable [8]. Explainable recommendation research seeks to accompany scores with understandable, preferably grounded reasons [14]."),
    ("p", "CareerFit does not attempt to outperform these research systems on a shared benchmark. Its engineering gap is different: connecting an inspectable lexical baseline and explicit relevance feedback to an end-to-end recruitment workflow with role-based interfaces, policy gating, actionable email, and auditable state changes. The resulting contribution is system integration and controlled automation within an IT-focused academic prototype."),
]


REFERENCES = [
    "[1] G. Salton, A. Wong, and C. S. Yang, “A vector space model for automatic indexing,” Communications of the ACM, vol. 18, no. 11, pp. 613–620, 1975, doi: 10.1145/361219.361220.",
    "[2] C. D. Manning, P. Raghavan, and H. Schütze, Introduction to Information Retrieval. Cambridge, U.K.: Cambridge University Press, 2008.",
    "[3] C. Buck and P. Koehn, “Quick and reliable document alignment via TF/IDF-weighted cosine distance,” in Proceedings of the First Conference on Machine Translation, vol. 2, Berlin, Germany, 2016, pp. 672–678, doi: 10.18653/v1/W16-2365.",
    "[4] J. J. Rocchio, “Relevance feedback in information retrieval,” in The SMART Retrieval System: Experiments in Automatic Document Processing, G. Salton, Ed. Englewood Cliffs, NJ, USA: Prentice-Hall, 1971, pp. 313–323.",
    "[5] K. Järvelin and J. Kekäläinen, “Cumulated gain-based evaluation of IR techniques,” ACM Transactions on Information Systems, vol. 20, no. 4, pp. 422–446, 2002, doi: 10.1145/582415.582418.",
    "[6] N. Drushchak and M. Romanyshyn, “Introducing the Djinni Recruitment Dataset: A corpus of anonymized CVs and job postings,” in Proceedings of the Third Ukrainian Natural Language Processing Workshop, Torino, Italy, 2024, pp. 8–13, doi: 10.63317/2dcvy45ws6yb.",
    "[7] X. Yu, R. Xu, C. Xue, J. Zhang, X. Ma, and Z. Yu, “ConFit v2: Improving resume-job matching using hypothetical resume embedding and runner-up hard-negative mining,” in Findings of the Association for Computational Linguistics: ACL 2025, Vienna, Austria, 2025, pp. 12775–12790, doi: 10.18653/v1/2025.findings-acl.661.",
    "[8] S. Vaishampayan et al., “Human and LLM-based resume matching: An observational study,” in Findings of the Association for Computational Linguistics: NAACL 2025, Albuquerque, NM, USA, 2025, pp. 4823–4838, doi: 10.18653/v1/2025.findings-naacl.270.",
    "[9] E. Tabassi, Artificial Intelligence Risk Management Framework (AI RMF 1.0), NIST AI 100-1. Gaithersburg, MD, USA: National Institute of Standards and Technology, 2023, doi: 10.6028/NIST.AI.100-1.",
    "[10] M. Raghavan, S. Barocas, J. Kleinberg, and K. Levy, “Mitigating bias in algorithmic hiring: Evaluating claims and practices,” in Proceedings of the 2020 Conference on Fairness, Accountability, and Transparency, Barcelona, Spain, 2020, pp. 469–481, doi: 10.1145/3351095.3372828.",
    "[11] R. T. Fielding, “Architectural styles and the design of network-based software architectures,” Ph.D. dissertation, University of California, Irvine, CA, USA, 2000.",
    "[12] M. Jones, J. Bradley, and N. Sakimura, “JSON Web Token (JWT),” Internet Engineering Task Force, RFC 7519, May 2015, doi: 10.17487/RFC7519.",
    "[13] K. Kent and M. Souppaya, Guide to Computer Security Log Management, NIST SP 800-92. Gaithersburg, MD, USA: National Institute of Standards and Technology, 2006, doi: 10.6028/NIST.SP.800-92.",
    "[14] A. Colas, J. Araki, Z. Zhou, B. Wang, and Z. Feng, “Knowledge-grounded natural language recommendation explanation,” in Proceedings of the 6th BlackboxNLP Workshop, Singapore, 2023, pp. 1–15, doi: 10.18653/v1/2023.blackboxnlp-1.1.",
]


def main():
    document = Document(PATH)
    if any(p.text.strip() == "2.1 Recruitment Information and CV–Job Description Matching" for p in document.paragraphs):
        raise SystemExit("Chapter 2 is already populated; no changes made.")

    placeholder = "NOTE: [Chapter 2 content will be written after academic sources are collected and verified. Planned figures: text-processing pipeline and Rocchio update illustration.]"
    for paragraph in list(document.paragraphs):
        if paragraph.text.strip() == placeholder:
            remove_paragraph(paragraph)

    anchor = find_paragraph(document, "CHAPTER 3. SYSTEM ANALYSIS AND DESIGN")
    for item_type, text in CONTENT:
        if item_type == "h2":
            insert_paragraph(document, anchor, text, level=2)
        elif item_type == "h3":
            insert_paragraph(document, anchor, text, level=3)
        elif item_type == "p":
            insert_paragraph(document, anchor, text)
        elif item_type == "eq":
            paragraph = insert_paragraph(document, anchor, text, italic=True, indent=False)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif item_type == "note":
            insert_note(document, anchor, text)

    insert_paragraph(document, anchor, "Table 2.1. Positioning of CareerFit against major approach categories", indent=False)
    insert_table(
        document,
        anchor,
        ["Approach", "Strength", "Limitation relative to CareerFit scope", "CareerFit position"],
        [
            ("Lexical TF-IDF/cosine", "Transparent, efficient, reproducible", "Limited semantic equivalence", "Implemented baseline with normalization, reasons, and validation"),
            ("Dense resume–job encoders", "Captures semantic relations", "Requires training data, governance, and benchmark validation", "Future hybrid/semantic extension; not claimed as implemented"),
            ("General job portal", "Strong search, publication, and application UX", "May not expose scoring, feedback, or policy internals", "Combines portal workflows with inspectable matching and automation"),
            ("Fully automated decision pipeline", "Reduces manual steps", "Higher control, accountability, and error-propagation risk", "Policy-gated HITL actions with confirmation and audit"),
            ("Explainable recommender", "Provides user-facing reasons", "Explanation may be unfaithful if not grounded", "Uses reasons tied to processed fields and lexical evidence"),
        ],
    )
    insert_paragraph(document, anchor, "The comparison shows that CareerFit deliberately accepts the semantic limitations of a lexical baseline in exchange for inspectability and implementation control. Its research value must therefore be assessed through reproducible behavior, workflow integration, and clearly stated limitations rather than through unsupported claims of general AI superiority.")
    insert_paragraph(document, anchor, "2.10 Chapter Summary", level=2)
    insert_paragraph(document, anchor, "This chapter established the foundations for CareerFit. Recruitment matching was framed as a family of ranking and decision-support tasks rather than a single hiring decision. TF-IDF and cosine similarity provide a transparent lexical baseline, while Rocchio relevance feedback enables bounded adaptation from explicit judgments. Ranking metrics evaluate ordered results under documented relevance assumptions. Human-in-the-Loop policy controls separate similarity evidence from consequential actions, and web security and audit mechanisms support controlled execution and traceability. The next chapter translates these principles into the requirements, architecture, data model, and workflows of the CareerFit system.")

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if "[CITATION REQUIRED]" in run.text:
                run.text = run.text.replace("[CITATION REQUIRED]", "[10]")

    reference_note = "NOTE: [References will be formatted consistently in IEEE style. Every entry must be cited in the text, and every in-text citation must have a corresponding reference entry.]"
    for paragraph in list(document.paragraphs):
        if paragraph.text.strip() == reference_note:
            remove_paragraph(paragraph)
    reference_anchor = find_paragraph(document, "APPENDICES")
    for reference in REFERENCES:
        paragraph = insert_paragraph(document, reference_anchor, reference, indent=False)
        paragraph.paragraph_format.left_indent = Inches(0.3)
        paragraph.paragraph_format.first_line_indent = Inches(-0.3)

    document.save(PATH)
    print(f"Updated {PATH} ({PATH.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
