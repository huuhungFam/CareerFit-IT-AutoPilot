from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = (
    ROOT
    / "Doc"
    / "working"
    / "CareerFit-Thesis-Report-before-20260811-chapter-restructure.docx"
)


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Replace visible text while retaining paragraph properties and child bookmarks."""
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def find_paragraph(doc: Document, exact_text: str) -> Paragraph:
    matches = [p for p in doc.paragraphs if p.text.strip() == exact_text]
    if len(matches) != 1:
        raise ValueError(f"Expected one paragraph {exact_text!r}, found {len(matches)}")
    return matches[0]


def replace_tokens(text: str, mapping: dict[str, str]) -> str:
    if not mapping:
        return text
    pattern = re.compile("|".join(re.escape(key) for key in sorted(mapping, key=len, reverse=True)))
    return pattern.sub(lambda match: mapping[match.group(0)], text)


if not DOCX.exists():
    raise FileNotFoundError(DOCX)
BACKUP.parent.mkdir(parents=True, exist_ok=True)
if not BACKUP.exists():
    shutil.copy2(DOCX, BACKUP)

doc = Document(DOCX)

# Move the existing solution-design block (2.11 through 2.15) so that it follows
# the Chapter 3 title and precedes the existing implementation overview.
start = find_paragraph(doc, "2.11 System Architecture")._p
end = find_paragraph(doc, "2.16 Chapter Summary")._p
anchor = find_paragraph(doc, "3.1 Implementation Overview")._p
children = list(doc.element.body.iterchildren())
start_index = children.index(start)
end_index = children.index(end)
anchor_index = children.index(anchor)
if not (start_index < end_index < anchor_index):
    raise ValueError("Unexpected Chapter 2/3 element order")
moving = children[start_index:end_index]
for element in moving:
    anchor.addprevious(element)

heading_map = {
    "CHAPTER 2. THEORETICAL BACKGROUND AND SOLUTION DESIGN": "CHAPTER 2. THEORETICAL BACKGROUND",
    "2.16 Chapter Summary": "2.11 Chapter Summary",
    "CHAPTER 3. SYSTEM IMPLEMENTATION": "CHAPTER 3. SYSTEM DESIGN AND IMPLEMENTATION",
    "2.11 System Architecture": "3.1 System Architecture",
    "2.12 Module Design": "3.2 Module Design",
    "2.13 Data Design": "3.3 Data Design",
    "2.13.1 Core Identity and Recruitment Data": "3.3.1 Core Identity and Recruitment Data",
    "2.13.2 Automation, Communication, and Operational Data": "3.3.2 Automation, Communication, and Operational Data",
    "2.14 Security, Failure, and Consistency Design": "3.4 Security, Failure, and Consistency Design",
    "2.14.1 Authentication and Authorization": "3.4.1 Authentication and Authorization",
    "2.14.2 Failure Handling": "3.4.2 Failure Handling",
    "2.14.3 Identified Design Risks": "3.4.3 Identified Design Risks",
    "2.15 Deployment Architecture": "3.5 Deployment Architecture",
    "3.1 Implementation Overview": "3.6 Implementation Overview",
    "3.2 Authentication and Authorization Implementation": "3.7 Authentication and Authorization Implementation",
    "3.2.1 Login and JWT Processing": "3.7.1 Login and JWT Processing",
    "3.2.2 URL Rules and Ownership Checks": "3.7.2 URL Rules and Ownership Checks",
    "3.3 CV Ingestion and Validation": "3.8 CV Ingestion and Validation",
    "3.3.1 Upload and Manual Creation": "3.8.1 Upload and Manual Creation",
    "3.3.2 File Validation and Extraction": "3.8.2 File Validation and Extraction",
    "3.4 Text Normalization and TF-IDF Implementation": "3.9 Text Normalization and TF-IDF Implementation",
    "3.4.1 Text Normalization": "3.9.1 Text Normalization",
    "3.4.2 Static Corpus and Vector Construction": "3.9.2 Static Corpus and Vector Construction",
    "3.5 Matching and Recommendation Implementation": "3.10 Matching and Recommendation Implementation",
    "3.5.1 Scoring Service": "3.10.1 Scoring Service",
    "3.5.2 Matching Orchestration and Persistence": "3.10.2 Matching Orchestration and Persistence",
    "3.5.3 Recommendation Service": "3.10.3 Recommendation Service",
    "3.6 Feedback Learning with Rocchio": "3.11 Feedback Learning with Rocchio",
    "3.7 Application and Recruiter Workflow": "3.12 Application and Recruiter Workflow",
    "3.8 AutoFit and Background Processing": "3.13 AutoFit and Background Processing",
    "3.8.1 Async Executor and Scheduler": "3.13.1 Async Executor and Scheduler",
    "3.8.2 Auto-Apply": "3.13.2 Auto-Apply",
    "3.8.3 Notification Guard and Delivery": "3.13.3 Notification Guard and Delivery",
    "3.9 Email Action and Audit Implementation": "3.14 Email Action and Audit Implementation",
    "3.10 Persistence and API Implementation": "3.15 Persistence and API Implementation",
    "3.11 Frontend Integration": "3.16 Frontend Integration",
    "3.12 Deployment and Observability Implementation": "3.17 Deployment and Observability Implementation",
    "3.13 Implementation Limitations": "3.18 Implementation Limitations",
    "3.14 Chapter Summary": "3.19 Chapter Summary",
    "PART 4. REFERENCES": "REFERENCES",
    "PART 5. APPENDICES": "APPENDICES",
}

paragraphs_by_text = {p.text.strip(): p for p in doc.paragraphs}
missing_headings = [text for text in heading_map if text not in paragraphs_by_text]
if missing_headings:
    raise ValueError(f"Missing headings: {missing_headings}")
for old, new in heading_map.items():
    set_paragraph_text(paragraphs_by_text[old], new)

figure_number_map = {
    "Figure 2.5": "Figure 3.1",
    "Figure 2.6": "Figure 3.2",
    "Figure 2.7": "Figure 3.3",
    "Figure 3.1": "Figure 3.4",
    "Figure 3.2": "Figure 3.5",
    "Figure 3.3": "Figure 3.6",
    "Figure 3.4": "Figure 3.7",
    "Figure 3.5": "Figure 3.8",
    "Figure 3.6": "Figure 3.9",
    "Figure 3.7": "Figure 3.10",
    "Figure 3.8": "Figure 3.11",
    "Figure 3.9": "Figure 3.12",
    "Figure 3.10": "Figure 3.13",
}
table_number_map = {
    "Table 2.2": "Table 3.1",
    "Table 2.3": "Table 3.2",
    "Table 2.4": "Table 3.3",
    "Table 3.1": "Table 3.4",
    "Table 3.2": "Table 3.5",
    "Table 3.3": "Table 3.6",
    "Table 3.4": "Table 3.7",
    "Table 3.5": "Table 3.8",
    "Table 3.6": "Table 3.9",
}
number_map = {**figure_number_map, **table_number_map}

# Update visible captions and any actor-level prose references, but leave cached
# TOC/figure/table lists to Microsoft Word's field refresh.
for paragraph in doc.paragraphs:
    if paragraph.style.name.lower().startswith("toc"):
        continue
    updated = replace_tokens(paragraph.text, number_map)
    if updated != paragraph.text:
        set_paragraph_text(paragraph, updated)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                updated = replace_tokens(paragraph.text, number_map)
                if updated != paragraph.text:
                    set_paragraph_text(paragraph, updated)

text_replacements = {
    "Chapter 2 presents the theoretical background and solution design.":
        "Chapter 2 presents the theoretical background.",
    "Chapter 3 explains the implementation.":
        "Chapter 3 presents the system design and implementation.",
    (
        "An append-oriented audit history supports traceability, but database permissions, retention, "
        "integrity, and access control determine whether that history is trustworthy. Operational logs "
        "used for debugging and domain audit records used to explain a business action may have different "
        "schemas and retention requirements. Chapter 2 defines these responsibilities at the design level, "
        "and Chapter 3 describes the mechanisms actually implemented in CareerFit."
    ): (
        "An append-oriented audit history supports traceability, but database permissions, retention, "
        "integrity, and access control determine whether that history is trustworthy. Operational logs "
        "used for debugging and domain audit records used to explain a business action may have different "
        "schemas and retention requirements. Chapter 3 presents these design responsibilities and the "
        "mechanisms implemented in CareerFit."
    ),
    (
        "This chapter connected the theoretical foundations to the CareerFit solution design. TF-IDF, "
        "cosine similarity, Rocchio feedback, ranking metrics, and Human-in-the-Loop principles explain "
        "the core decision-support behavior. The architecture, modules, data model, security boundaries, "
        "consistency rules, and deployment topology define how those ideas are separated into implementable "
        "components."
    ): (
        "This chapter presented the theoretical foundations used by CareerFit. It distinguished matching, "
        "recommendation, and recruitment action; described text preprocessing, TF-IDF, cosine similarity, "
        "Rocchio feedback, and ranking metrics; and discussed Human-in-the-Loop control, explainability, "
        "security, auditability, related work, and the research gap. These foundations guide the system "
        "design and implementation presented in Chapter 3."
    ),
    (
        "CareerFit is implemented as two application projects and a shared runtime environment. The backend "
        "is a Java 21 application based on Spring Boot 3.2.5. The frontend is a React 18 single-page "
        "application written in TypeScript and built with Vite. PostgreSQL is the transactional database, "
        "Flyway owns schema migration, and Docker Compose provides the reproducible local database and optional "
        "backend container. The implementation follows the modular-monolith design established in Chapter 2: "
        "domain packages share one backend process while retaining controllers, services, repositories, "
        "entities, and DTOs for each functional area."
    ): (
        "CareerFit is implemented as two application projects and a shared runtime environment. The backend "
        "is a Java 21 application based on Spring Boot 3.2.5. The frontend is a React 18 single-page "
        "application written in TypeScript and built with Vite. PostgreSQL is the transactional database, "
        "Flyway owns schema migration, and Docker Compose provides the reproducible local database and optional "
        "backend container. The implementation follows the modular-monolith design presented earlier in this "
        "chapter: domain packages share one backend process while retaining controllers, services, repositories, "
        "entities, and DTOs for each functional area."
    ),
    (
        "This chapter explained the current implementation in Spring Boot, React, PostgreSQL, and Flyway. "
        "It covered JWT security, CV review, direct matching, the Potential assessment, feedback learning, "
        "applications, Talent Pool, AutoFit, email actions, Job/CV reporting, administrator moderation, frontend "
        "integration, and monitoring. Chapter 4 evaluates the refreshed system."
    ): (
        "This chapter presented the CareerFit architecture, modules, data model, security and consistency "
        "design, deployment topology, and implementation in Spring Boot, React, PostgreSQL, and Flyway. It "
        "covered JWT security, CV review, direct matching, the Potential assessment, feedback learning, "
        "applications, Talent Pool, AutoFit, email actions, Job/CV reporting, administrator moderation, frontend "
        "integration, and monitoring. Chapter 4 evaluates the refreshed system."
    ),
}

paragraphs_by_text = {p.text.strip(): p for p in doc.paragraphs}
missing_text = [text for text in text_replacements if text not in paragraphs_by_text]
if missing_text:
    raise ValueError(f"Missing synchronization paragraphs: {missing_text}")
for old, new in text_replacements.items():
    set_paragraph_text(paragraphs_by_text[old], new)

# Synchronize title/description text for numbered diagrams. Descriptive actor-level
# alt text that does not depend on figure numbering is intentionally preserved.
paragraphs = doc.paragraphs
for index, paragraph in enumerate(paragraphs[:-1]):
    if "w:drawing" not in paragraph._p.xml:
        continue
    caption = paragraphs[index + 1]
    if caption.style.name != "Figure Caption":
        continue
    caption_text = caption.text.strip()
    for node in paragraph._p.iter():
        if not node.tag.endswith("}docPr"):
            continue
        for attr in ("title", "descr"):
            old_value = node.get(attr, "")
            if old_value.startswith("Figure ") or old_value.startswith("Current CareerFit diagram"):
                node.set(attr, caption_text)

doc.save(DOCX)

print(f"updated={DOCX}")
print(f"backup={BACKUP}")
print(f"moved_body_elements={len(moving)}")
print(f"heading_updates={len(heading_map)}")
print(f"figure_renumberings={len(figure_number_map)}")
print(f"table_renumberings={len(table_number_map)}")
print(f"summary_updates={len(text_replacements)}")
