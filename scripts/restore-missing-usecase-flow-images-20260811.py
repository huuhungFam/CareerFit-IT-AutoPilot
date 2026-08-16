from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Emu


ROOT = Path(__file__).resolve().parents[1]
SOURCE = (
    ROOT
    / "Doc"
    / "working"
    / "CareerFit-Thesis-Report-before-20260811-shortened-usecases.docx"
)
TARGET = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"


def find_caption(doc: Document, text: str):
    matches = [p for p in doc.paragraphs if p.text.strip() == text]
    if len(matches) != 1:
        raise ValueError(f"Expected one caption {text!r}, found {len(matches)}")
    return matches[0]


def image_before_caption(doc: Document, caption_text: str):
    caption = find_caption(doc, caption_text)
    index = next(i for i, paragraph in enumerate(doc.paragraphs) if paragraph._p is caption._p)
    paragraph = doc.paragraphs[index - 1]
    inlines = [node for node in paragraph._p.iter() if node.tag.endswith("}inline")]
    if len(inlines) != 1:
        raise ValueError(f"Expected one image before {caption_text!r}, found {len(inlines)}")
    inline = inlines[0]
    blips = [node for node in inline.iter() if node.tag.endswith("}blip")]
    extents = [node for node in inline.iter() if node.tag.endswith("}extent")]
    if len(blips) != 1 or not extents:
        raise ValueError(f"Cannot extract image geometry for {caption_text!r}")
    relationship_id = blips[0].get(qn("r:embed"))
    blob = doc.part.related_parts[relationship_id].blob
    extent = extents[0]
    return blob, int(extent.get("cx")), int(extent.get("cy"))


source = Document(SOURCE)
target = Document(TARGET)

jobs = [
    (
        "Figure 1.4. Feedback learning and recomputation sequence",
        "Figure 1.5. Feedback learning and recomputation sequence",
    ),
    (
        "Figure 1.5. AutoFit decision flow",
        "Figure 1.6. AutoFit decision flow",
    ),
]

for source_caption, target_caption in jobs:
    blob, width, height = image_before_caption(source, source_caption)
    caption = find_caption(target, target_caption)
    index = next(i for i, paragraph in enumerate(target.paragraphs) if paragraph._p is caption._p)
    image_paragraph = target.paragraphs[index - 1]
    if any(node.tag.endswith("}drawing") for node in image_paragraph._p.iter()):
        raise ValueError(f"Target image already exists before {target_caption!r}")
    run = image_paragraph.add_run()
    run.add_picture(BytesIO(blob), width=Emu(width), height=Emu(height))
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.keep_with_next = True
    image_paragraph.paragraph_format.keep_together = True
    for node in image_paragraph._p.iter():
        if node.tag.endswith("}docPr"):
            node.set("title", target_caption)
            node.set("descr", target_caption)

target.save(TARGET)
print(f"updated={TARGET}")
print(f"source={SOURCE}")
print(f"restored_images={len(jobs)}")
