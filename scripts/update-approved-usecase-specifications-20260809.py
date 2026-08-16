from __future__ import annotations

from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-before-20260809-approved-usecase-specifications.docx"
FIGURE = ROOT / "Doc" / "figures" / "fig-1-2-approved-usecases-20260809.png"
NOTIFICATION_FIGURE = ROOT / "Doc" / "figures" / "fig-3-8-notification-policy-guard-20260809.png"


def lines(*items: str) -> str:
    return "\n".join(items)


def flow(label: str, *steps: str) -> str:
    return label + "\n" + "\n".join(f"{i}. {step}" for i, step in enumerate(steps, 1))


CASES = [
    {
        "name": "Manage Career Profile",
        "caption": "Manage career profile",
        "primary": "Candidate",
        "secondary": "None",
        "description": "This use case allows a Candidate to maintain the career information used by CareerFit. The Candidate can manage their profile, Portfolio, and CV collection, including reviewing and confirming CV content for subsequent matching.",
        "pre": lines("P1. The Candidate account is active.", "P2. The Candidate is authenticated.", "P3. The Candidate is authorized to manage only their own profile, Portfolio, and CVs."),
        "trigger": "The Candidate opens the Career Profile area and selects an available profile, Portfolio, or CV management action.",
        "main": lines(
            "1. The Candidate opens the Career Profile area.",
            "2. The System displays the Candidate Profile, Portfolio, and owned CVs with their current statuses.",
            "3. The Candidate selects the option to upload a new CV.",
            "4. The System displays the CV upload option and supported file requirements.",
            "5. The Candidate selects and submits a CV file.",
            "6. The System validates the submitted file.",
            "7. The System creates the CV and extracts its information into reviewable sections.",
            "8. The System marks the CV as REVIEW_REQUIRED and displays the extracted information and detected warnings.",
            "9. The Candidate reviews the information and makes necessary corrections.",
            "10. The Candidate confirms the reviewed CV.",
            "11. The System validates the confirmed CV content.",
            "12. The System accepts the CV for matching processing and displays its processing status.",
            "13. After successful processing, the System marks the CV as SCORING_DONE.",
            "14. If no default CV exists, the System makes the first successfully processed CV the default.",
            "15. The System displays the updated CV collection. Generated Matching results are available through UC-05.",
        ),
        "alt": "\n\n".join([
            flow("A1 – Create a CV manually (from Step 3)", "The Candidate selects manual CV creation.", "The System displays the manual CV form.", "The Candidate enters and submits the required CV information.", "The System validates the information and creates the CV in REVIEW_REQUIRED status.", "The flow continues at Step 8 of the Main Flow."),
            flow("A2 – Save manual CV work as a draft (from A1, Step 3)", "The Candidate saves unfinished manual content.", "The System stores the CV in DRAFT status and displays it in the CV collection.", "The Use Case ends; the Candidate may resume the draft later."),
            flow("A3 – Update Candidate Profile (from Step 3)", "The Candidate opens the Candidate Profile section.", "The System displays the current profile.", "The Candidate changes editable information.", "The System validates, saves, and displays the updated profile."),
            flow("A4 – Manage Portfolio (from Step 3)", "The Candidate opens the Portfolio section.", "The System displays existing links and projects.", "The Candidate adds, updates, or removes a Portfolio item.", "The System validates, saves, and displays the change."),
            flow("A5 – Resume an unconfirmed CV (from Step 3)", "The Candidate selects an owned CV in DRAFT or REVIEW_REQUIRED status.", "The System displays its editable review content.", "The Candidate updates the content.", "If confirmed, the flow continues at Step 11; otherwise, the CV remains unconfirmed."),
            flow("A6 – Select a default CV (from Step 3)", "The Candidate selects an owned CV in SCORING_DONE status.", "The Candidate requests to make it the default.", "The System updates the default designation and displays the CV collection."),
            flow("A7 – Delete an eligible CV (from Step 3)", "The Candidate selects an owned, non-default CV for deletion.", "The System verifies deletion eligibility.", "The Candidate confirms deletion.", "The System removes the CV from the collection."),
        ]),
        "exc": "\n\n".join([
            flow("E1 – Unsupported or invalid CV file (from Step 6)", "The System rejects a file that does not meet supported requirements.", "The System informs the Candidate of the validation problem.", "No valid CV is created from the rejected file."),
            flow("E2 – CV extraction fails (from Step 7)", "The System cannot complete extraction.", "The System marks the CV as FAILED and records an available reason.", "No Matching result is made available for the failed CV."),
            flow("E3 – Confirmed CV content is invalid (from Step 11)", "The System detects missing or invalid confirmation information.", "The System rejects confirmation and displays the issues.", "The CV remains unconfirmed and may be corrected."),
            flow("E4 – CV is unavailable or unauthorized", "The System cannot find the CV or determines that it belongs to another Candidate.", "The System denies the operation.", "No CV information is changed."),
            flow("E5 – CV is not eligible to become the default (from A6, Step 2)", "The selected CV is not in SCORING_DONE status.", "The System rejects the request.", "The existing default CV remains unchanged."),
            flow("E6 – CV cannot be deleted (from A7, Step 2)", "The CV is default, BANNED, or has pending reports.", "The System rejects deletion and explains that the CV is not eligible.", "The CV remains in the collection."),
            flow("E7 – Processing fails after confirmation (from Step 12)", "The System cannot complete processing.", "The System marks the CV as FAILED and displays the failure status.", "No invalid Matching result is presented."),
        ]),
        "post": lines("Success Postconditions", "S1. Approved Candidate Profile or Portfolio changes are stored.", "S2. A draft CV remains in DRAFT; an unconfirmed CV remains REVIEW_REQUIRED.", "S3. A successfully processed CV reaches SCORING_DONE.", "S4. The Candidate has at most one default CV.", "S5. Matching generation may complete as a processing result; viewing results remains part of UC-05.", "S6. If no eligible Jobs are available, the CV may still reach SCORING_DONE without available Matching results.", "", "Minimal Guarantees", "F1. Invalid changes do not replace previously valid career information.", "F2. A failed CV does not expose invalid Matching results.", "F3. The Candidate cannot modify another Candidate's career information."),
        "related": lines("UC-03 – Manage Job Applications", "UC-05 – Review Personalized Career Insights", "UC-06 – Manage AutoFit"),
        "priority": "High",
    },
    {
        "name": "Explore Jobs", "caption": "Explore Jobs", "primary": "Candidate", "secondary": "Guest – alternative unauthenticated context",
        "description": "This use case allows a Candidate to discover and review available Jobs. Public exploration is also available to a Guest, while protected Candidate actions require authentication.",
        "pre": lines("P1. No authentication is required for public Job exploration.", "P2. Candidate-specific information and protected actions require an active, authenticated Candidate account."),
        "trigger": "The Candidate opens Job exploration or submits supported search and filter criteria.",
        "main": lines("1. The Candidate opens Job exploration.", "2. The System displays available ACTIVE Jobs and supported search, filter, sorting, and pagination controls.", "3. The Candidate enters a keyword or selects supported criteria.", "4. The System applies the criteria.", "5. The System displays a paginated list of matching ACTIVE Jobs.", "6. The Candidate reviews the Job summaries.", "7. The Candidate selects a Job.", "8. The System verifies that the Job is publicly available.", "9. The System displays Job details and available employer information.", "10. The System displays similar Jobs when available.", "11. The Candidate may return to the results or select another Job."),
        "alt": "\n\n".join([
            flow("A1 – Guest explores Jobs (from Step 1)", "A Guest opens the public Job page.", "The System displays publicly available ACTIVE Jobs.", "The Guest searches, filters, and views Job details.", "The System requests authentication before a protected Candidate action."),
            flow("A2 – View urgent Jobs (from Step 3)", "The Candidate selects the urgent-Job option.", "The System filters the results to Jobs marked urgent.", "The flow continues at Step 5."),
            flow("A3 – Browse without criteria (from Step 3)", "The Candidate enters no search or filter criteria.", "The System displays available ACTIVE Jobs using the default ordering.", "The flow continues at Step 6."),
            flow("A4 – No Jobs match (from Step 5)", "The System finds no matching Job.", "The System displays an empty result state.", "The Candidate may change or clear the criteria."),
            flow("A5 – Change sorting or page (from Step 5)", "The Candidate changes sorting or requests more results.", "The System displays the corresponding result set.", "The flow returns to Step 6."),
            flow("A6 – Open a similar Job (from Step 10)", "The Candidate selects a similar Job.", "The System loads it and returns to Step 8."),
        ]),
        "exc": "\n\n".join([
            flow("E1 – Job list cannot be loaded (from Step 2 or Step 4)", "The System cannot complete the Job request.", "The System informs the Candidate.", "The Candidate may retry."),
            flow("E2 – Selected Job is unavailable (from Step 8)", "The System cannot find the Job or it is no longer public.", "The System does not display the unavailable details.", "The Candidate may return to the Job list."),
            flow("E3 – Protected action requires authentication", "The System detects that valid Candidate authentication is missing.", "The System does not perform the action.", "The System requests authentication."),
        ]),
        "post": lines("Success Postconditions", "S1. The Candidate or Guest can view current public Job information.", "S2. The selected criteria determine the displayed result set.", "S3. Exploration alone does not create an Application.", "", "Minimal Guarantees", "F1. An unavailable or non-public Job is not exposed through public detail.", "F2. A protected action is not executed without authentication.", "F3. Job exploration does not modify Job or Candidate data."),
        "related": lines("UC-03 – Manage Job Applications", "UC-05 – Review Personalized Career Insights", "UC-07 – Report Suspicious Recruitment Content"), "priority": "High",
    },
    {
        "name": "Manage Job Applications", "caption": "Manage Job Applications", "primary": "Candidate", "secondary": "None",
        "description": "This use case allows a Candidate to submit and manage Job Applications. It also covers Application history, eligible withdrawal, and responses to Recruiter invitations.",
        "pre": lines("P1. The Candidate account is active.", "P2. The Candidate is authenticated.", "P3. The Candidate manages only their own Applications.", "P4. For a new Application, the selected Job is ACTIVE.", "P5. The Candidate owns the selected CV or has an owned default CV."),
        "trigger": "The Candidate applies for an ACTIVE Job or opens Application management to review an Application or invitation.",
        "main": lines("1. The Candidate opens an ACTIVE Job.", "2. The Candidate selects Apply.", "3. The System verifies the Candidate and Job.", "4. The System checks for an existing Application.", "5. The System identifies the selected owned CV or owned default CV.", "6. The System creates an Application in PENDING status.", "7. The System confirms successful submission.", "8. The Application becomes available in Candidate history and the Recruiter's applicant list.", "9. The Candidate reviews the submitted Application."),
        "alt": "\n\n".join([
            flow("A1 – Use an explicitly selected CV (from Step 5)", "The request identifies a specific CV.", "The System verifies Candidate ownership.", "The System uses the selected CV and continues at Step 6."),
            flow("A2 – Use the default CV (from Step 5)", "No specific CV is supplied.", "The System retrieves the Candidate's default CV.", "The flow continues at Step 6."),
            flow("A3 – View Application history (alternative from the Trigger)", "The Candidate opens Application management.", "The System displays owned Applications and statuses.", "The Candidate may filter and review an Application."),
            flow("A4 – Withdraw an eligible Application (from A3)", "The Candidate selects an Application that is not APPROVED or REJECTED.", "The Candidate requests withdrawal.", "The System verifies ownership and eligibility.", "The System changes the status to NOT_INTERESTED and displays it."),
            flow("A5 – View Recruiter invitations (alternative from the Trigger)", "The Candidate opens invitations.", "The System displays Applications in INVITED status.", "The Candidate selects an invitation."),
            flow("A6 – Accept an invitation (from A5)", "The Candidate accepts the invitation.", "The System verifies ownership and INVITED status.", "The System changes the status to PENDING and displays it in history."),
            flow("A7 – Decline an invitation (from A5)", "The Candidate declines the invitation.", "The System verifies ownership and INVITED status.", "The System changes the status to NOT_INTERESTED and displays the result."),
        ]),
        "exc": "\n\n".join([
            flow("E1 – Job is unavailable for application (from Step 3)", "The Job is missing or not ACTIVE.", "The System rejects the request.", "No Application is created."),
            flow("E2 – Application already exists (from Step 4)", "The System finds an existing Application for the Candidate and Job.", "The System rejects the duplicate request.", "The existing Application remains unchanged."),
            flow("E3 – No eligible owned CV is available (from Step 5)", "The System finds neither the requested owned CV nor an owned default CV.", "The System rejects the request.", "No Application is created."),
            flow("E4 – Another Candidate's CV is selected (from A1)", "The System detects invalid ownership.", "The System denies the request.", "No Application is created."),
            flow("E5 – Application cannot be withdrawn (from A4)", "The Application is APPROVED, REJECTED, unavailable, or unauthorized.", "The System rejects withdrawal.", "The status remains unchanged."),
            flow("E6 – Invitation cannot be processed (from A6 or A7)", "The invitation is missing, unauthorized, or no longer INVITED.", "The System rejects the response.", "The status remains unchanged."),
        ]),
        "post": lines("Success Postconditions", "S1. A manually submitted Application is PENDING.", "S2. It appears in Candidate history and the Recruiter's applicant list.", "S3. A withdrawn Application is NOT_INTERESTED.", "S4. An accepted invitation changes from INVITED to PENDING.", "S5. A declined invitation changes from INVITED to NOT_INTERESTED.", "", "Minimal Guarantees", "F1. Failed submission creates no duplicate or unauthorized Application.", "F2. Rejected withdrawal or response does not change status.", "F3. A Candidate cannot manage another Candidate's Application."),
        "related": lines("UC-01 – Manage Career Profile", "UC-02 – Explore Jobs", "UC-06 – Manage AutoFit", "UC-09 – Review and Process Applicants", "UC-14 – Respond Through Actionable Email"), "priority": "High",
    },
    {
        "name": "Provide Matching Feedback", "caption": "Provide Matching Feedback", "primary": "Candidate", "secondary": "None",
        "description": "This use case allows a Candidate to evaluate a CV–Job Matching result. The Feedback may be used to improve later rankings for that Candidate.",
        "pre": lines("P1. The Candidate account is active.", "P2. The Candidate is authenticated.", "P3. The Matching belongs to a CV owned by the Candidate."),
        "trigger": "The Candidate selects a Feedback option for a Matching result.",
        "main": lines("1. The Candidate opens a Matching result.", "2. The System displays the Job, Matching information, and supported Feedback options.", "3. The Candidate selects a Feedback type.", "4. The System verifies ownership.", "5. The System validates the Feedback type.", "6. The System records the Feedback.", "7. The System confirms receipt.", "8. When applicable, the System uses the Feedback for later personalized rankings."),
        "alt": "\n\n".join([flow("A1 – Change previous Feedback (from Step 3)", "The Candidate selects another type for a Matching with existing Feedback.", "The System validates and replaces the previous Feedback.", "The flow continues at Step 7."), flow("A2 – Mark the Job not interesting (from Step 3)", "The Candidate selects NOT_INTERESTED.", "The System records it without treating it as a preference-learning signal.", "The flow continues at Step 7.")]),
        "exc": "\n\n".join([flow("E1 – Matching is unavailable or unauthorized (from Step 4)", "The System cannot find the Matching or ownership is invalid.", "The System rejects the request.", "No Feedback changes."), flow("E2 – Unsupported Feedback type (from Step 5)", "The submitted value is unsupported.", "The System rejects it and informs the Candidate.", "Existing Feedback remains unchanged.")]),
        "post": lines("Success Postconditions", "S1. The Matching has the Candidate's latest Feedback.", "S2. Supported Feedback may influence later rankings.", "S3. NOT_INTERESTED is recorded without preference-vector learning.", "", "Minimal Guarantees", "F1. Invalid or unauthorized Feedback does not change the Matching.", "F2. Previous valid Feedback remains unchanged when an update is rejected."),
        "related": lines("UC-05 – Review Personalized Career Insights", "UC-14 – Respond Through Actionable Email"), "priority": "High",
    },
    {
        "name": "Review Personalized Career Insights", "caption": "Review Personalized Career Insights", "primary": "Candidate", "secondary": "None",
        "description": "This use case allows a Candidate to review personalized Job recommendations, available CV–Job Matching results, and career analytics. Matching and Job Recommendation remain separate workflows within the career-insights area.",
        "pre": lines("P1. The Candidate account is active.", "P2. The Candidate is authenticated."),
        "trigger": "The Candidate opens personalized career insights and selects recommendations, Matching results, or analytics.",
        "main": lines("1. The Candidate opens personalized career insights.", "2. The System displays the available insight views.", "3. The Candidate selects personalized Job recommendations.", "4. The System prepares recommendations using supported Candidate information and recommendation inputs.", "5. The System displays recommended ACTIVE Jobs in ranked order.", "6. The Candidate reviews available scores, labels, and matching-skill information.", "7. The Candidate selects a recommended Job.", "8. The System verifies that the Job remains available.", "9. The System displays current Job details.", "10. The Candidate may return and review other Jobs."),
        "alt": "\n\n".join([
            flow("A1 – Review CV–Job Matching results (from Step 3)", "The Candidate selects Matching results.", "The System identifies available results for the Candidate's CVs.", "The System displays ranked cards with available scores, labels, Potential indicators, and explanations.", "The Candidate may open a Job detail."),
            flow("A2 – Review Candidate analytics (from Step 3)", "The Candidate selects analytics.", "The System displays available overview, Matching trends, skill demand, and profile gaps.", "The Candidate reviews the information."),
            flow("A3 – Limited recommendation data (from Step 4)", "Some profile or preference information is incomplete.", "The System uses the remaining supported information.", "If no suitable result exists, the System displays an empty recommendation state.", "Matching results and analytics remain independently accessible."),
            flow("A4 – No Matching results are available (from A1)", "The System finds no available Matching result.", "The System displays an empty Matching state.", "Recommendations and analytics remain accessible."),
            flow("A5 – Limited analytics data (from A2)", "Data is insufficient for one or more analytical sections.", "The System displays available analytics and limited-data states for the rest."),
        ]),
        "exc": flow("E1 – Selected Job is unavailable (from Step 8 or A1)", "The System cannot find the Job or it is no longer Candidate-visible.", "The System does not display it.", "The Candidate may return to remaining results."),
        "post": lines("Success Postconditions", "S1. Recommendations and CV–Job Matching results remain separate review workflows.", "S2. Available Matching and analytics information can be reviewed without changing operational data.", "S3. Opening a Job does not create an Application.", "", "Minimal Guarantees", "F1. Missing data is represented as a limited-data or empty state.", "F2. An unavailable Job is not exposed as accessible detail.", "F3. Review does not modify Profile, CV, Feedback, or Application data."),
        "related": lines("UC-01 – Manage Career Profile", "UC-02 – Explore Jobs", "UC-03 – Manage Job Applications", "UC-04 – Provide Matching Feedback"), "priority": "High",
    },
    {
        "name": "Manage AutoFit", "caption": "Manage AutoFit", "primary": "Candidate", "secondary": "None",
        "description": "This use case allows a Candidate to configure and control AutoFit. CareerFit can use the approved configuration to submit a limited number of Applications for eligible matching Jobs.",
        "pre": lines("P1. The Candidate account is active.", "P2. The Candidate is authenticated."),
        "trigger": lines("Primary Trigger: The Candidate manually requests an AutoFit run or changes AutoFit configuration.", "Alternative Trigger: An AutoFit execution is initiated automatically according to the system-configured AutoFit schedule."),
        "main": lines("1. The Candidate opens AutoFit management.", "2. The System displays the current AutoFit status and supported configuration.", "3. The Candidate enables AutoFit and sets the required matching threshold.", "4. The System validates a threshold from 50 through 100 and saves the configuration.", "5. The Candidate requests a manual AutoFit run.", "6. The System verifies an owned default CV in SCORING_DONE status.", "7. The System identifies Matchings for ACTIVE Jobs that meet the threshold.", "8. The System excludes Jobs with an existing Candidate Application.", "9. The System creates eligible Applications in AUTO_APPLIED status within the supported per-run limit.", "10. The System displays the number of Applications created."),
        "alt": "\n\n".join([
            flow("A1 – Automatic AutoFit execution (alternative from the Trigger)", "CareerFit initiates AutoFit according to the system-configured schedule.", "The System verifies that AutoFit is enabled and not paused.", "The flow continues at Step 6."),
            flow("A2 – Pause AutoFit (from Step 3)", "The Candidate requests a pause.", "The System records it.", "Automatic execution is skipped while the pause remains effective."),
            flow("A3 – Resume AutoFit (from Step 3)", "The Candidate requests to resume AutoFit.", "The System removes the active pause.", "The System displays AutoFit as available."),
            flow("A4 – No eligible Matching result (from Step 7)", "No Matching satisfies the verified conditions.", "The System creates no Application.", "The System reports zero Applications."),
            flow("A5 – No eligible default CV (from Step 6)", "The System finds no owned default CV in SCORING_DONE status.", "The System creates no Application.", "The System reports zero Applications."),
            flow("A6 – Save configuration without running AutoFit (from Step 4)", "The Candidate does not request a manual run.", "The System retains the saved AutoFit configuration for future manual or automatic execution.", "The Use Case ends."),
        ]),
        "exc": flow("E1 – Invalid AutoFit threshold (from Step 4)", "The threshold is below 50 or above 100.", "The System rejects the change and informs the Candidate.", "The previous valid configuration remains unchanged."),
        "post": lines("Success Postconditions", "S1. Valid AutoFit configuration is stored.", "S2. An execution creates no more than three eligible Applications per run.", "S3. Created automatic Applications have AUTO_APPLIED status.", "S4. No duplicate Application is created.", "S5. A run with no eligible CV or Matching creates zero Applications.", "", "Minimal Guarantees", "F1. An invalid threshold does not replace valid configuration.", "F2. AutoFit does not use a non-default or non-SCORING_DONE CV.", "F3. AutoFit does not apply to a non-ACTIVE Job."),
        "related": lines("UC-01 – Manage Career Profile", "UC-03 – Manage Job Applications", "UC-05 – Review Personalized Career Insights"), "priority": "High",
    },
    {
        "name": "Report Suspicious Recruitment Content", "caption": "Report Suspicious Recruitment Content", "primary": "Candidate, Recruiter", "secondary": "None",
        "description": "This use case allows a Candidate to report a suspicious Job and a Recruiter to report a suspicious CV for later Administrator review.",
        "pre": lines("P1. The reporting account is active and authenticated.", "P2. A Candidate reports an ACTIVE Job.", "P3. A Recruiter reports a CV through an owned Job where it is visible through an Application or Matching."),
        "trigger": "A Candidate or Recruiter selects Report for eligible recruitment content.",
        "main": lines("1. The Candidate opens an ACTIVE Job.", "2. The Candidate selects Report.", "3. The System displays supported reasons and a comment field.", "4. The Candidate selects a reason and submits.", "5. The System verifies actor, target, reason, and duplicate state.", "6. The System creates a PENDING report.", "7. The System confirms submission for Administrator review."),
        "alt": "\n\n".join([flow("A1 – Recruiter reports a CV (alternative from the Trigger)", "The Recruiter opens a CV visible through an owned Job.", "The Recruiter selects Report and supplies the reason, comment, and Job.", "The System verifies ownership and CV visibility.", "The System creates a PENDING report and confirms submission."), flow("A2 – Use OTHER reason", "The actor selects OTHER.", "The actor provides a required explanatory comment.", "The flow returns to validation.")]),
        "exc": "\n\n".join([flow("E1 – Duplicate pending report", "The System finds the same actor's PENDING report for the target.", "The System rejects the duplicate.", "The existing report remains pending."), flow("E2 – Invalid reason or comment", "The reason is unsupported or OTHER lacks a comment.", "The System rejects the request.", "No report is created."), flow("E3 – Target is ineligible or invisible", "The Job is not ACTIVE, the CV is BANNED, or ownership/visibility fails.", "The System rejects the request.", "No report is created.")]),
        "post": lines("Success Postconditions", "S1. A valid report exists in PENDING status.", "S2. It is available in the Administrator queue.", "S3. Submission alone does not ban the target.", "", "Minimal Guarantees", "F1. Invalid, unauthorized, or duplicate requests create no extra pending report.", "F2. Submission does not change the target moderation status."),
        "related": lines("UC-02 – Explore Jobs", "UC-09 – Review and Process Applicants", "UC-10 – Manage Talent Pool and Invitations", "UC-13 – Review and Resolve Content Reports"), "priority": "Medium",
    },
    {
        "name": "Manage Employer Profile and Job Postings", "caption": "Manage Employer Profile and Job Postings", "primary": "Recruiter", "secondary": "None",
        "description": "This use case allows a Recruiter to maintain employer information and manage the lifecycle of owned Job postings.",
        "pre": lines("P1. The Recruiter account is active.", "P2. The Recruiter is authenticated.", "P3. Existing-Job operations apply only to Jobs owned by the Recruiter."),
        "trigger": "The Recruiter opens the employer or Job management workspace and selects an action.",
        "main": lines("1. The Recruiter opens Job management.", "2. The System displays Employer Profile status and owned Jobs.", "3. The Recruiter selects create and publish.", "4. The System verifies that an Employer Profile exists and displays the form.", "5. The Recruiter enters Job information.", "6. The Recruiter submits the Job for publication.", "7. The System validates Job information and quality requirements.", "8. The System creates the Job in ACTIVE status.", "9. The System displays the published Job."),
        "alt": "\n\n".join([
            flow("A1 – Create or update Employer Profile (from Step 2)", "The Recruiter opens Employer Profile.", "The System displays current or empty information.", "The Recruiter enters changes.", "The System validates, saves, and displays the profile."),
            flow("A2 – Save as draft (from Step 6)", "The Recruiter saves without publishing.", "The System stores available information in DRAFT status.", "The System displays the draft."),
            flow("A3 – Publish an existing draft (alternative from Step 3)", "The Recruiter selects an owned DRAFT Job and submits it for publication.", "The System validates complete information and quality.", "The System changes the Job to ACTIVE and displays it."),
            flow("A4 – Update a Job or supported status", "The Recruiter selects an owned Job and changes information or an allowed status.", "The System validates and saves the change.", "The System displays the updated Job."),
            flow("A5 – Delete an eligible Job", "The Recruiter selects deletion.", "The System verifies no Applications and no pending reports.", "The Recruiter confirms.", "The System removes the Job."),
            flow("A6 – Export Job information", "The Recruiter requests export.", "The System prepares permitted Job information as CSV.", "The System provides the export."),
        ]),
        "exc": "\n\n".join([flow("E1 – Employer Profile is missing (from Step 4)", "The System does not create the Job.", "The System directs the Recruiter to A1."), flow("E2 – Job validation fails", "The System displays missing, invalid, or blocking quality issues.", "Creation or publication is rejected.", "No invalid Job becomes ACTIVE."), flow("E3 – Job is unavailable or unauthorized", "The Job is missing or belongs to another Recruiter.", "The System denies the operation.", "The Job remains unchanged."), flow("E4 – Deletion is not allowed", "The Job has Applications or pending reports.", "The System rejects deletion.", "The Job remains available."), flow("E5 – Moderation controls the status", "The Job is HIDDEN_BY_ADMIN or BANNED, or a protected status is requested.", "The System rejects the change.", "The status remains unchanged.")]),
        "post": lines("Success Postconditions", "S1. Valid Employer Profile changes are stored.", "S2. A saved draft remains DRAFT.", "S3. A published Job is ACTIVE and Candidate-visible where applicable.", "S4. Valid changes appear in Job management.", "S5. An eligible deleted Job is unavailable.", "", "Minimal Guarantees", "F1. A Recruiter cannot modify another Recruiter's Job.", "F2. Invalid Job information is not published.", "F3. A Job with Applications or pending reports is not deleted."),
        "related": lines("UC-02 – Explore Jobs", "UC-09 – Review and Process Applicants", "UC-10 – Manage Talent Pool and Invitations"), "priority": "High",
    },
    {
        "name": "Review and Process Applicants", "caption": "Review and Process Applicants", "primary": "Recruiter", "secondary": "None",
        "description": "This use case allows a Recruiter to review Applications for an owned Job and record a supported recruitment decision.",
        "pre": lines("P1. The Recruiter account is active.", "P2. The Recruiter is authenticated.", "P3. The selected Job belongs to the Recruiter."),
        "trigger": "The Recruiter opens the applicant workspace for an owned Job.",
        "main": lines("1. The Recruiter selects an owned Job.", "2. The System displays its Applications.", "3. The Recruiter applies supported filters.", "4. The System displays matching applicants.", "5. The Recruiter selects an applicant.", "6. The System displays available Candidate, CV, Portfolio, Application, Matching, and Potential information.", "7. The Recruiter reviews the information.", "8. The Recruiter selects Approve or Reject.", "9. The System verifies ownership and the requested status.", "10. The System updates the Application.", "11. The System displays the result."),
        "alt": "\n\n".join([flow("A1 – Review without a decision (from Step 7)", "The Recruiter closes the detail without a decision.", "The Application remains unchanged."), flow("A2 – No applicants match (from Step 4)", "The System displays an empty applicant state.", "The Recruiter may change filters or select another Job.")]),
        "exc": "\n\n".join([flow("E1 – Job or Application is unavailable or unauthorized", "The System cannot find the record or ownership fails.", "The System denies access or rejects the decision.", "No status changes."), flow("E2 – Associated CV is banned (from Step 9)", "The Application is associated with a CV in BANNED status.", "The System rejects the decision.", "The Application status remains unchanged."), flow("E3 – Unsupported Application status (from Step 9)", "The requested status is unsupported.", "The System rejects the update.", "The existing status remains unchanged.")]),
        "post": lines("Success Postconditions", "S1. The Application reflects the supported Recruiter decision.", "S2. The updated result appears in the applicant workspace.", "S3. The Candidate can later observe the status.", "", "Minimal Guarantees", "F1. Review without a decision changes nothing.", "F2. A Recruiter cannot process another Recruiter's Application.", "F3. A rejected decision does not change status."),
        "related": lines("UC-03 – Manage Job Applications", "UC-07 – Report Suspicious Recruitment Content", "UC-08 – Manage Employer Profile and Job Postings", "UC-10 – Manage Talent Pool and Invitations"), "priority": "High",
    },
    {
        "name": "Manage Talent Pool and Invitations", "caption": "Manage Talent Pool and Invitations", "primary": "Recruiter", "secondary": "None",
        "description": "This use case allows a Recruiter to discover suitable Candidates for an owned Job, maintain Job-specific CV bookmarks, and manage Recruiter invitations.",
        "pre": lines("P1. The Recruiter account is active.", "P2. The Recruiter is authenticated.", "P3. The selected Job belongs to the Recruiter."),
        "trigger": "The Recruiter opens the Talent Pool for an owned Job.",
        "main": lines("1. The Recruiter selects an owned Job.", "2. The System displays Candidates with available Matchings for that Job.", "3. The System separates High and Potential groups.", "4. The Recruiter filters and reviews summaries.", "5. The Recruiter selects a Candidate.", "6. The System displays available CV, score, label, and explanation.", "7. The Recruiter selects Invite.", "8. The System verifies the ACTIVE Job and Candidate default CV.", "9. The System creates an Application in INVITED status.", "10. The System displays the Candidate in the invited group."),
        "alt": "\n\n".join([flow("A1 – Add or remove a bookmark (from Step 6)", "The Recruiter selects the bookmark action.", "The System adds or removes the Job-specific bookmark.", "The System updates the bookmark view."), flow("A2 – Withdraw a pending invitation (from Step 10)", "The Recruiter opens invited Candidates and selects an INVITED Application.", "The Recruiter confirms withdrawal.", "The System removes the invitation Application and updates the view."), flow("A3 – An Application already exists (from Step 8)", "The System finds an existing Candidate–Job Application.", "The System returns it instead of creating a duplicate.", "The System displays the existing state."), flow("A4 – Talent Pool is empty (from Step 2)", "The System finds no available Candidate.", "The System displays an empty state.", "No bookmark or invitation is created.")]),
        "exc": "\n\n".join([flow("E1 – Job is unavailable or unauthorized", "The Job is missing or ownership fails.", "The System denies the operation.", "No invitation is created."), flow("E2 – Job is not ACTIVE (from Step 8)", "The System rejects the invitation.", "No invitation is created."), flow("E3 – Candidate has no default CV (from Step 8)", "The System rejects the invitation.", "No invitation is created."), flow("E4 – Invitation cannot be withdrawn", "The Application is missing, unauthorized, or no longer INVITED.", "The System rejects withdrawal.", "The Application remains unchanged.")]),
        "post": lines("Success Postconditions", "S1. Job-specific bookmarks reflect the Recruiter's selections.", "S2. A new invitation is INVITED.", "S3. A withdrawn pending invitation is removed.", "S4. An existing Application is not duplicated.", "", "Minimal Guarantees", "F1. A Recruiter cannot manage another Recruiter's Talent Pool.", "F2. No invitation is created for a non-ACTIVE Job.", "F3. Invalid withdrawal changes nothing."),
        "related": lines("UC-03 – Manage Job Applications", "UC-07 – Report Suspicious Recruitment Content", "UC-08 – Manage Employer Profile and Job Postings", "UC-09 – Review and Process Applicants"), "priority": "Medium",
    },
    {
        "name": "Review Recruitment Analytics", "caption": "Review Recruitment Analytics", "primary": "Recruiter", "secondary": "None",
        "description": "This use case allows a Recruiter to review analytical information about owned Jobs, Applications, recruitment trends, and skill gaps.",
        "pre": lines("P1. The Recruiter account is active.", "P2. The Recruiter is authenticated.", "P3. Job-specific analytics apply only to a Job owned by the Recruiter."),
        "trigger": "The Recruiter opens recruitment analytics.",
        "main": lines("1. The Recruiter opens recruitment analytics.", "2. The System displays the Recruiter overview for the selected period.", "3. The System displays available Job and Application statistics.", "4. The Recruiter selects an owned Job.", "5. The System verifies ownership.", "6. The System displays the Job funnel.", "7. The System displays Job-specific skill gaps.", "8. The System displays recruitment trends.", "9. The Recruiter reviews the results."),
        "alt": "\n\n".join([flow("A1 – Review aggregate overview (from Step 4)", "The Recruiter selects no Job.", "The System continues to display Recruiter-level analytics."), flow("A2 – Limited analytics data (from Step 6)", "The Job or period has insufficient activity.", "The System displays available values and limited-data states.", "The Recruiter may select another Job or period.")]),
        "exc": flow("E1 – Job is unavailable or unauthorized (from Step 5)", "The Job is missing or belongs to another Recruiter.", "The System denies Job-specific analytics.", "The Recruiter overview remains available."),
        "post": lines("Success Postconditions", "S1. The Recruiter can review available analytics.", "S2. Limited data does not block other available sections.", "S3. Analytics review changes no Job or Application data.", "", "Minimal Guarantees", "F1. Another Recruiter's Job analytics are not exposed.", "F2. Analytics review does not modify operational records."),
        "related": lines("UC-08 – Manage Employer Profile and Job Postings", "UC-09 – Review and Process Applicants", "UC-10 – Manage Talent Pool and Invitations"), "priority": "Medium",
    },
    {
        "name": "Administer Platform Access and Job Visibility", "caption": "Administer Platform Access and Job Visibility", "primary": "Administrator", "secondary": "None",
        "description": "This use case allows an Administrator to control User-account access and public Job visibility.",
        "pre": lines("P1. The Administrator account is active.", "P2. The Administrator is authenticated with the Administrator role."),
        "trigger": "The Administrator opens User or Job administration and selects an action.",
        "main": lines("1. The Administrator opens User administration.", "2. The System displays searchable and filterable Users.", "3. The Administrator selects a User.", "4. The System displays account information and active state.", "5. The Administrator selects Suspend.", "6. The System verifies that the target is not the current Administrator.", "7. The System sets the account inactive.", "8. The System confirms suspension."),
        "alt": "\n\n".join([flow("A1 – Activate a User (from Step 5)", "The Administrator selects an inactive User.", "The Administrator selects Activate.", "The System sets the account active and confirms."), flow("A2 – Hide an ACTIVE Job (alternative from the Trigger)", "The Administrator opens Job administration and selects an ACTIVE Job.", "The Administrator selects Hide.", "The System changes the status to HIDDEN_BY_ADMIN and confirms."), flow("A3 – Restore a hidden Job (from A2)", "The Administrator selects a HIDDEN_BY_ADMIN Job.", "The Administrator selects Restore.", "The System changes the status to ACTIVE and confirms."), flow("A4 – No records match", "The System displays an empty result state.", "The Administrator may change the filters.")]),
        "exc": "\n\n".join([flow("E1 – Administrator attempts self-suspension (from Step 6)", "The System rejects the request.", "The Administrator remains active."), flow("E2 – Target is unavailable", "The System cannot find the User or Job.", "The System rejects the action.", "No state changes."), flow("E3 – Job is not eligible to be hidden", "The Job is not ACTIVE.", "The System rejects Hide.", "The status remains unchanged."), flow("E4 – Job is not eligible to be restored", "The Job is not HIDDEN_BY_ADMIN.", "The System rejects Restore.", "The status remains unchanged.")]),
        "post": lines("Success Postconditions", "S1. A suspended User is inactive; an activated User is active.", "S2. A hidden Job is HIDDEN_BY_ADMIN and excluded from applicable public views.", "S3. A restored Job is ACTIVE.", "", "Minimal Guarantees", "F1. An Administrator cannot suspend their own account.", "F2. Invalid requests do not change target state.", "F3. Job visibility follows allowed transitions."),
        "related": lines("UC-08 – Manage Employer Profile and Job Postings", "UC-13 – Review and Resolve Content Reports"), "priority": "Medium",
    },
    {
        "name": "Review and Resolve Content Reports", "caption": "Review and Resolve Content Reports", "primary": "Administrator", "secondary": "None",
        "description": "This use case allows an Administrator to review reports grouped by target and resolve them by dismissing reports or banning the reported Job or CV.",
        "pre": lines("P1. The Administrator account is active.", "P2. The Administrator is authenticated with the Administrator role."),
        "trigger": "The Administrator opens the content-report queue.",
        "main": lines("1. The Administrator opens the report queue.", "2. The System groups pending reports by Job or CV target.", "3. The Administrator selects a case.", "4. The System displays target content, reasons, comments, count, and status.", "5. The Administrator reviews the evidence.", "6. The Administrator selects Dismiss.", "7. The System verifies that pending reports remain.", "8. The System marks them DISMISSED and clears the target count.", "9. The System displays the resolved case."),
        "alt": "\n\n".join([flow("A1 – Ban the reported target (from Step 6)", "The Administrator selects Ban and may provide a note.", "The System verifies pending reports.", "The System changes the Job or CV to BANNED.", "For a CV, the System removes the default designation.", "The System marks reports ACTIONED, clears the count, and displays the case."), flow("A2 – Report queue is empty (from Step 2)", "The System finds no pending case.", "The System displays an empty queue.", "No moderation action occurs.")]),
        "exc": "\n\n".join([flow("E1 – Report target is unavailable", "The System cannot find the Job or CV.", "The System cannot resolve the case.", "No decision is recorded."), flow("E2 – No pending reports remain", "The case was already resolved or has no pending report.", "The System rejects repeated resolution.", "Existing states remain unchanged.")]),
        "post": lines("Success Postconditions", "S1. Dismissed reports are DISMISSED.", "S2. Reports for a banned case are ACTIONED.", "S3. A banned target has BANNED status.", "S4. The target has no pending-report count.", "", "Minimal Guarantees", "F1. Review alone does not change target status.", "F2. A case without pending reports is not resolved again.", "F3. Failed resolution does not partially change the case."),
        "related": lines("UC-07 – Report Suspicious Recruitment Content", "UC-12 – Administer Platform Access and Job Visibility"), "priority": "Medium",
    },
    {
        "name": "Respond Through Actionable Email", "caption": "Respond Through Actionable Email", "primary": "Candidate", "secondary": "None",
        "description": "This use case allows a Candidate to confirm and complete a supported CareerFit action from an actionable email. Successful behaviours are Matching Feedback and Recruiter-invitation responses.",
        "pre": lines("P1. The Candidate has received a CareerFit actionable email.", "P2. The email contains an actionable link associated with a CareerFit email action."),
        "trigger": "The Candidate opens the actionable link in the email.",
        "main": lines("1. The Candidate opens the link.", "2. The System identifies the associated email action.", "3. The System verifies that the action is pending and unexpired.", "4. The System displays the action and a confirmation page without changing business data.", "5. The Candidate reviews the action.", "6. The Candidate confirms it.", "7. The System validates the action and referenced record again.", "8. The System performs the supported Candidate action.", "9. The System marks the email action REDEEMED.", "10. The System displays the result."),
        "alt": "\n\n".join([flow("A1 – Submit Matching Feedback (from Step 8)", "The action is GOOD_MATCH, POTENTIAL, BAD_MATCH, or NOT_INTERESTED.", "The System submits the corresponding Feedback.", "The flow continues at Step 9."), flow("A2 – Respond to a Recruiter invitation (from Step 8)", "The action is INVITATION_ACCEPT or INVITATION_DECLINE.", "The System verifies INVITED status.", "Acceptance changes the Application to PENDING; decline changes it to NOT_INTERESTED.", "The flow continues at Step 9."), flow("A3 – Leave without confirming (from Step 5)", "The Candidate leaves the page.", "The System performs no business action.", "The email action remains pending until redemption or expiry.")]),
        "exc": "\n\n".join([flow("E1 – Action link is invalid (from Step 2)", "The System cannot identify an email action.", "The System displays an invalid-link message.", "No business data changes."), flow("E2 – Action was already redeemed (from Step 3 or Step 7)", "The System determines that the email action has already been successfully redeemed.", "The System informs the Candidate that the action was processed previously.", "The System does not execute the action again."), flow("E3 – Action has expired (from Step 3 or Step 7)", "The System determines that the email action has expired.", "The System rejects the action and informs the Candidate.", "If detected when confirmation is submitted, the System records the email action as EXPIRED.", "Referenced Feedback or Application state remains unchanged."), flow("E4 – Referenced Matching is unavailable (from A1)", "The System cannot identify an available Matching.", "The System rejects the action.", "No Feedback is submitted."), flow("E5 – Referenced invitation is unavailable (from A2)", "The Application is missing or no longer INVITED.", "The System rejects the response.", "The Application remains unchanged.")]),
        "post": lines("Success Postconditions", "S1. A supported action is applied to the referenced Matching or Application.", "S2. The email action becomes REDEEMED.", "S3. The redeemed action is not accepted again.", "", "Minimal Guarantees", "F1. Opening confirmation changes no business data.", "F2. Invalid, expired, or redeemed actions do not modify the referenced state.", "F3. An expired confirmed action may become EXPIRED without changing business state.", "F4. A failed action is not presented as successful."),
        "related": lines("UC-03 – Manage Job Applications", "UC-04 – Provide Matching Feedback"), "priority": "Medium",
    },
]


def set_font(run, size=13, bold=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), "Times New Roman")


def set_paragraph(paragraph, text, size=13, bold=False):
    paragraph.clear()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(text)
    set_font(run, size, bold)


def set_cell(cell, text, bold=False):
    cell.text = ""
    set_paragraph(cell.paragraphs[0], text, 13, bold)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = 1


def set_repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    element = trpr.find(qn("w:tblHeader"))
    if element is None:
        element = OxmlElement("w:tblHeader")
        trpr.append(element)
    element.set(qn("w:val"), "true")


def allow_row_split(row):
    trpr = row._tr.get_or_add_trPr()
    for node in list(trpr.findall(qn("w:cantSplit"))):
        trpr.remove(node)


def set_spec_table(table, index, case):
    rows = [
        ("Field", "Description"),
        ("Use Case ID", f"UC-{index:02d}"),
        ("Use Case Name", case["name"]),
        ("Primary Actor(s)", case["primary"]),
        ("Secondary Actor(s)", case["secondary"]),
        ("Description", case["description"]),
        ("Preconditions", case["pre"]),
        ("Trigger", case["trigger"]),
        ("Main Flow", case["main"]),
        ("Alternative Flows", case["alt"]),
        ("Exception Flows", case["exc"]),
        ("Postconditions", case["post"]),
        ("Related Use Cases", case["related"]),
        ("Priority", case["priority"]),
    ]
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    for r, values in zip(table.rows, rows):
        allow_row_split(r)
        set_cell(r.cells[0], values[0], bold=True)
        set_cell(r.cells[1], values[1], bold=(values[0] == "Field"))
    set_repeat_header(table.rows[0])


def replace_exact_or_prefix(doc, prefix, text):
    matches = [p for p in doc.paragraphs if p.text.strip().startswith(prefix) or prefix in p.text]
    if not matches and any(p.text.strip() == text for p in doc.paragraphs):
        return
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting {prefix!r}; found {len(matches)}")
    set_paragraph(matches[0], text, 13)


def set_simple_table(table, rows):
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    for ri, values in enumerate(rows):
        for ci, value in enumerate(values):
            set_cell(table.rows[ri].cells[ci], value, bold=(ri == 0))
        allow_row_split(table.rows[ri])
    set_repeat_header(table.rows[0])


def find_image_before_caption(doc, prefix):
    captions = [i for i, p in enumerate(doc.paragraphs)
                if p.style.name == "Figure Caption" and p.text.strip().startswith(prefix)]
    if len(captions) != 1:
        raise RuntimeError(f"Expected one caption {prefix}; found {len(captions)}")
    for i in range(captions[0] - 1, max(-1, captions[0] - 6), -1):
        if doc.paragraphs[i]._p.xpath(".//a:blip"):
            return doc.paragraphs[i]
    raise RuntimeError(f"No image before {prefix}")


def replace_picture(doc, caption_prefix, image_path, title="CareerFit use-case overview", description="CareerFit overview of the fourteen approved role-oriented use cases."):
    paragraph = find_image_before_caption(doc, caption_prefix)
    blip = paragraph._p.xpath(".//a:blip")[0]
    rid = blip.get(qn("r:embed"))
    doc.part.related_parts[rid]._blob = image_path.read_bytes()
    with Image.open(image_path) as source:
        ratio = source.width / source.height
    width_cm = min(15.0, 9.4 * ratio)
    height_cm = width_cm / ratio
    cx, cy = int(Cm(width_cm)), int(Cm(height_cm))
    for extent in paragraph._p.xpath(".//wp:extent") + paragraph._p.xpath(".//a:xfrm/a:ext"):
        extent.set("cx", str(cx)); extent.set("cy", str(cy))
    for prop in paragraph._p.xpath(".//wp:docPr"):
        prop.set("descr", description)
        prop.set("title", title)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def generate_figure():
    width, height = 2100, 1320
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    regular = ImageFont.truetype(r"C:\Windows\Fonts\arial.ttf", 25)
    bold = ImageFont.truetype(r"C:\Windows\Fonts\arialbd.ttf", 27)
    title = ImageFont.truetype(r"C:\Windows\Fonts\arialbd.ttf", 46)
    draw.text((60, 38), "CareerFit use-case overview", font=title, fill="#111827")
    lanes = [
        (135, 485, "Candidate", "#eef2ff", "#4f46e5", [(1,"Manage Career Profile"),(2,"Explore Jobs"),(3,"Manage Job Applications"),(4,"Provide Matching Feedback"),(5,"Review Personalized Career Insights"),(6,"Manage AutoFit")]),
        (515, 750, "Candidate / Recruiter", "#fff7ed", "#ea580c", [(7,"Report Suspicious Recruitment Content"),(14,"Respond Through Actionable Email")]),
        (780, 1045, "Recruiter", "#ecfeff", "#0891b2", [(8,"Manage Employer Profile and Job Postings"),(9,"Review and Process Applicants"),(10,"Manage Talent Pool and Invitations"),(11,"Review Recruitment Analytics")]),
        (1075, 1270, "Administrator", "#f0fdf4", "#15803d", [(12,"Administer Platform Access and Job Visibility"),(13,"Review and Resolve Content Reports")]),
    ]
    for top,bottom,label,fill,outline,cases in lanes:
        draw.rounded_rectangle((55,top,2045,bottom), radius=22, fill="#fafafa", outline="#cbd5e1", width=3)
        draw.rounded_rectangle((75,top+22,335,bottom-22), radius=18, fill=fill, outline=outline, width=4)
        lane_lines = ["Candidate /", "Recruiter"] if label == "Candidate / Recruiter" else [label]
        lane_font = regular if len(lane_lines) > 1 else bold
        lane_height = len(lane_lines) * 31
        lane_y = (top + bottom - lane_height) / 2
        for lane_line in lane_lines:
            box=draw.textbbox((0,0),lane_line,font=lane_font)
            draw.text((205-(box[2]-box[0])/2,lane_y),lane_line,font=lane_font,fill="#172554")
            lane_y += 31
        cols = 3 if len(cases) >= 3 else 2
        available = 1665
        gap = 18
        cw = (available - gap*(cols-1))//cols
        rows = (len(cases)+cols-1)//cols
        ch = min(105, (bottom-top-58-gap*(rows-1))//rows)
        for idx,(cid,name) in enumerate(cases):
            rr,cc=divmod(idx,cols); x=360+cc*(cw+gap); y=top+28+rr*(ch+gap)
            draw.rounded_rectangle((x,y,x+cw,y+ch),radius=14,fill=fill,outline=outline,width=3)
            draw.rounded_rectangle((x+10,y+12,x+92,y+50),radius=9,fill=outline)
            draw.text((x+20,y+18),f"UC-{cid:02d}",font=regular,fill="white")
            words=name.split(); lines=[]; current=""
            for word in words:
                test=(current+" "+word).strip()
                if draw.textbbox((0,0),test,font=bold)[2] < cw-125: current=test
                else: lines.append(current); current=word
            if current: lines.append(current)
            yy=y+(ch-len(lines)*32)//2
            for line in lines:
                draw.text((x+108,yy),line,font=bold,fill="#172554"); yy+=32
    FIGURE.parent.mkdir(parents=True, exist_ok=True)
    image.save(FIGURE, quality=95)


def generate_notification_figure():
    width, height = 1600, 900
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title = ImageFont.truetype(r"C:\Windows\Fonts\arialbd.ttf", 42)
    bold = ImageFont.truetype(r"C:\Windows\Fonts\arialbd.ttf", 27)
    footer = ImageFont.truetype(r"C:\Windows\Fonts\arial.ttf", 20)
    draw.text((70, 45), "Per-account notification policy guard", font=title, fill="#111827")
    nodes = [
        "Email enabled and policy active?", "Notification category allowed?",
        "Deduplication and interaction checks", "Daily quota and cooldown available?",
        "Outside quiet hours?", "Send, suppress, or defer notification",
        "Audit and delivery log",
    ]
    boxes = []
    for index, label in enumerate(nodes):
        row, col = divmod(index, 4)
        x1, y1 = 70 + col * 385, 170 + row * 270
        x2, y2 = x1 + 330, y1 + 130
        boxes.append((x1, y1, x2, y2))
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill="#eef2ff", outline="#4f46e5", width=4)
        words, lines, current = label.split(), [], ""
        for word in words:
            candidate = (current + " " + word).strip()
            if draw.textbbox((0, 0), candidate, font=bold)[2] <= 300:
                current = candidate
            else:
                lines.append(current); current = word
        if current:
            lines.append(current)
        yy = y1 + (130 - len(lines) * 34) / 2
        for line in lines:
            bounds = draw.textbbox((0, 0), line, font=bold)
            draw.text((x1 + (330 - (bounds[2] - bounds[0])) / 2, yy), line, font=bold, fill="#172554")
            yy += 34
    for index in range(len(boxes) - 1):
        left, right = boxes[index], boxes[index + 1]
        if index // 4 == (index + 1) // 4:
            start, end = (left[2], (left[1] + left[3]) // 2), (right[0], (right[1] + right[3]) // 2)
        else:
            start, end = ((left[0] + left[2]) // 2, left[3]), ((right[0] + right[2]) // 2, right[1])
        draw.line((start, end), fill="#4f46e5", width=5)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 18, ey - 10), (ex - 18, ey + 10)], fill="#4f46e5")
    draw.text((70, 845), "CareerFit IT AutoPilot — thesis diagram", font=footer, fill="#64748b")
    NOTIFICATION_FIGURE.parent.mkdir(parents=True, exist_ok=True)
    image.save(NOTIFICATION_FIGURE, quality=95)


def main():
    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(REPORT, BACKUP)
    generate_figure()
    doc = Document(REPORT)

    headings = [p for p in doc.paragraphs if p.style.name == "Heading 3" and p.text.strip().startswith("1.5.")]
    if len(headings) != 14:
        raise RuntimeError(f"Expected 14 use-case headings; found {len(headings)}")
    captions = [p for p in doc.paragraphs if p.style.name == "Table Caption" and p.text.strip().startswith("Table 1.")][4:18]
    if len(captions) != 14:
        raise RuntimeError(f"Expected 14 use-case captions; found {len(captions)}")
    for index, case in enumerate(CASES, 1):
        set_paragraph(headings[index-1], f"1.5.{index} {case['name']}", 13)
        set_paragraph(captions[index-1], f"Table 1.{index+4}. Use case - {case['caption']}", 11)
        set_spec_table(doc.tables[index+5], index, case)

    set_simple_table(doc.tables[3], [
        ["Actor", "Primary responsibilities", "Access boundary"],
        ["Guest", "Explore public Jobs and view public Job and employer details", "Public Job operations only; authentication is required for Candidate-specific actions"],
        ["Candidate", "Manage Career Profile; explore Jobs; manage Applications; review Matching, recommendations, and analytics; provide Feedback; manage AutoFit; report an ACTIVE Job; use actionable email", "Candidate-owned data and Candidate/report operations"],
        ["Recruiter", "Manage Employer Profile and Jobs; review applicants; manage Talent Pool and invitations; review recruitment analytics; report a visible CV", "Owned Jobs and CVs visible through an owned Job"],
        ["Administrator", "Manage User access and Job visibility; review and resolve reported Job/CV cases", "Administrative operations only"],
    ])
    set_simple_table(doc.tables[4], [
        ["Group", "Required capabilities", "Main actor"],
        ["Authentication and account", "Registration, login, account lookup, role enforcement, and supporting settings", "Candidate, Recruiter, Administrator"],
        ["Job exploration", "Public search, filters, sorting, pagination, urgent Jobs, Job/employer details, and similar Jobs", "Guest, Candidate"],
        ["Career Profile", "Candidate Profile and Portfolio editing; CV draft, upload, review, confirmation, default selection, and eligible deletion", "Candidate"],
        ["Matching, recommendation, and analytics", "Separate CV–Job Matching and personalized Job Recommendation views; Candidate and Recruiter analytics", "Candidate, Recruiter"],
        ["Employer and Job management", "Employer Profile; Job create, draft, publish, update, supported status, delete, and export", "Recruiter"],
        ["Applications and Talent Pool", "Apply, history, withdrawal, invitation response, applicant decisions, Candidate discovery, bookmarks, and invitations", "Candidate, Recruiter"],
        ["Feedback and AutoFit", "Matching Feedback; AutoFit enablement, threshold, pause/resume, manual run, and automatic execution", "Candidate"],
        ["Supporting notification settings", "Email-type preferences, quota, cooldown, quiet hours, digest, and related delivery controls", "Candidate, Recruiter"],
        ["Content reporting and moderation", "Candidate Job report; Recruiter visible-CV report; Administrator queue, dismiss, and ban", "Candidate, Recruiter, Administrator"],
        ["Actionable email", "Confirm-then-execute Matching Feedback and Recruiter invitation response", "Candidate"],
    ])

    # Objective and implementation summaries are aligned with verified execution rules.
    doc.tables[2].cell(4,1).text = "AutoFit threshold, default-CV, active-Job, duplicate, per-run, notification, and audit controls"
    for run in doc.tables[2].cell(4,1).paragraphs[0].runs: set_font(run,13)
    doc.tables[29].cell(6,1).text = "Auto-application and notification policy are implemented through separate execution paths; Application deadline is stored but not enforced by manual or automatic submission"
    doc.tables[29].cell(6,2).text = "Add deadline enforcement and focused tests; keep notification quota, cooldown, and quiet hours separate from auto-application eligibility"
    for cell in (doc.tables[29].cell(6,1),doc.tables[29].cell(6,2)):
        for run in cell.paragraphs[0].runs: set_font(run,13)

    replacements = {
        "The third motivation is controlled adaptation": "The third motivation is controlled adaptation and automation. AutoFit separates a matching score from an application action. Current auto-application checks the Candidate policy, pause state, default SCORING_DONE CV, ACTIVE Job status, score threshold, existing Application, and per-run limit. Notification quota, cooldown, time zone, quiet hours, and email preferences are handled separately for message delivery.",
        "Implement AutoFit policies that convert": "• Implement AutoFit policies that convert eligible Matching results into controlled Applications using enablement, pause state, a processed default CV, ACTIVE Job state, score threshold, duplicate checks, and a bounded per-run limit.",
        "ApplicationService.submit resolves": "ApplicationService.submit resolves the authenticated Candidate, requires an ACTIVE Job, rejects an existing Candidate–Job Application, and uses the requested owned CV or owned default CV. The current service does not enforce applicationDeadline or require a manually selected CV to be SCORING_DONE. If a Matching exists, it is attached as score context. The accepted Application is made available to the Candidate and owning Recruiter.",
        "AutomationScheduler coordinates": "AutomationScheduler coordinates matching recomputation, digests, high-match checks, auto-apply, token cleanup, and reminders using system-configured schedules. Per-item error handling prevents one failed record from stopping a full scan. Scheduled execution is internal CareerFit behavior and is not modeled as an external UML actor.",
        "AutoApplyService.runForPolicy": "AutoApplyService.runForPolicy resolves the policy owner and Candidate, then requires an owned default CV in SCORING_DONE status. It selects Matchings for ACTIVE Jobs at or above the configured threshold, skips an existing Candidate–Job Application, and creates at most three AUTO_APPLIED Applications per run. The current execution does not use applicationDeadline, notification quota, cooldown, quiet hours, or notification preferences as application-eligibility rules.",
        "NotificationPolicyGuard evaluates": "NotificationPolicyGuard evaluates email delivery separately from AutoFit application eligibility. It applies email enablement, email-type preferences, deduplication, daily quota, cooldown, and quiet hours, and records sent, skipped, or failed delivery outcomes. SMTP is used when enabled, while NoOpMailService supports local development without claiming production delivery evidence.",
        "EmailActionService generates": "EmailActionService generates a 32-character token derived from UUID text and stores the related email-action state with the recipient, optional Matching or Application, action type, and expiry. The current successful actionable-email scope covers Matching Feedback and Recruiter invitation responses. Although VIEW_JOB and UNSUBSCRIBE_DIGEST action types exist, the current VIEW_JOB branch does not redirect and the unsubscribe branch does not update the Candidate policy, so they are not documented as successful UC-14 behavior. Tokens remain valid for 72 hours.",
        "EmailActionController exposes": "EmailActionController exposes a public confirmation and redemption flow because possession of the high-entropy token is the credential. GET validates the token and displays a non-mutating confirmation page. POST validates it again, performs a supported action, marks a successful action REDEEMED, and records EXPIRED when an expired action is submitted for confirmation. Invalid or already redeemed actions do not change the referenced business state.",
        "This implementation removes the earlier": "This implementation removes the earlier state-changing GET and raw-token persistence findings. The remaining production concerns are rate limiting, deployment-specific link origin, secret rotation, mail-delivery monitoring, clearer user-facing error messages, and protected audit review.",
        "The core traceability chain is:": "The core traceability chain is: UC-01 to Career Profile, Portfolio, CV review, default, and deletion tests; UC-02 to public Job catalogue, detail, employer, urgent, and similar-Job contracts; UC-03 to Application history, submission, withdrawal, and invitation-response tests; UC-04 to Feedback and learning tests; UC-05 to separate Matching, recommendation, and Candidate analytics contracts; UC-06 to AutomationPolicy and AutoApply tests; UC-07 to Candidate Job and Recruiter visible-CV reporting; UC-08 to Employer Profile and Job lifecycle tests; UC-09 to applicant review and decision tests; UC-10 to Talent Pool, bookmark, invitation, and withdrawal flows; UC-11 to Recruiter analytics; UC-12 to administrative User and Job visibility tests; UC-13 to report moderation tests; and UC-14 to confirmation, Feedback/invitation redemption, replay, and expiry behavior.",
        "4. Configure Candidate AutoFit": "4. Configure Candidate AutoFit enablement and threshold, verify pause/resume, then use Run Now for a controlled check. Review notification quota, cooldown, quiet hours, and email preferences separately as supporting settings.",
        "This chapter defined CareerFit's actors": "This chapter defined CareerFit's actors, functional and non-functional requirements, and fourteen approved role-oriented use cases. The specifications distinguish actor goals from internal processing and keep CV–Job Matching, Job Recommendation, Applications, AutoFit eligibility, notification delivery, reporting, moderation, and actionable email behavior traceable to verified implementation.",
    }
    for prefix,text in replacements.items():
        replace_exact_or_prefix(doc,prefix,text)

    generate_figure()
    replace_picture(doc,"Figure 1.2.",FIGURE)
    figure_caption = [p for p in doc.paragraphs
                      if p.style.name == "Figure Caption" and p.text.strip().startswith("Figure 1.2.")]
    if len(figure_caption) != 1:
        raise RuntimeError(f"Expected one body Figure 1.2 caption; found {len(figure_caption)}")
    set_paragraph(figure_caption[0], "Figure 1.2. CareerFit use-case overview", 11)

    generate_notification_figure()
    replace_picture(
        doc,
        "Figure 3.8.",
        NOTIFICATION_FIGURE,
        title="Per-account notification policy guard",
        description="Notification delivery checks, kept separate from AutoFit application eligibility.",
    )
    notification_caption = [p for p in doc.paragraphs
                            if p.style.name == "Figure Caption" and p.text.strip().startswith("Figure 3.8.")]
    if len(notification_caption) != 1:
        raise RuntimeError(f"Expected one body Figure 3.8 caption; found {len(notification_caption)}")
    set_paragraph(notification_caption[0], "Figure 3.8. Per-account notification policy guard", 11)
    # Keep the cached List of Figures text synchronized while preserving its
    # hyperlink and PAGEREF field; Word will refresh the field again on open.
    for paragraph in doc.paragraphs:
        if "Figure 3.8. Per-account AutoFit policy guard" in paragraph.text:
            text_nodes = paragraph._p.xpath(".//w:t")
            if text_nodes:
                text_nodes[0].text = "Figure 3.8. Per-account notification policy guard"

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields"); settings.append(update)
    update.set(qn("w:val"), "true")
    doc.save(REPORT)
    print(REPORT)
    print(FIGURE)


if __name__ == "__main__":
    main()
