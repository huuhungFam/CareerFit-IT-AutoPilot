from pathlib import Path

from docx import Document


DOCX = Path(__file__).resolve().parents[1] / "Doc" / "CareerFit-Thesis-Report.docx"
EXPECTED = [
    "Figure 1.2. Overall CareerFit use-case diagram",
    "Figure 1.3. Candidate use-case diagram",
    "Figure 1.7. Recruiter use-case diagram",
    "Figure 1.8. Shared Candidate–Recruiter reporting use-case diagram",
    "Figure 1.9. Administrator use-case diagram",
]


doc = Document(DOCX)
paragraphs = doc.paragraphs
errors: list[str] = []

for caption in EXPECTED:
    matches = [i for i, p in enumerate(paragraphs) if p.text.strip() == caption]
    if len(matches) != 1:
        errors.append(f"caption_count={len(matches)} caption={caption}")
        continue
    index = matches[0]
    if paragraphs[index].style.name != "Figure Caption":
        errors.append(f"wrong_caption_style={paragraphs[index].style.name} caption={caption}")
    if index == 0 or "w:drawing" not in paragraphs[index - 1]._p.xml:
        errors.append(f"caption_not_immediately_below_image={caption}")
    elif paragraphs[index - 1].alignment != 1:
        errors.append(f"image_not_centered={caption}")

    if not any(p.text.strip().startswith(caption) for p in paragraphs if p.style.name == "toc 1"):
        errors.append(f"missing_list_of_figures_entry={caption}")

chapter_one = [
    p.text.strip()
    for p in paragraphs
    if p.style.name == "Figure Caption" and p.text.strip().startswith("Figure 1.")
]
expected_chapter_one = [
    "Figure 1.1. CareerFit system context",
    *EXPECTED[:2],
    "Figure 1.4. CV review, confirmation, and matching sequence",
    "Figure 1.5. Feedback learning and recomputation sequence",
    "Figure 1.6. AutoFit decision flow",
    *EXPECTED[2:],
]
if chapter_one != expected_chapter_one:
    errors.append(f"chapter_one_caption_sequence={chapter_one!r}")

empty_headings = [
    (i, p.style.name)
    for i, p in enumerate(paragraphs)
    if p.style.name.startswith("Heading") and not p.text.strip()
]
if empty_headings:
    errors.append(f"empty_headings={empty_headings!r}")

all_text = "\n".join(p.text for p in paragraphs)
for marker in ("Error! Bookmark not defined.", "\ufffd", "\x00"):
    if marker in all_text:
        errors.append(f"invalid_text_marker={marker!r}")

table_caption_errors = []
for table_index, table in enumerate(doc.tables):
    # The first two tables are borderless layout containers on the two cover pages.
    if table_index < 2:
        continue
    table_element = table._tbl
    previous = table_element.getprevious()
    while previous is not None and previous.tag.endswith("}p"):
        text = "".join(previous.itertext()).strip()
        if text:
            break
        previous = previous.getprevious()
    if previous is None or not previous.tag.endswith("}p"):
        table_caption_errors.append(table_index + 1)
        continue
    paragraph = next((p for p in paragraphs if p._p is previous), None)
    if paragraph is None or paragraph.style.name != "Table Caption":
        table_caption_errors.append(table_index + 1)
if table_caption_errors:
    errors.append(f"tables_without_caption_immediately_above={table_caption_errors!r}")

for section_index, section in enumerate(doc.sections):
    values = (
        round(section.page_width.cm, 1),
        round(section.page_height.cm, 1),
        round(section.top_margin.cm, 1),
        round(section.bottom_margin.cm, 1),
        round(section.left_margin.cm, 1),
        round(section.right_margin.cm, 1),
    )
    if values != (21.0, 29.7, 3.0, 3.0, 3.5, 2.0):
        errors.append(f"section_{section_index}_page_or_margins={values!r}")

normal = doc.styles["Normal"]
if normal.font.name != "Times New Roman" or round(normal.font.size.pt, 1) != 13.0:
    errors.append(f"normal_font={normal.font.name}/{normal.font.size.pt}")
if normal.paragraph_format.line_spacing != 1.2:
    errors.append(f"normal_line_spacing={normal.paragraph_format.line_spacing}")

print(f"docx={DOCX}")
print(f"paragraphs={len(paragraphs)} tables={len(doc.tables)} images={len(doc.inline_shapes)}")
print(f"chapter_1_figure_captions={len(chapter_one)}")
print(f"verified_new_usecase_diagrams={len(EXPECTED)}")
print(f"errors={len(errors)}")
for error in errors:
    print(error)
