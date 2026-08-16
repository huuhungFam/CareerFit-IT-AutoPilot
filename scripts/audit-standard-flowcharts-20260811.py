from __future__ import annotations

import hashlib
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-before-20260811-remove-placeholder-diagrams.docx"


def table_digest(document: Document) -> str:
    payload = []
    for table in document.tables:
        payload.append([[cell.text for cell in row.cells] for row in table.rows])
    return hashlib.sha256(repr(payload).encode("utf-8")).hexdigest()


def previous_paragraph_has_drawing(paragraph) -> bool:
    previous = paragraph._p.getprevious()
    while previous is not None and not previous.tag.endswith("}p"):
        previous = previous.getprevious()
    return previous is not None and bool(previous.xpath(".//w:drawing"))


def main() -> int:
    document = Document(DOCX)
    backup = Document(BACKUP)
    failures: list[str] = []

    if table_digest(document) != table_digest(backup):
        failures.append("table contents differ from the pre-edit backup")

    figure_captions = [
        p.text.strip()
        for p in document.paragraphs
        if p.style and p.style.name == "Figure Caption" and p.text.strip().startswith("Figure ")
    ]
    expected_ch1 = [
        "Figure 1.1. CareerFit system context",
        "Figure 1.2. Overall CareerFit use-case diagram",
        "Figure 1.3. Candidate use-case diagram",
        "Figure 1.4. Recruiter use-case diagram",
        "Figure 1.5. Shared Candidate-Recruiter reporting use-case diagram",
        "Figure 1.6. Administrator use-case diagram",
    ]
    expected_ch3 = [
        "Figure 3.1. CareerFit container and component architecture",
        "Figure 3.2. CareerFit core logical data model",
        "Figure 3.3. Local and containerized deployment topology",
        "Figure 3.4. Backend module structure and request path",
        "Figure 3.5. JWT authentication, role authorization, and ownership-check sequence",
        "Figure 3.6. CV ingestion, review, confirmation, and matching flowchart",
        "Figure 3.7. Seed-corpus initialization and TF-IDF construction flowchart",
        "Figure 3.8. CV-Job matching and Potential assessment flowchart",
        "Figure 3.9. Rocchio feedback-learning and recomputation flowchart",
        "Figure 3.10. Application and invitation UML state machine",
        "Figure 3.11. AutoFit eligibility and application decision flowchart",
        "Figure 3.12. Notification policy evaluation and delivery outcome flowchart",
        "Figure 3.13. Actionable-email confirmation and redemption flowchart",
        "Figure 3.14. Frontend request and API-response sequence",
    ]
    expected_ch4 = [
        "Figure 4.1. Baseline and Rocchio benchmark metrics at K = 5",
        "Figure 4.2. Observed local Job-search latency statistics",
    ]
    actual_ch1 = [caption for caption in figure_captions if caption.startswith("Figure 1.")]
    actual_ch2 = [caption for caption in figure_captions if caption.startswith("Figure 2.")]
    actual_ch3 = [caption for caption in figure_captions if caption.startswith("Figure 3.")]
    actual_ch4 = [caption for caption in figure_captions if caption.startswith("Figure 4.")]
    if actual_ch1 != expected_ch1:
        failures.append(f"Chapter 1 captions mismatch: {actual_ch1}")
    if actual_ch3 != expected_ch3:
        failures.append(f"Chapter 3 captions mismatch: {actual_ch3}")
    if actual_ch2:
        failures.append(f"Chapter 2 should not contain figure captions: {actual_ch2}")
    if actual_ch4 != expected_ch4:
        failures.append(f"Chapter 4 captions mismatch: {actual_ch4}")

    old_phrases = (
        "CV review, confirmation, and matching sequence",
        "Feedback learning and recomputation sequence",
        "AutoFit decision flow",
        "Direct score and Potential assessment flow",
        "Feedback processing and post-commit learning",
        "Distinction between matching, recommendation, and recruitment action",
        "TF-IDF vectorization and cosine-similarity pipeline",
        "Rocchio relevance-feedback vector update",
        "Human-in-the-Loop control cycle in CareerFit",
        "Evaluation environments and evidence sources",
        "P0 end-to-end workflow coverage",
        "CareerFit logical entity relationships",
        "Backend module structure and request flow",
        "JWT authentication and authorization boundaries",
        "Application and invitation state transitions",
        "Frontend routes, server-side catalogue, and API data flow",
        "Local Job-search latency distribution",
    )
    all_text = "\n".join(p.text for p in document.paragraphs)
    for phrase in old_phrases:
        if phrase in all_text:
            failures.append(f"obsolete caption text remains: {phrase}")

    for caption in document.paragraphs:
        if caption.style and caption.style.name == "Figure Caption" and caption.text.strip().startswith("Figure "):
            if not previous_paragraph_has_drawing(caption):
                failures.append(f"caption is not immediately preceded by a drawing: {caption.text}")

    use_case_headings = [
        p.text.strip()
        for p in document.paragraphs
        if p.style and p.style.name == "Heading 3" and p.text.strip().startswith("1.5.")
    ]
    if len(use_case_headings) != 14:
        failures.append(f"expected 14 Use Case headings, found {len(use_case_headings)}")

    if len(document.tables) != len(backup.tables):
        failures.append(f"table count changed: {len(backup.tables)} -> {len(document.tables)}")

    drawing_count = len(document._element.body.xpath(".//w:drawing"))
    print(f"physical drawings: {drawing_count}")
    print(f"tables: {len(document.tables)}; Use Case headings: {len(use_case_headings)}")
    toc_figure_entries = [
        p.text.strip().split("\t", 1)[0]
        for p in document.paragraphs
        if p.style and p.style.name.startswith("toc") and p.text.strip().startswith("Figure ")
    ]
    if toc_figure_entries != figure_captions:
        failures.append("List of Figures does not exactly match the current figure captions")

    print(
        f"Chapter figures: Ch1={len(actual_ch1)}, Ch2={len(actual_ch2)}, "
        f"Ch3={len(actual_ch3)}, Ch4={len(actual_ch4)}"
    )

    if failures:
        print("AUDIT FAILED")
        for failure in failures:
            print(f"- {failure}")
        return 1
    print("AUDIT PASSED")
    return 0


if __name__ == "__main__":
    sys.exit(main())
