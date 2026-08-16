from __future__ import annotations

from pathlib import Path
import shutil
import zipfile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-before-20260718-review.docx"
TEMP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-20260718-updated.docx"
SCREEN_DIR = ROOT / "Doc" / "screenshots"


PARAGRAPH_REPLACEMENTS = {
    "A Candidate or Recruiter submits feedback for a matching record. The system validates the actor and matching context, stores one current feedback judgment per actor and matching, and invokes the Rocchio service to rebuild the learned representation from the base vector and feedback history. Affected matching rows are marked needs_recompute. The scheduler or an explicit process rescans those rows and clears the marker after successful scoring. This two-stage design keeps feedback recording transactional while allowing ranking recomputation to run separately.":
        "A Candidate submits feedback for a Matching that belongs to the Candidate's CV. The endpoint rejects another role or an unrelated Matching, stores one current judgment per actor and Matching, and starts Rocchio learning only after the feedback transaction commits. Affected Matching rows are marked needs_recompute. The scheduler then rescans those rows and clears the marker after successful scoring. Recruiter decisions use the separate application and invitation workflow rather than this Candidate feedback endpoint.",
    "For match notifications and digests, the backend creates EmailAction records with random token strings and a 72-hour expiry. The recipient opens the public /api/email-action/redeem endpoint. The controller validates existence, pending status, and expiry; dispatches supported feedback actions; and marks the record redeemed. Reuse is rejected by status, and an expired pending record is changed to expired. A scheduled task expires old pending actions and later purges them.":
        "For match notifications and digests, the backend creates EmailAction records with a high-entropy raw token for the recipient, a persisted SHA-256 token hash, and a 72-hour expiry. The public /api/email-action/redeem endpoint validates the hash, pending state, and expiry. Reuse is rejected, expired pending actions are marked expired, and a scheduled task later removes old records.",
    "CV processing represents long-running work with statuses rather than holding a browser request open until every score completes. Extraction or validation failure records a failure state and reason. Batch scoring isolates individual pair failures so one invalid pair does not necessarily invalidate the entire batch. Scheduled recomputation keeps unsuccessful rows marked for later review instead of falsely marking them complete.":
        "CV processing represents long-running work with statuses rather than holding a browser request open until every score completes. Extraction or validation failure records a failure state and reason. CV and Job matching tasks are submitted through an after-commit executor so background work cannot read uncommitted rows. Batch scoring isolates individual pair failures, and scheduled recomputation keeps unsuccessful rows marked for later review instead of falsely marking them complete.",
    "After CV vectorization, MatchingService.scoreAllJobsForCv runs asynchronously. It obtains eligible active Jobs, calls ScoringService for each pair, and creates or updates the unique matching row. Reasons and potential metadata are serialized as JSON. Database uniqueness on (cv_id, job_id) prevents duplicate pair records, while conflict handling covers concurrent execution. The service updates CV status and audit information and may trigger high-match notification behavior.":
        "After CV vectorization commits, AfterCommitExecutor submits MatchingService.scoreAllJobsForCv by CV identifier to the bounded task executor. The service reloads the persisted CV, obtains eligible active Jobs, calls ScoringService for each pair, and creates or updates the unique Matching row. Reasons and potential metadata are serialized as JSON. Database uniqueness on (cv_id, job_id) prevents duplicate pairs, while conflict handling covers concurrent execution. The same boundary is used when a new or updated Job is scored against existing CVs.",
    "When a Recruiter creates or updates a Job, corresponding services build the Job vector and make the Job available to ranking queries. MatchingQueryService shapes Candidate cards and recruiter-oriented results, including pagination and metadata for empty, filtered, low-match, or tied-score states. Stable ordering is important because rounded scores can be equal.":
        "When a Recruiter creates or updates a Job, the service builds the Job vector and schedules post-commit scoring. MatchingQueryService shapes Candidate cards and recruiter-oriented results, including pagination and metadata for empty, filtered, low-match, or tied-score states. CandidatePortfolioVisibilityService adds portfolio data only when the Candidate enabled showPortfolioAfterApply and has actually applied or moved beyond an invitation-only state; otherwise the response contains an explicit hidden reason. Stable score ordering remains important because rounded scores can be equal.",
    "FeedbackService.submitFeedback resolves the Matching and actor, then upserts feedback by (matching_id, actor_id). Concurrent uniqueness conflicts are resolved by reloading and updating the existing row. Each event stores actor role, feedback type, and source channel and creates an audit entry. NOT_INTERESTED is stored but does not trigger Rocchio because it can represent preference rather than technical relevance. GOOD_MATCH, POTENTIAL, and BAD_MATCH trigger asynchronous Job-vector learning.":
        "FeedbackService.submitFeedback resolves the Matching and authenticated Candidate, verifies that the Candidate owns the associated CV, and then upserts feedback by (matching_id, actor_id). Concurrent uniqueness conflicts are resolved by reloading and updating the existing row. Each event stores actor role, feedback type, and source channel and creates an audit entry. NOT_INTERESTED is stored but does not trigger Rocchio because it can represent preference rather than technical relevance. GOOD_MATCH, POTENTIAL, and BAD_MATCH trigger post-commit Job-vector learning.",
    "Using the immutable base vector and the complete current feedback history makes recomputation idempotent for the same data and parameters. However, the asynchronous call is invoked from the feedback service, so transaction timing and executor behavior require integration tests to ensure that newly committed feedback is visible when learning begins. Job locking prevents simultaneous updates from writing independent learned vectors.":
        "Using the immutable base vector and the complete current feedback history makes recomputation idempotent for the same data and parameters. FeedbackService registers learning after transaction commit, so Rocchio reads the persisted judgment rather than racing the initiating transaction. Job locking prevents simultaneous updates from writing independent learned vectors, while integration tests cover the post-commit executor boundary.",
    "The response DTOs combine application state with selected Candidate, CV, and Matching fields. This lets the Recruiter view skills and score context without returning persistence entities directly. List metadata communicates states such as no match, no filtered results, and ready, allowing the frontend to display a meaningful empty state.":
        "The response DTOs combine application state with selected Candidate, CV, and Matching fields without returning persistence entities directly. Portfolio fields are privacy-gated: a Recruiter receives them only after a qualifying application state and when the Candidate has enabled portfolio visibility after applying. List metadata communicates states such as no match, no filtered results, and ready, allowing the frontend to display a meaningful empty state.",
    "The application enables asynchronous methods and scheduling at startup. AsyncConfig creates a bounded ThreadPoolTaskExecutor using configured core size, maximum size, queue capacity, and thread prefix. CV matching, mail sending, and Rocchio learning can therefore leave the request thread, while persisted statuses allow clients to observe completion.":
        "AsyncConfig creates a bounded ThreadPoolTaskExecutor using configured core size, maximum size, queue capacity, and thread prefix. AfterCommitExecutor registers CV and Job matching work with Spring transaction synchronization and submits it only after commit. Mail sending and Rocchio learning can also leave the request thread, while persisted CV and Matching states allow clients to observe completion.",
    "NotificationPolicyGuard evaluates whether a message is allowed and writes delivery outcomes in independent transactions. It supports deduplication, timing, quota, cooldown, and skipped/failed reasons. NotificationEmailService creates HTML content for lifecycle events. MailService sends asynchronously through SMTP when mail is enabled, while NoOpMailService supports local development without external delivery.":
        "NotificationPolicyGuard evaluates whether a message is allowed and writes delivery outcomes in independent transactions. It supports deduplication, timing, quota, cooldown, and skipped/failed reasons. After CV scoring, the backend can immediately send a high-match action email, a low-match notice, or a no-match notice; the scheduled scan remains a later safety path and deduplication prevents duplicate delivery. MailService sends asynchronously through SMTP when mail is enabled, while NoOpMailService supports local development without external delivery.",
    "JPA entities map the domain tables, while Spring Data repositories provide derived queries, pagination, locking queries, and explicit update/delete operations. Service methods use @Transactional for state changes and readOnly=true for queries where applicable. Flyway migrations V1–V14 create and harden the current schema; Hibernate runs with ddl-auto=validate. This prevents application startup from silently inventing schema changes outside migration history.":
        "JPA entities map the domain tables, while Spring Data repositories provide derived queries, pagination, locking queries, and explicit update/delete operations. Service methods use @Transactional for state changes and readOnly=true for queries where applicable. Flyway migrations V1–V15 create and harden the current schema, including hashed email-action tokens; Hibernate runs with ddl-auto=validate. This prevents application startup from silently inventing schema changes outside migration history.",
    "App.tsx defines public, Candidate, Recruiter, and Administrator routes. protectedRoute checks the locally loaded account role and redirects or blocks an incompatible workspace. This improves navigation but is not a security boundary; the backend remains authoritative. Public routes include the home page, Jobs, Job details, and employer details. Candidate routes cover upload, profile, recommendations, applications, analytics, automation, and settings. Recruiter routes cover dashboard, Job workspace, ranking/applicants/potential views, analytics, automation, and settings. Administrator routes cover dashboard, users, Jobs, audit logs, and email monitoring.":
        "App.tsx defines public, Candidate, Recruiter, and Administrator routes. protectedRoute checks the loaded account role, while reload restoration calls /api/auth/me and clears an invalid session. Passwordless login has a dedicated magic-link verification route. Public routes include Jobs and employer details; Candidate routes cover asynchronous CV processing, CV management, recommendations, applications, analytics, automation, and settings; Recruiter and Administrator routes provide their role-specific workspaces. These checks improve navigation, but the backend remains the security boundary.",
    "Frontend/src/lib/api.ts centralizes HTTP access. It reads VITE_API_BASE_URL with /api as the default, attaches the bearer token from local storage, selects JSON headers except for multipart FormData, unwraps the shared response envelope, and throws a normalized error for unsuccessful responses. DTO mapping functions convert backend shapes and labels into UI models. React Query hooks manage loading, refetching, caching, and error state for server data.":
        "Frontend/src/lib/api.ts centralizes HTTP access. It reads VITE_API_BASE_URL with /api as the default, attaches the bearer token from sessionStorage, selects JSON headers except for multipart FormData, unwraps the shared response envelope, and throws a normalized error for unsuccessful responses. DTO mapping functions convert backend shapes and labels into UI models. React Query hooks manage loading, refetching, caching, polling, and error state for server data.",
    "Static mock data remains imported for mapper fallback values in parts of api.ts. API-driven pages are intended to show loading, error, or empty states when the backend fails, but mapper fallback can still mask missing DTO fields with sample presentation values. This behavior should be tested and reduced for thesis evidence so that screenshots do not imply backend data that was not returned.":
        "API-driven pages no longer substitute mock Job records when requests fail; they expose loading, retryable error, or empty states. The July 18 integration pass connected magic-link completion, session revalidation, CV status polling and management, the dedicated recommendation feed, and the market dashboard to backend contracts. Remaining UI gaps are limited to selected recruiter drill-down, employer self-service, automation pause/resume, telemetry, and administrative maintenance operations.",
    "This chapter explained how the design is implemented in Spring Boot, React, PostgreSQL, and Flyway. It covered security, CV processing, TF-IDF scoring, Rocchio feedback, application workflows, AutoFit, email actions, audit records, frontend integration, and runtime monitoring. Chapter 5 evaluates these parts using fresh test and runtime evidence.":
        "This chapter explained how the design is implemented in Spring Boot, React, PostgreSQL, and Flyway. It covered security, post-commit CV and Job processing, TF-IDF scoring, Candidate feedback with Rocchio, privacy-gated portfolios, application workflows, AutoFit notifications, hardened email actions, frontend integration, and runtime monitoring. Chapter 5 evaluates these parts using refreshed test and runtime evidence.",
    "All results in this chapter were refreshed on July 3, 2026 (ICT). The observed Git commit was e92e3d847992d7b628bcfbbcef9a57ab32677547, but the worktree already contained modifications and untracked files. Therefore, the results describe the evaluated working tree and are not reproducible from the commit hash alone unless the complete working-tree changes are also preserved.":
        "Backend, algorithm, frontend build, and Chromium test results were refreshed on July 18, 2026 (ICT). The observed Git HEAD was 65318fb0e0978574c9a04d9e54aecca5ba1eb241, but the worktree contained additional modifications and untracked files. Therefore, the results describe the evaluated working tree and are not reproducible from the commit hash alone unless the complete working-tree state is preserved.",
    "The detailed commands and raw-result pointers are recorded in evidence/CHAPTER5_EVIDENCE_20260703.md. Surefire XML reports, Playwright output, backend runtime logs, benchmark logs, and evaluation/result.json provide the underlying artifacts.":
        "The original command catalog is recorded in evidence/CHAPTER5_EVIDENCE_20260703.md. For the July 18 refresh, current Surefire XML reports, Playwright output, the production-build output, backend runtime logs, and evaluation/result.json provide the underlying artifacts.",
    "The complete backend suite was executed through the Maven Wrapper using .\\mvnw.cmd test. It compiled 125 application source files and 17 test source files, started Testcontainers, validated and applied 14 Flyway migrations, and executed 15 Surefire test suites.":
        "The complete backend suite was executed through the Maven Wrapper using .\\mvnw.cmd test. It compiled 127 application source files and 22 test source files, started Testcontainers, validated and applied 15 Flyway migrations, and executed 20 Surefire test suites.",
    "The suites covered application context startup, API contracts, auto-apply, candidate profile, Job service, matching batch behavior, PDF/document extraction, quality validation, Rocchio, scoring, settings, TF-IDF, production configuration, and security hardening. AlgorithmEvaluatorTest was also included in the complete suite.":
        "The suites covered application context startup, API contracts, post-commit execution, authentication, auto-apply, Candidate profile and CV ingestion, feedback authorization, Job service, matching batches, document extraction, quality validation, Rocchio, scoring, settings, TF-IDF, production configuration, and security hardening. AlgorithmEvaluatorTest was included in the same complete suite.",
    "All 63 registered JUnit tests passed in the final run. Flyway 9.22.3 still warned that PostgreSQL 16.14 is newer than the highest version it reports as tested, and some negative tests intentionally logged handled exceptions. The final benchmark log contained no StaleObjectStateException or build failure.":
        "All 72 registered JUnit tests passed in the July 18 run, with no failures, errors, or skips. Flyway 9.22.3 still warned that PostgreSQL 16.14 is newer than the highest version it reports as tested, and some negative tests intentionally logged handled exceptions. The final benchmark log contained no StaleObjectStateException or build failure.",
    "The increase across position-sensitive metrics shows that the controlled positive feedback moved relevant holdout CVs toward the top of the ranking. Precision@5 remains 0.168 after feedback, which means the top-five lists are not composed entirely of labeled relevant items. Recall@5 and HitRate@5 reach 0.84 in the designed scenario, while 16 percent of query cases still do not place the expected holdout item within the first five results.":
        "The increase across position-sensitive metrics shows that the controlled positive feedback moved relevant holdout CVs toward the top of the ranking. Precision@5 reaches 0.172 after feedback, so the top-five lists are not composed entirely of labeled relevant items. Recall@5 and HitRate@5 reach 0.86 in the designed scenario, while 14 percent of query cases still do not place the expected holdout item within the first five results.",
    "Three additional executions used the same dataset without modifying data or labels. All returned exit code zero and exactly the same dataset hash, baseline nDCG@5 (0.037737056145), Rocchio nDCG@5 (0.817737056145), and delta (0.78). Their wall times were 62.67, 66.80, and 68.66 seconds. Metric equality supports deterministic benchmark output for the observed code and dataset.":
        "The July 18 full-suite run used dataset hash 6e935639ba6d3290dca8ad91a35d714c5e30c7e69a59af23ddbcf89fcc5cc2f2. It produced baseline nDCG@5 of 0.037737056145, Rocchio nDCG@5 of 0.837737056145, and a delta of 0.80. AlgorithmEvaluatorTest completed in 212.2 seconds inside the full suite. These values describe the current code and dataset; they replace the earlier July 3 metric snapshot.",
    "This result creates two separate conclusions. The ranking metrics are reproducible, because the recorded dataset and metric values are identical across runs. The benchmark runtime is not operationally clean, because a background persistence conflict occurs in every observed repetition. Reporting only the green test status would hide this reliability problem.":
        "Earlier benchmark repetitions exposed a background persistence conflict even when foreground assertions passed. The current run no longer logged that exception because feedback learning and matching begin after the initiating transaction commits and the benchmark clears prior Matching state before recomputation. This distinction remains important: test assertions and background operational cleanliness must both be checked.",
    "The required correction is to isolate scheduling from algorithm tests or control it explicitly, ensure asynchronous work starts after the initiating transaction commits, and define retry or conflict handling for optimistic locking. A future acceptance criterion should fail the evaluation when uncaught background exceptions appear, even if all foreground JUnit assertions pass.":
        "The implemented correction controls asynchronous transaction timing and checks the final logs for background failures. Production work should still define retry, conflict, timeout, and recovery policies, and CI should continue to fail when uncaught background exceptions appear even if foreground JUnit assertions pass.",
    "npm run build performs TypeScript checking with no output emission and then builds the production assets with Vite. The command completed successfully and transformed 2,475 modules.":
        "npm run build performs TypeScript checking with no output emission and then builds the production assets with Vite 6.4.3. The July 18 command completed successfully and transformed 2,417 modules.",
    "Manual Rollup chunking separated React, query, chart, and icon dependencies. The largest generated JavaScript chunk was 375.64 kB, below Vite's 500 kB warning threshold. No browser performance profile was collected, so bundle size is reported as build evidence rather than a direct page-load measurement.":
        "Manual Rollup chunking separated React, query, chart, and icon dependencies. The largest generated JavaScript chunk was 378.44 kB, below Vite's 500 kB warning threshold. No browser performance profile was collected, so bundle size is reported as build evidence rather than a direct page-load measurement.",
    "The Playwright Chromium project ran against the Vite frontend, host backend, and Compose PostgreSQL. Four P0 scenarios passed in 33.9 seconds.":
        "The Playwright Chromium project ran against the Vite frontend, host backend, and Compose PostgreSQL. All 20 tests across the P0 workflow, backend-contract, and resilience suites passed in 34.1 seconds.",
    "The result demonstrates integration for the four scripted paths in Chromium. It is not a full UAT study: no independent participants performed tasks, no usability scale or task time was collected, and seeded credentials were used. Firefox and WebKit projects were not run in this evidence snapshot. The Recruiter scenario creates a timestamped test Job but does not automatically remove it; repeated E2E therefore changes local database contents and requires an explicit cleanup policy.":
        "The result covers the four original role workflows together with passwordless request, settings persistence, recommendations, role-route rendering, CV polling, magic-link completion, session restoration, retryable errors, feedback contracts, employer details, similar Jobs, and desktop navigation. It is not a full UAT study: no independent participants performed tasks, seeded credentials were used, and Firefox and WebKit were not run. The Recruiter flow now removes the Job it creates, although the local database still contains earlier test and imported records.",
    "The benchmark runs inside a Spring context, so asynchronous work and persisted matching records can affect repeatability if they are not controlled. The final test registers learning after commit, clears previous matching state before recomputation, and checks the logs for background failures. The browser tests also avoid relying only on the first seeded job and delete the job created by the recruiter flow.":
        "The benchmark runs inside a Spring context, so asynchronous work and persisted Matching records can affect repeatability if they are not controlled. The final test registers learning after commit, clears previous Matching state before recomputation, and checks logs for background failures. Browser tests use explicit API and UI conditions, avoid relying only on the first seeded Job, and delete the Job created by the Recruiter flow.",
    "The benchmark dataset hash and generated JSON support algorithm reproduction. Maven Wrapper, Flyway, Testcontainers, package lock, and recorded commands support environment reconstruction. Reproducibility is weakened by the dirty working tree, mutable local database, timestamped E2E records, external web assets, and missing automatic cleanup. The final thesis release should be associated with a clean commit or archive and immutable evidence bundle.":
        "The benchmark dataset hash and generated JSON support algorithm reproduction. Maven Wrapper, Flyway, Testcontainers, package lock, and recorded commands support environment reconstruction. Recruiter E2E cleanup is now automatic, but reproducibility is still weakened by the dirty working tree, mutable local database, imported records, and external web assets. The final thesis release should be associated with a clean commit or archive and immutable evidence bundle.",
    "The final evaluation passed all 63 backend tests, completed the frontend production build, and passed four Chromium P0 workflows. The controlled Rocchio benchmark improved the holdout ranking deterministically on the synthetic dataset, and its final log contained no optimistic-lock exception. Aggregate health returned HTTP 200 with status UP. These results support a functioning academic prototype, but not production readiness or proven effectiveness on real recruitment data.":
        "The July 18 evaluation passed all 72 backend tests, completed the frontend production build, and passed all 20 Chromium workflow, contract, and resilience tests. The controlled Rocchio benchmark improved the holdout ranking on the synthetic dataset, and its final log contained no optimistic-lock exception. Aggregate health returned HTTP 200 with status UP. These results support a functioning academic prototype, but not production readiness or proven effectiveness on real recruitment data.",
    "The fresh evaluation provides evidence for several technical outcomes. The complete backend suite executed 63 registered tests across 15 suites with no reported JUnit failures, errors, or skips. The frontend passed TypeScript checking and Vite production compilation. Four Chromium P0 workflows passed end to end: public Job search and detail, Candidate application and withdrawal, Recruiter JD creation, and Administrator user suspension/reactivation. Selected authorization checks returned the expected public, unauthenticated, authenticated, and role-denied HTTP statuses.":
        "The refreshed evaluation provides evidence for several technical outcomes. The complete backend suite executed 72 registered tests across 20 suites with no failures, errors, or skips. The frontend passed TypeScript checking and Vite production compilation. All 20 Chromium tests passed across role workflows, backend contracts, and resilience checks, including public search, Candidate application, Recruiter JD creation, Administrator operations, passwordless flow, CV polling, session restoration, and error states.",
    "The controlled Rocchio benchmark demonstrated the intended causal behavior on its synthetic dataset. With 50 Jobs, 100 unique CVs, 300 training pairs, and 300 holdout pairs, nDCG@5 increased from 0.037737 to 0.817737, Recall@5 and HitRate@5 increased from 0.06 to 0.84, and MRR increased from 0.058755 to 0.823617. Three repeated runs returned the same dataset hash and the same observed metric values. These results demonstrate deterministic adaptation in the constructed latent-skill scenario; they do not estimate effectiveness on organic recruitment data.":
        "The controlled Rocchio benchmark demonstrated the intended causal behavior on its synthetic dataset. With 50 Jobs, 100 unique CVs, 300 training pairs, and 300 holdout pairs, nDCG@5 increased from 0.037737 to 0.837737, Recall@5 and HitRate@5 increased from 0.06 to 0.86, and MRR increased from 0.058755 to 0.842665. These results demonstrate adaptation in the constructed latent-skill scenario; they do not estimate effectiveness on organic recruitment data.",
    "The final remediation pass removed the benchmark concurrency exception, restored aggregate Actuator health to UP, split the frontend production bundle so the largest generated chunk is below 500 kB, removed synthetic Job fallbacks, and made the recruiter E2E flow delete the Job it creates. These results support a defensible local demonstration and thesis evaluation; they do not establish production readiness.":
        "The July 18 refresh retained the clean benchmark and aggregate health results, kept the largest frontend chunk below 500 kB, removed API-data fallbacks, expanded Chromium coverage to 20 tests, and refreshed the six interface screenshots. It also verified post-commit matching, Candidate-only feedback authorization, immediate match-result notifications, and portfolio privacy gating. These results support a defensible local demonstration; they do not establish production readiness.",
    "The controlled benchmark confirms the intended Rocchio behavior in the planted latent-skill scenario. The final run completed without the earlier optimistic-lock exception after learning was moved to an after-commit boundary and benchmark state was cleared before recomputation. This supports deterministic controlled behavior, although it does not prove reliability under production concurrency.":
        "The controlled benchmark confirms the intended Rocchio behavior in the planted latent-skill scenario. The final run completed without the earlier optimistic-lock exception after learning was moved to an after-commit boundary and benchmark state was cleared before recomputation. nDCG@5 and HitRate@5 each increased by 0.80, although this does not prove reliability under production concurrency.",
    "The result does not show that every real recruiter feedback event improves ranking by 0.78 nDCG@5. Organic feedback may be inconsistent, sparse, biased, delayed, or based on salary and location rather than technical relevance. Production evaluation would require independently labeled data, time-based splits, multiple recruiters, disagreement analysis, and online or offline comparison against a baseline without feedback.":
        "The result does not show that every real Candidate feedback event improves ranking by 0.80 nDCG@5. Organic feedback may be inconsistent, sparse, biased, delayed, or based on salary and location rather than technical relevance. Production evaluation would require independently labeled data, time-based splits, multiple reviewers, disagreement analysis, and online or offline comparison against a baseline without feedback.",
    "The 63 backend tests do not prove exhaustive path coverage. Browser evaluation covers four P0 cases in Chromium only, and no independent users participated. Security checks were targeted API observations rather than penetration testing, while the latency sample used 30 sequential warm requests on one workstation. Aggregate health passed in the final run, but concurrency, capacity, cross-browser behavior, and real-user usability remain unevaluated.":
        "The 72 backend tests do not prove exhaustive path coverage. Browser evaluation covers 20 automated cases in Chromium only, and no independent users participated. Security checks were targeted API observations rather than penetration testing, while the latency sample used 30 sequential warm requests on one workstation. Aggregate health passed in the final run, but concurrency, capacity, cross-browser behavior, and real-user usability remain unevaluated.",
}


def replace_paragraphs(doc: Document) -> None:
    remaining = dict(PARAGRAPH_REPLACEMENTS)
    for paragraph in doc.paragraphs:
        old = paragraph.text
        if old in remaining:
            paragraph.text = remaining.pop(old)
    if remaining:
        missing = "\n".join(text[:120] for text in remaining)
        raise RuntimeError(f"Paragraph replacements not found:\n{missing}")


def set_cell(table, row: int, col: int, text: str) -> None:
    cell = table.rows[row].cells[col]
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
            run.font.size = Pt(11)


def update_tables(doc: Document) -> None:
    # Functional requirements: the current web feedback endpoint is Candidate-only.
    set_cell(doc.tables[5], 7, 2, "Candidate")

    # Evaluation environment.
    set_cell(doc.tables[16], 6, 1, "Vite 6.4.3, React 18.3, TypeScript 5.9")

    # Backend suite.
    values = {
        1: "20",
        2: "72",
        3: "0",
        4: "0",
        5: "0",
        6: "244.659 s",
        7: "Approximately 266 s",
    }
    for row, value in values.items():
        set_cell(doc.tables[17], row, 1, value)

    # Controlled benchmark.
    metrics = [
        ("Precision@5", "0.012000", "0.172000", "+0.160000"),
        ("Recall@5", "0.060000", "0.860000", "+0.800000"),
        ("nDCG@3", "0.030000", "0.830000", "+0.800000"),
        ("nDCG@5", "0.037737", "0.837737", "+0.800000"),
        ("nDCG@10", "0.050424", "0.850424", "+0.800000"),
        ("MRR", "0.058755", "0.842665", "+0.783910"),
        ("HitRate@5", "0.060000", "0.860000", "+0.800000"),
        ("Coverage", "1.000000", "1.000000", "0"),
    ]
    for row, values in enumerate(metrics, start=1):
        for col, value in enumerate(values):
            set_cell(doc.tables[18], row, col, value)

    # Frontend build.
    artifacts = [
        ("index.html", "1.37 kB", "0.67 kB"),
        ("Main CSS", "65.09 kB", "12.89 kB"),
        ("Icons chunk", "19.25 kB", "4.48 kB"),
        ("Query chunk", "39.60 kB", "12.18 kB"),
        ("React chunk", "180.98 kB", "59.54 kB"),
        ("Application chunk", "217.39 kB", "58.05 kB"),
        ("Charts chunk", "378.44 kB", "112.04 kB"),
    ]
    for row, values in enumerate(artifacts, start=1):
        for col, value in enumerate(values):
            set_cell(doc.tables[19], row, col, value)

    # Add grouped rows for the expanded Chromium suite.
    e2e_table = doc.tables[20]
    if len(e2e_table.rows) == 5:
        row = e2e_table.add_row()
        row.cells[0].text = "Candidate and authentication contracts"
        row.cells[1].text = "Passwordless request, settings persistence, recommendations, CV polling, magic-link completion, and session restore"
        row.cells[2].text = "Passed"
        row = e2e_table.add_row()
        row.cells[0].text = "Resilience and role navigation"
        row.cells[1].text = "Retryable API errors, feedback contract, employer/similar-job data, desktop header, and role-route rendering"
        row.cells[2].text = "Passed"
        for row_index in (5, 6):
            for col in range(3):
                set_cell(e2e_table, row_index, col, e2e_table.rows[row_index].cells[col].text)

    # Evaluation answers and summaries.
    set_cell(doc.tables[23], 2, 1, "Yes: 72/72 registered tests passed, and the final benchmark log contained no optimistic-lock exception.")
    set_cell(doc.tables[23], 3, 1, "All 20 Chromium workflow, contract, and resilience tests passed; Firefox/WebKit and participant UAT remain incomplete.")
    set_cell(doc.tables[24], 1, 1, "72/72 registered tests passed")
    set_cell(doc.tables[24], 4, 1, "20/20 Chromium tests passed")
    set_cell(doc.tables[24], 4, 2, "Workflow, contract, and resilience coverage; no Firefox/WebKit or independent participant UAT")

    # Objective assessment: remove resolved/stale findings.
    set_cell(doc.tables[25], 4, 1, "Achieved in the controlled test; production concurrency unverified")
    set_cell(doc.tables[25], 4, 2, "Current synthetic benchmark, after-commit execution, and clean background log")
    set_cell(doc.tables[25], 5, 1, "Achieved as a configurable prototype")
    set_cell(doc.tables[25], 5, 2, "Auto-apply tests, scheduler, immediate notifications, and hashed confirm-then-POST email actions")
    set_cell(doc.tables[25], 7, 2, "Local health passed; scale, independent users, and production security evidence remain missing")

    set_cell(doc.tables[26], 5, 0, "Chromium-only 20-test automation suite")
    set_cell(doc.tables[26], 5, 1, "Limits cross-browser and real-user usability conclusions")


def update_images_and_alt_text(doc: Document) -> dict[str, Path]:
    screen_sources = {
        "Screen 4.1.": SCREEN_DIR / "screen-4-1-public-jobs.png",
        "Screen 4.2.": SCREEN_DIR / "screen-4-2-candidate-matching.png",
        "Screen 4.3.": SCREEN_DIR / "screen-4-3-cv-upload.png",
        "Screen 4.4.": SCREEN_DIR / "screen-4-4-recruiter-workspace.png",
        "Screen 4.5.": SCREEN_DIR / "screen-4-5-autofit-settings.png",
        "Screen 4.6.": SCREEN_DIR / "screen-4-6-admin-audit.png",
    }
    replacements: dict[str, Path] = {}
    paragraphs = doc.paragraphs
    for index, paragraph in enumerate(paragraphs):
        drawings = paragraph._p.xpath(".//w:drawing")
        if not drawings:
            continue
        caption = "Can Tho University logo"
        if index + 1 < len(paragraphs) and paragraphs[index + 1].text.strip():
            caption = paragraphs[index + 1].text.strip()
        for doc_pr in paragraph._p.xpath(".//wp:docPr"):
            doc_pr.set("descr", caption)
            doc_pr.set("title", caption)
        blips = paragraph._p.xpath(".//a:blip")
        if blips:
            rel_id = blips[0].get(qn("r:embed"))
            if rel_id:
                target = str(doc.part.rels[rel_id].target_ref).replace("\\", "/")
                archive_name = "word/" + target.lstrip("/")
                for prefix, source in screen_sources.items():
                    if caption.startswith(prefix):
                        replacements[archive_name] = source
                        if prefix == "Screen 4.1.":
                            for crop in paragraph._p.xpath(".//a:srcRect"):
                                crop.getparent().remove(crop)
                        break
    if len(replacements) != 6:
        raise RuntimeError(f"Expected six screenshot relationships, found {len(replacements)}: {replacements}")
    return replacements


def replace_archive_members(docx: Path, replacements: dict[str, Path]) -> None:
    patched = docx.with_suffix(".patched.docx")
    with zipfile.ZipFile(docx, "r") as source, zipfile.ZipFile(patched, "w") as target:
        for info in source.infolist():
            data = replacements[info.filename].read_bytes() if info.filename in replacements else source.read(info.filename)
            target.writestr(info, data)
    patched.replace(docx)


def main() -> None:
    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)
    doc = Document(DOCX)
    replace_paragraphs(doc)
    update_tables(doc)
    screenshot_members = update_images_and_alt_text(doc)
    doc.save(TEMP)
    replace_archive_members(TEMP, screenshot_members)
    shutil.copy2(TEMP, DOCX)
    print(f"Updated: {DOCX}")
    print(f"Backup: {BACKUP}")
    print("Replaced screenshot members:")
    for member, source in screenshot_members.items():
        print(f"- {member} <- {source.name}")


if __name__ == "__main__":
    main()
