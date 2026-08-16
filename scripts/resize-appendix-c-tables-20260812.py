from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
OUTPUT = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-final-resized-20260812.docx"
FINAL = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"


def set_widths(table: Table, widths: list[float]) -> None:
    table.autofit = False
    total_twips = int(Inches(sum(widths)).twips)
    tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
    if tbl_w is None:
        from docx.oxml import OxmlElement

        tbl_w = OxmlElement("w:tblW")
        table._tbl.tblPr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_twips))
    layout = table._tbl.tblPr.find(qn("w:tblLayout"))
    if layout is None:
        from docx.oxml import OxmlElement

        layout = OxmlElement("w:tblLayout")
        table._tbl.tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for grid_col, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(int(Inches(width).twips)))
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)


def main() -> None:
    doc = Document(SOURCE)
    for marker in doc._element.xpath(".//w:lastRenderedPageBreak"):
        marker.getparent().remove(marker)
    active = False
    count = 0
    preceding_caption = ""
    for child in doc._element.body.iterchildren():
        if child.tag == qn("w:p"):
            text = Paragraph(child, doc._body).text.strip()
            preceding_caption = text
            if text == "Appendix C. Full Data Dictionary":
                active = True
            elif text == "Appendix D. Evaluation Summary":
                active = False
        elif child.tag == qn("w:tbl"):
            table = Table(child, doc._body)
            selected = False
            if preceding_caption.startswith("Table App.A.1."):
                set_widths(table, [0.80, 2.20, 3.75])
                selected = True
            elif preceding_caption.startswith("Table App.B.1."):
                set_widths(table, [0.90, 2.55, 3.30])
                selected = True
            elif active:
                count += 1
                selected = True
                if len(table.columns) == 7:
                    set_widths(table, [0.35, 1.35, 1.55, 0.45, 0.50, 2.15, 3.10])
                elif len(table.columns) == 4:
                    set_widths(table, [0.55, 2.00, 1.75, 5.15])
            if not selected:
                preceding_caption = ""
                continue
            for row_index, row in enumerate(table.rows):
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = (
                            WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                        )
            if len(table.rows) >= 4:
                for cell in table.rows[-3].cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.keep_with_next = True
            if len(table.rows) >= 3:
                for cell in table.rows[-2].cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.keep_with_next = True
            preceding_caption = ""

    if count != 25:
        raise RuntimeError(f"Expected 25 Appendix C tables, found {count}")
    doc.save(OUTPUT)
    shutil.copy2(OUTPUT, FINAL)
    print(f"Updated {count} Appendix C tables")
    print(FINAL)


if __name__ == "__main__":
    main()
