from __future__ import annotations

import hashlib
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "Doc" / "mau-luan-van" / "Thesis template_SE_English_updated.docx"
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-before-20260811-template-cover-border.docx"


def table_digest(document: Document) -> str:
    payload = []
    for table in document.tables:
        payload.append([[cell.text for cell in row.cells] for row in table.rows])
    return hashlib.sha256(repr(payload).encode("utf-8")).hexdigest()


def drawing_count(document: Document) -> int:
    return len(document._element.body.xpath(".//w:drawing"))


def page_border(section):
    return section._sectPr.find(qn("w:pgBorders"))


def border_signature(border):
    if border is None:
        return None
    result = []
    for side in ("top", "left", "bottom", "right"):
        edge = border.find(qn(f"w:{side}"))
        if edge is None:
            result.append((side, None))
        else:
            result.append(
                (
                    side,
                    edge.get(qn("w:val")),
                    edge.get(qn("w:sz")),
                    edge.get(qn("w:space")),
                    edge.get(qn("w:color")),
                )
            )
    return tuple(result)


def main() -> None:
    if not TEMPLATE.exists():
        raise FileNotFoundError(TEMPLATE)
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)

    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(DOCX, BACKUP)

    template = Document(TEMPLATE)
    document = Document(DOCX)
    before_tables = table_digest(document)
    before_paragraphs = len(document.paragraphs)
    before_drawings = drawing_count(document)
    before_sections = len(document.sections)

    source_border = page_border(template.sections[0])
    if source_border is None:
        raise RuntimeError("The template's first section has no page border")

    target_sect_pr = document.sections[0]._sectPr
    old_border = page_border(document.sections[0])
    if old_border is not None:
        insert_at = target_sect_pr.index(old_border)
        target_sect_pr.remove(old_border)
        target_sect_pr.insert(insert_at, deepcopy(source_border))
    else:
        # Insert after page margins when no existing border is available.
        page_margin = target_sect_pr.find(qn("w:pgMar"))
        insert_at = target_sect_pr.index(page_margin) + 1 if page_margin is not None else len(target_sect_pr)
        target_sect_pr.insert(insert_at, deepcopy(source_border))

    if table_digest(document) != before_tables:
        raise RuntimeError("Table contents changed during the cover-border update")
    if len(document.paragraphs) != before_paragraphs:
        raise RuntimeError("Paragraph count changed during the cover-border update")
    if drawing_count(document) != before_drawings:
        raise RuntimeError("Drawing count changed during the cover-border update")
    if len(document.sections) != before_sections:
        raise RuntimeError("Section count changed during the cover-border update")

    actual = border_signature(page_border(document.sections[0]))
    expected = border_signature(source_border)
    if actual != expected:
        raise RuntimeError(f"Copied border differs from template: {actual!r} != {expected!r}")

    for section in document.sections[1:]:
        if page_border(section) is not None:
            raise RuntimeError("A non-cover section unexpectedly has a page border")

    document.save(DOCX)
    print(f"Updated: {DOCX}")
    print(f"Backup:  {BACKUP}")
    print(f"Border signature: {actual}")
    print(f"Tables unchanged: {before_tables}")
    print(f"Paragraphs/drawings/sections: {before_paragraphs}/{before_drawings}/{before_sections}")


if __name__ == "__main__":
    main()
