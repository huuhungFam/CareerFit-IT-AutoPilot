from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
doc = Document(DOCX)

targets = {
    "2. Background and Related Work",
    "2.9 Related Work and Research Gap",
    "REFERENCES",
}

for index, paragraph in enumerate(doc.paragraphs):
    if " ".join(paragraph.text.split()) in targets:
        print(f"\n--- {index}: {paragraph.text!r} [{paragraph.style.name}] ---")
        for following in doc.paragraphs[index + 1 : index + 12]:
            print(f"{following.text!r} [{following.style.name}]")
