from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shape import InlineShape
from docx.shared import Cm, Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
TEMPLATE = ROOT / "Doc" / "Thesis-Report.docx"
LOGO = ROOT / "Doc" / "working" / "ctu-template-logo.png"


def set_cell(cell, text):
    cell.text = text


def table_by_first_header(doc, header):
    for table in doc.tables:
        if table.rows and table.rows[0].cells and table.rows[0].cells[0].text.strip() == header:
            return table
    raise RuntimeError(f"Table not found: {header}")


def replace_table_rows(table, rows):
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        tr_pr.append(repeat)


def set_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


doc = Document(DOCX)

# Page and base paragraph settings follow the approved template.
section = doc.sections[0]
section.page_width = Cm(21.01)
section.page_height = Cm(29.69)
section.left_margin = Cm(3.0)
section.right_margin = Cm(2.01)
section.top_margin = Cm(2.01)
section.bottom_margin = Cm(2.01)
section.header_distance = Cm(1.27)
section.footer_distance = Cm(1.27)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(13)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(0)

h1 = doc.styles["Heading 1"]
h1.font.name = "Times New Roman"
h1.font.size = Pt(16)
h1.font.bold = True
h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
h1.paragraph_format.line_spacing = 1.5
h1.paragraph_format.space_before = Pt(18)
h1.paragraph_format.space_after = Pt(6)
h1.paragraph_format.keep_with_next = True
h1.paragraph_format.keep_together = True
num_pr = h1.element.pPr.find(qn("w:numPr"))
if num_pr is not None:
    h1.element.pPr.remove(num_pr)

h2 = doc.styles["Heading 2"]
h2.font.name = "Times New Roman"
h2.font.size = Pt(14)
h2.font.bold = True
h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h2.paragraph_format.line_spacing = 1.5
h2.paragraph_format.space_before = Pt(12)
h2.paragraph_format.space_after = Pt(6)
h2.paragraph_format.keep_with_next = True
h2.paragraph_format.keep_together = True
num_pr = h2.element.pPr.find(qn("w:numPr"))
if num_pr is not None:
    h2.element.pPr.remove(num_pr)

h3 = doc.styles["Heading 3"]
h3.font.name = "Times New Roman"
h3.font.size = Pt(13)
h3.font.bold = True
h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h3.paragraph_format.line_spacing = 1.5
h3.paragraph_format.space_before = Pt(8)
h3.paragraph_format.space_after = Pt(4)
h3.paragraph_format.keep_with_next = True
h3.paragraph_format.keep_together = True
num_pr = h3.element.pPr.find(qn("w:numPr"))
if num_pr is not None:
    h3.element.pPr.remove(num_pr)

list_style = doc.styles["List Paragraph"]
list_style.font.name = "Times New Roman"
list_style.font.size = Pt(13)
list_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
list_style.paragraph_format.line_spacing = 1.5
list_style.paragraph_format.left_indent = Cm(0.75)
list_style.paragraph_format.first_line_indent = Cm(-0.75)
list_style.paragraph_format.space_after = Pt(0)

for name in ("Figure Caption", "Table Caption"):
    style = doc.styles[name]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.font.italic = True
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.keep_together = True
doc.styles["Figure Caption"].paragraph_format.space_before = Pt(3)
doc.styles["Figure Caption"].paragraph_format.space_after = Pt(8)
doc.styles["Table Caption"].paragraph_format.space_before = Pt(8)
doc.styles["Table Caption"].paragraph_format.space_after = Pt(3)
doc.styles["Table Caption"].paragraph_format.keep_with_next = True

# Major divisions start on new pages. Appendix items are subordinate to APPENDICES.
major = {
    "ACKNOWLEDGEMENTS", "ABSTRACT", "TABLE OF CONTENTS", "LIST OF FIGURES",
    "LIST OF TABLES", "LIST OF ABBREVIATIONS", "CHAPTER 1. INTRODUCTION",
    "CHAPTER 2. THEORETICAL BACKGROUND AND RELATED WORK",
    "CHAPTER 3. SYSTEM ANALYSIS AND DESIGN", "CHAPTER 4. SYSTEM IMPLEMENTATION",
    "CHAPTER 5. EXPERIMENTAL EVALUATION",
    "CHAPTER 6. RESULTS, DISCUSSION AND CONCLUSION", "REFERENCES", "APPENDICES",
}
for paragraph in doc.paragraphs:
    if paragraph.text in major:
        paragraph.style = doc.styles["Heading 1"]
        paragraph.paragraph_format.page_break_before = True
    elif paragraph.text.startswith("Appendix "):
        paragraph.style = doc.styles["Heading 2"]
        paragraph.paragraph_format.page_break_before = False

# Restore the CTU logo from the approved template if it is not present.
with ZipFile(TEMPLATE) as archive:
    media = next(name for name in archive.namelist() if name.startswith("word/media/"))
    LOGO.write_bytes(archive.read(media))
logo_paragraph = doc.paragraphs[5]
if "w:drawing" not in logo_paragraph._p.xml:
    logo_paragraph.clear()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_paragraph.add_run().add_picture(str(LOGO), width=Inches(1.39236), height=Inches(1.39514))

# Crop the repeated full-page public-job capture to its first 712-pixel viewport.
caption = next(p for p in doc.paragraphs if p.text == "Screen 4.1. Public Job search")
caption_index = next(i for i, p in enumerate(doc.paragraphs) if p.text == "Screen 4.1. Public Job search")
image_paragraph = doc.paragraphs[caption_index - 1]
inline = image_paragraph._p.xpath(".//wp:inline")
if len(inline) != 1:
    raise RuntimeError("Screen 4.1 inline image was not found")
shape = InlineShape(inline[0])
shape.width = Inches(6.15)
shape.height = Inches(6.15 * 712 / 1265)
blip_fill = image_paragraph._p.find(".//" + qn("pic:blipFill"))
src_rect = blip_fill.find(qn("a:srcRect"))
if src_rect is None:
    src_rect = OxmlElement("a:srcRect")
    blip_fill.insert(1, src_rect)
src_rect.set("b", str(round((1 - 712 / 8902) * 100000)))

# Keep every image with its caption and center both elements.
for i, paragraph in enumerate(doc.paragraphs[:-1]):
    if "w:drawing" in paragraph._p.xml and doc.paragraphs[i + 1].style.name == "Figure Caption":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(0)

# Table 1.1 was the only table whose caption followed the table; move it above.
objective_table = table_by_first_header(doc, "Thesis objective")
objective_caption = next(p for p in doc.paragraphs if p.text.startswith("Table 1.1."))
objective_table._element.addprevious(objective_caption._p)

# Standardize table typography, header repetition, and row pagination.
for table_index, table in enumerate(doc.tables):
    table.autofit = True
    if table_index > 0:
        set_repeat_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        set_cant_split(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11 if table_index > 0 else 13)
                    if row_index == 0 and table_index > 0:
                        run.bold = True

# References use a hanging indent and compact spacing.
for paragraph in doc.paragraphs:
    if paragraph.text.startswith("[") and len(paragraph.text) > 3 and paragraph.text[1].isdigit():
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.first_line_indent = Cm(-0.75)
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(4)

# Update stale result tables and prose that predated the final remediation run.
build_table = table_by_first_header(doc, "Artifact")
replace_table_rows(build_table, [
    ("index.html", "1.37 kB", "0.68 kB"),
    ("Main CSS", "61.91 kB", "12.27 kB"),
    ("Icons chunk", "18.96 kB", "4.45 kB"),
    ("Query chunk", "39.51 kB", "12.14 kB"),
    ("React chunk", "179.74 kB", "58.98 kB"),
    ("Application chunk", "188.05 kB", "51.84 kB"),
    ("Charts chunk", "375.64 kB", "110.88 kB"),
])
for paragraph in doc.paragraphs:
    if paragraph.text.startswith("Vite emitted a warning because"):
        paragraph.clear()
        paragraph.add_run("Manual Rollup chunking separated React, query, chart, and icon dependencies. The largest generated JavaScript chunk was 375.64 kB, below Vite's 500 kB warning threshold. No browser performance profile was collected, so bundle size is reported as build evidence rather than a direct page-load measurement.")

risk = table_by_first_header(doc, "Risk")
set_cell(risk.rows[3].cells[2], "Implemented hashed storage, GET confirmation, and POST execution; retain replay and expiry tests")
set_cell(risk.rows[4].cells[1], "GET is non-mutating and displays confirmation")
set_cell(risk.rows[4].cells[2], "Add rate limiting and deployment-specific origin controls")
set_cell(risk.rows[5].cells[1], "Scheduler timing reads from app.scheduler properties")
set_cell(risk.rows[5].cells[2], "Maintain configuration and schedule-boundary tests")

limitations = table_by_first_header(doc, "Area")
set_cell(limitations.rows[1].cells[1], "Hashed token storage with GET confirmation and POST execution")
set_cell(limitations.rows[1].cells[2], "Add rate limiting, origin controls, secret rotation, and delivery monitoring")
set_cell(limitations.rows[5].cells[1], "All scheduler expressions use app.scheduler properties")
set_cell(limitations.rows[5].cells[2], "Add schedule-boundary and time-zone integration tests")
set_cell(limitations.rows[6].cells[1], "Learning is registered after transaction commit")
set_cell(limitations.rows[6].cells[2], "Add production retry, conflict, and failure-recovery policy")
set_cell(limitations.rows[7].cells[1], "JWT and account summary use sessionStorage")
set_cell(limitations.rows[7].cells[2], "Evaluate HttpOnly cookie or refresh-token architecture and strengthen XSS controls")
set_cell(limitations.rows[8].cells[1], "API-driven views no longer substitute mock Job fields")
set_cell(limitations.rows[8].cells[2], "Maintain contract tests to prevent fallback data from returning")

runtime = table_by_first_header(doc, "Endpoint or check")
set_cell(runtime.rows[2].cells[1], "200, UP")

questions = table_by_first_header(doc, "Question")
set_cell(questions.rows[2].cells[1], "Yes: 63/63 registered tests passed, and the final benchmark log contained no optimistic-lock exception.")
set_cell(questions.rows[5].cells[1], "Yes in the evaluated local profile: aggregate health, liveness, readiness, and the core API returned HTTP 200. This is not production monitoring validation.")

status = table_by_first_header(doc, "Evaluation area")
set_cell(status.rows[1].cells[2], "Final logs were also checked for background exceptions")
set_cell(status.rows[3].cells[2], "Largest JavaScript chunk was 375.64 kB; no Vite size warning")
set_cell(status.rows[6].cells[1], "Core Job API and aggregate/liveness/readiness health returned HTTP 200/UP")
set_cell(status.rows[6].cells[2], "Local profile only; no external production monitoring")

principal = table_by_first_header(doc, "Limitation")
set_cell(principal.rows[3].cells[0], "Production concurrency not evaluated")
set_cell(principal.rows[3].cells[1], "Final benchmark is clean, but production conflict and retry behavior remain unproven")
set_cell(principal.rows[4].cells[0], "Local-only health verification")
set_cell(principal.rows[4].cells[1], "Health passed locally but does not establish deployed availability")
set_cell(principal.rows[6].cells[0], "Email delivery and abuse controls not production-tested")
set_cell(principal.rows[6].cells[1], "Hashed confirm-then-POST flow is implemented; rate limiting and real delivery remain unverified")

# Ask Word to refresh all generated lists when the file opens.
settings = doc.settings._element
update = settings.find(qn("w:updateFields"))
if update is None:
    update = OxmlElement("w:updateFields")
    settings.append(update)
update.set(qn("w:val"), "true")

doc.save(DOCX)
print(DOCX)
