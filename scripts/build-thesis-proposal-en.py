from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"C:\CODING\Thesis")
REFERENCE = ROOT / "Doc" / "Đề cương LV Mau.docx"
OUTPUT = ROOT / "Doc" / "CareerFit-Thesis-Proposal-English.docx"


def clear_body(doc: Document) -> None:
    body = doc._body._element
    sect_pr = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sect_pr = deepcopy(child)
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "808080", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        element = margins.find(qn("w:" + key))
        if element is None:
            element = OxmlElement("w:" + key)
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def format_run(run, *, bold=False, italic=False, size=13):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_paragraph(doc, text="", *, bold=False, italic=False, size=13, align=None, before=0, after=3, first_line=True):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = 1.15
    if first_line:
        fmt.first_line_indent = Inches(0.3)
    run = p.add_run(text)
    format_run(run, bold=bold, italic=italic, size=size)
    return p


def add_section(doc, title):
    return add_paragraph(doc, title, bold=True, size=13, before=8, after=4, first_line=False)


def add_subsection(doc, title):
    return add_paragraph(doc, title, bold=True, size=13, before=4, after=3, first_line=False)


def add_label_para(doc, label, text):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(3)
    fmt.line_spacing = 1.15
    fmt.first_line_indent = Inches(0.3)
    r1 = p.add_run(label)
    format_run(r1, bold=True, size=13)
    r2 = p.add_run(" " + text)
    format_run(r2, size=13)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = widths[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "D9EAF7")
        set_cell_border(cell)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        format_run(run, bold=True, size=12)
    for row in rows:
        docx_row = table.add_row()
        set_row_cant_split(docx_row)
        cells = docx_row.cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.width = widths[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            format_run(run, size=11)
    add_paragraph(doc, "", after=2, first_line=False)
    return table


def build():
    doc = Document(REFERENCE)
    clear_body(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)

    add_paragraph(doc, "THESIS PROPOSAL", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, after=8, first_line=False)
    add_paragraph(
        doc,
        "DESIGN AND IMPLEMENTATION OF A HUMAN-IN-THE-LOOP AI-ASSISTED RECRUITMENT AUTOMATION PLATFORM FOR CV-JD EVALUATION AND RECOMMENDATION IN IT",
        bold=True,
        size=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=3,
        first_line=False,
    )
    add_paragraph(
        doc,
        "CareerFit IT AutoPilot",
        bold=True,
        italic=True,
        size=13,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=12,
        first_line=False,
    )
    add_label_para(doc, "Student:", "Pham Huu Hung - Student ID: B2203557")
    add_label_para(doc, "Supervisor 1:", "Ph.D. Nguyen Thanh Khoa - Staff ID: [to be added]")
    add_label_para(doc, "Supervisor 2:", "None")

    add_section(doc, "I. INTRODUCTION")
    add_subsection(doc, "1. Rationale")
    add_paragraph(
        doc,
        "In information technology recruitment, candidates and recruiters must process a large amount of heterogeneous information from curricula vitae, job descriptions, skills, experience, work locations, salary expectations, and language requirements. Common job portals support posting, searching, and applying for jobs, but CV-JD screening and candidate discovery are still often performed manually and depend heavily on subjective judgement.",
    )
    add_paragraph(
        doc,
        "Simple keyword filtering can miss important relationships between skills and job requirements, while fully automated decision-making can reduce transparency and user control. A practical recruitment platform should therefore combine automated recommendation, explainable matching scores, user feedback, and Human-in-the-Loop safeguards for consequential actions such as applying, inviting, rejecting, or marking a candidate as potential.",
    )
    add_paragraph(
        doc,
        "This thesis proposes CareerFit, a web-based recruitment platform that integrates a Job Portal, CV-JD Matching, Job Recommendation, AutoFit Automation, an Email Action Channel, and an Audit Trail. Instead of only ranking CVs or jobs, the system demonstrates a controlled recruitment workflow following the loop of Perception - Decision - Action - Learning - Audit.",
    )

    add_subsection(doc, "2. Research objectives")
    add_label_para(
        doc,
        "Overall objective:",
        "To design and implement a web platform that supports CV-JD fit evaluation, job recommendation, recruiter-side candidate discovery, and policy-driven recruitment automation with Human-in-the-Loop control, explainability, and auditability.",
    )
    add_label_para(
        doc,
        "Specific objectives:",
        "The thesis aims to build role-based experiences for Guest, Candidate, Recruiter, and Admin users; ingest CV data from PDF files or manual forms; validate and normalize recruitment text; implement separate workflows for CV-JD matching and candidate-profile recommendation; apply TF-IDF, cosine similarity, score normalization, and reason generation; incorporate feedback learning using Rocchio; implement AutoFit policies, scheduled processing, email magic links, and audit logging; and evaluate the solution using focused unit tests, integration tests, E2E/UAT scenarios, and ranking metrics where appropriate.",
    )

    add_subsection(doc, "3. Research object and scope")
    add_label_para(
        doc,
        "Research object:",
        "The research object is the IT recruitment workflow on a web platform, including CV/JD textual data, ranking and recommendation algorithms based on lexical vector representation, user feedback, and controlled automation policies.",
    )
    add_label_para(
        doc,
        "Scope:",
        "The thesis focuses on an internal job portal, CV/JD management, matching, recommendation, application, feedback, AutoFit policy, actionable email, analytics, and audit workflows for candidate, recruiter, and admin roles. Full ATS processes such as interviews, offers, and payroll, automatic submission to third-party websites, autonomous LLM-based decision-making, and large-scale commercial effectiveness validation are outside the core scope.",
    )

    add_section(doc, "II. THEORETICAL BACKGROUND")
    add_subsection(doc, "1. Recruitment information and CV-JD matching")
    add_paragraph(
        doc,
        "CVs and job descriptions in IT recruitment contain both structured and unstructured information. Structured fields include job title, skills, years of experience, seniority, location, language, and salary. Unstructured text includes project experience, responsibilities, achievements, and requirement descriptions. The matching task transforms these data sources into comparable representations and ranks the degree of fit rather than making the final hiring decision.",
    )
    add_subsection(doc, "2. Text preprocessing, TF-IDF, and cosine similarity")
    add_paragraph(
        doc,
        "Text preprocessing includes case normalization, noise removal, tokenization, skill normalization, and vocabulary construction. Term Frequency-Inverse Document Frequency represents documents by weighting terms according to their frequency within a document and their rarity across the corpus. Cosine similarity measures the angle between two vectors and is suitable for ranking the lexical similarity between CVs, job descriptions, and desired candidate profiles.",
    )
    add_subsection(doc, "3. Relevance feedback and the Rocchio algorithm")
    add_paragraph(
        doc,
        "Rocchio is a relevance feedback method from information retrieval that adjusts a query or profile vector based on positive and negative feedback. In CareerFit, feedback labels such as Good Match, Potential, Bad Match, and Not Interested act as signals for updating the learned vector and influencing later recommendation rankings.",
    )
    add_subsection(doc, "4. Human-in-the-Loop and explainable automation")
    add_paragraph(
        doc,
        "Human-in-the-Loop design allows a system to automate computation, filtering, and recommendation while keeping important actions under human confirmation, approval, or intervention. In CareerFit, the AutoFit policy layer converts matching scores into actions only when threshold, consent, quota, state, and notification timing conditions are satisfied. Email actions use magic links, a confirmation screen, and audit logging to avoid unintended execution.",
    )
    add_subsection(doc, "5. Web, security, and data foundations")
    add_paragraph(
        doc,
        "The system follows a client-server architecture with REST APIs, a React and TypeScript frontend, a Spring Boot backend, and PostgreSQL persistence. Relevant foundations include JWT authentication, role-based authorization, Flyway database migration, asynchronous processing, scheduled jobs, one-time email action tokens, audit logging, and consistent error handling.",
    )
    add_paragraph(
        doc,
        "The theoretical references include information retrieval, recommender systems, relevance feedback, Human-AI interaction, and technical documentation for Spring Boot, React, and PostgreSQL. Internal project documents such as the SRS, architecture notes, test cases, and existing thesis report are used to keep the proposal aligned with the implemented system.",
    )

    add_section(doc, "III. RESEARCH METHODOLOGY")
    add_subsection(doc, "1. Research approach")
    add_paragraph(
        doc,
        "This thesis follows an applied software engineering research approach that combines requirement analysis, system design, software implementation, and empirical evaluation. First, the IT recruitment problem is modeled into workflows for candidates, recruiters, administrators, and background workers. Next, the system is designed into data, matching, recommendation, feedback, automation, email action, and audit modules. Finally, the solution is evaluated through unit tests, integration tests, E2E/UAT scenarios, and ranking metrics on controlled data.",
    )
    add_subsection(doc, "2. Implementation process")
    add_table(
        doc,
        ["No.", "Activity", "Expected result"],
        [
            ["1", "Study IT recruitment workflows, job portals, ATS concepts, CV-JD matching, Human-in-the-Loop automation, and feedback learning.", "A background report, problem scope, and identified research gap."],
            ["2", "Analyze requirements for Guest, Candidate, Recruiter, Admin, and background/email workflows.", "SRS, use cases, functional requirements, and non-functional requirements."],
            ["3", "Design the system architecture, data model, CV/JD processing pipeline, scoring algorithm, and AutoFit policy layer.", "Architecture diagrams, ERD, sequence flows, and module design."],
            ["4", "Implement backend, frontend, database migrations, matching/recommendation, Rocchio feedback, email action, and audit logging.", "A working CareerFit web prototype for the main roles."],
            ["5", "Conduct unit, integration, E2E/UAT, algorithm evaluation, and security/operation review.", "Test reports, acceptance scripts, evidence artifacts, and evaluation results."],
            ["6", "Synthesize results, analyze limitations, finalize the thesis report, and prepare the defense materials.", "Completed thesis report, future development plan, and demo documentation."],
        ],
        [Inches(0.55), Inches(4.35), Inches(2.1)],
    )

    add_section(doc, "IV. PRACTICAL APPLICATION AND DATA SOURCES")
    add_subsection(doc, "1. Practical application")
    add_paragraph(
        doc,
        "CareerFit can be used as an internal recruitment system or as a research prototype for technology companies, career centers, and HR-tech product teams. Candidates can search for jobs, manage CVs, view fit scores, receive recommendations, and provide feedback. Recruiters can create job descriptions, view applicants, discover high-potential candidates, send invitations, and teach the system through feedback. Administrators can monitor users, jobs, email activity, and audit logs.",
    )
    add_paragraph(
        doc,
        "The main practical value of the thesis is to demonstrate a controlled recruitment workflow. The system can recommend and automate actions when conditions are satisfied, but it also records reasons, states, policies, and action history so that users can inspect and review system behavior. This design is suitable for contexts that need to reduce repetitive work without handing recruitment decisions entirely to an automated black box.",
    )
    doc.add_page_break()
    add_subsection(doc, "2. Experimental data sources")
    add_table(
        doc,
        ["Data source", "Description", "Location"],
        [
            ["CareerFit seed data", "Demo accounts, jobs, employers, default CV data, and matching cards for candidate, recruiter, and admin workflow testing.", "Backend/careerfit-backend/src/main/resources/db/migration"],
            ["Crawled job data", "Normalized IT job postings imported into PostgreSQL for testing the job portal, search, filtering, and source metadata.", "scraped-data/jobs_for_careerfit_import.json"],
            ["Controlled evaluation dataset", "A controlled dataset for evaluating ranking, recommendation, and algorithm metrics.", "evaluation/controlled-dataset.json"],
            ["Evidence and test artifacts", "Backend, frontend, E2E, UAT, and demo evidence used to cross-check the thesis report.", "evidence/, TEST_CASES.md, CAREERFIT_E2E_TEST_SCRIPT.md"],
        ],
        [Inches(1.65), Inches(3.65), Inches(1.7)],
    )

    add_section(doc, "V. RELATED PROJECTS")
    add_paragraph(
        doc,
        "Related systems are reviewed according to job portal functionality, ATS/recruitment workflow support, matching and recommendation, explainability, feedback learning, Human-in-the-Loop policy control, and auditability. CareerFit is not intended to replace large commercial platforms; instead, it focuses on demonstrating an explainable and controlled integrated architecture within the thesis scope."
    )
    add_table(
        doc,
        ["Project/System", "Description", "Relevance to CareerFit"],
        [
            ["LinkedIn Jobs / Indeed", "Popular job portal platforms that support job posting, searching, applying, and large-scale job recommendation.", "Reference for job feed, search, job detail, and recommendation experiences; CareerFit focuses on an auditable prototype with explicit policies."],
            ["Greenhouse / Workable", "ATS platforms that support recruitment pipeline management, candidate tracking, interview workflows, and team collaboration.", "Reference for recruiter workspace and applicant management; CareerFit keeps a narrower scope around CV-JD matching, discovery, invitation, and audit."],
            ["Recommender Systems", "Systems that use user/item representations, ranking metrics, and feedback to personalize results.", "Theoretical foundation for candidate-job recommendation and metrics such as Precision@K, Recall@K, MRR, and nDCG."],
            ["Human-AI Interaction Guidelines", "Design principles for transparency, controllability, confirmation, error recovery, and appropriate reliance in human-AI systems.", "Foundation for AutoFit policy, email confirmation before action, reason display, and audit trail."],
        ],
        [Inches(1.6), Inches(3.1), Inches(2.3)],
    )

    add_section(doc, "VI. EXPECTED RESULTS AND CONTRIBUTIONS")
    add_paragraph(
        doc,
        "The expected result is a working CareerFit web platform with Guest, Candidate, Recruiter, and Admin roles; a Spring Boot backend; a React and TypeScript frontend; and a PostgreSQL database managed through Flyway migrations. The system includes CV/JD processing, validation, matching, recommendation, feedback learning, AutoFit policy evaluation, email action handling, and audit logging.",
    )
    add_paragraph(
        doc,
        "Technically, the thesis contributes an integrated recruitment architecture consisting of a Job Portal, Matching Engine, Recommendation Engine, AutoFit Automation, and Human-in-the-Loop Email Action Channel. The scoring pipeline uses TF-IDF and cosine similarity to keep results explainable at the skill/keyword level. Rocchio feedback adjusts learned representations, while the policy engine converts scores into controlled actions rather than unconditional automation.",
    )
    add_paragraph(
        doc,
        "Empirically, the thesis is expected to provide unit tests, integration tests, E2E/UAT scenarios, controlled evaluation data, demo scripts, and evaluation reports. The results will be presented within the correct scope: algorithm evaluation on controlled data, system evaluation through tests and UAT, and explicit limitations such as lexical TF-IDF representation, limited experimental data, and the need for human final decisions in recruitment.",
    )
    add_paragraph(
        doc,
        "In terms of future application, the system can be extended with semantic embeddings, learning-to-rank models, advanced analytics dashboards, real email integration, external ATS integration, and evaluation with real users.",
    )

    add_section(doc, "VII. REFERENCES")
    references = [
        "Manning, C. D., Raghavan, P., & Schutze, H. (2008). Introduction to Information Retrieval. Cambridge University Press.",
        "Salton, G., & Buckley, C. (1988). Term-weighting approaches in automatic text retrieval. Information Processing & Management, 24(5), 513-523.",
        "Rocchio, J. J. (1971). Relevance feedback in information retrieval. In The SMART Retrieval System: Experiments in Automatic Document Processing.",
        "Ricci, F., Rokach, L., & Shapira, B. (Eds.). (2015). Recommender Systems Handbook. Springer.",
        "Aggarwal, C. C. (2016). Recommender Systems: The Textbook. Springer.",
        "Amershi, S., et al. (2019). Guidelines for Human-AI Interaction. Proceedings of CHI 2019.",
        "Spring Boot Documentation. https://docs.spring.io/spring-boot/",
        "React Documentation. https://react.dev/",
        "PostgreSQL Documentation. https://www.postgresql.org/docs/",
        "CareerFit repository documents: proposal.md, srs.md, architecture.md, TEST_CASES.md, CAREERFIT_E2E_TEST_SCRIPT.md, Doc/THESIS_REPORT_CONTENT_PLAN.md.",
    ]
    for ref in references:
        add_paragraph(doc, ref, size=12, first_line=False, after=2)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
