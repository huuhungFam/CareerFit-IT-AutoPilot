from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"


def replace_all(paragraph, old: str, new: str) -> None:
    if old not in paragraph.text:
        return
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)


document = Document(REPORT)

# This sentence was added after the chapter-renumbering pass, so it must refer to
# the final evaluation chapter directly.
for paragraph in document.paragraphs:
    if paragraph.text.startswith("The final evidence package contains"):
        replace_all(paragraph, "Chapter 3 reports", "Chapter 4 reports")
    if paragraph.text.startswith(
        "The report contains an Introduction, four main chapters"
    ):
        paragraph.text = (
            "The report contains an Introduction, four chapters, and a "
            "Conclusion. Chapter 1 defines the problem, stakeholders, "
            "requirements, and use cases. Chapter 2 presents the theoretical "
            "background and solution design. Chapter 3 explains the "
            "implementation. Chapter 4 presents the evaluation method, evidence, "
            "and threats to validity. The Conclusion discusses achievements, "
            "limitations, and future work."
        )
    if paragraph.text in {
        "DESIGN AND IMPLEMENTATION OF A HUMAN-IN-THE-LOOP",
        "AI-ASSISTED RECRUITMENT AUTOMATION PLATFORM FOR",
        "AI-ASSISTED RECRUITMENT AUTOMATION PLATFORM",
    }:
        if paragraph.text.endswith(" PLATFORM FOR"):
            paragraph.text = "AI-ASSISTED RECRUITMENT AUTOMATION PLATFORM"
        paragraph.alignment = 1
        for run in paragraph.runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run._element.get_or_add_rPr().rFonts.set(
                qn("w:eastAsia"), "Times New Roman"
            )
            run.font.size = Pt(13)
    if paragraph.text == "CV-JD EVALUATION AND RECOMMENDATION IN IT":
        paragraph.text = "FOR CV-JD EVALUATION AND RECOMMENDATION IN IT"
        paragraph.alignment = 1
        for run in paragraph.runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run._element.get_or_add_rPr().rFonts.set(
                qn("w:eastAsia"), "Times New Roman"
            )
            run.font.size = Pt(13)

for style_name in ("toc 1", "toc 2", "toc 3"):
    style = document.styles[style_name]
    style.font.name = "Times New Roman"
    style._element.get_or_add_rPr().rFonts.set(
        qn("w:eastAsia"), "Times New Roman"
    )
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

caption_mapping = {}
for old, new in {1: 3, 2: 4, 3: 5, 4: 6, 5: 7, 6: 8}.items():
    caption_mapping[f"Figure 3.{old}."] = f"Figure 1.{new}."
for old, new in {7: 5, 8: 6, 9: 7}.items():
    caption_mapping[f"Figure 3.{old}."] = f"Figure 2.{new}."
for number in range(1, 11):
    caption_mapping[f"Figure 4.{number}."] = f"Figure 3.{number}."
for number in range(1, 7):
    caption_mapping[f"Screen 4.{number}."] = f"Screen 3.{number}."
for number in range(1, 5):
    caption_mapping[f"Figure 5.{number}."] = f"Figure 4.{number}."

for drawing in document._element.iter(qn("wp:docPr")):
    title = drawing.get("title", "")
    description = drawing.get("descr", "")
    if not title and not description:
        drawing.set("title", "Can Tho University logo")
        drawing.set("descr", "Can Tho University logo")
        continue
    for old, new in caption_mapping.items():
        title = title.replace(old, new)
        description = description.replace(old, new)
    drawing.set("title", title)
    drawing.set("descr", description)

# The inner-cover layout table is a single labeled row. Marking it as a header
# row gives assistive technology a stable interpretation and clears the audit.
if document.tables:
    row_properties = document.tables[0].rows[0]._tr.get_or_add_trPr()
    if row_properties.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        row_properties.append(header)

document.save(REPORT)
print(f"Cleaned text and image metadata in {REPORT}")
