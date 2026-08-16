from collections import Counter, defaultdict
from pathlib import Path
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
FILES = [ROOT / "Doc" / "Thesis-Report.docx", ROOT / "Doc" / "CareerFit-Thesis-Report.docx"]


def cm(value):
    return None if value is None else round(value.cm, 2)


def pt(value):
    return None if value is None else round(value.pt, 1)


for path in FILES:
    doc = Document(path)
    print(f"\n===== {path.name} =====")
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} images={len(doc.inline_shapes)} sections={len(doc.sections)}")
    for i, section in enumerate(doc.sections):
        print(
            f"section {i}: page={cm(section.page_width)}x{cm(section.page_height)} "
            f"margins L/R/T/B={cm(section.left_margin)}/{cm(section.right_margin)}/{cm(section.top_margin)}/{cm(section.bottom_margin)} "
            f"header/footer={cm(section.header_distance)}/{cm(section.footer_distance)} start={section.start_type} "
            f"differentFirst={section.different_first_page_header_footer}"
        )

    print("STYLES")
    wanted = ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "List Paragraph", "Caption", "Figure Caption", "Table Caption"]
    for name in wanted:
        if name not in doc.styles:
            continue
        s = doc.styles[name]
        pf = s.paragraph_format
        print(
            f"{name}: font={s.font.name} size={pt(s.font.size)} bold={s.font.bold} italic={s.font.italic} "
            f"align={pf.alignment} line={pf.line_spacing} before/after={pt(pf.space_before)}/{pt(pf.space_after)} "
            f"keepNext={pf.keep_with_next} keepTogether={pf.keep_together} pageBefore={pf.page_break_before}"
        )

    style_counts = Counter(p.style.name for p in doc.paragraphs)
    print("STYLE COUNTS", style_counts.most_common(15))
    direct_fonts = Counter()
    direct_sizes = Counter()
    mixed = []
    for i, p in enumerate(doc.paragraphs):
        names = {r.font.name for r in p.runs if r.text.strip() and r.font.name}
        sizes = {pt(r.font.size) for r in p.runs if r.text.strip() and r.font.size}
        direct_fonts.update(names)
        direct_sizes.update(sizes)
        if len(names) > 1 or len(sizes) > 1:
            mixed.append((i, p.style.name, names, sizes, p.text[:100]))
    print("DIRECT FONTS", direct_fonts)
    print("DIRECT SIZES", direct_sizes)
    print("MIXED RUN FORMAT", len(mixed))
    for row in mixed[:30]:
        print(row)

    print("HEADINGS")
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith("Heading"):
            print(i, p.style.name, repr(p.text), "pageBefore", p.paragraph_format.page_break_before, "keepNext", p.paragraph_format.keep_with_next)

    explicit_breaks = []
    for i, p in enumerate(doc.paragraphs):
        xml = p._p.xml
        if 'w:type="page"' in xml or "w:lastRenderedPageBreak" in xml:
            explicit_breaks.append((i, p.style.name, p.text[:100], xml.count('w:type="page"'), xml.count("w:lastRenderedPageBreak")))
    print("PAGE BREAK MARKERS", len(explicit_breaks))
    for row in explicit_breaks[:80]:
        print(row)

    print("EMPTY OR CONSECUTIVE HEADINGS")
    for i, p in enumerate(doc.paragraphs[:-1]):
        if p.style.name.startswith("Heading"):
            nxt = doc.paragraphs[i + 1]
            if not nxt.text.strip() or nxt.style.name.startswith("Heading"):
                print(i, p.style.name, repr(p.text), "->", nxt.style.name, repr(nxt.text))

    print("IMAGE/CAPTION PAIRS")
    for i, p in enumerate(doc.paragraphs):
        if "w:drawing" in p._p.xml:
            following = doc.paragraphs[i + 1].text if i + 1 < len(doc.paragraphs) else ""
            print(i, "image paragraph", "align", p.alignment, "next", repr(following[:120]))

    print("TOC/LIST COUNTS")
    print("toc1", style_counts.get("toc 1", 0), "toc2", style_counts.get("toc 2", 0), "toc3", style_counts.get("toc 3", 0))
    print("figure entries", sum(p.style.name == "toc 1" and (p.text.startswith("Figure ") or p.text.startswith("Screen ")) for p in doc.paragraphs))
    print("table entries", sum(p.style.name == "toc 1" and p.text.startswith("Table ") for p in doc.paragraphs))
