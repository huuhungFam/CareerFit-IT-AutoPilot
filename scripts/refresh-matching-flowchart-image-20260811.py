from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.text.paragraph import Paragraph

from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
IMAGE = ROOT / "Doc" / "figures" / "flowchart-matching-potential-20260811.png"


def find_exact(document, text):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one caption; found {len(matches)}")
    return matches[0]


def previous_paragraph(paragraph):
    previous = paragraph._p.getprevious()
    while previous is not None and not previous.tag.endswith("}p"):
        previous = previous.getprevious()
    return Paragraph(previous, paragraph._parent)


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def configure_image_paragraph(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.page_break_before = True
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0


def add_picture(paragraph, image_path, alt_text):
    inline = paragraph.add_run().add_picture(str(image_path), width=Inches(5.2))
    inline._inline.docPr.set("title", alt_text)
    inline._inline.docPr.set("descr", alt_text)


caption_text = "Figure 3.8. CV-Job matching and Potential assessment flowchart"
document = Document(DOCX)
caption = find_exact(document, caption_text)
old_image = previous_paragraph(caption)
if not old_image._p.xpath(".//w:drawing"):
    raise RuntimeError("The Figure 3.8 caption is not preceded by an image")
remove_paragraph(old_image)
image_paragraph = caption.insert_paragraph_before()
configure_image_paragraph(image_paragraph)
add_picture(image_paragraph, IMAGE, caption_text)
document.save(DOCX)
print(f"Refreshed Figure 3.8 in {DOCX}")
