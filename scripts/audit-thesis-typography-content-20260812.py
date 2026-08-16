from __future__ import annotations

import json
import re
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
OUTPUT = ROOT / "Doc" / "working" / "thesis-typography-content-audit-20260812.json"


def clean(text: str) -> str:
    return " ".join(text.replace("\u00a0", " ").split())


def direct_font(run) -> str | None:
    if run.font.name:
        return run.font.name
    rpr = run._r.rPr
    if rpr is None or rpr.rFonts is None:
        return None
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        value = rpr.rFonts.get(qn(f"w:{attr}"))
        if value:
            return value
    return None


def direct_size(run) -> float | None:
    if run.font.size:
        return round(run.font.size.pt, 2)
    return None


def paragraph_alignment(paragraph) -> str:
    alignment = paragraph.alignment
    if alignment is None:
        alignment = paragraph.style.paragraph_format.alignment
    names = {
        WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
        WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
        WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
        WD_ALIGN_PARAGRAPH.DISTRIBUTE: "DISTRIBUTE",
        WD_ALIGN_PARAGRAPH.JUSTIFY_MED: "JUSTIFY_MED",
        WD_ALIGN_PARAGRAPH.JUSTIFY_HI: "JUSTIFY_HI",
        WD_ALIGN_PARAGRAPH.JUSTIFY_LOW: "JUSTIFY_LOW",
        WD_ALIGN_PARAGRAPH.THAI_JUSTIFY: "THAI_JUSTIFY",
    }
    return names.get(alignment, str(alignment))


def paragraph_features(paragraph) -> dict:
    xml = paragraph._p.xml
    text = paragraph.text
    visible_runs = [r for r in paragraph.runs if r.text.strip()]
    fonts = sorted({font for r in visible_runs if (font := direct_font(r))})
    sizes = sorted({size for r in visible_runs if (size := direct_size(r)) is not None})
    return {
        "text": clean(text),
        "style": paragraph.style.name,
        "alignment": paragraph_alignment(paragraph),
        "fonts": fonts,
        "sizes": sizes,
        "tabs": xml.count("<w:tab"),
        "line_breaks": xml.count("<w:br") - xml.count('w:type="page"'),
        "page_breaks": xml.count('w:type="page"'),
        "multiple_spaces": bool(re.search(r" {2,}", text)),
        "nbsp": text.count("\u00a0"),
        "longest_token": max((len(token) for token in text.split()), default=0),
        "word_count": len(re.findall(r"\b[\w'-]+\b", text, flags=re.UNICODE)),
        "character_spacing": "w:spacing" in xml,
        "fit_text": "w:fitText" in xml,
    }


def iter_all_paragraphs(document: Document):
    for index, paragraph in enumerate(document.paragraphs):
        yield f"body.p{index}", paragraph, False
    for table_index, table in enumerate(document.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            for cell_index, cell in enumerate(row.cells, 1):
                for paragraph_index, paragraph in enumerate(cell.paragraphs, 1):
                    yield f"table{table_index}.r{row_index}.c{cell_index}.p{paragraph_index}", paragraph, True
    for section_index, section in enumerate(document.sections, 1):
        for area_name, area in (("header", section.header), ("footer", section.footer), ("first_header", section.first_page_header), ("first_footer", section.first_page_footer)):
            for paragraph_index, paragraph in enumerate(area.paragraphs, 1):
                yield f"section{section_index}.{area_name}.p{paragraph_index}", paragraph, False


def main() -> None:
    document = Document(DOCX)
    paragraph_records = []
    font_counts = Counter()
    size_counts = Counter()
    alignment_counts = Counter()
    style_counts = Counter()
    suspects = defaultdict(list)

    for location, paragraph, in_table in iter_all_paragraphs(document):
        record = paragraph_features(paragraph)
        record["location"] = location
        record["in_table"] = in_table
        paragraph_records.append(record)
        style_counts[record["style"]] += 1
        alignment_counts[record["alignment"]] += 1
        font_counts.update(record["fonts"])
        size_counts.update(record["sizes"])

        if record["text"]:
            non_tnr = [font for font in record["fonts"] if font not in {"Times New Roman", "Cambria Math", "Symbol"}]
            if non_tnr:
                suspects["non_times_fonts"].append({k: record[k] for k in ("location", "style", "text", "fonts", "sizes")})
            if record["alignment"] in {"JUSTIFY", "DISTRIBUTE", "JUSTIFY_MED", "JUSTIFY_HI", "JUSTIFY_LOW"}:
                if record["tabs"] or record["line_breaks"] or record["longest_token"] >= 35 or in_table or record["word_count"] <= 8:
                    suspects["wide_word_spacing_risk"].append({k: record[k] for k in ("location", "style", "text", "alignment", "tabs", "line_breaks", "longest_token", "word_count", "in_table")})
            if record["multiple_spaces"] or record["nbsp"]:
                suspects["literal_spacing"].append({k: record[k] for k in ("location", "style", "text", "multiple_spaces", "nbsp")})
            if record["character_spacing"] or record["fit_text"]:
                suspects["character_spacing_or_fit_text"].append({k: record[k] for k in ("location", "style", "text", "character_spacing", "fit_text")})

    body = document.paragraphs
    headings = []
    current_h1 = None
    current_h2 = None
    section_stats = defaultdict(lambda: {"paragraphs": 0, "words": 0, "citations": 0, "long_paragraphs": 0})
    for index, paragraph in enumerate(body):
        text = clean(paragraph.text)
        style = paragraph.style.name
        if style == "Heading 1":
            current_h1 = text
            current_h2 = None
            headings.append({"index": index, "level": 1, "text": text})
        elif style == "Heading 2":
            current_h2 = text
            headings.append({"index": index, "level": 2, "text": text})
        elif style.startswith("Heading "):
            headings.append({"index": index, "level": int(style.split()[-1]), "text": text})

        if text and not style.startswith("Heading"):
            key = current_h2 or current_h1 or "FRONT MATTER"
            words = len(re.findall(r"\b[\w'-]+\b", text, flags=re.UNICODE))
            section_stats[key]["paragraphs"] += 1
            section_stats[key]["words"] += words
            section_stats[key]["citations"] += len(re.findall(r"\[\d+\]", text))
            section_stats[key]["long_paragraphs"] += int(words >= 140)

    sentence_map = defaultdict(list)
    for index, paragraph in enumerate(body):
        if paragraph.style.name.startswith("Heading"):
            continue
        for sentence in re.split(r"(?<=[.!?])\s+", clean(paragraph.text)):
            normalized = re.sub(r"\W+", " ", sentence.lower()).strip()
            if len(normalized.split()) >= 10:
                sentence_map[normalized].append(index)
    repeated_sentences = [
        {"sentence": sentence, "paragraphs": indexes}
        for sentence, indexes in sentence_map.items()
        if len(indexes) > 1
    ]

    text_all = "\n".join(p.text for p in body)
    result = {
        "document": str(DOCX),
        "paragraph_count_body": len(document.paragraphs),
        "table_count": len(document.tables),
        "drawing_count": len(document.element.body.xpath(".//w:drawing")),
        "style_counts": style_counts,
        "alignment_counts": alignment_counts,
        "direct_font_counts": font_counts,
        "direct_size_counts": {str(key): value for key, value in sorted(size_counts.items())},
        "headings": headings,
        "section_stats": section_stats,
        "suspects": suspects,
        "repeated_sentences": repeated_sentences,
        "invalid_characters": {
            "replacement": text_all.count("\ufffd"),
            "nul": text_all.count("\x00"),
            "soft_hyphen": text_all.count("\u00ad"),
            "nbsp": text_all.count("\u00a0"),
        },
    }
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT.write_text(json.dumps(result, ensure_ascii=False, indent=2, default=dict), encoding="utf-8")

    print(f"output={OUTPUT}")
    print(f"paragraphs={result['paragraph_count_body']} tables={result['table_count']} drawings={result['drawing_count']}")
    print(f"fonts={dict(font_counts)}")
    print(f"alignments={dict(alignment_counts)}")
    print(f"non_times_fonts={len(suspects['non_times_fonts'])}")
    print(f"wide_word_spacing_risk={len(suspects['wide_word_spacing_risk'])}")
    print(f"literal_spacing={len(suspects['literal_spacing'])}")
    print(f"character_spacing_or_fit_text={len(suspects['character_spacing_or_fit_text'])}")
    print(f"repeated_sentences={len(repeated_sentences)}")
    print(f"invalid_characters={result['invalid_characters']}")


if __name__ == "__main__":
    main()
