from __future__ import annotations

import hashlib
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"


def clean(text: str) -> str:
    return " ".join(text.split())


def configure_run(run) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)


def set_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    configure_run(run)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.2


def find_start(document: Document, prefix: str) -> Paragraph:
    matches = [p for p in document.paragraphs if clean(p.text).startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting {prefix!r}; found {len(matches)}")
    return matches[0]


def find_exact(document: Document, text: str) -> Paragraph:
    matches = [p for p in document.paragraphs if clean(p.text) == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph {text!r}; found {len(matches)}")
    return matches[0]


def replace_start(document: Document, prefix: str, text: str) -> None:
    set_text(find_start(document, prefix), text)


def delete_start(document: Document, prefix: str) -> None:
    paragraph = find_start(document, prefix)
    paragraph._p.getparent().remove(paragraph._p)


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    node = OxmlElement("w:p")
    paragraph._p.addnext(node)
    created = Paragraph(node, paragraph._parent)
    created.style = "Normal"
    set_text(created, text)
    return created


def section_payload(document: Document, heading_text: str) -> tuple[tuple[str, str], ...]:
    heading = find_exact(document, heading_text)
    level = int(heading.style.name.split()[-1])
    paragraphs = document.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p._p is heading._p)
    result = []
    for paragraph in paragraphs[start + 1 :]:
        if paragraph.style.name.startswith("Heading "):
            next_level = int(paragraph.style.name.split()[-1])
            if next_level <= level:
                break
        result.append((paragraph.style.name, paragraph.text))
    return tuple(result)


def payload_hash(payload: tuple[tuple[str, str], ...]) -> str:
    raw = repr(payload).encode("utf-8")
    return hashlib.sha256(raw).hexdigest()


def main() -> None:
    document = Document(DOCX)
    protected = {
        title: payload_hash(section_payload(document, title))
        for title in (
            "2.2.2 Bag-of-Words Representation",
            "2.3.1 Term Frequency and Inverse Document Frequency",
            "2.3.2 Cosine Similarity",
        )
    }

    # Introduction: separate problem, motivation, scope, and contribution roles.
    replace_start(
        document,
        "IT recruitment requires candidates and recruiters to compare many details",
        "IT recruitment requires candidates and recruiters to compare skills, experience, location, and working conditions across many Jobs and CVs. As these records grow, manual comparison becomes slow and inconsistent [10].",
    )
    replace_start(
        document,
        "Most job portals support posting, searching, and applying",
        "Most Job portals support posting, search, and application but provide limited explanation of fit. CareerFit therefore separates CV–JD Matching, which compares one CV and Job, from profile-based Recommendation, which reflects broader Candidate preferences.",
    )
    replace_start(
        document,
        "Automation can reduce repetitive work, but a similarity score",
        "A Matching score alone should not trigger a recruitment action. CareerFit combines consent, policy conditions, existing state, and review points in a Human-in-the-Loop workflow.",
    )
    replace_start(
        document,
        "This thesis develops CareerFit IT AutoPilot, an IT-focused web platform",
        "This thesis develops CareerFit IT AutoPilot as an IT-focused decision-support platform with explainable Matching, feedback learning, controlled AutoFit, actionable email, Recruiter discovery, moderation, and auditability. It supports rather than replaces recruitment judgment.",
    )
    replace_start(
        document,
        "The first motivation is to reduce fragmentation",
        "CareerFit reduces fragmentation by combining Candidate Job, CV, and Application workflows with Recruiter Job, applicant, and Talent Pool workflows in one role-based system. Email provides an additional action channel.",
    )
    replace_start(
        document,
        "The second motivation is explainability",
        "TF-IDF and cosine similarity have less semantic capacity than dense models, but their terms and weights are inspectable. They provide a reproducible academic baseline for explaining scores, limitations, and feedback effects.",
    )
    replace_start(
        document,
        "The third motivation is controlled adaptation and automation",
        "AutoFit separates relevance from action. It checks Candidate consent and policy state, an eligible default CV, ACTIVE Job status, threshold, duplicates, and a per-run limit; notification delivery applies separate quota, cooldown, timezone, quiet-hour, and preference rules.",
    )
    replace_start(
        document,
        "Related work includes classical lexical retrieval",
        "Related work covers lexical retrieval, learned resume–Job representations, explainable recommendation, and human-centered governance [6]–[10], [14]. CareerFit addresses an integration gap by connecting transparent scoring, feedback learning, policy-controlled actions, and audit records in one academic prototype.",
    )
    replace_start(
        document,
        "This thesis focuses on an academic prototype for IT recruitment. Its functional scope",
        "The functional scope includes public Job discovery; Candidate, Recruiter, and Administrator workspaces; reviewed CV/Profile and employer/Job management; Matching and Recommendation; Applications, Talent Pool and invitations; feedback; AutoFit; notifications; analytics; reporting, moderation, and audit.",
    )
    replace_start(
        document,
        "The thesis does not attempt to implement a complete enterprise applicant tracking system",
        "The prototype excludes interview scheduling, offer, payroll, and employee lifecycle management; external Job-board application submission; unrestricted large-language-model decisions; distributed microservices; large message brokers; and semantic embedding models. These may be future work but are not presented as implemented features.",
    )

    # Chapter 2: remove generic explanation while retaining project-specific theory.
    replace_start(
        document,
        "Recruitment matching is an information-retrieval and decision-support problem",
        "Recruitment matching is an information-retrieval and decision-support task that combines structured conditions with inconsistent free text. Exact filters can enforce location, experience, or employment conditions, but they cannot by themselves rank overall fit between a CV and a Job.",
    )
    replace_start(
        document,
        "The same information supports different user tasks",
        "The same information supports different goals: Candidates rank Jobs, while Recruiters review applicants or discover Potential Candidates. These workflows may share representations and similarity functions, but a Recommendation is not an Application and a high score is not a hiring decision. Scores only support prioritization.",
    )
    replace_start(
        document,
        "Recruitment data also has domain-specific ambiguity",
        "Technical terms, titles, and skill families vary across companies and seniority levels, while public CV corpora are limited by privacy. The Djinni Recruitment Dataset shows the scale and value of anonymized recruitment data [6]. CareerFit uses a narrower IT scope, explicit fields, and normalized text so its behavior can be inspected.",
    )
    replace_start(
        document,
        "Before vectorization, text must be converted into a consistent form",
        "Before vectorization, text is extracted and normalized into deterministic tokens [2]. Recruitment preprocessing must preserve technical names, versions, and abbreviations because excessive filtering can remove important evidence.",
    )
    replace_start(
        document,
        "Bilingual data adds further complexity",
        "Vietnamese diacritics and multi-syllable expressions coexist with English technical terms. Structured attributes such as location, language, seniority, and salary mode remain available for validation or filtering instead of being forced into one text vector.",
    )
    replace_start(
        document,
        "Preprocessing quality directly affects every later score",
        "Missing OCR text, malformed files, very short descriptions, and duplicated boilerplate can distort scores. CareerFit therefore validates content before scoring: blocking errors stop unusable input, while warnings allow review and correction.",
    )
    replace_start(
        document,
        "Ranking quality cannot be described well by ordinary classification accuracy",
        "Ranking evaluation focuses on the ordered results users inspect. It therefore measures relevance at cutoff K, the position of useful items, and graded relevance when available [2].",
    )
    replace_start(
        document,
        "Precision@K is the proportion of the top K results",
        "Precision@K measures how much of the visible top-K list is relevant, Recall@K measures how much known relevant content it retrieves, and HitRate@K records whether the list contains at least one relevant result [2].",
    )
    replace_start(
        document,
        "For a query, reciprocal rank is 1/r",
        "Reciprocal rank is 1/r for the first relevant result at rank r, or zero when none is retrieved. MRR averages this value across queries, emphasizing the first useful result rather than additional relevant items [2].",
    )
    replace_start(
        document,
        "Discounted Cumulative Gain (DCG) supports graded relevance",
        "Discounted Cumulative Gain (DCG) supports graded relevance and discounts useful items at lower ranks. Normalized DCG divides the observed score by the ideal ordering, allowing comparison across queries [5]. One common formulation is:",
    )
    replace_start(
        document,
        "Metric values are meaningful only when the relevance labels",
        "Metric values require documented labels, candidate sets, cutoff K, and calculation rules. Synthetic improvement demonstrates only the designed scenario, so Chapter 4 reports the data source, holdout logic, repeated-run behavior, and validity limits with the scores.",
    )
    replace_start(
        document,
        "CareerFit uses a web client and a server-side application connected through HTTP APIs",
        "CareerFit separates presentation and navigation in the web client from authentication, authorization, validation, transactions, persistence, and protected business rules in the backend. Browser state is therefore not treated as authority for protected actions.",
    )
    replace_start(
        document,
        "REST is an architectural style characterized by constraints",
        "The backend exposes a REST-oriented HTTP API with stateless interaction and resource representations [11]. Chapter 3 documents the actual endpoints and state transitions rather than assuming that JSON over HTTP satisfies every REST constraint.",
    )
    replace_start(
        document,
        "Authentication establishes the identity associated with a request",
        "CareerFit combines authenticated identity, role authorization, and resource ownership checks. Sharing the Candidate role, for example, must not permit one account to access another Candidate's private CV.",
    )
    replace_start(
        document,
        "JSON Web Token is a compact, URL-safe format",
        "JWT carries signed claims such as subject, role, and expiry [12]. CareerFit must still validate the expected algorithm and claims, reload account state, and apply business authorization. Access tokens and purpose-specific email-action tokens are not interchangeable credentials.",
    )
    replace_start(
        document,
        "An actionable email link may be opened by the intended user",
        "An actionable link can be forwarded or opened by a mail scanner. CareerFit therefore uses a non-mutating GET confirmation followed by deliberate POST execution. Only the SHA-256 token hash is stored, and expiry and redeemed-state checks prevent normal reuse. Production still requires rate limiting, origin controls, secret rotation, and delivery monitoring.",
    )
    replace_start(
        document,
        "Audit logging records security-relevant and business-relevant events",
        "CareerFit audit records identify the actor, action, target, source channel, result, timestamp, and relevant policy or score context. Sensitive CV content and credentials are excluded unless a documented need exists. This supports investigation without treating ordinary debug logs as business evidence [13].",
    )
    replace_start(
        document,
        "An append-oriented audit history supports traceability",
        "Traceability still depends on database permissions, retention, integrity, and access control. Operational debug logs and domain audit records may therefore use different schemas and retention rules; Chapter 3 documents the implemented CareerFit mechanisms.",
    )

    # Chapter 3: add missing model interpretation, then remove design/implementation repetition.
    insert_after(
        find_exact(document, "3.3.2 Conceptual Data Model"),
        "The conceptual model groups CareerFit information into identity, Career Profile, recruitment, Feedback, Automation, communication, and governance. A Candidate owns Profile and CV information, while a Recruiter owns an employer profile and Jobs. Matching links a CV and Job as decision-support evidence; Application records a separate recruitment action, so a score is never treated as an Application or final decision.",
    )
    insert_after(
        find_exact(document, "3.3.3 Logical Data Model"),
        "The logical model resolves these concepts into normalized entities and explicit cardinalities. Four diagrams keep the main ownership and lifecycle relationships readable: identity and Career Profile; recruitment and Matching; Automation and communication; and governance and analytics. Unique and foreign-key relationships protect one-account policies, one Candidate–Job Application, one CV–Job Matching, and traceable Feedback, invitation, Report, notification, and audit records.",
    )
    delete_start(document, "Appendix C lists every active column")
    replace_start(
        document,
        "The backend is stateless at the HTTP session layer",
        "CareerFit uses stateless bearer authentication and combines role checks with service-level ownership rules. Public access is limited to explicitly permitted operations, while Candidate, Recruiter, Administrator, and Automation workflows require the appropriate authenticated context.",
    )
    replace_start(
        document,
        "The security configuration disables CSRF because the application uses stateless bearer authentication",
        "Security headers, CORS restrictions, and HTTPS deployment reduce browser and transport risks but do not replace ownership checks or production controls. Development access to API documentation and metrics must be restricted in production.",
    )
    replace_start(
        document,
        "The optional backend container exposes port 8080",
        "The optional backend container exposes port 8080, waits for PostgreSQL health, and mounts persistent CV storage. The React/Vite frontend runs separately during development. Environment variables provide database, JWT, URL/CORS, mail, OCR, and storage configuration; the backend image includes the project's Vietnamese/English Tesseract setup.",
    )
    replace_start(
        document,
        "CareerFit is implemented as two application projects and a shared runtime environment",
        "CareerFit consists of a Java 21/Spring Boot 3.2.5 backend and a React 18/TypeScript frontend built with Vite. PostgreSQL, Flyway, and Docker Compose provide persistence and the reproducible local runtime. Domain packages share the modular-monolith process described in Section 3.1.",
    )
    replace_start(
        document,
        "The backend package root is com.careerfit.backend",
        "The backend package root is com.careerfit.backend. Domain packages cover authentication, career data, recruitment, Matching and Recommendation, Feedback, Automation and Notification, Reporting and Administration, Audit, and Settings. Shared configuration, security, validation, exceptions, and text utilities remain cross-cutting.",
    )
    replace_start(
        document,
        "RecommendationService serves personalized Job results",
        "RecommendationService produces Candidate Job recommendations from Profile/preferences and available Matching context, while similar-Job retrieval remains public. Results are bounded to eligible visible Jobs and may return a limited-data or empty state when suitable information is unavailable. Recommendation stays separate from Application: viewing or skipping a result changes no recruitment state, and creation requires an explicit request or a separately enabled AutoFit policy.",
    )
    replace_start(
        document,
        "JPA entities map the domain tables",
        "Flyway V1–V25 defines the schema, and Hibernate validates the mappings. V25 adds content_report, moderation indexes and counters, the unique pending-report rule, and BANNED Job/CV states.",
    )
    replace_start(
        document,
        "Controllers return shared API envelopes and DTOs",
        "Controllers return validated DTOs in shared API envelopes. Service exceptions distinguish not found, forbidden, bad request, and conflict outcomes; springdoc publishes OpenAPI. Catalogue endpoints use pagination or bounded queries instead of unbounded payloads.",
    )
    replace_start(
        document,
        "API-driven pages do not replace failed responses with mock Job data",
        "API-driven pages use server-side filtering and pagination and display real errors instead of mock Job data. Reporting uses fixed reasons, optional comments, pending-state badges, and an Administrator moderation queue. Detailed interfaces appear in Section 3.4 and their workflows are evaluated in Chapter 4.",
    )
    replace_start(
        document,
        "The default local runtime starts PostgreSQL through Docker Compose",
        "The implementation follows the deployment topology in Section 3.6. Environment variables select host or container database addresses and configure JWT, CORS, application URL, mail, OCR, and CV storage. The Vite development server calls /api or a configured backend base URL.",
    )

    # Chapter 4: retain evidence and remove repeated descriptions.
    replace_start(
        document,
        "The suites covered application context, API contracts, security",
        "Coverage includes security, CV and Job lifecycles, Matching and Recommendation, Feedback, Automation, Applications, Talent Pool, analytics, reporting, moderation, and the controlled algorithm experiment.",
    )
    replace_start(
        document,
        "The final isolated benchmark completed without StaleObjectStateException",
        "The final isolated benchmark completed without StaleObjectStateException. Feedback learning starts after transaction commit, and benchmark setup removes prior Matchings and clears the persistence context before recomputation. Earlier runs showed a background conflict even when foreground assertions passed, so both test results and asynchronous logs must be checked.",
    )
    delete_start(document, "Earlier benchmark runs showed a background database conflict")
    replace_start(
        document,
        "The public Job API returned HTTP 200",
        "The public Job API, aggregate health, liveness, and readiness endpoints returned HTTP 200 with status UP, and Prometheus metrics were available in the local profile. Mail health is disabled when application mail is disabled, preventing an intentionally absent provider from making aggregate health misleading. Production monitoring still requires protected component details, alerting, and an external mail check when email is enabled.",
    )
    delete_start(document, "The final runtime check returned HTTP 200")
    replace_start(
        document,
        "The August 7 evaluation passed all 141 backend tests",
        "All backend, frontend, and integrated Chrome checks passed, and the controlled Rocchio benchmark reproduced its expected synthetic improvement. The evidence supports a functioning academic prototype, not production readiness or hiring effectiveness.",
    )

    # Conclusion: keep objective traceability and interpretation, remove repeated inventories.
    replace_start(
        document,
        "CareerFit IT AutoPilot was implemented as an IT recruitment prototype with public Job discovery",
        "CareerFit IT AutoPilot was implemented as an IT recruitment prototype that combines public Job discovery, role-based workflows, reviewed CV processing, separate Matching and Recommendation, controlled Automation, moderation, and auditability.",
    )
    replace_start(
        document,
        "The refreshed evaluation provides evidence for several technical outcomes",
        "The August 7 evaluation passed 141 backend tests, frontend type, lint, build, and bundle checks, and 46 integrated Chrome tests across the principal role workflows and moderation contracts.",
    )
    replace_start(
        document,
        "The controlled Rocchio benchmark showed the expected behavior",
        "On the synthetic benchmark, Rocchio increased nDCG@5 from 0.037737 to 0.837737 and Recall@5 and HitRate@5 from 0.06 to 0.86. This demonstrates the designed adaptation behavior, not expected performance on real recruitment data.",
    )
    replace_start(
        document,
        "The August 7 refresh synchronized the new content-reporting feature",
        "The refreshed prototype includes content reporting, Flyway V25, BANNED Job/CV states, and Administrator moderation. It remains suitable for local thesis demonstration rather than production release.",
    )
    replace_start(
        document,
        "The role-based workflow objective was achieved at prototype level",
        "The role-based platform objective was achieved at prototype level. Guests browse public Jobs; Candidates manage CVs, results, Applications, Feedback, AutoFit, and Reports; Recruiters manage Jobs, applicants, Talent Pool, invitations, and visible-CV Reports; Administrators manage access and moderation. The 46 Chrome tests cover representative flows, not every route, browser, or usability condition.",
    )
    replace_start(
        document,
        "The project implements uploaded and manually created CVs",
        "CareerFit supports uploaded and manual CVs, format and magic-byte validation, PDF/DOCX extraction, OCR fallback, review and confirmation, processing states, quality signals, and structured Job validation. Tests cover extraction, validation, Job behavior, and integration contracts. Local storage, external Tesseract, processing limits, and deterministic tokenization remain prototype constraints.",
    )
    delete_start(document, "The implementation is limited by local storage")
    replace_start(
        document,
        "CareerFit implements a static-corpus TF-IDF vectorizer",
        "The project implements a static-corpus TF-IDF/cosine baseline, direct labels and reasons, a separate Potential assessment, and profile-oriented Recommendation. Tests confirm scoring and API behavior, but punctuation-sensitive technologies, synonyms, Vietnamese phrases, and career context remain limitations. Displayed percentages represent normalized lexical similarity, not hiring probability.",
    )
    delete_start(document, "The objective was achieved as an interpretable lexical baseline")
    replace_start(
        document,
        "The system records GOOD_MATCH, POTENTIAL, BAD_MATCH, and NOT_INTERESTED feedback",
        "CareerFit records typed Feedback and uses Rocchio (α=1.0, β=0.75, γ=0.15) to rebuild learned Job vectors from the original vector and current Feedback. Stale Matchings are recomputed asynchronously. The controlled benchmark reproduced the intended improvement without the earlier background conflict, but it does not prove production concurrency or ranking quality.",
    )
    delete_start(document, "The controlled benchmark confirms the intended Rocchio behavior")
    replace_start(
        document,
        "AutoFit policies store thresholds, enablement, notification preferences",
        "AutoFit separates Matching evidence from Application authorization and checks policy state, default CV, ACTIVE Job, threshold, duplicates, and a three-Application per-run limit. Notification delivery uses separate policy guards. Actionable email stores hashed expiring tokens and follows GET confirmation with POST execution. Rate limiting, secret rotation, provider testing, and deployment controls remain future work.",
    )
    delete_start(document, "This objective is achieved as configurable prototype automation")
    replace_start(
        document,
        "Major application, feedback, automation, content-report, moderation",
        "Structured audit records cover major Application, Feedback, Automation, moderation, and administrative events. Unit, integration, security, algorithm, and E2E evidence supports the prototype, although complete audit coverage and metadata consistency have not been formally established.",
    )
    delete_start(document, "Audit coverage has not been formally proven")
    replace_start(
        document,
        "The earlier background StaleObjectStateException showed why test counts",
        "The earlier background StaleObjectStateException showed that passing assertions do not guarantee clean asynchronous execution. Reliable evidence must include logs, health components, output files, and side effects. CI should control schedulers, wait for background work, fail on uncaught exceptions, and archive logs while protecting operational details.",
    )
    delete_start(document, "Future CI should continue to fail on uncaught background exceptions")
    replace_start(
        document,
        "The controlled dataset is synthetic and designed to show how feedback changes ranking",
        "The synthetic dataset demonstrates Feedback-driven ranking behavior but lacks organic language, Recruiter disagreement, demographic analysis, misleading CVs, and changing market terms. TF-IDF uses a static corpus and simple tokenization, while the versioned Potential knowledge base still depends on manually defined aliases, weights, families, and thresholds.",
    )
    replace_start(
        document,
        "The 141 backend tests do not prove complete path coverage",
        "The 141 backend tests do not prove complete path coverage. Browser evaluation covers 46 desktop-Chrome cases without independent users. Concurrency, capacity, accessibility, cross-browser behavior, real moderation operations, and hiring outcomes remain unevaluated.",
    )
    replace_start(
        document,
        "Notification action tokens are hashed",
        "Email-action tokens are hashed and state changes require POST confirmation, but frontend access tokens remain in sessionStorage. Local CV storage lacks proven malware scanning, encryption, retention, and recovery controls. Public management endpoints, privacy, fairness, and personal-data governance remain production concerns.",
    )
    replace_start(
        document,
        "The current implementation includes after-commit matching",
        "The implementation includes after-commit processing, CV review, versioned Potential rules, server-side catalogues, policy guards, and moderation. Audit metadata is not fully centralized, the evaluated worktree is not a fixed release, and moderation has no appeal or restoration workflow.",
    )
    replace_start(
        document,
        "Future security work should add rate limiting",
        "Future security work should add rate limiting, stronger sessions, protected management endpoints, malware scanning, encrypted CV storage, retention rules, backup recovery, moderation appeals, and real-provider email testing.",
    )
    replace_start(
        document,
        "The matching baseline can then be extended using a hybrid approach",
        "Matching can progress from domain-aware tokenization and aliases to hybrid lexical, embedding, and structured-constraint retrieval. New models should be compared with the current baseline on a frozen, independently labeled dataset rather than assumed superior because they are more complex.",
    )
    replace_start(
        document,
        "Evaluation should expand to recruiter-labeled CV–JD pairs",
        "Evaluation should add Recruiter-labeled pairs, temporal holdouts, agreement analysis, hard negatives, fairness checks, calibrated labels, user studies, cross-browser E2E, load and failure testing, backup/restore, email delivery, and security testing.",
    )
    replace_start(
        document,
        "Operational development should move CVs to protected object storage",
        "Operational work should move CVs to protected object storage, centralize audit generation and redaction, strengthen token/session architecture, and bind each release to a clean commit and evidence archive. External ATS or Job-board integration requires explicit consent, ownership, recovery, and audit contracts.",
    )
    replace_start(
        document,
        "CareerFit demonstrates controlled IT recruitment automation by connecting job discovery",
        "CareerFit demonstrates controlled IT recruitment automation by connecting Job discovery, explainable Matching, profile-based Recommendation, Feedback learning, policy-driven actions, moderation, and audit records. Its local evaluation supports a tested Human-in-the-Loop prototype, not a system ready to make real hiring decisions.",
    )
    delete_start(document, "Local evaluation passed 141 backend tests")

    # Confirm the three explicitly protected theoretical sections are byte-for-byte
    # identical at the paragraph/style level.
    for title, before_hash in protected.items():
        after_hash = payload_hash(section_payload(document, title))
        if after_hash != before_hash:
            raise RuntimeError(f"Protected section changed: {title}")

    # Request Word to refresh generated fields at the next open/automation pass.
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    document.save(DOCX)
    print(f"updated={DOCX}")
    print("protected_sections=3")


if __name__ == "__main__":
    main()
