from pathlib import Path
from docx import Document


FILES = [
    Path(r"C:\Users\PhamHuuHwng\Downloads\mau-ve-data.docx"),
    Path(r"C:\Users\PhamHuuHwng\Downloads\mau-ve-data2.docx"),
    Path(r"C:\CODING\Thesis\Doc\CareerFit-Thesis-Report.docx"),
]


def clean(text: str) -> str:
    return " ".join(text.replace("\u00a0", " ").split())


for path in FILES:
    doc = Document(path)
    print("\n" + "=" * 100)
    print(path)
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} inline_shapes={len(doc.inline_shapes)} sections={len(doc.sections)}")
    print("-- paragraphs likely related to data/database/model/entity/schema --")
    for i, paragraph in enumerate(doc.paragraphs):
        text = clean(paragraph.text)
        style = paragraph.style.name if paragraph.style else ""
        haystack = f"{style} {text}".lower()
        if text and any(term in haystack for term in (
            "data", "database", "conceptual", "logical", "physical", "cdm", "ldm", "pdm",
            "entity", "relationship", "schema", "dictionary", "table design", "data model",
        )):
            print(f"P{i:04d} [{style}] {text}")
    print("-- tables: dimensions and first non-empty rows --")
    for ti, table in enumerate(doc.tables):
        preview = []
        for row in table.rows[:4]:
            values = [clean(cell.text) for cell in row.cells]
            if any(values):
                preview.append(" | ".join(values))
        joined = " || ".join(preview)
        print(f"T{ti:03d} rows={len(table.rows)} cols={len(table.columns)} :: {joined[:900]}")
