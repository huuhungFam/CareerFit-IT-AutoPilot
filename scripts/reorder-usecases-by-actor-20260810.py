from __future__ import annotations

from pathlib import Path
import re
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-before-20260810-usecase-reorder.docx"

NEW_ORDER = [1, 2, 3, 4, 5, 6, 14, 8, 9, 10, 11, 7, 12, 13]
NEW_ID_BY_OLD = {
    1: 1, 2: 2, 3: 3, 4: 4, 5: 5, 6: 6,
    14: 7,
    8: 8, 9: 9, 10: 10, 11: 11,
    7: 12,
    12: 13,
    13: 14,
}
NAME_BY_OLD = {
    1: "Manage Career Profile",
    2: "Explore Jobs",
    3: "Manage Job Applications",
    4: "Provide Matching Feedback",
    5: "Review Personalized Career Insights",
    6: "Manage AutoFit",
    7: "Report Suspicious Recruitment Content",
    8: "Manage Employer Profile and Job Postings",
    9: "Review and Process Applicants",
    10: "Manage Talent Pool and Invitations",
    11: "Review Recruitment Analytics",
    12: "Administer Platform Access and Job Visibility",
    13: "Review and Resolve Content Reports",
    14: "Respond Through Actionable Email",
}


def set_paragraph_text(paragraph, text, size=13, italic=False):
    for child in list(paragraph._p):
        if child.tag.endswith("}pPr"):
            continue
        paragraph._p.remove(child)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.italic = italic


def simultaneous_id_replace(text):
    placeholders = {f"UC-{old:02d}": f"__CAREERFIT_UC_{old:02d}__" for old in NEW_ID_BY_OLD}
    for old_id, placeholder in placeholders.items():
        text = text.replace(old_id, placeholder)
    for old, new in NEW_ID_BY_OLD.items():
        text = text.replace(placeholders[f"UC-{old:02d}"], f"UC-{new:02d}")
    return text


def replace_ids_in_text_nodes(document):
    for text_node in document._element.xpath(".//w:t"):
        if text_node.text and "UC-" in text_node.text:
            text_node.text = simultaneous_id_replace(text_node.text)


def paragraph_for_element(document, element):
    for paragraph in document.paragraphs:
        if paragraph._p is element:
            return paragraph
    return None


def table_for_element(document, element):
    for table in document.tables:
        if table._tbl is element:
            return table
    return None


def find_use_case_blocks(document):
    body = document._element.body
    children = list(body.iterchildren())
    starts = []
    for index, element in enumerate(children):
        paragraph = paragraph_for_element(document, element)
        if paragraph is None or paragraph.style.name != "Heading 3":
            continue
        match = re.match(r"1\.5\.(\d+)\s+", paragraph.text.strip())
        if match:
            starts.append((int(match.group(1)), index))
    if [old for old, _ in starts] != list(range(1, 15)):
        raise RuntimeError(f"Unexpected use-case headings: {[old for old, _ in starts]}")

    summary_index = None
    for index, element in enumerate(children):
        paragraph = paragraph_for_element(document, element)
        if paragraph is not None and paragraph.style.name == "Heading 2" and paragraph.text.strip().startswith("1.6 "):
            summary_index = index
            break
    if summary_index is None:
        raise RuntimeError("Could not find the 1.6 Chapter Summary anchor")

    blocks = {}
    for position, (old_id, start) in enumerate(starts):
        end = starts[position + 1][1] if position + 1 < len(starts) else summary_index
        blocks[old_id] = children[start:end]
    return blocks, children[summary_index]


def update_block(document, block, old_id, new_id):
    headings = []
    captions = []
    tables = []
    for element in block:
        paragraph = paragraph_for_element(document, element)
        if paragraph is not None:
            if paragraph.style.name == "Heading 3" and paragraph.text.strip().startswith("1.5."):
                headings.append(paragraph)
            elif paragraph.style.name == "Table Caption" and "Use case -" in paragraph.text:
                captions.append(paragraph)
        table = table_for_element(document, element)
        if table is not None:
            tables.append(table)
    if len(headings) != 1 or len(captions) != 1 or len(tables) != 1:
        raise RuntimeError(f"Unexpected block structure for old UC-{old_id:02d}")

    name = NAME_BY_OLD[old_id]
    set_paragraph_text(headings[0], f"1.5.{new_id} {name}", 13)
    set_paragraph_text(captions[0], f"Table 1.{new_id + 4}. Use case - {name}", 11, italic=True)
    captions[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    captions[0].paragraph_format.keep_with_next = True
    captions[0].paragraph_format.keep_together = True

    table = tables[0]
    for row in table.rows:
        if row.cells[0].text.strip() == "Use Case ID":
            set_paragraph_text(row.cells[1].paragraphs[0], f"UC-{new_id:02d}", 13)
            break
    else:
        raise RuntimeError(f"Missing Use Case ID row for old UC-{old_id:02d}")


def add_or_update_figure_note(document):
    note_text = (
        "NOTE: Figure 1.2 is temporarily retained without redrawing. When it is manually redrawn, "
        "use the revised ordering: UC-01–UC-07 for Candidate functions, UC-08–UC-11 for Recruiter "
        "functions, UC-12 for the shared Candidate–Recruiter reporting function, and UC-13–UC-14 "
        "for Administrator functions."
    )
    captions = [paragraph for paragraph in document.paragraphs
                if paragraph.style.name == "Figure Caption"
                and paragraph.text.strip().startswith("Figure 1.2.")]
    if len(captions) != 1:
        raise RuntimeError(f"Expected one Figure 1.2 caption, found {len(captions)}")
    caption = captions[0]
    next_element = caption._p.getnext()
    next_paragraph = paragraph_for_element(document, next_element) if next_element is not None else None
    if next_paragraph is not None and next_paragraph.text.strip().startswith("NOTE: Figure 1.2"):
        note = next_paragraph
        set_paragraph_text(note, note_text, 13, italic=True)
    else:
        note_element = OxmlElement("w:p")
        caption._p.addnext(note_element)
        note = paragraph_for_element(document, note_element)
        if note is None:
            # python-docx refreshes the paragraph collection lazily only through a wrapper.
            from docx.text.paragraph import Paragraph
            note = Paragraph(note_element, caption._parent)
        note.style = document.styles["Normal"]
        set_paragraph_text(note, note_text, 13, italic=True)
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note.paragraph_format.keep_together = True


def main():
    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(REPORT, BACKUP)

    document = Document(REPORT)
    replace_ids_in_text_nodes(document)
    blocks, summary_anchor = find_use_case_blocks(document)

    for old_id, block in blocks.items():
        update_block(document, block, old_id, NEW_ID_BY_OLD[old_id])

    for block in blocks.values():
        for element in block:
            element.getparent().remove(element)
    for old_id in NEW_ORDER:
        for element in blocks[old_id]:
            summary_anchor.addprevious(element)

    traceability = [paragraph for paragraph in document.paragraphs
                    if paragraph.text.strip().startswith("The core traceability chain is:")]
    if len(traceability) != 1:
        raise RuntimeError(f"Expected one traceability paragraph, found {len(traceability)}")
    set_paragraph_text(
        traceability[0],
        "The core traceability chain is: UC-01 to Career Profile, Portfolio, CV review, default, and deletion tests; "
        "UC-02 to public Job catalogue, detail, employer, urgent, and similar-Job contracts; UC-03 to Application "
        "history, submission, withdrawal, and invitation-response tests; UC-04 to Feedback and learning tests; "
        "UC-05 to separate Matching, recommendation, and Candidate analytics contracts; UC-06 to AutomationPolicy "
        "and AutoApply tests; UC-07 to confirmation, Feedback/invitation redemption, replay, and expiry behavior; "
        "UC-08 to Employer Profile and Job lifecycle tests; UC-09 to applicant review and decision tests; UC-10 to "
        "Talent Pool, bookmark, invitation, and withdrawal flows; UC-11 to Recruiter analytics; UC-12 to Candidate Job "
        "and Recruiter visible-CV reporting; UC-13 to administrative User and Job visibility tests; and UC-14 to report "
        "moderation tests.",
        13,
    )

    add_or_update_figure_note(document)

    settings = document.settings._element
    update_fields = settings.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}updateFields")
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "true")

    document.save(REPORT)
    print(REPORT)
    print(BACKUP)


if __name__ == "__main__":
    main()
