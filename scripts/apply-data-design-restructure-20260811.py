from __future__ import annotations

import hashlib
import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-before-20260811-data-design.docx"
FIGURES = ROOT / "Doc" / "figures"
COLUMNS_FILE = ROOT / "Doc" / "working" / "careerfit-schema-physical-columns-20260811.txt"
CONSTRAINTS_FILE = ROOT / "Doc" / "working" / "careerfit-schema-constraints-20260811.txt"


TABLE_GROUPS = [
    ("Identity, Candidate, CV, and Portfolio", [
        "user_account", "user_settings", "candidate", "candidate_portfolio_link",
        "candidate_portfolio_project", "cv", "skills",
    ]),
    ("Employer, Job, Matching, and Recruitment", [
        "employer_profile", "job", "matching", "feedback", "application",
        "recommendation_interaction", "recruiter_cv_bookmark",
    ]),
    ("Automation, Email, and Notification", [
        "automation_policy", "email_action", "email_action_token", "notification_job",
        "notification_delivery_log",
    ]),
    ("Reporting, Analytics, and Audit", [
        "content_report", "analytics_event", "audit_log", "job_trend_snapshot",
        "job_market_snapshot",
    ]),
]

TABLE_PURPOSES = {
    "analytics_event": "Records product and recruitment events used by the implemented analytics views.",
    "application": "Stores a Candidate's application or Recruiter invitation state for a Job.",
    "audit_log": "Stores append-oriented evidence for important user, system, moderation, and automated actions.",
    "automation_policy": "Stores one account-level AutoFit, digest, email, and notification policy.",
    "candidate": "Stores the Candidate profile, preferences, experience, and compatibility AutoFit fields.",
    "candidate_portfolio_link": "Stores external portfolio links owned by a Candidate.",
    "candidate_portfolio_project": "Stores structured project entries in a Candidate portfolio.",
    "content_report": "Stores Candidate or Recruiter reports about a Job or visible CV and the moderation outcome.",
    "cv": "Stores CV metadata, extracted and reviewed content, processing state, and reporting state.",
    "email_action": "Tracks creation and delivery lifecycle metadata for an outbound actionable email.",
    "email_action_token": "Stores the hash and lifecycle of a one-time actionable-email token.",
    "employer_profile": "Stores the Recruiter's company profile and public employer information.",
    "feedback": "Stores Candidate or Recruiter judgments about a CV-Job Matching result.",
    "job": "Stores the Job Description, structured recruitment fields, vectors, publication state, and counters.",
    "job_market_snapshot": "Stores date-based aggregate Job-market counts and distributions.",
    "job_trend_snapshot": "Stores date-based view and application counters for an individual Job.",
    "matching": "Stores one scored relationship between a CV and Job, including labels, reasons, and Potential data.",
    "notification_delivery_log": "Records notification-policy outcomes used for quota, deduplication, cooldown, and audit evidence.",
    "notification_job": "Stores queued notification work and its retry lifecycle.",
    "recommendation_interaction": "Stores Candidate interactions with personalized Job recommendations.",
    "recruiter_cv_bookmark": "Stores a Recruiter's saved Candidate/CV for a specific owned Job.",
    "skills": "Stores the searchable and normalized skill catalogue used by profile and Job inputs.",
    "user_account": "Stores authentication identity, role, activation state, verification state, and language.",
    "user_settings": "Stores supporting account settings that do not require their own domain table.",
}


SPECIAL_DESCRIPTIONS = {
    "email": "Normalized email address used as the account login identifier.",
    "password_hash": "BCrypt password hash; the original password is not stored.",
    "full_name": "Display name of the user.",
    "email_verified": "Indicates whether the account email has been verified.",
    "preferred_language": "Preferred interface or content language code.",
    "settings": "JSON object containing supporting account settings.",
    "desired_title": "Job title preferred by the Candidate.",
    "desired_seniority": "Seniority level preferred by the Candidate.",
    "desired_skills": "JSON list of skills preferred by the Candidate.",
    "desired_work_model": "Preferred onsite, hybrid, or remote work model.",
    "desired_salary_min": "Lower bound of the Candidate's preferred salary range.",
    "desired_salary_max": "Upper bound of the Candidate's preferred salary range.",
    "desired_salary_currency": "Currency used for the Candidate's preferred salary range.",
    "years_of_experience": "Candidate's declared total years of professional experience.",
    "about_me": "Candidate's profile introduction.",
    "avatar_url": "Location of the Candidate avatar image.",
    "display_name": "User-facing name of the CV.",
    "raw_text": "Current CV text used by the review and processing workflow.",
    "original_raw_text": "Original extracted or manually entered CV text before review edits.",
    "parsed_summary": "Summary derived from the reviewed CV content.",
    "top_skills": "JSON list of the main skills extracted or confirmed for the CV.",
    "extracted_terms": "JSON representation of normalized terms used by scoring.",
    "review_sections": "JSON object containing editable CV review sections.",
    "review_issues": "JSON list of validation warnings or issues presented during review.",
    "file_path": "Protected storage path of the uploaded CV file.",
    "file_original_name": "Original client filename of the uploaded CV.",
    "failure_reason": "Actor-visible reason recorded when CV processing fails.",
    "last_scored_at": "Timestamp of the most recent completed CV scoring operation.",
    "pending_report_count": "Number of unresolved reports currently associated with the Job or CV.",
    "canonical_key": "Stable canonical identifier of a skill.",
    "normalized_name": "Normalized skill name used for uniqueness and search.",
    "search_text": "Searchable text assembled for skill lookup.",
    "popularity": "Observed usage count used for skill ordering.",
    "company_name": "Public company name shown with the Employer Profile.",
    "slug": "Unique URL-friendly identifier of the Employer Profile.",
    "logo_url": "Location of the company logo image.",
    "cover_url": "Location of the company cover image.",
    "company_size": "Declared size category of the employer.",
    "benefits": "JSON list of employer benefits displayed to users.",
    "is_featured": "Indicates whether the employer is highlighted in public views.",
    "original_text": "Original Job Description text used for display and processing.",
    "required_skills": "JSON list of skills required by the Job.",
    "nice_to_have_skills": "JSON list of preferred but non-mandatory skills.",
    "seniority_level": "Seniority level requested by the Job.",
    "employment_type": "Employment arrangement associated with the Job.",
    "remote_type": "Onsite, hybrid, or remote arrangement of the Job.",
    "domain": "IT domain or specialization assigned to the Job.",
    "salary_mode": "Mode controlling how salary information is stored and displayed.",
    "salary_min": "Lower bound of the Job salary range when applicable.",
    "salary_max": "Upper bound of the Job salary range when applicable.",
    "salary_currency": "Currency of the Job salary values.",
    "salary_type": "Time basis of salary values, such as monthly or yearly.",
    "salary_is_visible": "Indicates whether salary information is visible to users.",
    "salary_display_text": "Preformatted salary text used by the interface when available.",
    "learned_profile_vector": "JSON sparse Job vector updated from relevance feedback.",
    "tfidf_vector": "JSON sparse TF-IDF vector built from the Job Description.",
    "deadline": "Legacy or date-only Job application deadline.",
    "application_deadline": "Timestamp after which the Job no longer accepts applications.",
    "applicant_count": "Stored number of applicants associated with the Job.",
    "application_count": "Stored application count used by Job popularity and analytics flows.",
    "view_count": "Stored number of Job-detail views.",
    "source_platform": "External platform from which the Job was imported, when applicable.",
    "source_url": "Original URL of an imported Job.",
    "scraped_at": "Timestamp when an external Job record was collected.",
    "external_hash": "Stable hash used to detect duplicate imported Jobs.",
    "is_urgent": "Indicates that the Recruiter marked the Job for urgent hiring.",
    "raw_score": "Cosine-similarity score in the implemented zero-to-one range.",
    "normalized_score": "User-facing Matching score normalized to zero through one hundred.",
    "label": "Implemented LOW, MEDIUM, HIGH, or POTENTIAL Matching label.",
    "is_potential": "Indicates that the separate Potential assessment applies.",
    "match_reasons": "JSON reasons explaining the direct CV-Job score.",
    "potential_reason": "JSON explanation produced by the Potential assessment.",
    "needs_recompute": "Indicates that the persisted Matching should be recalculated.",
    "recruiter_label": "Recruiter decision label attached to the Matching when present.",
    "feedback_type": "Business judgment submitted for the Matching.",
    "is_auto_applied": "Indicates that the Application was created by AutoFit rather than manual submission.",
    "cover_letter": "Optional cover letter submitted with the Application.",
    "recruiter_notes": "Recruiter notes associated with the Application decision workflow.",
    "applied_at": "Timestamp when the Application was created.",
    "action": "Candidate interaction recorded for a recommended Job.",
    "token_hash": "SHA-256 hash of the one-time email action token.",
    "expires_at": "Timestamp after which the email action token is rejected.",
    "redeemed_at": "Timestamp when the email action token was successfully redeemed.",
    "recipient_user_id": "Account that receives the tracked email or notification.",
    "recipient_id": "Account allowed to redeem the one-time email action.",
    "subject": "Subject line of the outbound actionable email.",
    "template_name": "Template identifier used to construct the outbound email.",
    "sent_at": "Timestamp when the email was marked as sent.",
    "opened_at": "Timestamp when the email lifecycle recorded an open event.",
    "executed_at": "Timestamp when the referenced email action completed.",
    "auto_apply_enabled": "Indicates whether AutoFit auto-application is enabled for the account.",
    "auto_apply_threshold": "Minimum normalized score configured for AutoFit application eligibility.",
    "auto_invite_enabled": "Stored policy flag for automated invitation support.",
    "daily_digest_enabled": "Indicates whether the daily digest is enabled.",
    "daily_digest_time": "Local time at which the digest is intended to run.",
    "user_timezone": "Time zone used to interpret digest and quiet-hour settings.",
    "quiet_hours_enabled": "Indicates whether notification quiet hours are active.",
    "quiet_hours_start": "Local start time of the configured quiet-hours interval.",
    "quiet_hours_end": "Local end time of the configured quiet-hours interval.",
    "job_scan_enabled": "Indicates whether scheduled Job scanning is enabled.",
    "job_scan_frequency_hours": "Configured interval between scheduled Job scans.",
    "high_match_email_enabled": "Indicates whether high-match email notifications are enabled.",
    "high_match_threshold": "Minimum score used by the high-match email policy.",
    "max_email_per_day": "Maximum number of policy-controlled emails allowed per day.",
    "notification_cooldown_hours": "Minimum cooldown between related notification categories.",
    "replacement_after_skip_enabled": "Indicates whether a replacement recommendation may follow a skipped item.",
    "replacement_delay_minutes": "Delay before a configured replacement notification is eligible.",
    "email_action_enabled": "Indicates whether actionable email links are enabled.",
    "email_notifications_enabled": "Global account switch for email notifications.",
    "paused_until": "Timestamp until which account-level automation is paused.",
    "job_type": "Type of notification work stored in the queue.",
    "payload": "JSON payload required to execute the queued notification work.",
    "retry_count": "Number of processing retries already attempted.",
    "next_retry_at": "Timestamp when the failed notification job may be retried.",
    "email_type": "Notification or email category evaluated by the delivery policy.",
    "context_key": "Optional business context used for notification deduplication.",
    "reason": "Business reason associated with the report or delivery outcome.",
    "comment": "Optional explanatory comment supplied by the reporter.",
    "resolution_note": "Administrator note explaining the report resolution.",
    "resolved_at": "Timestamp when the content report was resolved.",
    "event_type": "Analytics event category.",
    "actor_role": "Role of the user that produced the event or feedback.",
    "subject_type": "Type of business object associated with the analytics event.",
    "subject_id": "Identifier of the business object associated with the analytics event.",
    "occurred_at": "Timestamp when the analytics event occurred.",
    "actor_type": "Distinguishes user-generated and system-generated audit actions.",
    "action_type": "Business or system action recorded by this row.",
    "target_type": "Type of Job, CV, or other target referenced by the record.",
    "target_id": "Identifier of the polymorphic target referenced by the record.",
    "result": "Recorded success, failure, or denial outcome of the audited action.",
    "ip_address": "Request IP address retained as audit metadata when available.",
    "user_agent": "Browser or client identification retained as audit metadata when available.",
    "snapshot_date": "Date represented by the aggregate or per-Job snapshot.",
    "total_posted_jobs": "Total number of posted Jobs counted for the snapshot.",
    "active_jobs": "Number of active Jobs counted for the snapshot.",
    "new_jobs": "Number of newly posted Jobs counted for the snapshot.",
    "employer_count": "Number of employers represented in the snapshot.",
    "distribution_by_role": "JSON distribution of Jobs by role category.",
    "distribution_by_salary": "JSON distribution of Jobs by salary category.",
}


def clean_text(value: str) -> str:
    return " ".join(value.replace("\u00a0", " ").split())


def load_columns():
    tables = {}
    for raw in COLUMNS_FILE.read_text(encoding="utf-8-sig").splitlines():
        table, column, sql_type, nullable, default = raw.split("|", 4)
        tables.setdefault(table, []).append({
            "name": column,
            "type": sql_type,
            "nullable": nullable,
            "default": default,
            "keys": [],
            "checks": [],
        })
    return tables


def load_constraints(tables):
    seen = set()
    for raw in CONSTRAINTS_FILE.read_text(encoding="utf-8-sig").splitlines():
        parts = raw.split("|", 6)
        if len(parts) != 7:
            continue
        table, column, kind, name, ref_table, ref_column, definition = parts
        key = (table, column, kind, name)
        if key in seen or table not in tables:
            continue
        seen.add(key)
        target = next((c for c in tables[table] if c["name"] == column), None)
        if target is None:
            continue
        if kind == "p":
            target["keys"].append("PK")
        elif kind == "f":
            target["keys"].append(f"FK -> {ref_table}.{ref_column}")
            target["ref_table"] = ref_table
        elif kind == "u":
            target["keys"].append("UQ")
        elif kind == "c":
            target["checks"].append((name, definition))


def simplify_default(default: str) -> str:
    value = default.strip()
    if not value:
        return ""
    value = re.sub(r"::(?:character varying|jsonb|text)$", "", value)
    if value.upper() == "CURRENT_TIMESTAMP":
        return "now()"
    return value


def simplify_check(name: str, definition: str) -> str:
    values = []
    if "ARRAY[" in definition:
        values = list(dict.fromkeys(re.findall(r"'([^']+)'::character varying", definition)))
    if values:
        return "Allowed: " + ", ".join(values)
    if "normalized_score_range" in name or "threshold" in name:
        return "Range 0-100"
    if "raw_score_range" in name:
        return "Range 0-1"
    if "non_negative" in name or "email_limits" in name:
        return "Must be non-negative"
    if "salary_range" in name:
        return "Minimum must not exceed maximum"
    if "salary_mode_fields" in name:
        return "Salary fields must match salary mode"
    if "frequency" in name:
        return "Frequency must be positive; delay non-negative"
    return name


def constraint_summary(column) -> str:
    parts = []
    default = simplify_default(column["default"])
    if default:
        parts.append("Default: " + default)
    for name, definition in column["checks"]:
        value = simplify_check(name, definition)
        if value not in parts:
            parts.append(value)
    return "; ".join(parts) if parts else "-"


def description(table: str, column) -> str:
    name = column["name"]
    if name == "id":
        noun = TABLE_PURPOSES[table].split(".")[0].replace("Stores ", "").replace("Records ", "record for ")
        return "Unique identifier for this record."
    if "ref_table" in column:
        return f"References the related {column['ref_table']} record."
    if name in SPECIAL_DESCRIPTIONS:
        return SPECIAL_DESCRIPTIONS[name]
    if name == "status":
        return f"Current lifecycle state of the {table.replace('_', ' ')} record."
    if name == "role":
        return "Role used by authorization and role-specific workflows."
    if name == "source":
        return "Channel or origin from which the record was created."
    if name == "metadata":
        return "Structured JSON metadata associated with the record."
    if name == "created_at":
        return "Timestamp when the record was created."
    if name == "updated_at":
        return "Timestamp when the record was last updated."
    if name == "version":
        return "Optimistic-lock version used to detect conflicting updates."
    if name.startswith("is_"):
        return f"Indicates whether {name[3:].replace('_', ' ')} applies to the record."
    if name.endswith("_enabled"):
        return f"Indicates whether {name[:-8].replace('_', ' ')} is enabled."
    if name.endswith("_count"):
        return f"Stored count of {name[:-6].replace('_', ' ')} items."
    if name.endswith("_url"):
        return f"URL associated with the {name[:-4].replace('_', ' ')}."
    if name.endswith("_at"):
        return f"Timestamp associated with {name[:-3].replace('_', ' ')}."
    if name.endswith("_date"):
        return f"Date associated with {name[:-5].replace('_', ' ')}."
    if name.endswith("_name"):
        return f"Stored name of the {name[:-5].replace('_', ' ')}."
    return f"Stores the {name.replace('_', ' ')} value for this record."


def find_exact(document: Document, text: str) -> Paragraph:
    matches = [p for p in document.paragraphs if p.text.strip() == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph for {text!r}; found {len(matches)}")
    return matches[0]


def set_cell_text(cell, text: str, bold: bool = False, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill="E7E7E7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.allow_autofit = False
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width.inches * 1440)))
        grid.append(col)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(sum(w.inches for w in widths) * 1440)))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
            set_cell_margins(cell)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def insert_paragraph_before(anchor: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    paragraph = anchor.insert_paragraph_before(text)
    if style:
        paragraph.style = style
    return paragraph


def configure_heading(paragraph: Paragraph, page_break=False):
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    if page_break:
        paragraph.paragraph_format.page_break_before = True


def configure_body(paragraph: Paragraph):
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def configure_caption(paragraph: Paragraph, style="Figure Caption"):
    paragraph.style = style
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = style == "Table Caption"
    paragraph.paragraph_format.keep_together = True


def insert_figure(anchor: Paragraph, image_name: str, caption_text: str):
    image_p = insert_paragraph_before(anchor)
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.paragraph_format.page_break_before = True
    image_p.paragraph_format.keep_with_next = True
    image_p.paragraph_format.space_before = Pt(0)
    image_p.paragraph_format.space_after = Pt(0)
    inline = image_p.add_run().add_picture(str(FIGURES / image_name), width=Inches(6.0))
    inline._inline.docPr.set("title", caption_text)
    inline._inline.docPr.set("descr", caption_text)
    caption = insert_paragraph_before(anchor, caption_text)
    configure_caption(caption)


def request_field_refresh(document: Document):
    settings = document.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def table_digest(document: Document):
    payload = []
    for table in document.tables:
        payload.append([[clean_text(cell.text) for cell in row.cells] for row in table.rows])
    return hashlib.sha256(repr(payload).encode("utf-8")).hexdigest()


def replace_data_design(document: Document):
    heading = find_exact(document, "3.3 Data Design")
    next_heading = find_exact(document, "3.4 Security, Failure, and Consistency Design")
    constraint_caption = find_exact(document, "Table 3.2. Key integrity constraints")
    constraint_table = next(
        table for table in document.tables
        if table.rows and table.rows[0].cells[0].text.strip() == "Constraint"
        and table.rows[0].cells[1].text.strip() == "Purpose"
    )
    preserved_table = deepcopy(constraint_table._tbl)

    node = heading._p.getnext()
    while node is not None and node is not next_heading._p:
        following = node.getnext()
        node.getparent().remove(node)
        node = following

    h = insert_paragraph_before(next_heading, "3.3.1 Data Design Overview", "Heading 3")
    configure_heading(h)
    p = insert_paragraph_before(next_heading,
        "CareerFit uses PostgreSQL as its transactional data store. Flyway migrations V1-V25 are the authoritative schema history, while Hibernate ddl-auto=validate checks that the application mappings match the migrated schema. The final schema contains 24 active application tables and 293 columns; the earlier email_token table was removed by V16. JSONB is retained for variable collections and snapshots such as skills, vectors, reasons, review sections, notification payloads, and analytic distributions.")
    configure_body(p)
    p = insert_paragraph_before(next_heading,
        "The models below separate three levels of abstraction. The conceptual model shows business information without implementation detail. The logical models show normalized entities and cardinalities. The physical models use the final PostgreSQL table and column names. This separation prevents a high-level domain diagram from being mistaken for the deployed schema.")
    configure_body(p)

    h = insert_paragraph_before(next_heading, "3.3.2 Conceptual Data Model", "Heading 3")
    configure_heading(h)
    p = insert_paragraph_before(next_heading,
        "The conceptual data model focuses on the information needed by Candidates, Recruiters, Administrators, and the CareerFit decision-support workflows. It intentionally omits SQL types, indexes, implementation logs, and low-level token storage.")
    configure_body(p)
    insert_figure(next_heading, "data-cdm-careerfit-20260811.png", "Figure 3.2. CareerFit conceptual data model")

    h = insert_paragraph_before(next_heading, "3.3.3 Logical Data Model", "Heading 3")
    configure_heading(h, page_break=True)
    p = insert_paragraph_before(next_heading,
        "The logical model expands the conceptual objects into normalized entities. It is divided into four views so that keys, attributes, and relationships remain readable on A4 pages. Dashed relationships denote a logical or polymorphic association that is not represented by a direct foreign key in the final schema.")
    configure_body(p)
    insert_figure(next_heading, "data-ldm-identity-20260811.png", "Figure 3.3. CareerFit logical data model - identity and Career Profile")
    insert_figure(next_heading, "data-ldm-recruitment-20260811.png", "Figure 3.4. CareerFit logical data model - recruitment and Matching")
    insert_figure(next_heading, "data-ldm-communication-20260811.png", "Figure 3.5. CareerFit logical data model - automation and communication")
    insert_figure(next_heading, "data-ldm-governance-20260811.png", "Figure 3.6. CareerFit logical data model - governance and analytics")

    h = insert_paragraph_before(next_heading, "3.3.4 Physical Data Model", "Heading 3")
    configure_heading(h, page_break=True)
    p = insert_paragraph_before(next_heading,
        "The physical model reflects the final PostgreSQL schema after all successful Flyway migrations. The figures show physical table names, PostgreSQL types, primary and foreign-key columns, and representative business columns. Tables repeated as external references are shaded and are defined in their main domain view. Appendix C lists every active column, including nullability, defaults, keys, constraints, and a short description.")
    configure_body(p)
    insert_figure(next_heading, "data-pdm-identity-20260811.png", "Figure 3.7. CareerFit physical data model - identity and Career Profile tables")
    insert_figure(next_heading, "data-pdm-recruitment-20260811.png", "Figure 3.8. CareerFit physical data model - recruitment and Matching tables")
    insert_figure(next_heading, "data-pdm-communication-20260811.png", "Figure 3.9. CareerFit physical data model - automation and communication tables")
    insert_figure(next_heading, "data-pdm-governance-20260811.png", "Figure 3.10. CareerFit physical data model - governance and analytics tables")

    caption = insert_paragraph_before(next_heading, "Table 3.2. Key integrity constraints")
    configure_caption(caption, "Table Caption")
    next_heading._p.addprevious(preserved_table)

    p = insert_paragraph_before(next_heading,
        "The constraints above are enforced at the database boundary in addition to service-level validation. Indexes support active Job search, ownership queries, Matching order, application history, feedback lookup, token redemption, notification policy checks, audit chronology, and the pending content-report queue.")
    configure_body(p)

    h = insert_paragraph_before(next_heading, "3.3.5 Data Dictionary", "Heading 3")
    configure_heading(h)
    p = insert_paragraph_before(next_heading,
        "Appendix C contains the physical table catalogue and field-level data dictionary. Each entry is derived from the running PostgreSQL schema produced by Flyway V1-V25 and records the column name, PostgreSQL type, nullability, key role, default or check constraint, and business meaning. The dictionary documents persisted columns only; DTO-only fields, transient properties, repositories, and internal service variables are excluded.")
    configure_body(p)


def renumber_later_figures(document: Document):
    for paragraph in document.paragraphs:
        if paragraph.style.name != "Figure Caption":
            continue
        match = re.match(r"Figure 3\.(\d+)\.(.*)", paragraph.text.strip())
        if not match:
            continue
        number = int(match.group(1))
        if 3 <= number <= 14:
            new_text = f"Figure 3.{number + 8}.{match.group(2).strip()}"
            paragraph.text = new_text
            configure_caption(paragraph)
            previous = paragraph._p.getprevious()
            while previous is not None and not previous.tag.endswith("}p"):
                previous = previous.getprevious()
            if previous is not None:
                for prop in previous.xpath(".//wp:docPr"):
                    prop.set("title", new_text)
                    prop.set("descr", new_text)


def add_catalogue_table(document: Document, anchor: Paragraph, tables):
    caption = insert_paragraph_before(anchor, "Table App.C.1. CareerFit physical table catalogue")
    configure_caption(caption, "Table Caption")
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["No.", "Table name", "Data group", "Purpose"]
    for index, value in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], value, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[index])
    repeat_header(table.rows[0])
    group_lookup = {name: group for group, names in TABLE_GROUPS for name in names}
    for number, name in enumerate(sorted(tables), start=1):
        cells = table.add_row().cells
        values = [str(number), name, group_lookup[name], TABLE_PURPOSES[name]]
        for index, value in enumerate(values):
            set_cell_text(cells[index], value, align=WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT)
    set_table_widths(table, [Inches(0.38), Inches(1.25), Inches(1.55), Inches(2.92)])
    anchor._p.addprevious(table._tbl)


def add_dictionary_table(document: Document, anchor: Paragraph, table_name: str, columns, caption_no: int):
    caption = insert_paragraph_before(anchor, f"Table App.C.{caption_no}. Data dictionary for {table_name}")
    configure_caption(caption, "Table Caption")
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["No.", "Column", "PostgreSQL type", "Null", "Key", "Default / constraint", "Description"]
    for index, value in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], value, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[index])
    repeat_header(table.rows[0])
    for number, column in enumerate(columns, start=1):
        cells = table.add_row().cells
        key = "; ".join(dict.fromkeys(column["keys"])) if column["keys"] else "-"
        values = [
            str(number), column["name"], column["type"], "Yes" if column["nullable"] == "YES" else "No",
            key, constraint_summary(column), description(table_name, column),
        ]
        for index, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if index in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[index], value, align=align)
    set_table_widths(table, [
        Inches(0.33), Inches(0.92), Inches(0.92), Inches(0.48), Inches(0.92), Inches(1.18), Inches(1.35)
    ])
    anchor._p.addprevious(table._tbl)


def replace_appendix_c(document: Document, tables):
    heading = find_exact(document, "Appendix C. Data Model Summary")
    heading.text = "Appendix C. Full Data Dictionary"
    heading.style = "Heading 2"
    next_heading = find_exact(document, "Appendix D. Evaluation Summary")
    node = heading._p.getnext()
    while node is not None and node is not next_heading._p:
        following = node.getnext()
        node.getparent().remove(node)
        node = following

    intro = insert_paragraph_before(next_heading,
        "This appendix documents the final CareerFit PostgreSQL schema after Flyway V1-V25. It contains 24 active tables and 293 persisted columns. The earlier email_token table is excluded because migration V16 removes it. Keys and constraints are reported from the running migrated schema; descriptions summarize the observable role of each stored value.")
    configure_body(intro)

    h = insert_paragraph_before(next_heading, "C.1 Physical Table Catalogue", "Heading 3")
    configure_heading(h)
    add_catalogue_table(document, next_heading, tables)

    h = insert_paragraph_before(next_heading, "C.2 Field-Level Data Dictionary", "Heading 3")
    configure_heading(h, page_break=True)
    caption_no = 2
    section_no = 1
    for group_name, table_names in TABLE_GROUPS:
        group_h = insert_paragraph_before(next_heading, f"C.2.{section_no} {group_name}", "Heading 3")
        configure_heading(group_h, page_break=True)
        section_no += 1
        for table_name in table_names:
            table_h = insert_paragraph_before(next_heading, table_name, "Heading 4")
            configure_heading(table_h, page_break=True)
            purpose = insert_paragraph_before(next_heading, TABLE_PURPOSES[table_name])
            configure_body(purpose)
            add_dictionary_table(document, next_heading, table_name, tables[table_name], caption_no)
            caption_no += 1


def audit(document: Document, tables):
    captions = [p.text.strip() for p in document.paragraphs if p.style.name == "Figure Caption"]
    expected = [f"Figure 3.{i}." for i in range(1, 23)]
    for prefix in expected:
        if sum(c.startswith(prefix) for c in captions) != 1:
            raise RuntimeError(f"Missing or duplicate figure caption: {prefix}")
    if any(c.startswith("Figure 3.23.") for c in captions):
        raise RuntimeError("Unexpected Figure 3.23")

    dictionary_captions = [
        p.text.strip() for p in document.paragraphs
        if p.style.name == "Table Caption" and p.text.startswith("Table App.C.")
    ]
    if len(dictionary_captions) != 25:
        raise RuntimeError(f"Expected 25 Appendix C table captions; found {len(dictionary_captions)}")
    if len(tables) != 24 or sum(len(cols) for cols in tables.values()) != 293:
        raise RuntimeError("Schema count changed unexpectedly")
    all_cell_paragraphs = [p for table in document.tables for row in table.rows for cell in row.cells for p in cell.paragraphs]
    for paragraph in all_cell_paragraphs:
        for run in paragraph.runs:
            if run.text and run.font.size is not None and round(run.font.size.pt, 2) != 13:
                raise RuntimeError(f"Non-13pt table run found: {run.text[:40]!r} {run.font.size.pt}")
    text = "\n".join(p.text for p in document.paragraphs)
    if any(p.style.name == "Heading 2" and p.text.strip() == "Appendix C. Data Model Summary" for p in document.paragraphs):
        raise RuntimeError("Old Appendix C title remains")
    for bad in ("\x00", "\ufffd"):
        if bad in text:
            raise RuntimeError("Invalid character found in document text")


def main():
    tables = load_columns()
    load_constraints(tables)
    expected_order = [name for _, names in TABLE_GROUPS for name in names]
    if set(tables) != set(expected_order):
        raise RuntimeError(f"Schema table mismatch: {sorted(set(tables) ^ set(expected_order))}")

    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(DOCX, BACKUP)
    document = Document(DOCX)
    before_tables = len(document.tables)
    before_drawings = len(document.inline_shapes)

    renumber_later_figures(document)
    replace_data_design(document)
    replace_appendix_c(document, tables)
    request_field_refresh(document)
    audit(document, tables)

    document.save(DOCX)
    print(f"Updated: {DOCX}")
    print(f"Backup: {BACKUP}")
    print(f"Tables: {before_tables} -> {len(document.tables)}")
    print(f"Inline shapes: {before_drawings} -> {len(document.inline_shapes)}")
    print(f"Schema: {len(tables)} tables, {sum(len(cols) for cols in tables.values())} columns")
    print(f"Table digest after edit: {table_digest(document)}")


if __name__ == "__main__":
    main()
