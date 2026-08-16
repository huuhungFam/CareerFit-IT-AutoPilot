from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.text.paragraph import Paragraph
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"


def clean(text: str) -> str:
    return " ".join(text.split())


def find_start(document: Document, prefix: str) -> Paragraph:
    matches = [p for p in document.paragraphs if clean(p.text).startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting {prefix!r}; found {len(matches)}")
    return matches[0]


def set_font(run) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)


def set_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    set_font(run)


def clone_before(template: Paragraph, anchor: Paragraph) -> Paragraph:
    cloned = deepcopy(template._p)
    anchor._p.addprevious(cloned)
    paragraph = Paragraph(cloned, anchor._parent)
    paragraph.clear()
    return paragraph


def add_hyperlink(paragraph: Paragraph, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), "Times New Roman")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "26")
    size_complex = OxmlElement("w:szCs")
    size_complex.set(qn("w:val"), "26")
    properties.extend((fonts, size, size_complex))
    text = OxmlElement("w:t")
    text.text = url
    run.extend((properties, text))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_reference(template: Paragraph, reference_anchor: Paragraph, number: int, organization: str, title: str, url: str, venue: str | None = None) -> None:
    paragraph = clone_before(template, reference_anchor)
    prefix = f'[{number}] {organization}, "{title},"'
    if venue:
        prefix += f" {venue},"
    prefix += " [Online]. Available: "
    run = paragraph.add_run(prefix)
    set_font(run)
    add_hyperlink(paragraph, url)
    suffix = paragraph.add_run(". [Accessed: Aug. 12, 2026].")
    set_font(suffix)


def main() -> None:
    document = Document(DOCX)

    background = find_start(document, "CareerFit reduces fragmentation by combining Candidate Job")
    set_text(
        background,
        "Vietnamese recruitment portals provide the practical baseline for online Job discovery, CV support, and employer access. "
        "ITviec and TopDev emphasize technology recruitment, while CareerViet and TopCV provide broader Job-search and recruitment services [15]-[18]. "
        "CareerFit reduces fragmentation by combining Candidate Job, CV, and Application workflows with Recruiter Job, applicant, and Talent Pool workflows in one role-based system. Email provides an additional action channel.",
    )

    gap = find_start(document, "CareerFit does not attempt to outperform these research systems")
    industry = gap.insert_paragraph_before()
    industry.style = "Normal"
    set_text(
        industry,
        "A CareerViet industry article describes AI Matching as comparing Candidate CV information with Job requirements, presenting a matching score, and supporting suitable-Job and CV-improvement guidance [19]. "
        "This article was a practical source of inspiration for CareerFit, but it is treated as an industry example rather than academic evidence of algorithmic effectiveness.",
    )

    appendix = next(p for p in document.paragraphs if clean(p.text) == "APPENDICES")
    reference_template = next(p for p in document.paragraphs if clean(p.text).startswith("[14]"))
    if any(clean(p.text).startswith("[15]") for p in document.paragraphs):
        raise RuntimeError("Reference [15] already exists")

    add_reference(reference_template, appendix, 15, "ITviec", "ITviec | Top IT Jobs for You", "https://itviec.com/")
    add_reference(reference_template, appendix, 16, "CareerViet", "Fast Recruitment & Job Search at CareerViet.vn", "https://careerviet.vn/en")
    add_reference(reference_template, appendix, 17, "TopDev", "TopDev - Việc Làm Lương Cao Hàng Đầu", "https://topdev.vn/")
    add_reference(reference_template, appendix, 18, "TopCV Vietnam", "TopCV Pro - Nơi kết nối Ứng viên chất & Doanh nghiệp hàng đầu", "https://www.topcv.vn/pro")
    add_reference(
        reference_template,
        appendix,
        19,
        "CareerViet",
        "Tăng Tỉ Lệ Ứng Tuyển Thành Công Với AI Matching - Cách Biến CV Của Bạn Thành 'Ứng Viên Sáng Giá'",
        "https://careerviet.vn/en/talentcommunity/tang-ti-le-ung-tuyen-thanh-cong-voi-ai-matching-cach-bien-cv-cua-ban-thanh-ung-vien-sang-gia.35A52D7D.html",
        "CareerViet Talent Community",
    )

    document.save(DOCX)
    print(f"updated={DOCX}")
    print("references_added=5")


if __name__ == "__main__":
    main()
