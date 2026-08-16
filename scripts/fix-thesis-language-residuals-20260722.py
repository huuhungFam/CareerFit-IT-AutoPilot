from pathlib import Path

from docx import Document


REPORT = Path(r"C:\CODING\Thesis\Doc\CareerFit-Thesis-Report.docx")

REPLACEMENTS = {
    "An actionable email link can be opened by the intended user":
        "An actionable email link may be opened by the intended user, forwarded by mistake, or visited by a mail-security scanner. CareerFit uses a two-step flow to reduce this risk. The GET request checks the high-entropy token and shows a confirmation page without changing recruitment data. A deliberate POST request performs the action and marks the token as redeemed. Only the SHA-256 token hash is stored, and status and expiry checks prevent normal reuse. Production deployment should still add rate limiting, origin controls, secret rotation, and mail-delivery monitoring.",
    "The July 18 refresh retained the clean benchmark":
        "The July 18 refresh retained the clean benchmark and aggregate health results, kept the largest frontend chunk below 500 kB, removed API-data fallbacks, expanded Chromium coverage to 20 tests, and refreshed the six interface screenshots. It also verified post-commit matching, Candidate-only feedback authorization, immediate match-result notifications, and portfolio privacy gating. These results support a well-tested local demonstration; they do not establish production readiness.",
    "Notification action tokens are hashed":
        "Notification action tokens are hashed, and state changes require POST after confirmation. Access tokens are stored in frontend sessionStorage rather than persistent localStorage. Local CV storage lacks documented malware scanning, encryption, retention enforcement, and recovery testing. Public Swagger and Prometheus access are suitable for local demonstration but should be restricted in production. Privacy, fairness, and management of users' personal data have not been validated with a real deployment population.",
    "The final remediation pass resolved":
        "The final remediation pass resolved the observed optimistic-lock benchmark exception, scheduler and configuration drift, mock Job fallback values, persistent localStorage authentication, the oversized frontend bundle warning, and missing E2E Job cleanup. Audit records are still written directly by multiple services, and the evaluated working tree is not a fixed release commit. These limits make production operation and exact reproduction of the results more difficult.",
    "CareerFit demonstrates an end-to-end approach":
        "CareerFit demonstrates controlled IT recruitment automation by connecting job discovery, explainable matching, profile-based recommendations, feedback learning, policy-driven actions, and audit records. It keeps scores, business state, and user decisions visible.",
    "The evaluated local prototype passed":
        "Local evaluation passed the backend tests, frontend build, selected browser workflows, health checks, and controlled Rocchio benchmark. This supports a tested Human-in-the-Loop workflow, not a system ready for real hiring decisions.",
}

STYLE_REPLACEMENTS = {
    ("Figure Caption", "Figure 4.5. Scoring and Matching persistence flow"):
        "Figure 4.5. Scoring and matching persistence flow",
    ("Heading 3", "6.3.4 Green Assertions versus Operational Cleanliness"):
        "6.3.4 Test Results versus Background Errors",
}

document = Document(REPORT)
remaining = dict(REPLACEMENTS)
for prefix, replacement in list(remaining.items()):
    if any(paragraph.text == replacement for paragraph in document.paragraphs):
        del remaining[prefix]
for paragraph in document.paragraphs:
    for prefix, replacement in list(remaining.items()):
        if paragraph.text.startswith(prefix):
            paragraph.text = replacement
            del remaining[prefix]
            break

if remaining:
    raise RuntimeError(f"Missing residual paragraphs: {list(remaining)}")

style_remaining = dict(STYLE_REPLACEMENTS)
for key, replacement in list(style_remaining.items()):
    style_name, _ = key
    if any(
        paragraph.style.name == style_name and paragraph.text == replacement
        for paragraph in document.paragraphs
    ):
        del style_remaining[key]
for paragraph in document.paragraphs:
    key = (paragraph.style.name, paragraph.text)
    if key in style_remaining:
        paragraph.text = style_remaining.pop(key)

if style_remaining:
    raise RuntimeError(f"Missing style-specific paragraphs: {list(style_remaining)}")

document.save(REPORT)
print(f"Fixed {len(REPLACEMENTS) + len(STYLE_REPLACEMENTS)} residual paragraphs")
