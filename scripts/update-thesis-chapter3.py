from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"C:\CODING\Thesis")
DOCX = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
SOURCE = ROOT / "Doc" / "working" / "CHAPTER_3_SOURCE.md"


def set_font(run, size=13, bold=None, italic=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def find(document, text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(text)


def remove(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def insert_paragraph(document, anchor, text, level=None, note=False):
    paragraph = document.add_paragraph(style=f"Heading {level}" if level else "Normal")
    run = paragraph.add_run(text)
    set_font(run, 14 if level == 2 else 13, bool(level) if level else None, note)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    if level or note:
        paragraph.paragraph_format.first_line_indent = Inches(0)
    else:
        paragraph.paragraph_format.first_line_indent = Inches(0.5)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if note:
        paragraph.paragraph_format.left_indent = Inches(0.25)
    anchor._element.addprevious(paragraph._element)


def insert_table(document, anchor, markdown_lines):
    rows = [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in markdown_lines]
    rows = [rows[0], *rows[2:]]  # discard Markdown separator row
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, values in enumerate(rows):
        cells = table.rows[0].cells if row_index == 0 else table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Inches(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_font(run, 9.5, row_index == 0)
    anchor._element.addprevious(table._element)


def parse_and_insert(document, anchor):
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    paragraph_lines = []

    def flush_paragraph():
        nonlocal paragraph_lines
        if paragraph_lines:
            insert_paragraph(document, anchor, " ".join(line.strip() for line in paragraph_lines))
            paragraph_lines = []

    while index < len(lines):
        line = lines[index]
        if not line.strip():
            flush_paragraph()
            index += 1
            continue
        if line.startswith("## "):
            flush_paragraph()
            insert_paragraph(document, anchor, line[3:], level=2)
            index += 1
            continue
        if line.startswith("### "):
            flush_paragraph()
            insert_paragraph(document, anchor, line[4:], level=3)
            index += 1
            continue
        if line.startswith("NOTE: "):
            flush_paragraph()
            insert_paragraph(document, anchor, line, note=True)
            index += 1
            continue
        if line.startswith("Table "):
            flush_paragraph()
            insert_paragraph(document, anchor, line)
            index += 1
            table_lines = []
            while index < len(lines) and (not lines[index].strip() or lines[index].startswith("|")):
                if lines[index].startswith("|"):
                    table_lines.append(lines[index])
                index += 1
            insert_table(document, anchor, table_lines)
            continue
        paragraph_lines.append(line)
        index += 1
    flush_paragraph()


def replace_claims(document):
    replacements = {
        "Important actions can be reviewed through web interfaces or confirmed through single-use magic links, and audit records are maintained to support traceability.":
            "Important actions can be reviewed through web interfaces, while the current notification-email flow provides expiring one-click action links. Audit records and action statuses support traceability, but the one-click GET design is identified as a limitation requiring confirmation-before-execution for production use.",
        "An actionable email mechanism that uses confirmation pages, expiring single-use tokens, state-changing POST operations, and audit records to provide a fast but controlled interaction channel.":
            "An actionable email mechanism with expiring one-click tokens and action-state tracking, together with an identified requirement to replace state-changing GET redemption with confirmation followed by POST before production use.",
        "A safer pattern lets a GET request display a confirmation page without changing business state and requires a deliberate POST request to execute the action.":
            "A safer target pattern lets a GET request display a confirmation page without changing business state and requires a deliberate POST request to execute the action; the current CareerFit email-action implementation does not yet satisfy this target.",
        "Support actionable email through expiring, single-use magic links, confirmation before state-changing operations, and audit logging of significant decisions and actions.":
            "Support expiring email-action links, action-state tracking, and audit logging, while documenting confirmation-before-execution as a required security improvement.",
    }
    for paragraph in document.paragraphs:
        for old, new in replacements.items():
            if old in paragraph.text:
                updated = paragraph.text.replace(old, new)
                paragraph.clear()
                set_font(paragraph.add_run(updated))


def main():
    document = Document(DOCX)
    if any(p.text.strip() == "3.1 System Analysis Context" for p in document.paragraphs):
        replace_claims(document)
        document.save(DOCX)
        raise SystemExit("Chapter 3 was already populated; claim corrections were refreshed.")
    placeholder = "NOTE: [Chapter 3 content will be written from the verified SRS, source code, API contracts, and Flyway schema. Planned diagrams: use case, system architecture, component diagram, ERD, sequence diagrams, AutoFit decision flow, and deployment architecture.]"
    for paragraph in list(document.paragraphs):
        if paragraph.text.strip() == placeholder:
            remove(paragraph)
    anchor = find(document, "CHAPTER 4. SYSTEM IMPLEMENTATION")
    parse_and_insert(document, anchor)
    replace_claims(document)
    document.save(DOCX)
    print(f"Updated {DOCX} ({DOCX.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
