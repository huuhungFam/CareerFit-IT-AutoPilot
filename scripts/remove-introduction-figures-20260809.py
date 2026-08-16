from pathlib import Path
import shutil

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-before-20260809-remove-introduction-figures.docx"

REMOVE = {
    "Figure 1.1. CareerFit problem context and principal information flows",
    "Figure 1.2. Scope boundary of the CareerFit thesis",
}

RENUMBER = {
    "Figure 1.3.": "Figure 1.1.",
    "Figure 1.4.": "Figure 1.2.",
    "Figure 1.5.": "Figure 1.3.",
    "Figure 1.6.": "Figure 1.4.",
    "Figure 1.7.": "Figure 1.5.",
    "Figure 1.8.": "Figure 1.6.",
}


def remove_element(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def format_caption(paragraph):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.font.italic = True
        rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
        rfonts.set(qn("w:eastAsia"), "Times New Roman")


def main():
    if not BACKUP.exists():
        BACKUP.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(REPORT, BACKUP)

    document = Document(REPORT)
    paragraphs = document.paragraphs
    removed = []

    for index, paragraph in enumerate(paragraphs):
        if paragraph.style.name != "Figure Caption" or paragraph.text.strip() not in REMOVE:
            continue
        if index == 0 or not paragraphs[index - 1]._p.xpath(".//w:drawing"):
            raise RuntimeError(f"No drawing immediately before {paragraph.text!r}")
        remove_element(paragraphs[index - 1]._p)
        remove_element(paragraph._p)
        removed.append(paragraph.text.strip())

    if set(removed) != REMOVE:
        raise RuntimeError(f"Expected to remove {sorted(REMOVE)}, removed {sorted(removed)}")

    renumbered = []
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name != "Figure Caption":
            continue
        old_prefix = next((old for old in RENUMBER if paragraph.text.startswith(old)), None)
        if old_prefix is None:
            continue
        new_prefix = RENUMBER[old_prefix]
        old_text = paragraph.text
        new_text = new_prefix + old_text[len(old_prefix):]

        if index == 0 or not document.paragraphs[index - 1]._p.xpath(".//w:drawing"):
            raise RuntimeError(f"No drawing immediately before {old_text!r}")
        drawing_paragraph = document.paragraphs[index - 1]
        for prop in drawing_paragraph._p.xpath(".//wp:docPr"):
            for attribute in ("title", "descr"):
                value = prop.get(attribute)
                if value:
                    prop.set(attribute, value.replace(old_prefix.rstrip("."), new_prefix.rstrip(".")))

        paragraph.text = new_text
        format_caption(paragraph)
        renumbered.append((old_text, new_text))

    if len(renumbered) != len(RENUMBER):
        raise RuntimeError(f"Expected {len(RENUMBER)} renumbered captions, found {len(renumbered)}")

    document.save(REPORT)
    print(f"updated={REPORT}")
    print(f"backup={BACKUP}")
    print(f"removed={len(removed)} renumbered={len(renumbered)}")
    for old, new in renumbered:
        print(f"{old} -> {new}")


if __name__ == "__main__":
    main()
