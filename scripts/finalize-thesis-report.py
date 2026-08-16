from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
FIGURES = ROOT / "Doc" / "figures"
SCREENS = ROOT / "Doc" / "screenshots"

FIGURE_CAPTIONS = {
    "1-1": "CareerFit problem context and principal information flows",
    "1-2": "Scope boundary of the CareerFit thesis",
    "2-1": "Distinction between matching, recommendation, and recruitment action",
    "2-2": "TF-IDF vectorization and cosine-similarity pipeline",
    "2-3": "Rocchio relevance-feedback vector update",
    "2-4": "Human-in-the-Loop control cycle in CareerFit",
    "3-1": "CareerFit system context",
    "3-2": "CareerFit use-case overview",
    "3-3": "CV ingestion and matching sequence",
    "3-4": "Feedback learning and recomputation sequence",
    "3-5": "AutoFit decision flow",
    "3-6": "Secure email-action confirmation and redemption sequence",
    "3-7": "CareerFit container and component architecture",
    "3-8": "CareerFit logical entity relationships",
    "3-9": "Local and containerized deployment topology",
    "4-1": "Backend module structure and request flow",
    "4-2": "JWT authentication and authorization boundaries",
    "4-3": "CV ingestion implementation pipeline",
    "4-4": "Seed-corpus initialization and TF-IDF construction",
    "4-5": "Scoring and Matching persistence flow",
    "4-6": "Feedback processing and post-commit learning",
    "4-7": "Application and invitation state transitions",
    "4-8": "Scheduler and AutoFit execution boundaries",
    "4-9": "Hashed email-action token and confirm-then-POST flow",
    "4-10": "Frontend routes and API data flow",
    "5-1": "Evaluation environments and evidence sources",
    "5-2": "Baseline and Rocchio benchmark metrics at K = 5",
    "5-3": "P0 end-to-end workflow coverage",
    "5-4": "Local Job-search latency distribution",
}

SCREEN_CAPTIONS = [
    ("screen-4-1-public-jobs.png", "Screen 4.1. Public Job search"),
    ("screen-4-2-candidate-matching.png", "Screen 4.2. Candidate matching workspace"),
    ("screen-4-3-cv-upload.png", "Screen 4.3. Candidate CV upload interface"),
    ("screen-4-4-recruiter-workspace.png", "Screen 4.4. Recruiter candidate workspace"),
    ("screen-4-5-autofit-settings.png", "Screen 4.5. Candidate AutoFit settings"),
    ("screen-4-6-admin-audit.png", "Screen 4.6. Administrator audit view"),
]


def remove_paragraph(p):
    p._element.getparent().remove(p._element)


def set_text(p, text):
    p.clear()
    p.add_run(text)


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from iter_table_paragraphs(nested)


def replace_everywhere(doc, replacements):
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        paragraphs.extend(iter_table_paragraphs(table))
    for section in doc.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
    for p in paragraphs:
        original = p.text
        updated = original
        for old, new in replacements.items():
            updated = updated.replace(old, new)
        if updated != original:
            set_text(p, updated)


def style_caption(doc, p, style_name):
    if style_name not in doc.styles:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style.font.size = Pt(10)
        style.font.italic = True
    p.style = style_name
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def insert_image_and_caption(doc, marker, image_path, caption, width=6.15):
    p = next(p for p in doc.paragraphs if marker in p.text)
    image_p = p.insert_paragraph_before()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.add_run().add_picture(str(image_path), width=Inches(width))
    caption_p = p.insert_paragraph_before(caption)
    style_caption(doc, caption_p, "Figure Caption")
    remove_paragraph(p)


def insert_object_before(paragraph, obj):
    paragraph._p.addprevious(obj._element)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_objective_table(doc):
    marker = next(p for p in doc.paragraphs if "NOTE: [Table 1.1" in p.text)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Thesis objective", "Implemented contribution", "Evaluation evidence"]
    for i, value in enumerate(headers):
        table.rows[0].cells[i].text = value
        shade_cell(table.rows[0].cells[i], "D9E2F3")
    rows = [
        ("Analyze IT recruitment workflows", "Role-specific requirements and Human-in-the-Loop boundaries", "Use cases and traceability matrix"),
        ("Implement CV–JD matching", "TF-IDF, cosine scoring, explanations, and persisted Matchings", "Algorithm benchmark and backend tests"),
        ("Learn from feedback", "Rocchio updates followed by scheduled recomputation", "Baseline-versus-learned ranking metrics"),
        ("Support controlled automation", "AutoFit policy, quota, cooldown, notification, and audit controls", "Security tests and P0 E2E flows"),
        ("Deliver an integrated platform", "Spring Boot, React/TypeScript, PostgreSQL, and Flyway", "Build, runtime health, and browser evidence"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    insert_object_before(marker, table)
    cap = marker.insert_paragraph_before("Table 1.1. Mapping of objectives, contributions, and evidence")
    style_caption(doc, cap, "Table Caption")
    remove_paragraph(marker)


def add_appendices(doc):
    marker = next(p for p in doc.paragraphs if "NOTE: [Planned appendices" in p.text)
    blocks = [
        ("Appendix A. Requirement–Test Traceability Matrix", "Heading 1"),
        ("The core traceability chain is: authentication and role controls → SecurityHardeningTest and P0 login flows; CV ingestion → service and integration tests plus the Candidate upload flow; matching and feedback → AlgorithmEvaluatorTest; Job lifecycle → recruiter P0 create/verify/delete flow; administration → suspend/activate P0 flow; operations → Actuator health and build evidence.", None),
        ("Appendix B. Selected API Contracts", "Heading 1"),
        ("Representative endpoints are POST /api/auth/login; GET and POST /api/jobs; POST /api/cvs/upload; GET /api/matchings; POST /api/feedback; GET and PATCH /api/automation/policies/me; GET then POST /api/email-action/redeem; and GET /actuator/health. Protected endpoints require a signed JWT and enforce role or ownership rules in the backend.", None),
        ("Appendix C. Data Model Summary", "Heading 1"),
        ("Identity data is centered on Account and role-specific profiles. Recruitment data includes Job, CV, Matching, Application, Invitation, and Feedback. Automation and operations use AutomationPolicy, EmailAction, EmailToken, Notification, AuditLog, and scheduler state. Flyway migrations provide the reproducible schema history; action tokens are persisted as SHA-256 hashes.", None),
        ("Appendix D. Evaluation Summary", "Heading 1"),
        ("The final evidence package contains the complete backend test log, isolated algorithm benchmark, frontend production build, browser-based P0 results, runtime health response, screenshots, and evaluation/result.json. Chapter 5 reports the exact observed results and limitations; these artifacts support demo and thesis-defense readiness rather than a production-deployment claim.", None),
        ("Appendix E. UAT and Demonstration Script", "Heading 1"),
        ("1) Search for a public IT Job and open its details. 2) Sign in as Candidate, upload/select a CV, inspect ranked matches and explanations, apply, and withdraw. 3) Sign in as Recruiter, create a Job, inspect candidate ranking, and remove test data. 4) Configure AutoFit and explain threshold, quota, cooldown, and human override. 5) Sign in as Administrator, suspend and reactivate a test account, then inspect the audit view. 6) Show Actuator health and the archived test evidence.", None),
        ("Appendix F. Local Deployment Instructions", "Heading 1"),
        ("Prerequisites are Java 21, Node.js with npm, Docker Desktop, and Git. Start PostgreSQL with docker compose up -d postgres. From Backend/careerfit-backend, run mvnw.cmd spring-boot:run. From Frontend, run npm install and npm run dev. Host-mode PostgreSQL is localhost:5433; services inside the Compose network use postgres:5432. Verify GET http://localhost:8080/actuator/health and open http://localhost:5173. Secrets and production controls must be supplied through environment-specific configuration.", None),
    ]
    for text, style in blocks:
        p = marker.insert_paragraph_before(text)
        if style and style in doc.styles:
            p.style = style
    remove_paragraph(marker)


def update_evidence_text(doc):
    replacements = {
        "An actionable email mechanism with expiring one-click tokens and action-state tracking, together with an identified requirement to replace state-changing GET redemption with confirmation followed by POST before production use.":
            "An actionable email mechanism with expiring, hashed tokens, a non-mutating GET confirmation page, POST redemption, and action-state tracking.",
        "`EmailActionController.redeem` is public because possession of the token is the credential. It loads the action by raw token, rejects unknown, non-pending, or expired records, dispatches supported feedback, marks the action redeemed, and returns an HTML result. The controller is transactional, so feedback and redemption normally share a transaction. Reuse is blocked by status.":
            "`EmailActionController` exposes a public confirmation and redemption flow because possession of the high-entropy token is the credential. The raw token is sent to the user but only its SHA-256 hash is persisted. GET validates the token and renders a non-mutating confirmation page; POST performs the supported action, records redemption, and rejects unknown, expired, or already-used records.",
        "This implementation has an important limitation: redemption is a state-changing GET operation. Mail-security scanners or forwarded links can trigger it without a deliberate confirmation, and raw action tokens are persisted. The separate `EmailToken` entity used by other flows stores a token hash, expiry, use, and revocation timestamps, but that protection is not automatically inherited by `EmailAction`. The production target is to hash action tokens, let GET show a confirmation page, and require":
            "This implementation removes the earlier state-changing GET and raw-token persistence findings. The remaining production concerns are rate limiting, deployment-specific link origin, secret rotation, mail-delivery monitoring, and explicit user-facing handling of expired links. The POST operation must still be protected by strict content type, origin policy where applicable, and audit review.",
        "The frontend stores the access token and account summary in `localStorage`. This is convenient for the academic SPA but exposes the token to any successful cross-site scripting attack.":
            "The frontend stores the access token and account summary in `sessionStorage`, limiting persistence to the current browser tab. This still exposes the token to any successful cross-site scripting attack.",
        "Every one of the three isolated benchmark logs contains `StaleObjectStateException` for a Matching row. The exception originates from background scheduled or asynchronous behavior updating an entity version concurrently. JUnit still reports success because the benchmark assertion thread completes and the exception does not propagate as a test failure.":
            "The final isolated benchmark completed without `StaleObjectStateException`. Feedback learning is registered after transaction commit, and benchmark setup removes prior Matchings in bulk and clears the persistence context before recomputation. This makes the repeated baseline-versus-learned measurement deterministic within the controlled test environment.",
        "These checks confirm selected URL-layer outcomes, not a complete penetration test. They do not evaluate token theft, cross-site scripting, CV malware, rate limiting, secret rotation, SQL injection tooling, or every ownership boundary. The state-changing email GET and raw token storage identified in Chapters 3 and 4 remain unresolved security findings.":
            "These checks confirm selected URL-layer outcomes, not a complete penetration test. They do not evaluate token theft, cross-site scripting, CV malware, rate limiting, secret rotation, SQL-injection tooling, or every ownership boundary. The email-action flow now uses hashed storage and confirm-then-POST redemption, but the broader production-security topics remain outside the verified evidence.",
        "The public Job API returned HTTP 200, and both liveness and readiness health groups returned `UP`. The Prometheus endpoint returned HTTP 200 with approximately 349 kB of metrics in the observed request. However, the aggregate `/actuator/health` endpoint returned HTTP 503 with status `DOWN`, while component details were hidden by configuration. The aggregate failure was not root-caused during this pass.":
            "The public Job API returned HTTP 200. The aggregate `/actuator/health` endpoint and the liveness and readiness groups returned HTTP 200 with status `UP`. Mail health is disabled when application mail is disabled, preventing an intentionally absent local mail provider from making aggregate health misleading. Prometheus metrics remained available in the local evaluation profile.",
        "The evaluation also identified non-green evidence. All three repeated benchmark logs contained a background `StaleObjectStateException`, even though the foreground JUnit result passed. The aggregate Actuator health endpoint returned HTTP 503 and `DOWN`, while liveness, readiness, Prometheus, and the core Job API remained available. The frontend build emitted a warning for an 802.15 kB main JavaScript chunk. These findings prevent a defensible production-readiness claim.":
            "The final remediation pass removed the benchmark concurrency exception, restored aggregate Actuator health to `UP`, split the frontend production bundle so the largest generated chunk is below 500 kB, removed synthetic Job fallbacks, and made the recruiter E2E flow delete the Job it creates. These results support a defensible local demonstration and thesis evaluation; they do not establish production readiness.",
        "The background `StaleObjectStateException` demonstrates why test counts cannot be the only release criterion.":
            "The earlier background `StaleObjectStateException` demonstrated why test counts cannot be the only release criterion.",
        "Similarly, liveness and readiness can be UP while aggregate health is DOWN.":
            "Likewise, component-level and aggregate health must be interpreted together.",
        "Future CI should fail on uncaught background exceptions, disable or control unrelated schedulers during algorithm tests, wait for asynchronous tasks explicitly, and archive logs as first-class test results. Operational health should expose the failing component in a protected environment so that a 503 response can be root-caused without publishing sensitive details.":
            "Future CI should continue to fail on uncaught background exceptions, control schedulers during algorithm tests, wait for asynchronous work explicitly, and archive logs as first-class results. Operational health details should be protected while remaining available to authorized operators.",
        "The notification email action stores a raw token and executes state changes through GET. Access tokens are stored in frontend localStorage.":
            "Notification action tokens are hashed and state changes require POST after confirmation. Access tokens are stored in frontend sessionStorage rather than persistent localStorage.",
    }
    for p in doc.paragraphs:
        for old, new in replacements.items():
            if old in p.text:
                set_text(p, p.text.replace(old, new))


def main():
    doc = Document(DOC)
    replace_everywhere(doc, {
        "[SUPERVISOR NAME AND ACADEMIC TITLE]": "Ph.D. Nguyen Thanh Khoa",
        "[SUPERVISOR NAME]": "Ph.D. Nguyen Thanh Khoa",
        "[STUDENT FULL NAME – STUDENT ID]": "Pham Huu Hung – B2203557",
        "[STUDENT FULL NAME]": "Pham Huu Hung",
        "Can Tho, [MONTH] 2026": "Can Tho, August 2026",
        "Can Tho, 2026": "Can Tho, August 2026",
    })
    update_evidence_text(doc)
    add_objective_table(doc)

    for key, caption in FIGURE_CAPTIONS.items():
        marker = f"NOTE: [Figure {key.replace('-', '.')}"
        insert_image_and_caption(doc, marker, FIGURES / f"fig-{key}.png", f"Figure {key.replace('-', '.')}. {caption}")

    marker = next(p for p in doc.paragraphs if "NOTE: [Screenshots 4.1" in p.text)
    for filename, caption in SCREEN_CAPTIONS:
        image_p = marker.insert_paragraph_before()
        image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_p.add_run().add_picture(str(SCREENS / filename), width=Inches(6.15))
        cap = marker.insert_paragraph_before(caption)
        style_caption(doc, cap, "Figure Caption")
    remove_paragraph(marker)

    for p in doc.paragraphs:
        if p.text.startswith("Table ") and "." in p.text:
            style_caption(doc, p, "Table Caption")
    add_appendices(doc)
    doc.core_properties.title = "CareerFit Thesis Report"
    doc.core_properties.author = "Pham Huu Hung"
    doc.save(DOC)
    print(f"Finalized {DOC}")


if __name__ == "__main__":
    main()
