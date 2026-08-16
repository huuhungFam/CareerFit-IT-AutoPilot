from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import math
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-before-20260809-usecase-expansion.docx"
FIGURE = ROOT / "Doc" / "figures" / "fig-3-2-usecases-20260809.png"

FONT = Path(r"C:\Windows\Fonts\arial.ttf")
BOLD = Path(r"C:\Windows\Fonts\arialbd.ttf")


USE_CASES = [
    {
        "title": "Candidate Manages CV, Profile, and Portfolio and Receives Match Results",
        "caption": "Candidate manages CV, profile, portfolio, and match results",
        "actor": "Candidate",
        "pre": "The Candidate account is active and authenticated. Candidate-owned profile, CV, and portfolio APIs are available.",
        "trigger": "The Candidate uploads or drafts a CV, edits profile/portfolio data, manages an existing CV, or opens its match results.",
        "flow": "Create or upload a CV; validate and extract content; review and edit sections; confirm the CV; update profile and portfolio links/projects; list owned CVs; set an eligible SCORING_DONE CV as default or delete a permitted CV; vectorize confirmed content; score active Jobs; show the score, label, Potential explanation, and reasons.",
        "alt": "Save incomplete manual content as DRAFT; keep extracted content in REVIEW_REQUIRED until confirmation; reject unsafe or unsupported files; record FAILED processing; reject a default change for an unscored CV; protect portfolio data according to the Candidate's privacy settings.",
        "post": "The owned profile, portfolio, and CV collection reflect the accepted changes, and confirmed CVs expose current matching results when scoring completes.",
    },
    {
        "title": "Candidate Searches and Reviews Job Descriptions",
        "caption": "Candidate searches and reviews Job descriptions",
        "actor": "Guest or Candidate",
        "pre": "Public Job discovery requires no account. Personalized information and recruitment actions require an authenticated Candidate.",
        "trigger": "The user searches, applies filters, selects an urgent or suggested Job, or opens a Job/employer detail page.",
        "flow": "Send server-side search, filter, sort, and pagination parameters; display active Jobs; open the selected JD; show structured Job and employer information; load similar Jobs; retain the selected filters when the user returns to the catalogue.",
        "alt": "Show a clear empty state when no Job matches; show a retryable error when the API fails; do not expose closed, hidden, banned, or otherwise ineligible Jobs through the public catalogue.",
        "post": "The user receives current Job information without creating an Application; authentication is requested before a protected action.",
    },
    {
        "title": "Candidate Manages Applications and Recruiter Invitations",
        "caption": "Candidate manages applications and Recruiter invitations",
        "actor": "Candidate",
        "pre": "The Candidate is authenticated. Applying requires an active Job before its deadline and a confirmed selected or default CV.",
        "trigger": "The Candidate applies, opens application history, withdraws an application, or accepts/declines a Recruiter invitation.",
        "flow": "Validate the Candidate, Job, deadline, CV, and duplicate state; create the Application with available Matching context; list and filter owned Applications; withdraw when the state permits; display pending invitations; accept or decline an invitation; notify the related parties and record important changes.",
        "alt": "Reject missing authentication/CV, duplicate application, inactive or expired Job, another Candidate's Application, or an unsupported withdrawal/invitation transition.",
        "post": "The Application or invitation response is stored with a valid lifecycle state and remains visible in the Candidate's history.",
    },
    {
        "title": "Candidate Provides Match Feedback and Updates Later Ranking",
        "caption": "Candidate provides match feedback and updates later ranking",
        "actor": "Candidate",
        "pre": "The Candidate owns the CV and Matching connected to the feedback.",
        "trigger": "The Candidate submits GOOD_MATCH, POTENTIAL, BAD_MATCH, or NOT_INTERESTED.",
        "flow": "Resolve the Matching and authenticated actor; verify ownership; upsert one current judgment; store audit evidence; commit the transaction; start Rocchio learning for supported positive/negative feedback; mark affected Matchings for recomputation; rescore and clear the stale marker after success.",
        "alt": "Reject another role, another Candidate's Matching, missing evidence, or invalid feedback; store NOT_INTERESTED without treating it as a technical relevance signal; retain retry visibility when asynchronous recomputation fails.",
        "post": "The feedback is stored, and later rankings use the rebuilt learned Job vector when recomputation completes.",
    },
    {
        "title": "Candidate Reviews Personalized Recommendations and Analytics",
        "caption": "Candidate reviews personalized recommendations and analytics",
        "actor": "Candidate",
        "pre": "The Candidate is authenticated. Personalized results require available profile, preference, CV, Matching, or analytics data.",
        "trigger": "The Candidate opens Recommendations or Advanced Analytics.",
        "flow": "Load personalized Job recommendations separately from the Matching feed; open, skip, or apply from a recommendation; load matching trends, skill demand, and profile-gap summaries; present the results as decision support; allow the Candidate to return to profile/CV management when improvement is needed.",
        "alt": "Show an empty or limited-data state when the profile, confirmed CV, or analytics history is insufficient; show a retry action for API failure; do not create an Application only because a Job was recommended.",
        "post": "The Candidate can use the recommendation and analytics evidence to choose a Job or improve the profile without an automatic recruitment decision.",
    },
    {
        "title": "Candidate Configures and Runs AutoFit",
        "caption": "Candidate configures and runs AutoFit",
        "actor": "Candidate and background scheduler",
        "pre": "The Candidate owns an AutomationPolicy. Execution also requires an eligible confirmed default CV and completed matching data.",
        "trigger": "The Candidate changes AutoFit settings, pauses/resumes automation, selects Run Now, or a scheduled scan evaluates the policy.",
        "flow": "Save enablement, thresholds, categories, quota, cooldown, time zone, quiet hours, and notification preferences; resolve the Candidate, CV, Jobs, Matchings, consent, and previous interactions; select only eligible actions; create permitted Applications or notifications; write delivery and audit evidence.",
        "alt": "Skip disabled, paused, ineligible, duplicate, stale, quota-exceeded, cooldown, quiet-hour, invalid-deadline, or invalid-state items; isolate a failed item instead of stopping the full scan.",
        "post": "Only policy-allowed actions are recorded; skipped items remain unchanged and explainable through stored outcomes.",
    },
    {
        "title": "Candidate or Recruiter Reports Suspicious Recruitment Content",
        "caption": "Candidate or Recruiter reports suspicious recruitment content",
        "actor": "Candidate or Recruiter",
        "pre": "A Candidate may report an ACTIVE Job. A Recruiter may report a CV only through an owned Job that provides visibility by Application or Matching.",
        "trigger": "The actor opens Report on an eligible Job or CV, selects a supported reason, and optionally enters a comment.",
        "flow": "Resolve the actor and target; apply the role, Job ownership, CV visibility, and target-state checks; validate the reason/comment; prevent another pending report from the same reporter for the same target; create the PENDING ContentReport; increase the target counter; add audit evidence; expose the case to Administrator moderation.",
        "alt": "Require a comment for OTHER; reject an inactive Job, banned CV, unrelated Recruiter, invisible CV, unsupported reason, or duplicate pending report.",
        "post": "The target has a pending report count, and the report appears in the Administrator queue until it is dismissed or actioned.",
    },
    {
        "title": "Recruiter Manages Company Profile and Job Lifecycle",
        "caption": "Recruiter manages company profile and Job lifecycle",
        "actor": "Recruiter",
        "pre": "The Recruiter account is active and authenticated. A company profile is required before publishing a Job.",
        "trigger": "The Recruiter completes onboarding, edits the company profile, creates/edits a Job, changes its status, deletes it, or exports the owned Job list.",
        "flow": "Create or update the company profile; list owned Jobs; enter structured and free-text JD data; preview quality signals; save DRAFT; review and publish; set urgent and deadline fields; edit an owned Job; change among supported Recruiter lifecycle states; delete a permitted Job or export the list; start vectorization and scoring after a valid publish commit.",
        "alt": "Reject missing company data, blocking JD/deadline errors, another Recruiter's Job, an unsupported transition, or deletion that violates current state rules; retain incomplete work as DRAFT.",
        "post": "The company and owned Job state reflect the accepted operation, and published eligible Jobs are available for matching and applications.",
    },
    {
        "title": "Recruiter Reviews and Processes Applicants",
        "caption": "Recruiter reviews and processes applicants",
        "actor": "Recruiter",
        "pre": "The Recruiter is authenticated and owns the selected Job.",
        "trigger": "The Recruiter opens the ranking/applicants workspace or chooses an application decision.",
        "flow": "Load and filter applicants by valid states; show Candidate, CV, Matching, label, Potential, and permitted portfolio information; open a Candidate detail; approve or reject an eligible Application; update the list; send configured lifecycle notifications; add audit evidence.",
        "alt": "Reject a Job owned by another Recruiter, an unknown Application, hidden private data, or an invalid status transition; show empty states for Jobs without applicants.",
        "post": "The Application contains the valid Recruiter decision and the Candidate can observe the updated state.",
    },
    {
        "title": "Recruiter Manages Talent Pool and Candidate Invitations",
        "caption": "Recruiter manages Talent Pool and Candidate invitations",
        "actor": "Recruiter",
        "pre": "The Recruiter is authenticated and owns the Job used for Candidate discovery.",
        "trigger": "The Recruiter opens Discover, Bookmarked CVs, or Invited CVs for an owned Job.",
        "flow": "Load matching and Potential Candidates with filters; view permitted CV and explanation data; bookmark or unbookmark a Candidate for the Job; send an invitation; list invitations awaiting a response; withdraw a pending invitation when permitted; update the corresponding tabs and notifications.",
        "alt": "Reject unrelated Job ownership, an invisible/missing Candidate, duplicate invitation, invalid invitation withdrawal, or an unsupported Candidate state; show a clear empty state for each tab.",
        "post": "Bookmark and invitation records reflect the Recruiter's choices and remain scoped to the owned Job.",
    },
    {
        "title": "Recruiter Reviews Analytics and Configures Recruitment Notifications",
        "caption": "Recruiter reviews analytics and configures recruitment notifications",
        "actor": "Recruiter",
        "pre": "The Recruiter is authenticated; private recruitment analytics depend on available owned-Job activity.",
        "trigger": "The Recruiter opens Analytics/Advanced Analytics or recruitment notification settings.",
        "flow": "Load recruitment overview and trends; show owned-Job/application outcomes and market, skill, or salary summaries with their scope; enable or disable high-match CV alerts, new-application alerts, invitation-response alerts, approval digests, and Job-closing reminders; set the high-match threshold; save the preferences.",
        "alt": "Show limited-data or empty states when activity is insufficient; distinguish platform-wide market data from private company results; retain the previous settings if an update fails.",
        "post": "The Recruiter can monitor current evidence, and later notification decisions use the stored preferences.",
    },
    {
        "title": "Administrator Manages User Accounts and Job Visibility",
        "caption": "Administrator manages user accounts and Job visibility",
        "actor": "Administrator",
        "pre": "The Administrator is authenticated and authorized for administrative APIs.",
        "trigger": "The Administrator searches users/Jobs and selects Suspend, Activate, Hide, or Restore.",
        "flow": "Load paginated user or Job data; apply filters; suspend an ACTIVE account or activate a suspended account; hide an ACTIVE Job using HIDDEN_BY_ADMIN; restore only a Job hidden by Administrator; refresh the result; record the administrative action and target in audit history.",
        "alt": "Reject missing targets, invalid roles, or unsupported state changes; never use Restore to publish a DRAFT, PAUSED, CLOSED, or BANNED Job.",
        "post": "The selected account state or Job visibility is updated without changing unrelated recruitment records.",
    },
    {
        "title": "Administrator Reviews and Resolves Content Reports",
        "caption": "Administrator reviews and resolves content reports",
        "actor": "Administrator",
        "pre": "The Administrator is authenticated. At least one report case may be pending for a JOB or CV target.",
        "trigger": "The Administrator opens Report moderation, selects a queue/case, and chooses Dismiss or Ban.",
        "flow": "Load grouped pending JOB/CV cases; show target details, reasons, comments, reporters, and timestamps; lock and reload the target; dismiss all pending reports when no action is needed or set the target to BANNED when action is required; clear the target counter; resolve the pending report records; add audit evidence; refresh the queue.",
        "alt": "Show an empty queue when no case is pending; reject a missing target, a case with no pending reports, or an invalid target type; keep the case unchanged if the transaction fails.",
        "post": "Pending reports become DISMISSED or ACTIONED, and a banned target is excluded from normal use according to its domain rules.",
    },
    {
        "title": "Recipient Reviews and Completes an Actionable Email",
        "caption": "Recipient reviews and completes an actionable email",
        "actor": "Email recipient",
        "pre": "A pending, unexpired EmailAction exists for the recipient, and only the SHA-256 token hash is stored.",
        "trigger": "The recipient opens an action from MATCH_NOTIFICATION, DAILY_DIGEST, or RECRUITER_INVITATION and deliberately confirms it.",
        "flow": "GET hashes and validates the token and shows a non-mutating confirmation page; POST repeats validation and performs the selected action. Match notification supports GOOD_MATCH, POTENTIAL, or NOT_INTERESTED; Daily Digest supports per-Job feedback and UNSUBSCRIBE_DIGEST; Recruiter Invitation supports INVITATION_ACCEPT or INVITATION_DECLINE. The service records redemption and audit/feedback/application effects.",
        "alt": "Reject missing, malformed, expired, redeemed, or unsupported tokens/actions; do not change recruitment state from GET; show a clear result for unsuccessful redemption.",
        "post": "The requested action is executed at most once under normal validation. Other one-way lifecycle emails remain notifications rather than separate use cases.",
    },
]


def font(size: int, bold: bool = False):
    return ImageFont.truetype(str(BOLD if bold else FONT), size)


def wrap(draw: ImageDraw.ImageDraw, text: str, used_font, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=used_font)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_centered_text(draw, bounds, text, used_font, color="#172554", line_gap=4):
    x1, y1, x2, y2 = bounds
    lines = wrap(draw, text, used_font, x2 - x1 - 24)
    line_height = used_font.size + line_gap
    total = len(lines) * line_height - line_gap
    y = y1 + (y2 - y1 - total) / 2
    for line in lines:
        box = draw.textbbox((0, 0), line, font=used_font)
        x = x1 + (x2 - x1 - (box[2] - box[0])) / 2
        draw.text((x, y), line, font=used_font, fill=color)
        y += line_height


def draw_case(draw, bounds, case_id, label, fill, outline):
    x1, y1, x2, y2 = bounds
    draw.rounded_rectangle(bounds, radius=15, fill=fill, outline=outline, width=4)
    draw.rounded_rectangle((x1 + 9, y1 + 9, x1 + 78, y1 + 47), radius=10, fill=outline)
    draw_centered_text(draw, (x1 + 9, y1 + 8, x1 + 78, y1 + 48), case_id, font(24, True), color="white")
    draw_centered_text(draw, (x1 + 82, y1 + 5, x2 - 5, y2 - 5), label, font(25, True), color="#172554")


def generate_use_case_figure():
    width, height = 2000, 1180
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 40), "CareerFit role-oriented use cases", font=font(48, True), fill="#111827")
    draw.text((70, 102), "Fourteen user goals represented by current frontend, email actions, and protected backend workflows", font=font(25), fill="#475569")

    lanes = [
        (145, 505, "Guest / Candidate", "Search, profile, matching, applications, learning, and automation", "#eef2ff", "#4f46e5"),
        (535, 785, "Recruiter", "Company, Jobs, applicants, Talent Pool, analytics, and alerts", "#ecfeff", "#0891b2"),
        (815, 1050, "Admin / Email", "Moderation and secure recipient actions", "#f0fdf4", "#15803d"),
    ]
    for top, bottom, title, subtitle, fill, outline in lanes:
        draw.rounded_rectangle((55, top, 1945, bottom), radius=22, fill="#fafafa", outline="#cbd5e1", width=3)
        draw.rounded_rectangle((70, top + 18, 315, bottom - 18), radius=18, fill=fill, outline=outline, width=4)
        draw_centered_text(draw, (82, top + 35, 303, top + 125), title, font(31, True), color="#0f172a")
        draw_centered_text(draw, (84, top + 125, 301, bottom - 35), subtitle, font(22), color="#475569")

    candidate_cases = [
        ("UC-01", "CV, profile, portfolio, matches"),
        ("UC-02", "Search and review Jobs"),
        ("UC-03", "Applications and invitations"),
        ("UC-04", "Feedback and Rocchio"),
        ("UC-05", "Recommendations and analytics"),
        ("UC-06", "Configure and run AutoFit"),
        ("UC-07", "Report suspicious content"),
    ]
    recruiter_cases = [
        ("UC-07", "Report a visible CV"),
        ("UC-08", "Company and Job lifecycle"),
        ("UC-09", "Review applicants"),
        ("UC-10", "Talent Pool and invitations"),
        ("UC-11", "Analytics and notifications"),
    ]
    final_cases = [
        ("UC-12", "Manage users and Job visibility"),
        ("UC-13", "Resolve content reports"),
        ("UC-14", "Complete an email action"),
    ]

    start_x, gap = 340, 18
    case_width = 380
    for index, (case_id, label) in enumerate(candidate_cases):
        row = index // 4
        col = index % 4
        x1 = start_x + col * (case_width + gap)
        y1 = 175 + row * 150
        draw_case(draw, (x1, y1, x1 + case_width, y1 + 120), case_id, label, "#eef2ff", "#4f46e5")
    recruiter_width = 305
    for index, (case_id, label) in enumerate(recruiter_cases):
        x1 = start_x + index * (recruiter_width + 15)
        draw_case(draw, (x1, 600, x1 + recruiter_width, 720), case_id, label, "#ecfeff", "#0891b2")
    final_width = 500
    for index, (case_id, label) in enumerate(final_cases):
        x1 = start_x + index * (final_width + 30)
        draw_case(draw, (x1, 875, x1 + final_width, 995), case_id, label, "#f0fdf4", "#15803d")

    draw.text((70, 1105), "CareerFit IT AutoPilot — updated use-case scope, August 2026", font=font(22), fill="#64748b")
    FIGURE.parent.mkdir(parents=True, exist_ok=True)
    image.save(FIGURE, quality=95)


def set_run_font(run, size=13):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "Times New Roman")


def set_paragraph_text(paragraph, text, size=13):
    paragraph.text = text
    for run in paragraph.runs:
        set_run_font(run, size)


def replace_start(document, prefix, text):
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting {prefix!r}, found {len(matches)}")
    set_paragraph_text(matches[0], text)


def normalize_cell(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.right_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    set_run_font(run, 13)


def set_table(table, rows):
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    for row_index, row in enumerate(rows):
        if len(row) != len(table.columns):
            raise RuntimeError("Table shape mismatch")
        for column_index, value in enumerate(row):
            normalize_cell(table.cell(row_index, column_index), value, bold=(row_index == 0))


def insert_paragraph(anchor, text, style, size):
    paragraph = anchor.insert_paragraph_before(text, style=style)
    for run in paragraph.runs:
        set_run_font(run, size)
    return paragraph


def use_case_rows(index, item):
    return [
        ["Field", "Description"],
        ["Use-case ID", f"UC-{index:02d}"],
        ["Primary actor", item["actor"]],
        ["Preconditions", item["pre"]],
        ["Trigger", item["trigger"]],
        ["Main flow", item["flow"]],
        ["Alternative/exception flows", item["alt"]],
        ["Postconditions", item["post"]],
    ]


def image_paragraph_before(document, caption_prefix):
    captions = [i for i, p in enumerate(document.paragraphs)
                if p.style.name == "Figure Caption" and p.text.strip().startswith(caption_prefix)]
    if len(captions) != 1:
        raise RuntimeError(f"Expected one body caption {caption_prefix}, found {len(captions)}")
    caption_index = captions[0]
    for index in range(caption_index - 1, max(-1, caption_index - 5), -1):
        paragraph = document.paragraphs[index]
        if paragraph._p.xpath(".//a:blip"):
            return paragraph
    raise RuntimeError(f"No image before {caption_prefix}")


def replace_picture(document, caption_prefix, image_path, alt_text, max_width_cm=15.0, max_height_cm=9.4):
    paragraph = image_paragraph_before(document, caption_prefix)
    blips = paragraph._p.xpath(".//a:blip")
    if len(blips) != 1:
        raise RuntimeError(f"Expected one image before {caption_prefix}")
    relation_id = blips[0].get(qn("r:embed"))
    document.part.related_parts[relation_id]._blob = Path(image_path).read_bytes()
    with Image.open(image_path) as source:
        ratio = source.width / source.height
    width_cm = min(max_width_cm, max_height_cm * ratio)
    height_cm = width_cm / ratio
    cx, cy = int(Cm(width_cm)), int(Cm(height_cm))
    for extent in paragraph._p.xpath(".//wp:extent"):
        extent.set("cx", str(cx))
        extent.set("cy", str(cy))
    for extent in paragraph._p.xpath(".//a:xfrm/a:ext"):
        extent.set("cx", str(cx))
        extent.set("cy", str(cy))
    for prop in paragraph._p.xpath(".//wp:docPr"):
        prop.set("descr", alt_text)
        prop.set("title", alt_text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(REPORT, BACKUP)
    generate_use_case_figure()

    document = Document(REPORT)

    # Preserve the four existing sequence diagrams while rebuilding the use-case block.
    figures = {}
    for number in (5, 6, 7, 8):
        caption_prefix = f"Figure 1.{number}."
        caption = [p for p in document.paragraphs
                   if p.style.name == "Figure Caption" and p.text.strip().startswith(caption_prefix)]
        if len(caption) != 1:
            raise RuntimeError(f"Missing body {caption_prefix}")
        image_paragraph = image_paragraph_before(document, caption_prefix)
        figures[number] = (deepcopy(image_paragraph._p), deepcopy(caption[0]._p))

    usecase_heading = [p for p in document.paragraphs
                       if p.style.name == "Heading 2" and p.text.strip() == "1.5 Use-Case Analysis"]
    summary_heading = [p for p in document.paragraphs
                       if p.style.name == "Heading 2" and p.text.strip() == "1.6 Chapter Summary"]
    if len(usecase_heading) != 1 or len(summary_heading) != 1:
        raise RuntimeError("Use-case block anchors are not unique")
    start = usecase_heading[0]._p
    anchor = summary_heading[0]
    end = anchor._p
    template_table = deepcopy(document.tables[6]._tbl)

    current = start.getnext()
    while current is not None and current is not end:
        following = current.getnext()
        current.getparent().remove(current)
        current = following

    intro = insert_paragraph(
        anchor,
        "The following fourteen use cases represent the main role-based goals available through the current frontend, actionable email pages, and protected backend workflows. Related screens are grouped into one user goal instead of treating every button or passive notification as a separate use case.",
        "Normal",
        13,
    )
    intro.paragraph_format.keep_with_next = True

    figure_after_case = {1: 5, 4: 6, 6: 7, 14: 8}
    for index, item in enumerate(USE_CASES, start=1):
        heading = insert_paragraph(anchor, f"1.5.{index} {item['title']}", "Heading 3", 13)
        # Avoid leaving UC-05's heading, caption, and only the repeated table
        # header below Figure 1.6 at the bottom of a page.
        if index == 5:
            heading.paragraph_format.page_break_before = True
        insert_paragraph(anchor, f"Table 1.{index + 4}. Use case - {item['caption']}", "Table Caption", 11)
        table_element = deepcopy(template_table)
        anchor._p.addprevious(table_element)
        table = Table(table_element, document._body)
        set_table(table, use_case_rows(index, item))
        if index == 14:
            note = insert_paragraph(
                anchor,
                "Other lifecycle emails are treated as notifications rather than separate use cases because they inform the recipient without adding another token-based user goal or a separate state-changing workflow.",
                "Normal",
                13,
            )
            note.paragraph_format.keep_with_next = True
        if index in figure_after_case:
            figure_number = figure_after_case[index]
            anchor._p.addprevious(deepcopy(figures[figure_number][0]))
            anchor._p.addprevious(deepcopy(figures[figure_number][1]))

    # Actor and functional-requirement tables now reflect the expanded use-case set.
    set_table(document.tables[3], [
        ["Actor", "Primary responsibilities", "Access boundary"],
        ["Guest", "Browse and filter public Jobs, open Job/employer details, and review similar Jobs or public analytics", "Public GET operations only; authentication is required for personalized data or recruitment actions"],
        ["Candidate", "Manage profile, portfolio, and CVs; review matches/recommendations/analytics; apply or answer invitations; give feedback; configure AutoFit; report an active Job", "Candidate-owned resources and Candidate/report APIs"],
        ["Recruiter", "Maintain company and Jobs; review applicants; manage Talent Pool bookmarks/invitations; inspect analytics; configure alerts; report a visible CV", "Recruiter APIs, owned Jobs, and CVs visible through an owned Job"],
        ["Administrator", "Suspend/activate users; hide/restore Jobs; monitor audit/email state; review and resolve reported Job/CV cases", "Administrative APIs only; all moderation actions are audited"],
        ["Background scheduler", "Recompute stale Matchings, send digests/reminders, clean expired actions, and run eligible automation", "Internal service calls under system control"],
        ["Mail provider", "Deliver notification and signed action emails generated by the backend", "SMTP boundary; it does not own CareerFit business state"],
    ])
    set_table(document.tables[4], [
        ["Group", "Required capabilities", "Main actor"],
        ["Authentication and account", "Registration, email/password login, current-account lookup, role enforcement, and settings", "Candidate, Recruiter, Administrator"],
        ["Public Job discovery", "Server-side search, filters, urgent state, pagination, Job/employer details, similar Jobs, and public analytics", "Guest and authenticated users"],
        ["Candidate profile, CV, and portfolio", "Profile/portfolio editing; CV draft, upload, OCR, review/edit/confirm, default/delete, quality signals, and match results", "Candidate"],
        ["Recommendation and analytics", "Personalized Jobs, interaction handling, match trends, skill demand, profile gaps, and scoped Recruiter/market analytics", "Candidate and Recruiter"],
        ["Job and employer management", "Company onboarding/profile; JD quality; Job draft, publish, edit, status, urgent/deadline, delete, export, and counts", "Recruiter"],
        ["Application and Talent Pool", "Apply, history, withdrawal, invitation response, applicant decisions, Potential discovery, bookmarks, invitations, and withdrawal", "Candidate and Recruiter"],
        ["Feedback and automation", "Rocchio feedback/recomputation; account policy, threshold, quota, cooldown, quiet hours, pause/resume, auto-apply, and reminders", "Candidate and Scheduler"],
        ["Content reporting", "Candidate reports active Job; Recruiter reports visible CV with owned Job context; prevent duplicate pending reports", "Candidate and Recruiter"],
        ["Administration", "User suspend/activate, Job hide/restore, report queue/detail, dismiss/ban, email monitoring, maintenance, and audit logs", "Administrator"],
        ["Actionable email", "Match feedback, Daily Digest feedback/unsubscribe, and Recruiter invitation accept/decline through confirm-then-POST tokens", "Email recipient"],
    ])

    replace_start(
        document,
        "This chapter defined CareerFit's actors",
        "This chapter defined CareerFit's actors, functional and non-functional requirements, and fourteen representative use cases. The use cases follow real role-based goals across the frontend, actionable email flow, and protected services. They keep matching evidence, application state, user policy, report moderation, automated actions, and audit records separate so later design and implementation decisions can be traced to expected behavior.",
    )
    replace_start(
        document,
        "The core traceability chain is:",
        "The core traceability chain is: UC-01 to CV/profile/portfolio and ingestion tests; UC-02 to public catalogue, detail, employer, and similar-Job contracts; UC-03 to Application and invitation lifecycle tests; UC-04 to feedback and AlgorithmEvaluatorTest; UC-05 to recommendation and analytics contracts; UC-06 to AutomationPolicy/AutoApply tests; UC-07 to ContentReportServiceTest and Job/CV reporting browser contracts; UC-08 to company and Job lifecycle tests; UC-09 and UC-10 to applicants, Talent Pool, bookmark, and invitation flows; UC-11 to analytics and settings contracts; UC-12 and UC-13 to administrative account/Job/report moderation tests; UC-14 to hashed token, confirmation, redemption, replay, and expiry behavior.",
    )

    replace_picture(
        document,
        "Figure 1.4.",
        FIGURE,
        "CareerFit role-oriented overview mapping fourteen use cases to Guest/Candidate, Recruiter, Administrator, and Email Recipient actors.",
    )

    document.save(REPORT)
    print(REPORT)
    print(FIGURE)


if __name__ == "__main__":
    main()
