from __future__ import annotations

from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
BACKUP = ROOT / "Doc" / "working" / "CareerFit-Thesis-Report-before-20260807-project-sync-and-table-font.docx"


def set_run_font(run, size: float = 13) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "Times New Roman")


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = text
    for run in paragraph.runs:
        set_run_font(run)


def replace_start(document, prefix: str, text: str) -> None:
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting {prefix!r}, found {len(matches)}")
    set_paragraph_text(matches[0], text)


def normalize_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.right_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    set_run_font(run, 13)


def set_table(table, rows: list[list[str]]) -> None:
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    if any(len(row) != len(table.columns) for row in rows):
        raise RuntimeError("Table shape mismatch")
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            normalize_cell(table.cell(row_index, column_index), value, bold=(row_index == 0))


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    for old in tr_pr.findall(qn("w:tblHeader")):
        tr_pr.remove(old)
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def format_all_table_content(document) -> None:
    for table_index, table in enumerate(document.tables):
        if table_index > 0 and table.rows:
            set_repeat_header(table.rows[0])
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        set_run_font(run, 13)


def main() -> None:
    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(REPORT, BACKUP)

    document = Document(REPORT)

    replacements = {
        "Methods: Recruitment text": "Methods: Recruitment text is normalized and represented with Term Frequency-Inverse Document Frequency vectors. Cosine similarity ranks CV-job pairs and candidate-job recommendations, while Rocchio feedback updates learned representations from explicit feedback. A separate skill-transfer model explains Potential candidates without changing the direct cosine score. AutoFit checks consent and policy limits before an action. The platform also has a report moderation flow for suspicious Jobs and CVs. CareerFit uses a Spring Boot backend, a React and TypeScript frontend, and PostgreSQL with Flyway migrations.",
        "Results: On August 3, 2026": "Results: On August 7, 2026, the refreshed backend suite passed 141 of 141 tests across 35 suites. TypeScript checking, ESLint, the production frontend build, and the bundle check also passed. All 46 Chrome workflow, contract, and resilience tests passed against the integrated application, including Job/CV reporting and administrator moderation cases. In the controlled synthetic benchmark, nDCG@5 increased from 0.037737 to 0.837737 after Rocchio feedback. These results support a tested academic prototype, but they do not prove production readiness or real-world hiring effectiveness.",
        "• Design a role-based web platform": "• Design a role-based web platform for guests, candidates, recruiters, and administrators, with protected workflows and a content-reporting path for suspicious Jobs and CVs.",
        "This thesis focuses on an academic prototype": "This thesis focuses on an academic prototype for IT recruitment. Its functional scope includes a public job portal; authenticated Candidate, Recruiter, and Administrator workspaces; reviewed CV and profile management; recruiter company profiles; Job drafts and publishing; CV–JD matching; personalized recommendation; applications; Talent Pool bookmarks and invitations; explicit feedback; AutoFit policies; notifications; analytics; content reporting and administrator moderation; and audit records.",
        "• An integrated recruitment platform": "• An integrated recruitment platform that combines public job discovery, reviewed CV processing, CV–JD matching, personalized recommendation, application management, Recruiter Jobs, Talent Pool, and a report-moderation workflow.",
        "The implementation exposes these groups": "The implementation exposes these groups through REST-oriented endpoints. Examples include /api/jobs/search, /api/cv/upload, /api/cv/{cvId}/review, /api/matches/me/cards, /api/recommendations/jobs, /api/applications, /api/recruiter/jobs/{jobId}/applicants, /api/automation/policy, /api/reports/jobs/{jobId}, /api/reports/cvs/{cvId}, /api/admin/reports, and /api/admin/audit-logs. Chapter 3 explains the main behavior.",
        "This chapter defined CareerFit's actors": "This chapter defined CareerFit's actors, functional and non-functional requirements, and six representative use cases. The requirements separate matching evidence, application state, user policy, reported-content state, automated actions, and audit records so later design and implementation decisions can be traced to expected behavior.",
        "user_account stores identity": "user_account stores identity, role, activation, verification, language, and password hash. A Candidate owns a profile, multiple CVs, reviewable CV sections, extracted skills, portfolio links, and portfolio projects. A Recruiter owns an employer profile and Jobs. A Job stores its JD, structured recruitment fields, lifecycle status, report count, urgent flag, deadline, application count, source details, and vector data. A CV also stores its lifecycle status and pending report count.",
        "automation_policy stores one policy": "automation_policy stores one policy per account, including enable/pause state, thresholds, quotas, cooldowns, quiet hours, category guards, and notification preferences. email_action stores expiring one-time actions using a token hash. content_report stores the reporter, JOB or CV target, reason, comment, pending/resolved state, resolving administrator, note, and timestamps. Notification delivery records support policy decisions, while audit_log records important user, moderation, and automated actions.",
        "Flyway migrations are the authoritative": "Flyway migrations are the authoritative schema history. Hibernate uses ddl-auto=validate, so the application checks entity-to-schema compatibility instead of silently generating a different database. Indexes support active Jobs, ownership, matching, applications, feedback, tokens, audit chronology, and the pending report queue. A partial unique index prevents one reporter from creating duplicate pending reports for the same target.",
        "CV processing represents long-running": "CV processing represents long-running work with statuses rather than keeping a browser request open until every score completes. Extraction or validation failure records a failure state and reason. Matching tasks start after commit so background work cannot read uncommitted rows. Report moderation uses transactions and row locking to update the target, report counter, report resolutions, and audit state together. A banned Job is excluded from active flows, while a banned CV can no longer remain the default CV.",
        "The backend package root is": "The backend package root is com.careerfit.backend. Domain packages include auth, candidate, cv, skill, job, employer, matching, recommendation, application, feedback, automation, notification, analytics, report, admin, audit, and settings. Shared responses, exceptions, text utilities, validation, security filters, and configuration are placed in common or configuration packages.",
        "SecurityConfig uses a stateless": "SecurityConfig uses a stateless session policy. Public access is limited to selected authentication operations, public Job/employer/analytics GET routes, similar-job retrieval, health/metrics/OpenAPI routes, and email-action redemption. Candidate and Recruiter report endpoints require authentication and are checked again by role and ownership in ContentReportService. Administrative report queue, detail, ban, and dismiss endpoints require the Administrator role.",
        "A Candidate views owned applications": "A Candidate views owned applications through a paginated API and can filter them by status. A Recruiter can list applicants only for an owned Job, update valid states, invite a Candidate, and report a visible CV in the context of an owned Job. The Talent Pool adds Job-based Candidate discovery, a Potential-only view, and CV bookmarks. A Candidate can report an active Job from the Job interface.",
        "Audit records are written directly": "Audit records are written directly through AuditLogRepository in major services. Entries can contain actor type/ID, action, target, result, source channel, and JSON metadata. Examples include CV failure, application submission, invitation, feedback, auto-apply, JOB_REPORTED, CV_REPORTED, report dismissal, and banning a Job or CV. Direct repository calls are simple, but a centralized audit facade could better standardize metadata, redaction, and failure behavior.",
        "JPA entities map the domain tables": "JPA entities map the domain tables, and Spring Data repositories provide pagination, sorting, locking, and explicit update operations. Flyway migrations V1–V25 create and update the current schema. V25 adds content_report, report queue/target indexes, pending report counters, a unique pending-report rule, and the BANNED status for Job and CV. Hibernate uses ddl-auto=validate.",
        "Database constraints are treated": "Database constraints are treated as correctness boundaries rather than only documentation. Unique indexes protect email identity, default CV, Matching, Application, imported Job hash, and duplicate pending reports by the same reporter and target. Check constraints restrict roles, labels, lifecycle states, report target/reason/status, source channels, and salary modes. Version columns and explicit row locks help protect mutable core entities.",
        "App.tsx defines public": "App.tsx defines public, Candidate, Recruiter, and Administrator routes. Protected routes check the loaded account role, while session restoration calls /api/auth/me and clears invalid state. Candidate pages include Job reporting. Recruiter applicant and Talent Pool views can report a visible CV with the owned Job ID. The Administrator workspace includes a Report moderation tab for the pending queue, case detail, ban, and dismiss actions.",
        "API-driven pages do not replace": "API-driven pages do not replace failed responses with mock Job data. The current frontend uses server-side filters, pagination, sorting, CV review, recruiter onboarding, Job draft/publish, urgent and deadline controls, Talent Pool actions, AutoFit settings, and analytics. Reporting modals use fixed reasons and an optional comment; badges show pending report state, and the admin moderation tab groups pending reports by target.",
        "This chapter explained the current implementation": "This chapter explained the current implementation in Spring Boot, React, PostgreSQL, and Flyway. It covered JWT security, CV review, direct matching, the Potential assessment, feedback learning, applications, Talent Pool, AutoFit, email actions, Job/CV reporting, administrator moderation, frontend integration, and monitoring. Chapter 4 evaluates the refreshed system.",
        "Backend, algorithm, frontend build": "Backend, algorithm, frontend build, and integrated Chrome results were refreshed on August 7, 2026 (ICT). The observed Git HEAD was 242e13a8f7d16fc9ebcab9780264c2c2b2b4ef06, but the worktree contained 209 modified or untracked entries. The results therefore describe this working tree and cannot be reproduced from the commit alone unless the full source state is preserved.",
        "Earlier commands are recorded": "Earlier commands are recorded in evidence/CHAPTER5_EVIDENCE_20260703.md. For the August 7 refresh, Surefire XML, Maven output, TypeScript/ESLint/Vite results, Playwright output, the current report, and evaluation/result.json are the main evidence. The report states the working-tree limitation instead of treating the commit hash as a complete release identifier.",
        "The complete backend suite was executed": "The complete backend suite was executed with .\\mvnw.cmd clean verify. It compiled 148 application source files and 37 test source files, started PostgreSQL 16 Testcontainers, applied Flyway migrations V1–V25, executed 35 Surefire suites, and built the backend JAR.",
        "The suites covered application context": "The suites covered application context, API contracts, security, CV drafts and review, OCR/extraction, skill rules, Job lifecycle, catalogues, matching, feedback, automation, deadlines, bookmarks, invitations, analytics, settings, content reporting and moderation, and the controlled AlgorithmEvaluatorTest.",
        "All 131 registered JUnit tests": "All 141 registered JUnit tests passed on August 7, with no failures, errors, or skips. The aggregated Surefire suite time was 102.453 seconds. Negative tests still produced expected handled errors in their own cases, but the Maven build completed successfully and no final test report contained a failure.",
        "The August 3 full-suite run reproduced": "The August 7 full-suite run reproduced dataset hash 6e935639ba6d3290dca8ad91a35d714c5e30c7e69a59af23ddbcf89fcc5cc2f2. It produced baseline nDCG@5 of 0.037737056145, Rocchio nDCG@5 of 0.837737056145, and a delta of 0.80. These values describe the controlled dataset and should not be read as expected improvement on real recruitment data.",
        "The frontend verification ran": "The frontend verification ran TypeScript checking with no output, ESLint, Vite 6.4.3 production build, and the bundle check. The August 7 build completed successfully and transformed 2,362 modules.",
        "Manual Rollup chunking": "Manual Rollup chunking separated React, query, chart, and icon dependencies. The largest generated JavaScript chunk was charts at 387.39 kB (112.63 kB gzip), below Vite's 500 kB warning threshold. This is build evidence only; no browser performance profile was collected.",
        "The Playwright suite ran": "The Playwright suite ran with the Vite frontend, current backend, PostgreSQL, and installed desktop Chrome. All 46 tests across workflow, backend-contract, account-state, catalogue, and resilience coverage passed, including the three report-moderation contracts.",
        "The result covers Guest search": "The result covers Guest search and details; email/password login; Candidate CV, matching, applications, settings, AutoFit, and active-Job reporting; Recruiter onboarding, Jobs, applicants, Talent Pool, bookmarks, invitations, and visible-CV reporting; Administrator account operations and report moderation; session restoration; catalogues; retryable errors; dark mode; and role navigation. It is not an independent UAT study, and Firefox/WebKit were not run in this refresh.",
        "The August 3 evaluation passed": "The August 7 evaluation passed all 141 backend tests, completed TypeScript, ESLint, production-build, and bundle checks, and passed all 46 integrated Chrome tests. The controlled Rocchio benchmark kept its expected synthetic improvement. These results support a functioning academic prototype, but not production readiness or proven hiring effectiveness.",
        "CareerFit IT AutoPilot was implemented": "CareerFit IT AutoPilot was implemented as an IT recruitment prototype with public Job discovery, role workspaces, reviewed CV processing, direct matching, a separate Potential assessment, recommendations, Rocchio feedback, applications, Talent Pool features, AutoFit policies, email actions, content reporting, administrator moderation, analytics, and audit monitoring.",
        "The refreshed evaluation provides evidence": "The refreshed evaluation provides evidence for several technical outcomes. The backend passed 141 tests across 35 suites. TypeScript checking, ESLint, Vite production compilation, and the bundle check passed. All 46 integrated Chrome tests passed across role workflows, backend contracts, account state, catalogues, and resilience, including Candidate Job reporting, Recruiter CV reporting, and Administrator banning from the moderation tab.",
        "The August 3 refresh verified": "The August 7 refresh synchronized the new content-reporting feature, Flyway V25, BANNED Job/CV states, administrator moderation, and the related browser contracts. The largest JavaScript chunk remained below 500 kB. The result is suitable for a local thesis demonstration, not a production release.",
        "The 131 backend tests": "The 141 backend tests do not prove complete path coverage. Browser evaluation covers 46 automated cases in desktop Chrome only, and no independent users participated. The new moderation flow has functional and contract tests, but it has not been tested with real abuse volume, moderator teams, appeals, or fairness review. Concurrency, capacity, cross-browser behavior, accessibility with users, and real hiring outcomes remain unevaluated.",
        "The current implementation improves earlier": "The current implementation includes after-commit matching, CV review before scoring, versioned Potential rules, server-side catalogues, recruiter company onboarding, Job drafts, bookmarks, policy guards, and content-report moderation. Audit records are still written directly by several services, and the evaluated worktree is not a fixed release commit. Moderation currently supports ban or dismiss but not an appeal or restoration workflow.",
        "Future security work should add": "Future security work should add rate limiting, stronger session handling, protected management endpoints, malware scanning, encrypted CV storage, retention rules, and tested backup recovery. Content moderation should add an appeal/restoration workflow, moderator assignment, evidence retention, abuse-rate limits, and tests with concurrent reports. Email delivery should be tested with a real provider, including expiry, replay, scanner behavior, and failure recovery.",
        "CareerFit demonstrates controlled": "CareerFit demonstrates controlled IT recruitment automation by connecting job discovery, explainable matching, profile-based recommendations, feedback learning, policy-driven actions, content reporting, administrator moderation, and audit records. It keeps scores, business state, reported-content state, and user decisions visible.",
        "Local evaluation passed 131": "Local evaluation passed 141 backend tests, the frontend type/lint/build/bundle checks, 46 integrated Chrome tests, health checks, and the controlled Rocchio benchmark. This supports a tested Human-in-the-Loop prototype, not a system ready to make real hiring decisions.",
        "The core traceability chain is": "The core traceability chain is: authentication and role controls to SecurityHardeningTest and P0 login flows; CV ingestion to service and integration tests; matching and feedback to AlgorithmEvaluatorTest; Job lifecycle to Recruiter flows; content reporting to ContentReportServiceTest and the three report-moderation browser contracts; administration to moderation and account-state workflows; and operations to Actuator health, build, and runtime evidence.",
        "Representative endpoints are POST": "Representative endpoints are POST /api/auth/login; GET /api/jobs/search; POST /api/cv/upload; GET/PATCH /api/cv/{cvId}/review; POST /api/matches/{matchingId}/feedback; POST /api/applications; POST /api/reports/jobs/{jobId}; POST /api/reports/cvs/{cvId}; GET /api/admin/reports; POST /api/admin/reports/{type}/{targetId}/ban or /dismiss; PATCH /api/automation/policy; and GET /actuator/health. Protected routes require a signed JWT and backend role/ownership checks.",
        "Identity data is centered": "Identity data is centered on UserAccount and role profiles. Recruitment data includes Job, CV with review and report fields, Matching, Application, RecruiterCvBookmark, Invitation state, Feedback, and Skill. Automation and operations use AutomationPolicy, EmailAction, Notification, DeliveryLog, ContentReport, AuditLog, and scheduler state. Flyway V1–V25 records schema changes, and one-time email action tokens are stored as SHA-256 hashes.",
        "The final evidence package contains": "The final evidence package contains the 141-test backend reports, controlled algorithm benchmark, frontend type/lint/build results, 46-test Chrome output, runtime health evidence, the updated report, and evaluation/result.json. Chapter 4 reports what this evidence supports and what it does not support.",
        "2. As a Candidate, sign in": "2. As a Candidate, sign in, upload or draft a CV, review and confirm it, inspect matches, apply or withdraw, submit feedback, and report one designated active test Job with a clear reason.",
        "3. As a Recruiter, complete": "3. As a Recruiter, complete the company profile, publish a test Job, inspect applicants and Talent Pool, bookmark and invite a Candidate, update application status, and report one visible test CV using the owned Job context.",
        "5. As an Administrator": "5. As an Administrator, inspect users, audit records, and the Report moderation tab; open the designated test report case; dismiss it or ban only disposable test content; then verify status, counters, and audit evidence.",
        "Sign in, open My CVs": "Sign in, open My CVs, upload a PDF/DOCX/image or save a manual draft. Review and confirm the sections before scoring. Then choose the default CV, inspect match and Potential reasons, manage applications, provide feedback, configure AutoFit, and use Report on an active Job only when there is a real issue.",
        "Sign in to the Recruiter workspace": "Sign in to the Recruiter workspace, complete the company profile, and create or publish Jobs. Use applicants and Talent Pool to review visible CVs, bookmark or invite a Candidate, and update valid application states. A visible CV can be reported only with an owned Job that gives the Recruiter access to it.",
        "Use the administrative workspace": "Use the administrative workspace to inspect users, Jobs, audit records, notifications, email actions, and Report moderation. Review the report reasons and target details before choosing Dismiss or Ban. Banning is a state-changing action and should be tested only on disposable records.",
    }

    for prefix, text in replacements.items():
        replace_start(document, prefix, text)

    set_table(document.tables[3], [
        ["Actor", "Primary responsibilities", "Access boundary"],
        ["Guest", "Browse public Jobs, details, employers, urgent items, and public market analytics", "Public GET operations only; authentication is required for personalized data or applications"],
        ["Candidate", "Maintain profile/CVs; view matches; apply; give feedback; configure AutoFit; report an active Job", "Candidate-owned resources and Candidate/report APIs"],
        ["Recruiter", "Maintain company/Jobs; review applicants and Talent Pool; invite Candidates; report a visible CV", "Recruiter APIs, owned Jobs, and CVs visible through an owned Job"],
        ["Administrator", "Monitor users, Jobs, audit records, and report cases; dismiss reports or ban reported Jobs/CVs", "Administrative APIs only; moderation actions are audited"],
        ["Background scheduler", "Recompute Matchings, send reminders, clean actions, and run eligible automation", "Internal service calls under system control"],
        ["Mail provider", "Deliver notification and signed action emails generated by the backend", "SMTP boundary; it does not own CareerFit business state"],
    ])

    set_table(document.tables[4], [
        ["Group", "Required capabilities", "Main actor"],
        ["Authentication and account", "Registration, email/password login, current-account lookup, role enforcement, and settings", "Candidate, Recruiter, Administrator"],
        ["Public Job portal", "Server-side search, filters, urgent state, pagination, Job details, employers, and public analytics", "Guest and authenticated users"],
        ["Candidate profile and CV", "Profile/portfolio, CV drafts, upload, OCR cleanup, review/edit/confirm, default CV, and quality signals", "Candidate"],
        ["Job and employer management", "Company profile; JD quality; Job draft, publish, update, urgent/deadline fields, status, export, and counts", "Recruiter"],
        ["Matching and recommendation", "Direct CV–JD score, reasons, Potential assessment, Candidate cards, ranking, recommendations, and similar Jobs", "Candidate and Recruiter"],
        ["Application and Talent Pool", "Apply, history, withdrawal, applicant review, Potential discovery, bookmark, invitation, and status update", "Candidate and Recruiter"],
        ["Feedback and automation", "Rocchio feedback, recomputation, account policy, quota, cooldown, quiet hours, auto-apply, and reminders", "Candidate and Scheduler"],
        ["Content reporting", "Candidate reports active Job; Recruiter reports visible CV with owned Job context; prevent duplicate pending reports", "Candidate and Recruiter"],
        ["Administration", "Report queue/detail, dismiss or ban Job/CV, user/Job monitoring, notifications, email actions, and audit logs", "Administrator"],
    ])

    set_table(document.tables[13], [
        ["Module", "Main responsibility", "Representative services"],
        ["Auth and Security", "Password login, JWT validation, account resolution, role and route rules", "AuthService, JwtService, security filters"],
        ["Candidate, CV, and Skill", "Profile, CV draft/review, storage, OCR, validation, and skill suggestions", "CandidateProfileService, CvIngestionService, PdfExtractionService, SkillService"],
        ["Job and Employer", "Public catalogue, company onboarding, Job quality, draft/publish, urgent/deadline, and lifecycle", "JobService, EmployerService, QualityValidationService"],
        ["Matching and Recommendation", "TF-IDF/cosine, reasons, Potential, batch matching, filters, ranking, and Job recommendation", "ScoringService, SkillTransferService, MatchingService, RecommendationService"],
        ["Application and Talent", "Apply/withdraw, applicants, invitations, status changes, Talent Pool, and bookmarks", "ApplicationService, RecruiterTalentService"],
        ["Feedback and Automation", "Rocchio update, recomputation, account policy, auto-apply, reminders, and scheduler", "FeedbackService, RocchioService, AutoApplyService, AutomationScheduler"],
        ["Notification", "Policy guard, delivery log, SMTP/no-op mail, and signed email actions", "NotificationPolicyGuard, NotificationEmailService, EmailActionService"],
        ["Content Report", "Job/CV reports, access checks, queue grouping, ban/dismiss, counters, and audit", "ContentReportService, ContentReportController, AdminReportController"],
        ["Analytics, Admin, and Audit", "Live analytics, account/Job administration, monitoring, and audit persistence", "Analytics services, administrative services, audit/common components"],
    ])

    set_table(document.tables[14], [
        ["Constraint", "Purpose"],
        ["Unique normalized user email", "Prevent duplicate identities that differ only by case"],
        ["One role profile and one policy per account", "Preserve ownership and account-level settings"],
        ["Partial unique default CV per Candidate", "Ensure at most one active default CV"],
        ["Unique Matching per CV and Job", "Prevent duplicate direct scores for the same pair"],
        ["Unique Application per Candidate and Job", "Prevent duplicate applications or invitations"],
        ["Unique bookmark and feedback keys", "Keep one bookmark/judgment for each defined actor-target relation"],
        ["Unique pending report per reporter and target", "Prevent duplicate unresolved abuse reports from the same account"],
        ["Check constraints on roles, states, labels, report values, salary, and channels", "Reject invalid domain values at the database boundary"],
        ["Indexes for Jobs, ownership, ranking, audit, and report queue", "Support catalogues, reminders, matching, and moderation queries"],
    ])

    set_table(document.tables[15], [
        ["Risk", "Current control", "Remaining treatment"],
        ["Unauthorized cross-account access", "URL roles and service ownership/visibility checks", "Keep negative integration tests for identifier-based operations"],
        ["Duplicate applications, Matchings, or reports", "Service checks plus unique constraints/indexes", "Convert database conflicts into stable API errors"],
        ["False or abusive content reports", "Fixed reasons, one pending report per reporter-target, and Administrator review", "Add rate limits, evidence rules, moderator assignment, and abuse monitoring"],
        ["Incorrect moderation action", "Case detail, explicit Ban/Dismiss, transactions, and audit events", "Add appeal, restoration, and dual-review policy for production"],
        ["Replayed email action", "Hashed token, pending/redeemed state, expiry, and POST execution", "Retain replay, expiry, scanner, and rate-limit tests"],
        ["Scheduler/configuration drift", "Scheduler timing reads application properties", "Maintain configuration and schedule-boundary tests"],
        ["Sensitive data in logs", "Structured audit metadata and DTO boundaries", "Review retention, redaction, access, and exported evidence"],
        ["Lexical scoring bias or omission", "Reasons, validation, feedback, and HITL", "Evaluate representative data; never treat score as hiring probability"],
        ["Local file-storage loss or exposure", "Configured path and Docker volume", "Define backup, encryption, malware scanning, and retention"],
    ])

    set_table(document.tables[17], [
        ["State", "Meaning", "Typical transition"],
        ["DRAFT", "Manual content is saved but not submitted for scoring", "Draft edit → review/confirm"],
        ["UPLOADED", "Uploaded metadata and source are accepted", "Upload → validation/extraction"],
        ["VALIDATING", "File, content, and extraction checks are running", "Valid source → review; hard error → failed"],
        ["REVIEW_REQUIRED", "Extracted/manual sections are ready for Candidate review", "Edit sections → confirm"],
        ["PROCESSING", "Confirmed content is normalized and vectorized", "Vector saved → asynchronous scoring"],
        ["SCORING_DONE", "The confirmed CV vector is available for matching", "Use cards, feedback, applications, and AutoFit"],
        ["FAILED", "Processing cannot continue", "Failure reason is stored and audited"],
        ["BANNED", "An Administrator acted on a report case", "Removed as default and excluded from normal use"],
    ])

    set_table(document.tables[20], [
        ["API operation", "Controller/service responsibility", "Persistence effect"],
        ["POST /api/cv/upload; GET/PATCH /api/cv/{id}/review", "Extract content and expose/edit review sections", "CV, review data, quality signals, audit"],
        ["POST /api/cv/{id}/review/confirm", "Confirm owned CV and start matching", "CV vector, Matchings, audit"],
        ["GET /api/matches/me/cards; POST /api/matches/{id}/feedback", "Query Candidate cards and submit feedback", "Feedback, audit, learned Job vector, stale flags"],
        ["POST /api/jobs/drafts; POST /api/jobs/{id}/publish", "Validate company/JD/deadline and manage owned Job", "Job, vectors, Matchings, audit"],
        ["PUT /api/recruiter/talent/.../bookmark", "Validate Job ownership and save Talent Pool choice", "Recruiter CV bookmark"],
        ["PATCH /api/automation/policy", "Update the account's automation policy", "Automation policy and version"],
        ["POST /api/reports/jobs/{id}; POST /api/reports/cvs/{id}", "Validate role, target state, Job ownership/CV visibility, reason, and duplicate pending report", "ContentReport, pending report count, audit"],
        ["GET /api/admin/reports; POST .../{type}/{id}/ban or /dismiss", "Load report cases and resolve all pending reports for a target", "Target status/counter, report resolutions, audit"],
    ])

    set_table(document.tables[21], [
        ["Area", "Current limitation", "Consequence or required improvement"],
        ["Content moderation", "Ban/Dismiss exists, but appeal and restoration do not", "Add moderator workflow, appeal, evidence retention, and abuse-rate controls"],
        ["Email actions", "Hashed token and confirm-then-POST flow", "Add rate limiting, delivery monitoring, and deployment origin controls"],
        ["CV/OCR", "Preprocessing and cleanup are heuristic", "Scanned layouts and low-quality images still need user correction"],
        ["Text processing", "Simple tokenization remains the direct-score baseline", "Technology spelling and Vietnamese phrases may lose information"],
        ["Potential model", "Aliases and transfer weights are manually defined", "Validate rules with recruiter-labeled data and review drift"],
        ["Automation", "Many policy guards and schedules interact", "Add time-zone, quota, cooldown, deadline, and concurrency tests"],
        ["Frontend session", "JWT and account summary use sessionStorage", "Evaluate refresh/cookie architecture and XSS controls"],
        ["File storage", "Local path or Docker volume", "Add malware scanning, encryption, backup, retention, and object storage"],
        ["Audit", "Several services write audit rows directly", "Centralize redaction, request context, schema rules, and integrity policy"],
    ])

    set_table(document.tables[22], [
        ["Component", "Observed version or configuration"],
        ["Host operating environment", "Windows, PowerShell"],
        ["Java", "21.0.1"],
        ["Backend framework", "Spring Boot 3.2.5"],
        ["Maven", "Repository Maven Wrapper; clean verify"],
        ["Node.js / npm", "Bundled Node runtime with repository npm scripts"],
        ["Frontend build", "Vite 6.4.3, React 18.3, TypeScript 5.9, ESLint"],
        ["Docker / Compose", "Docker Desktop with Compose"],
        ["Test database", "PostgreSQL 16 through Testcontainers; local Compose PostgreSQL for E2E"],
        ["Database migrations", "Flyway V1–V25"],
        ["Browser E2E", "Playwright 1.61 using installed desktop Chrome"],
        ["Integrated backend", "Compose backend on port 8080"],
        ["Integrated frontend", "Vite development server on 127.0.0.1:5173"],
    ])

    set_table(document.tables[23], [
        ["Measure", "Result"],
        ["Test suites", "35"],
        ["Tests", "141"],
        ["Failures", "0"],
        ["Errors", "0"],
        ["Skipped", "0"],
        ["Aggregated Surefire suite time", "102.453 s"],
        ["Maven lifecycle", "clean verify and JAR build passed"],
    ])

    set_table(document.tables[25], [
        ["Artifact", "Uncompressed size", "Gzip size"],
        ["index.html", "1.37 kB", "0.67 kB"],
        ["Main CSS", "137.80 kB", "25.57 kB"],
        ["Icons chunk", "22.10 kB", "5.03 kB"],
        ["Query chunk", "39.78 kB", "12.20 kB"],
        ["React chunk", "181.35 kB", "59.62 kB"],
        ["Application chunk", "326.95 kB", "88.93 kB"],
        ["Charts chunk", "387.39 kB", "112.63 kB"],
    ])

    set_table(document.tables[26], [
        ["Scenario group", "Main verification", "Result"],
        ["Guest and authentication", "Public catalogue/detail, login, session restoration, and role routes", "Passed"],
        ["Candidate CV and matching", "Draft, upload/review/confirm, matching, filters, and feedback", "Passed"],
        ["Candidate recruitment", "Urgent Jobs, apply/withdraw, settings, AutoFit, and active-Job report", "Passed"],
        ["Recruiter workflow", "Company, Job draft/publish, applicants, Talent Pool, bookmark, invitation, and visible-CV report", "Passed"],
        ["Administration", "Suspend/reactivate, audit, report queue/detail, and ban action", "Passed"],
        ["Resilience and interface", "API errors, navigation, account state, status labels, and dark mode", "Passed"],
    ])

    set_table(document.tables[29], [
        ["Question", "Evidence-supported answer"],
        ["Does Rocchio change holdout ranking in the intended direction?", "Yes for the controlled synthetic dataset; the same large metric delta was reproduced."],
        ["Is the backend automated suite passing?", "Yes. All 141 tests across 35 suites passed, and clean verify built the JAR."],
        ["Do the main browser workflows execute?", "Yes. All 46 integrated Chrome tests passed; Firefox/WebKit and participant UAT remain incomplete."],
        ["Does report moderation work in the tested contracts?", "Yes for Candidate Job report, Recruiter visible-CV report, and Administrator Job ban; production abuse handling is not evaluated."],
        ["Are selected authorization boundaries enforced?", "Observed public, authentication, role, ownership, and report-visibility checks passed; this is not comprehensive security validation."],
        ["Is the runtime healthy?", "Yes in the evaluated local profile: health groups and the core API returned HTTP 200."],
        ["Is performance production-ready?", "Not evaluated. Only a small sequential localhost latency sample was collected."],
    ])

    set_table(document.tables[30], [
        ["Evaluation area", "Supported result", "Qualification"],
        ["Backend automated tests", "141/141 tests across 35 suites passed", "clean verify and JAR build completed"],
        ["Controlled feedback benchmark", "Large, deterministic Rocchio improvement", "Synthetic design; not production effectiveness"],
        ["Frontend checks", "TypeScript, ESLint, Vite build, and bundle check passed", "Largest JS chunk 387.39 kB; no Vite size warning"],
        ["Browser workflows", "46/46 integrated Chrome tests passed", "No Firefox/WebKit or independent participant UAT"],
        ["Report moderation", "Job/CV reporting and Administrator Job ban contracts passed", "No real moderation load, appeal, or restoration study"],
        ["Authorization spot checks", "Observed 200/401/403 outcomes matched expectations", "Not a complete security assessment"],
        ["Runtime APIs", "Core Job API and health endpoints returned HTTP 200/UP", "Local profile only; no production monitoring"],
        ["Local latency sample", "Mean 61.79 ms and p95 85.11 ms for 30 sequential Job queries", "No concurrency, controlled load, or production network"],
    ])

    set_table(document.tables[31], [
        ["Objective", "Assessment", "Primary evidence"],
        ["Role-based web platform", "Achieved for prototype scope", "API contracts and 46 integrated Chrome tests"],
        ["Reviewed CV/JD ingestion", "Achieved for supported formats and configured OCR", "Draft/review/confirm, extraction, quality, Job, and integration tests"],
        ["Explainable matching and Potential", "Achieved as rule-based baseline", "TF-IDF/scoring tests and skill-transfer reasons"],
        ["Rocchio feedback adaptation", "Achieved in controlled test", "Synthetic benchmark, after-commit execution, and clean reports"],
        ["AutoFit and communication", "Achieved as configurable prototype", "Policies, scheduler, reminders, notifications, and signed email actions"],
        ["Recruiter operational workflow", "Achieved for prototype scope", "Company, Job, applicants, Talent Pool, bookmark, and invitation tests"],
        ["Content reporting and moderation", "Achieved for the implemented basic flow", "ContentReportServiceTest and three browser contracts"],
        ["Production readiness", "Not demonstrated", "Local health passed; scale, real users, appeals, and production security remain missing"],
    ])

    set_table(document.tables[32], [
        ["Limitation", "Impact on conclusions"],
        ["Synthetic benchmark", "Supports causal behavior only, not real hiring effectiveness"],
        ["Lexical direct score", "Limits semantic and contextual understanding"],
        ["Manual skill-transfer knowledge base", "Potential rules require labeled validation and maintenance"],
        ["Basic moderation workflow only", "Ban/Dismiss is implemented, but appeal, restoration, and real abuse load are not evaluated"],
        ["Production concurrency not evaluated", "Conflict, retry, and recovery behavior remain unproven at scale"],
        ["Local-only health verification", "Does not establish deployed availability"],
        ["Chrome-only 46-test automation suite", "Limits cross-browser and real-user usability conclusions"],
        ["Email delivery and abuse controls not production-tested", "Signed actions and reports exist, but real delivery and rate limits remain unverified"],
        ["Dirty worktree and mutable E2E database", "Weakens exact experiment reproduction"],
    ])

    format_all_table_content(document)
    document.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    main()
