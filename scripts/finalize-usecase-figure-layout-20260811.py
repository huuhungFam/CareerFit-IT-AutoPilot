from pathlib import Path
import hashlib

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "Doc" / "CareerFit-Thesis-Report.docx"
UC12_SOURCE = Path(r"C:\Users\PhamHuuHwng\Downloads\uc12.png")
NOTE_PREFIX = "NOTE: In Figure 1.8, the left actor must be changed from Guest to Recruiter"


def remove_paragraph(paragraph):
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)


def main():
    document = Document(REPORT)

    notes = [p for p in document.paragraphs if p.text.strip().startswith(NOTE_PREFIX)]
    if len(notes) != 1:
        raise RuntimeError(f"Expected one obsolete UC-12 note; found {len(notes)}")
    remove_paragraph(notes[0])

    source_hash = hashlib.sha256(UC12_SOURCE.read_bytes()).hexdigest()
    found = 0
    for paragraph in document.paragraphs:
        blips = paragraph._p.xpath(".//a:blip")
        if not blips:
            continue
        rid = blips[0].get(qn("r:embed"))
        if hashlib.sha256(document.part.related_parts[rid].blob).hexdigest() != source_hash:
            continue
        found += 1
        for prop in paragraph._p.xpath(".//wp:docPr"):
            prop.set("title", "Shared Candidate–Recruiter reporting use-case diagram")
            prop.set("descr", "Shared UC-12 reporting use-case diagram initiated by Recruiter and Candidate.")
    if found != 1:
        raise RuntimeError(f"Expected one UC-12 image; found {found}")

    captions = [p for p in document.paragraphs if p.text.strip() == "Figure 1.2. Overall CareerFit use-case diagram"]
    if len(captions) != 1:
        raise RuntimeError(f"Expected one Figure 1.2 caption; found {len(captions)}")
    captions[0].paragraph_format.space_after = Pt(6)
    captions[0].paragraph_format.keep_with_next = False

    headings = [p for p in document.paragraphs if p.text.strip() == "1.3 Functional Requirements"]
    if len(headings) != 1:
        raise RuntimeError(f"Expected one 1.3 heading; found {len(headings)}")
    headings[0].paragraph_format.page_break_before = True

    document.save(REPORT)
    print("obsolete_uc12_note_removed=true")
    print("uc12_alt_text_corrected=true")
    print("section_1_3_starts_new_page=true")
    print(REPORT)


if __name__ == "__main__":
    main()
